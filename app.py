"""
UrbanLex - Aplicacao Flask Principal v3.5
Parametros Urbanisticos + Biblioteca Legislativa + Monitoramento IA

CORRECOES APLICADAS (25/02/2026):
  BUG 1: DELETE legislacao usava id() em vez de leg_id + foreign keys
  BUG 2: Login API tinha sessao como codigo morto (indentacao)
  BUG 3: Login HTML mostrava erro no GET
  BUG 4: Rotas diagnostico sem autenticacao
  BUG 5: Rotas FIX 6-10 movidas para antes do if __name__
  BUG 6: inicializar() chamada no nivel do modulo
  BUG 7: Query alteracoes usa COALESCE para municipio_nome
  BUG 8: Adicionada rota GET /logout
  BUG 9: Login HTML usa qry() em vez de conexao manual
  BUG 10: Rota diagnostico R2 adicionada
  BUG 11: Rota re-upload de arquivo adicionada
"""

import os
from dotenv import load_dotenv
load_dotenv()
import io, sys, json, hashlib, threading, time
from pathlib import Path
from datetime import datetime, timedelta
from functools import wraps

from flask import Flask, render_template, request, jsonify, redirect, url_for, send_file, session

sys.path.insert(0, str(Path(__file__).parent.parent))

import psycopg2
import psycopg2.extras
from werkzeug.utils import secure_filename

# -- Auth helpers (adapters para compatibilidade) --
try:
    import bcrypt as _bcrypt
    def hash_senha(s): return _bcrypt.hashpw(s.encode(), _bcrypt.gensalt(12)).decode()
    def verificar_senha(s, h):
        try: return _bcrypt.checkpw(s.encode(), h.encode())
        except: return False
except ImportError:
    import hashlib
    def hash_senha(s): return hashlib.sha256(s.encode()).hexdigest()
    def verificar_senha(s, h): return hash_senha(s) == h

import secrets as _secrets
def gerar_token(): return _secrets.token_urlsafe(32)

def _get_email_cfg():
    import smtplib, ssl
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    host  = os.getenv('EMAIL_HOST','smtp.gmail.com')
    port  = int(os.getenv('EMAIL_PORT','587'))
    user  = os.getenv('EMAIL_USER','')
    pwd   = os.getenv('EMAIL_PASS','')
    frm   = os.getenv('EMAIL_FROM', user)
    return host, port, user, pwd, frm

def enviar_email_generico(para, assunto, html):
    try:
        import smtplib
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        host, port, user, pwd, frm = _get_email_cfg()
        if not user: return
        msg = MIMEMultipart('alternative')
        msg['Subject'] = assunto; msg['From'] = frm; msg['To'] = para
        msg.attach(MIMEText(html, 'html'))
        with smtplib.SMTP(host, port) as srv:
            srv.starttls(); srv.login(user, pwd); srv.sendmail(frm, para, msg.as_string())
    except Exception as e:
        print(f"[EMAIL ERROR] {e}")


def get_app_url():
    """Retorna a URL base da aplicacao, com fallback para variavel de ambiente."""
    url = os.getenv('APP_URL', '')
    if not url:
        try:
            from flask import request
            url = request.host_url.rstrip('/')
        except:
            url = 'http://localhost:5000'
    return url.rstrip('/')

def enviar_email_ativacao(user, token):
    url = f"{get_app_url()}/ativar/{token}"
    enviar_email_generico(user['email'], 'Ative sua conta UrbanLex',
        f'<p>Ola {user["nome"]},</p><p>Clique para ativar sua conta:</p>'
        f'<p><a href="{url}">{url}</a></p><p>Link valido por 24 horas.</p>')

def enviar_email_aprovacao_admin(user):
    app_url = get_app_url()
    token_apr = gerar_token()
    exp = datetime.now() + timedelta(days=7)
    qry("INSERT INTO aprovacao_tokens (user_id,token,tipo,expira_em) VALUES (%s,%s,'aprovacao',%s)",
        (user['id'], token_apr, exp), commit=True, fetch=None)
    admin_email = os.getenv('ADMIN_EMAIL','')
    if not admin_email: return
    url_apr = f"{app_url}/admin/aprovar/{token_apr}"
    url_rej = f"{app_url}/admin/rejeitar/{token_apr}"
    enviar_email_generico(admin_email, f'Novo cadastro: {user["nome"]}',
        f'<p>Novo usuario aguardando aprovacao: <strong>{user["nome"]}</strong> ({user["email"]})</p>'
        f'<p><a href="{url_apr}">Aprovar</a> &nbsp; <a href="{url_rej}">Rejeitar</a></p>')

def enviar_email_boas_vindas(user):
    enviar_email_generico(user['email'], 'Bem-vindo ao UrbanLex!',
        f'<p>Ola {user["nome"]}, sua conta foi aprovada! Acesse: {get_app_url()}/login</p>')

def enviar_email_rejeicao(user):
    enviar_email_generico(user['email'], 'Cadastro UrbanLex',
        f'<p>Ola {user["nome"]}, seu cadastro nao foi aprovado. Entre em contato com o administrador.</p>')

def enviar_email_reset(user, token):
    url = f"{get_app_url()}/reset-senha/{token}"
    enviar_email_generico(user['email'], 'Redefinicao de senha - UrbanLex',
        f'<p>Ola {user["nome"]},</p><p>Clique para redefinir sua senha:</p>'
        f'<p><a href="{url}">{url}</a></p><p>Link valido por 1 hora.</p>')

try:
    from modulos.scheduler_integrado import iniciar_scheduler
    SCHEDULER_OK = True
except ImportError:
    SCHEDULER_OK = False

# -- Cloudflare R2 storage (opcional) --
try:
    from modulos.storage_r2 import upload_arquivo as r2_upload, deletar_arquivo as r2_delete, \
                                    gerar_url_assinada as r2_url_assinada, r2_disponivel, \
                                    download_arquivo as r2_download
    _R2_IMPORTADO = True
except ImportError:
    _R2_IMPORTADO = False
    def r2_disponivel(): return False
    def r2_upload(*a, **kw): return None
    def r2_delete(*a, **kw): return False
    def r2_url_assinada(*a, **kw): return None
    def r2_download(*a, **kw): return None

# Conversao de arquivos
try: import PyMuPDF as fitz; PDF_OK = True
except:
    try: import fitz; PDF_OK = True
    except: PDF_OK = False

try: from docx import Document as DocxDoc; DOCX_OK = True
except: DOCX_OK = False

try: import pytesseract; from PIL import Image; OCR_OK = True
except: OCR_OK = False

try: import pandas as pd; PANDAS_OK = True
except: PANDAS_OK = False

# -- App --
app = Flask(__name__, template_folder='templates')
app.secret_key = os.getenv('SECRET_KEY', 'urbanlex-dev-key-change-in-prod')
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB
app.config['SESSION_COOKIE_SECURE'] = False  # HTTP only
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_HTTPONLY'] = True
ADMIN_EMAIL = os.getenv('ADMIN_EMAIL', '')

# -- Debug: screenshots do navegador universal --
@app.route('/api/captcha/saldo')
def api_captcha_saldo():
    import requests as _req, os as _os
    key = _os.getenv('TWOCAPTCHA_API_KEY', '6d9468712a495eae8e97cbfa4c855191')
    try:
        r = _req.get(f'https://2captcha.com/res.php?key={key}&action=getbalance', timeout=8)
        saldo = r.text.strip(); float(saldo)
        return jsonify({'saldo': saldo})
    except Exception as e:
        return jsonify({'saldo': 'erro', 'msg': str(e)})

@app.route('/debug/screenshots')
def debug_screenshots():
    """Lista screenshots de debug do navegador universal."""
    import glob
    debug_dir = '/tmp/nav_screenshots'
    if not os.path.isdir(debug_dir):
        return '<h3>Nenhum screenshot ainda</h3>'
    
    files = sorted(glob.glob(f'{debug_dir}/*.png'), key=os.path.getmtime, reverse=True)
    html = '<h2>Screenshots do Navegador Universal</h2>'
    html += f'<p>{len(files)} arquivo(s)</p>'
    html += '<a href="/debug/screenshots/clear">Limpar todos</a><hr>'
    for f_path in files[:30]:
        fname = os.path.basename(f_path)
        html += f'<div style="margin:10px 0;border:1px solid #ccc;padding:10px;">'
        html += f'<b>{fname}</b><br>'
        html += f'<img src="/debug/screenshots/img/{fname}" style="max-width:100%;border:1px solid red;">'
        html += f'</div>'
    return html

@app.route('/debug/screenshots/img/<path:filename>')
def debug_screenshot_img(filename):
    """Serve um screenshot individual."""
    # Sanitizar: só permitir caracteres seguros
    filename = os.path.basename(filename)
    if not filename.endswith('.png'):
        return 'Invalid filename', 400
    f_path = f'/tmp/nav_screenshots/{filename}'
    if os.path.isfile(f_path):
        return send_file(f_path, mimetype='image/png')
    return 'Not found', 404

@app.route('/debug/screenshots/clear')
def debug_screenshots_clear():
    """Limpa screenshots de debug."""
    import glob
    for f in glob.glob('/tmp/nav_screenshots/*.png'):
        try:
            os.remove(f)
        except Exception:
            pass
    return '<p>Screenshots limpos. <a href="/debug/screenshots">Voltar</a></p>'

# -- DB --
def get_db():
    conn = psycopg2.connect(os.environ['DATABASE_URL'])
    conn.autocommit = False
    return conn

def qry(sql, params=None, fetch='all', commit=False):
    conn = get_db()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(sql, params or ())
        result = None
        if fetch == 'all': result = [dict(r) for r in cur.fetchall()]
        elif fetch == 'one': row = cur.fetchone(); result = dict(row) if row else None
        elif fetch == 'id':
            row = cur.fetchone()
            result = list(row.values())[0] if row else None
        if commit: conn.commit()
        return result
    finally:
        conn.close()

# -- Auth helpers --
def login_required(f):
    @wraps(f)
    def dec(*a, **k):
        if 'user_id' not in session:
            if request.path.startswith('/api/'): return jsonify({'error':'Nao autenticado'}), 401
            return redirect('/login')
        return f(*a, **k)
    return dec

def admin_required(f):
    @wraps(f)
    def dec(*a, **k):
        if 'user_id' not in session:
            if request.path.startswith('/api/'): return jsonify({'error':'Nao autenticado'}), 401
            return redirect('/login')
        if session.get('role') != 'admin':
            if request.path.startswith('/api/'): return jsonify({'error':'Acesso negado'}), 403
            return redirect('/')
        return f(*a, **k)
    return dec

def editor_required(f):
    @wraps(f)
    def dec(*a, **k):
        if 'user_id' not in session:
            if request.path.startswith('/api/'): return jsonify({'error':'Nao autenticado'}), 401
            return redirect('/login')
        if session.get('role') not in ('admin','editor'):
            if request.path.startswith('/api/'): return jsonify({'error':'Acesso negado'}), 403
            return redirect('/')
        return f(*a, **k)
    return dec

def tmpl_ctx():
    return {
        'username': session.get('nome',''),
        'role': session.get('role',''),
        'is_admin': session.get('role') == 'admin'
    }

def extrair_texto_arquivo(arquivo_bytes, nome_arquivo):
    ext = Path(nome_arquivo).suffix.lower()
    texto = ''
    try:
        if ext == '.pdf' and PDF_OK:
            doc = fitz.open(stream=arquivo_bytes, filetype='pdf')
            texto = '\n'.join(p.get_text() for p in doc)
        elif ext in ('.doc','.docx') and DOCX_OK:
            doc = DocxDoc(io.BytesIO(arquivo_bytes))
            texto = '\n'.join(p.text for p in doc.paragraphs)
        elif ext in ('.jpg','.jpeg','.png') and OCR_OK:
            img = Image.open(io.BytesIO(arquivo_bytes))
            texto = pytesseract.image_to_string(img, lang='por')
        elif ext in ('.xls','.xlsx') and PANDAS_OK:
            df = pd.read_excel(io.BytesIO(arquivo_bytes))
            texto = df.to_string()
        elif ext == '.txt':
            texto = arquivo_bytes.decode('utf-8', errors='ignore')
    except Exception as e:
        texto = f'[Erro ao extrair texto: {e}]'
    return texto


# -- Funcao auxiliar de busca com IA --
def _buscar_legislacao_internet(consulta: str) -> dict:
    """Tenta encontrar legislacao via GROQ ou busca simples."""
    try:
        from groq import Groq
        client = Groq(api_key=os.getenv('GROQ_API_KEY',''))
        prompt = (f"Encontre a seguinte legislacao urbanistica brasileira: '{consulta}'. "
                  "Retorne APENAS um JSON com: titulo, estado, municipio, numero, ano, url. "
                  "Se nao encontrar, retorne {{}}.")
        resp = client.chat.completions.create(
            model="llama3-8b-8192",
            messages=[{"role":"user","content":prompt}],
            max_tokens=300
        )
        text = resp.choices[0].message.content.strip()
        import re
        m = re.search(r'\{.*\}', text, re.DOTALL)
        if m:
            result = json.loads(m.group(0))
            if result.get('titulo') or result.get('url'):
                return result
    except Exception as e:
        print(f"[BUSCA IA] {e}")
    return {}


# -------------------------------------------------------------------
# PAGINAS HTML
# -------------------------------------------------------------------

# -- FIX 3: Login corrigido -- erro so aparece no POST com falha --
@app.route('/login', methods=['GET','POST'])
def login_page():
    if 'user_id' in session: return redirect('/')
    error = None
    if request.method == 'POST':
        email = request.form.get('email','').strip().lower()
        senha = request.form.get('senha','')
        user_exists = qry("SELECT * FROM users WHERE email=%s", (email,), 'one')
        user = qry("SELECT * FROM users WHERE email=%s AND ativo=TRUE AND aprovado=TRUE", (email,), 'one')
        if user and verificar_senha(senha, user['senha_hash']):
            session['user_id'] = user['id']
            session['nome'] = user['nome']
            session['email'] = user['email']
            session['role'] = user['role']
            # FIX 9: usar qry() em vez de conexao manual
            qry("UPDATE users SET ultimo_acesso=NOW() WHERE id=%s", (user['id'],), commit=True, fetch=None)
            return redirect('/')
        if not user_exists:
            error = 'Usuário não cadastrado. Faça o cadastro para acessar.'
        elif user_exists and not user_exists.get('aprovado'):
            error = 'Seu cadastro ainda não foi aprovado. Aguarde a liberação pelo administrador.'
        else:
            error = 'E-mail ou senha incorretos.'
    return render_template('login.html', error=error, **tmpl_ctx())

@app.route('/cadastro')
def pagina_cadastro(): return render_template('cadastro.html')

@app.route('/esqueci-senha')
def pagina_esqueci_senha(): return render_template('esqueci_senha.html', **tmpl_ctx())

@app.route('/reset-senha/<token>')
def pagina_reset_senha(token):
    tk = qry("SELECT * FROM password_reset_tokens WHERE token=%s AND usado=FALSE AND expira_em>NOW()", (token,), 'one')
    if not tk:
        return render_template('reset_senha.html', token=token,
            erro='Este link e invalido ou ja expirou. Solicite um novo.', **tmpl_ctx())
    return render_template('reset_senha.html', token=token, erro=None, **tmpl_ctx())

@app.route('/ativar/<token>')
def pagina_ativar(token):
    tk = qry("SELECT * FROM aprovacao_tokens WHERE token=%s AND tipo='ativacao' AND usado=FALSE AND expira_em>NOW()", (token,), 'one')
    if not tk: return render_template('conta_ativada.html', sucesso=False, msg='Link invalido ou expirado.', **tmpl_ctx())
    qry("UPDATE users SET ativo=TRUE WHERE id=%s", (tk['user_id'],), commit=True, fetch=None)
    qry("UPDATE aprovacao_tokens SET usado=TRUE WHERE id=%s", (tk['id'],), commit=True, fetch=None)
    user = qry("SELECT * FROM users WHERE id=%s", (tk['user_id'],), 'one')
    if ADMIN_EMAIL: enviar_email_aprovacao_admin(user)
    return render_template('conta_ativada.html', sucesso=True,
        msg='Conta ativada! Aguardando aprovacao do administrador.',
        nome=user.get('nome',''), email=user.get('email',''), **tmpl_ctx())

@app.route('/admin/aprovar/<token>')
def admin_aprovar(token):
    tk = qry("SELECT * FROM aprovacao_tokens WHERE token=%s AND tipo='aprovacao' AND usado=FALSE", (token,), 'one')
    if not tk: return render_template('resultado_aprovacao.html', sucesso=False, msg='Link invalido.', **tmpl_ctx())
    qry("UPDATE users SET aprovado=TRUE WHERE id=%s", (tk['user_id'],), commit=True, fetch=None)
    qry("UPDATE aprovacao_tokens SET usado=TRUE WHERE id=%s", (tk['id'],), commit=True, fetch=None)
    user = qry("SELECT * FROM users WHERE id=%s", (tk['user_id'],), 'one')
    if not user:
        return render_template('resultado_aprovacao.html', sucesso=False, msg='Usuario nao encontrado.', nome='', email='', **tmpl_ctx())
    enviar_email_boas_vindas(user)
    return render_template('resultado_aprovacao.html', sucesso=True,
        msg=f'Usuario {user["nome"]} aprovado com sucesso!',
        nome=user.get('nome',''), email=user.get('email',''), **tmpl_ctx())

@app.route('/admin/rejeitar/<token>')
def admin_rejeitar(token):
    tk = qry("SELECT * FROM aprovacao_tokens WHERE token=%s AND tipo='aprovacao' AND usado=FALSE", (token,), 'one')
    if not tk: return render_template('resultado_aprovacao.html', sucesso=False, msg='Link invalido.', **tmpl_ctx())
    qry("UPDATE users SET aprovado=FALSE, ativo=FALSE WHERE id=%s", (tk['user_id'],), commit=True, fetch=None)
    qry("UPDATE aprovacao_tokens SET usado=TRUE WHERE id=%s", (tk['id'],), commit=True, fetch=None)
    user = qry("SELECT * FROM users WHERE id=%s", (tk['user_id'],), 'one')
    if user: enviar_email_rejeicao(user)
    nome_rej = user.get('nome','usuario') if user else 'usuario'
    return render_template('resultado_aprovacao.html', sucesso=False,
        msg=f'Cadastro de {nome_rej} rejeitado.',
        nome=user.get('nome','') if user else '', email=user.get('email','') if user else '', **tmpl_ctx())

@app.route('/')
@login_required
def index(): return render_template('dashboard.html', active_page='dashboard', active_group='', **tmpl_ctx())

@app.route('/legislacoes')
@login_required
def pagina_legislacoes(): return render_template('legislacoes.html', active_page='legislacoes', active_group='biblioteca', **tmpl_ctx())

@app.route('/legislacoes/arvores')
@login_required
def pagina_arvores(): return render_template('legislacoes.html', active_page='arvores', active_group='biblioteca', **tmpl_ctx())

@app.route('/legislacoes/pendentes')
@login_required
def pagina_leg_pendentes(): return render_template('legislacoes.html', active_page='leg-pendentes', active_group='biblioteca', **tmpl_ctx())

@app.route('/legislacoes/arvore/<int:leg_id>')
@login_required
def pagina_arvore(leg_id): return render_template('arvore.html', legislacao_id=leg_id, active_page='arvores', active_group='biblioteca', **tmpl_ctx())

@app.route('/parametros')
@login_required
def pagina_parametros(): return render_template('parametros.html', active_page='parametros', active_group='parametros', **tmpl_ctx())

@app.route('/parametros/calculadora')
@login_required
def pagina_calculadora(): return render_template('parametros.html', active_page='calculadora', active_group='parametros', **tmpl_ctx())

@app.route('/parametros/importar')
@login_required
def pagina_importar(): return render_template('parametros.html', active_page='importar', active_group='parametros', **tmpl_ctx())

@app.route('/parametros/pendentes')
@login_required
def pagina_param_pendentes(): return render_template('parametros.html', active_page='param-pendentes', active_group='parametros', **tmpl_ctx())

@app.route('/monitoramento')
@login_required
def pagina_monitoramento(): return render_template('monitoramento.html', active_page='municipios', active_group='monitoramento', **tmpl_ctx())

@app.route('/monitoramento/historico')
@login_required
def pagina_historico(): return render_template('monitoramento.html', active_page='historico', active_group='monitoramento', **tmpl_ctx())

@app.route('/monitoramento/alteracoes')
@login_required
def pagina_alteracoes(): return render_template('monitoramento.html', active_page='alteracoes', active_group='monitoramento', **tmpl_ctx())

@app.route('/monitoramento/scheduler')
@login_required
def pagina_scheduler(): return render_template('monitoramento.html', active_page='scheduler', active_group='monitoramento', **tmpl_ctx())

@app.route('/integracao')
@login_required
def pagina_integracao(): return render_template('integracao.html', active_page='integracao-fila', active_group='integracoes', **tmpl_ctx())

@app.route('/integracao/aprovadas')
@login_required
def pagina_integracao_aprovadas(): return render_template('integracao.html', active_page='integracao-aprovadas', active_group='integracoes', **tmpl_ctx())

@app.route('/integracao/rejeitadas')
@login_required
def pagina_integracao_rejeitadas(): return render_template('integracao.html', active_page='integracao-rejeitadas', active_group='integracoes', **tmpl_ctx())

@app.route('/usuarios')
@admin_required
def pagina_usuarios(): return render_template('usuarios.html', active_page='todos-usuarios', active_group='usuarios', **tmpl_ctx())

@app.route('/usuarios/pendentes')
@admin_required
def pagina_users_pendentes(): return render_template('usuarios.html', active_page='usuarios-pendentes', active_group='usuarios', **tmpl_ctx())

@app.route('/config')
@app.route('/config/tipos-legislacao')
@admin_required
def pagina_config_tipos(): return render_template('configuracoes.html', active_page='tipos-leg', active_group='config', **tmpl_ctx())

@app.route('/config/assuntos')
@admin_required
def pagina_config_assuntos(): return render_template('configuracoes.html', active_page='assuntos', active_group='config', **tmpl_ctx())

@app.route('/config/email')
@admin_required
def pagina_config_email(): return render_template('configuracoes.html', active_page='config-email', active_group='config', **tmpl_ctx())

@app.route('/config/perfil')
@login_required
def pagina_perfil(): return render_template('perfil.html', active_page='perfil', active_group='config', **tmpl_ctx())

# -------------------------------------------------------------------
# API: AUTH
# -------------------------------------------------------------------

# -- FIX 2: Login API corrigido -- sessao agora e criada corretamente --
@app.route('/api/auth/login', methods=['POST'])
def api_login():
    d = request.json or {}
    email = d.get('email','').strip().lower()
    senha = d.get('senha','')
    user = qry("SELECT * FROM users WHERE email=%s AND ativo=TRUE AND aprovado=TRUE", (email,), 'one')
    if not user or not verificar_senha(senha, user['senha_hash']):
        return jsonify({'success':False,'error':'E-mail ou senha incorretos'}), 401
    session['user_id'] = user['id']
    session['nome'] = user['nome']
    session['email'] = user['email']
    session['role'] = user['role']
    qry("UPDATE users SET ultimo_acesso=NOW() WHERE id=%s", (user['id'],), commit=True, fetch=None)
    return jsonify({'success':True,'role':user['role']})

@app.route('/api/status')
def api_status():
    return jsonify({'ok': True})

@app.route('/api/auth/logout', methods=['POST'])
def api_logout(): session.clear(); return jsonify({'success':True})

# -- FIX 8: Rota GET para logout --
@app.route('/logout')
def logout():
    session.clear()
    return redirect('/login')


@app.route('/api/auth/cadastrar', methods=['POST'])
def api_cadastrar():
    d = request.json or {}
    nome = d.get('nome','').strip()
    email = d.get('email','').strip().lower()
    senha = d.get('senha','')
    if not nome or not email or not senha: return jsonify({'success':False,'error':'Campos obrigatorios'}), 400
    if qry("SELECT id FROM users WHERE email=%s", (email,), 'one'): return jsonify({'success':False,'error':'E-mail ja cadastrado'}), 400
    import re
    if not re.match(r'^(?=.*[A-Z])(?=.*[a-z])(?=.*\d)(?=.*[@$!%*?&\-_#])[A-Za-z\d@$!%*?&\-_#]{8,20}$', senha):
        return jsonify({'success':False,'error':'Senha fraca. Use 6-15 chars com maiuscula, minuscula, numero e especial.'}), 400
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("INSERT INTO users (nome,email,senha_hash,role,ativo,aprovado) VALUES (%s,%s,%s,'apenas_leitura',FALSE,FALSE) RETURNING id",
                    (nome, email, hash_senha(senha)))
        uid = cur.fetchone()[0]
        token_ativ = gerar_token()
        exp = datetime.now() + timedelta(hours=24)
        cur.execute("INSERT INTO aprovacao_tokens (user_id,token,tipo,expira_em) VALUES (%s,%s,'ativacao',%s)", (uid,token_ativ,exp))
        conn.commit()
    finally: conn.close()
    enviar_email_ativacao({'email':email,'nome':nome}, token_ativ)
    return jsonify({'success':True})

@app.route('/api/auth/esqueci-senha', methods=['POST'])
def api_esqueci_senha():
    d = request.json or {}
    email = d.get('email','').strip().lower()
    user = qry("SELECT * FROM users WHERE email=%s", (email,), 'one')
    if user:
        token = gerar_token()
        exp = datetime.now() + timedelta(hours=1)
        qry("INSERT INTO password_reset_tokens (user_id,token,expira_em) VALUES (%s,%s,%s)", (user['id'],token,exp), commit=True, fetch=None)
        enviar_email_reset(user, token)
    return jsonify({'success':True})

@app.route('/api/auth/reset-senha', methods=['POST'])
def api_reset_senha():
    d = request.json or {}
    token = d.get('token','')
    senha = d.get('nova_senha','') or d.get('senha','')
    tk = qry("SELECT * FROM password_reset_tokens WHERE token=%s AND usado=FALSE AND expira_em>NOW()", (token,), 'one')
    if not tk: return jsonify({'success':False,'error':'Link invalido ou expirado'}), 400
    qry("UPDATE users SET senha_hash=%s WHERE id=%s", (hash_senha(senha), tk['user_id']), commit=True, fetch=None)
    qry("UPDATE password_reset_tokens SET usado=TRUE WHERE id=%s", (tk['id'],), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/auth/alterar-senha', methods=['POST'])
@login_required
def api_alterar_senha():
    d = request.json or {}
    user = qry("SELECT * FROM users WHERE id=%s", (session['user_id'],), 'one')
    if not user: return jsonify({'success':False,'error':'Usuario nao encontrado'}), 404
    if not verificar_senha(d.get('senha_atual',''), user['senha_hash']):
        return jsonify({'success':False,'error':'Senha atual incorreta'}), 400
    qry("UPDATE users SET senha_hash=%s WHERE id=%s", (hash_senha(d['senha_nova']), session['user_id']), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/auth/perfil', methods=['GET','POST'])
@login_required
def api_perfil():
    if request.method == 'GET':
        user = qry("SELECT id,nome,email,role,criado_em FROM users WHERE id=%s", (session['user_id'],), 'one')
        return jsonify({'success':True,'data':user})
    from werkzeug.security import generate_password_hash, check_password_hash
    d = request.json or {}
    nome = d.get('nome','').strip()
    email = d.get('email','').strip().lower()
    senha_atual = d.get('senha_atual','').strip()
    senha_nova = d.get('senha_nova','').strip()
    if not nome: return jsonify({'success':False,'error':'Nome obrigatorio'}), 400
    if not email: return jsonify({'success':False,'error':'E-mail obrigatorio'}), 400
    # Verificar se email ja existe em outro usuario
    outro = qry("SELECT id FROM users WHERE email=%s AND id!=%s", (email, session['user_id']), 'one')
    if outro: return jsonify({'success':False,'error':'E-mail ja cadastrado por outro usuario'}), 400
    # Atualizar nome e email
    qry("UPDATE users SET nome=%s, email=%s WHERE id=%s", (nome, email, session['user_id']), commit=True, fetch=None)
    session['nome'] = nome
    session['email'] = email
    # Atualizar senha se fornecida
    if senha_nova:
        user = qry("SELECT senha_hash FROM users WHERE id=%s", (session['user_id'],), 'one')
        if not senha_atual: return jsonify({'success':False,'error':'Informe a senha atual para alterar a senha'}), 400
        if not user['senha_hash'] or not verificar_senha(senha_atual, user['senha_hash']):
            return jsonify({'success':False,'error':'Senha atual incorreta'}), 400
        import re
        if not re.match(r'^(?=.*[A-Z])(?=.*[a-z])(?=.*\d)(?=.*[@$!%*?&\-_#])[A-Za-z\d@$!%*?&\-_#]{8,20}$', senha_nova):
            return jsonify({'success':False,'error':'Nova senha fraca. Use 8-20 chars com maiuscula, minuscula, numero e especial'}), 400
        qry("UPDATE users SET senha_hash=%s WHERE id=%s", (hash_senha(senha_nova), session['user_id']), commit=True, fetch=None)
    return jsonify({'success':True})

# -------------------------------------------------------------------
# API: ADMIN USUARIOS
# -------------------------------------------------------------------

@app.route('/api/admin/usuarios', methods=['GET'])
@admin_required
def api_listar_usuarios():
    users = qry("SELECT id,nome,email,role,ativo,aprovado,criado_em,ultimo_acesso FROM users ORDER BY criado_em DESC")
    return jsonify({'success':True,'data':users})

@app.route('/api/admin/usuarios/<int:uid>/role', methods=['POST'])
@admin_required
def api_alterar_role(uid):
    role = (request.json or {}).get('role')
    if role not in ('admin','editor','apenas_leitura'): return jsonify({'success':False,'error':'Role invalida'}), 400
    qry("UPDATE users SET role=%s WHERE id=%s", (role, uid), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/admin/usuarios/<int:uid>', methods=['DELETE'])
@admin_required
def api_excluir_usuario(uid):
    if uid == session['user_id']: return jsonify({'success':False,'error':'Nao pode excluir a si mesmo'}), 400
    qry("DELETE FROM users WHERE id=%s", (uid,), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/admin/usuarios/<int:uid>/acesso', methods=['POST'])
@admin_required
def api_atualizar_acesso(uid):
    d = request.json or {}
    acesso = d.get('acesso')  # 'liberado' ou 'negado'
    if acesso not in ('liberado', 'negado'):
        return jsonify({'success': False, 'error': 'Valor invalido'})
    aprovado = acesso == 'liberado'
    ativo = acesso == 'liberado'
    user = qry("SELECT * FROM users WHERE id=%s", (uid,), 'one')
    if not user:
        return jsonify({'success': False, 'error': 'Usuario nao encontrado'})
    era_aprovado = user.get('aprovado', False)
    qry("UPDATE users SET aprovado=%s, ativo=%s WHERE id=%s", (aprovado, ativo, uid), commit=True, fetch=None)
    # Enviar e-mail de boas-vindas apenas quando liberar pela primeira vez
    if aprovado and not era_aprovado:
        try:
            enviar_email_generico(user['email'], 'Acesso Liberado — UrbanLex',
                f'<p>Ola {user["nome"]},</p>'
                f'<p>Seu acesso a plataforma <strong>UrbanLex</strong> foi liberado!</p>'
                f'<p>Acesse agora: <a href="{get_app_url()}/login">{get_app_url()}/login</a></p>'
                f'<p>Bem-vindo!</p>')
        except Exception as e:
            print(f"[EMAIL] Erro ao enviar email de acesso liberado: {e}")
    return jsonify({'success': True})

@app.route('/api/admin/verificar-senha', methods=['POST'])
@admin_required
def api_verificar_senha():
    from werkzeug.security import check_password_hash
    d = request.json or {}
    senha = d.get('senha', '')
    user = qry("SELECT senha_hash FROM users WHERE id=%s", (session['user_id'],), 'one')
    if not user:
        return jsonify({'success': False})
    ok = check_password_hash(user['senha_hash'], senha)
    return jsonify({'success': ok})

# -------------------------------------------------------------------
# API: BIBLIOTECA DE LEGISLACOES
# -------------------------------------------------------------------

@app.route('/api/legislacoes', methods=['GET'])
@login_required
def api_listar_legislacoes():
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 50))
    offset = (page - 1) * per_page
    where = ['l.pendente_aprovacao = FALSE']
    params = []
    for campo, col in [('estado','l.estado'),('municipio','l.municipio_nome'),('numero','l.numero'),
                       ('status','l.status'),('tipo_id','l.tipo_id'),('assunto_id','l.assunto_id')]:
        v = request.args.get(campo)
        if v: where.append(f"{col} = %s"); params.append(v)
    if request.args.get('em_monitoramento') in ('0','1'):
        where.append("l.em_monitoramento = %s"); params.append(request.args['em_monitoramento'] == '1')
    if request.args.get('keywords'):
        where.append("(l.ementa ILIKE %s OR l.palavras_chave ILIKE %s)")
        kw = f"%{request.args['keywords']}%"; params += [kw, kw]
    if request.args.get('ano'):
        where.append("l.ano = %s"); params.append(int(request.args['ano']))
    where_sql = ' AND '.join(where)
    total = qry(f"SELECT COUNT(*) as n FROM legislacoes l WHERE {where_sql}", params, 'one')['n']
    order = 'l.criado_em DESC' if request.args.get('order') == 'recente' else 'l.estado, l.municipio_nome, l.ano DESC'
    data = qry(f"""SELECT l.*, COALESCE(tl.nome, l.tipo_nome) as tipo_nome, COALESCE(al.nome, l.assunto_nome) as assunto_nome,
                   COALESCE(fa.qtd_arquivos, 0) as qtd_arquivos
                   FROM legislacoes l
                   LEFT JOIN tipos_legislacao tl ON l.tipo_id=tl.id
                   LEFT JOIN assuntos_legislacao al ON l.assunto_id=al.id
                   LEFT JOIN (SELECT legislacao_id, COUNT(*) as qtd_arquivos FROM legislacao_arquivos GROUP BY legislacao_id) fa ON fa.legislacao_id=l.id
                   WHERE {where_sql} ORDER BY {order} LIMIT %s OFFSET %s""",
               params + [per_page, offset])
    return jsonify({'success':True,'data':data,'total':total,'page':page})

@app.route('/api/legislacoes', methods=['POST'])
@editor_required
def api_criar_legislacao():
    arquivo = request.files.get('arquivo')
    if not arquivo: return jsonify({'success':False,'error':'Nenhum arquivo enviado'}), 400
    nome_arquivo = secure_filename(arquivo.filename)
    arquivo_bytes = arquivo.read()
    texto = extrair_texto_arquivo(arquivo_bytes, nome_arquivo)
    hash_c = hashlib.sha256(arquivo_bytes).hexdigest()
    ano = request.form.get('ano')
    tipo_id = request.form.get('tipo_id') or None
    assunto_id = request.form.get('assunto_id') or None
    estado = request.form.get('estado') or None
    esfera = request.form.get('esfera','municipal')
    municipio_nome = request.form.get('municipio') or None
    numero = request.form.get('numero')
    ementa = request.form.get('ementa') or None
    data_pub = request.form.get('data_publicacao') or None
    em_mon = request.form.get('em_monitoramento','0') == '1'
    kw = request.form.get('palavras_chave')
    kw_json = json.dumps([k.strip() for k in kw.split(',') if k.strip()]) if kw else None

    tipo_row = qry("SELECT nome FROM tipos_legislacao WHERE id=%s", (tipo_id,), 'one') if tipo_id else None
    assunto_row = qry("SELECT nome FROM assuntos_legislacao WHERE id=%s", (assunto_id,), 'one') if assunto_id else None

    leg_row = qry("""
        INSERT INTO legislacoes (pais,esfera,estado,municipio_nome,tipo_id,tipo_nome,numero,ano,data_publicacao,
            ementa,assunto_id,assunto_nome,palavras_chave,conteudo_texto,arquivo_nome,arquivo_tipo,
            hash_conteudo,em_monitoramento,origem,pendente_aprovacao,criado_em)
        VALUES ('BR',%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'manual',FALSE,NOW()) RETURNING id
    """, (esfera,estado,municipio_nome,tipo_id,tipo_row['nome'] if tipo_row else None,
          numero,int(ano) if ano else None,data_pub,ementa,assunto_id,
          assunto_row['nome'] if assunto_row else None,kw_json,texto,nome_arquivo,
          Path(nome_arquivo).suffix.lower()[1:],hash_c,em_mon), 'one', commit=True)

    novo_id = leg_row['id'] if leg_row else None

    url_r2 = None
    if novo_id and r2_disponivel():
        url_r2 = r2_upload(arquivo_bytes, nome_arquivo, leg_id=novo_id)
        if url_r2:
            qry("UPDATE legislacoes SET arquivo_url=%s WHERE id=%s", (url_r2, novo_id), commit=True, fetch=None)

    # Inserir tambem na tabela de arquivos (multi-arquivo)
    if novo_id:
        qry("""INSERT INTO legislacao_arquivos
               (legislacao_id, nome_arquivo, arquivo_tipo, arquivo_url,
                tamanho_bytes, hash_conteudo, conteudo_texto, criado_por)
               VALUES (%s,%s,%s,%s,%s,%s,%s,%s)""",
            (novo_id, nome_arquivo, Path(nome_arquivo).suffix.lower().lstrip('.'),
             url_r2, len(arquivo_bytes), hash_c, texto, session['user_id']),
            commit=True, fetch=None)

    return jsonify({'success':True,'id':novo_id})

@app.route('/api/legislacoes/<int:leg_id>', methods=['GET'])
@login_required
def api_get_legislacao(leg_id):
    l = qry("SELECT l.*, COALESCE(tl.nome, l.tipo_nome) as tipo_nome, COALESCE(al.nome, l.assunto_nome) as assunto_nome FROM legislacoes l LEFT JOIN tipos_legislacao tl ON l.tipo_id=tl.id LEFT JOIN assuntos_legislacao al ON l.assunto_id=al.id WHERE l.id=%s", (leg_id,), 'one')
    if not l: return jsonify({'success':False,'error':'Nao encontrada'}), 404
    arquivos = qry("""SELECT id, nome_arquivo, arquivo_tipo, tamanho_bytes, criado_em
                      FROM legislacao_arquivos WHERE legislacao_id=%s ORDER BY criado_em""", (leg_id,))
    l['arquivos'] = arquivos or []
    l['total_arquivos_bytes'] = sum(a.get('tamanho_bytes', 0) or 0 for a in (arquivos or []))
    # Incluir relacoes categorizadas
    rels = qry("""
        SELECT r.id as relacao_id, r.tipo_relacao,
               r.legislacao_pai_id, r.legislacao_filha_id,
               lp.numero as pai_numero, lp.ano as pai_ano, tlp.nome as pai_tipo,
               lf.numero as filha_numero, lf.ano as filha_ano, tlf.nome as filha_tipo
        FROM legislacao_relacoes r
        JOIN legislacoes lp ON r.legislacao_pai_id = lp.id
        JOIN legislacoes lf ON r.legislacao_filha_id = lf.id
        LEFT JOIN tipos_legislacao tlp ON lp.tipo_id = tlp.id
        LEFT JOIN tipos_legislacao tlf ON lf.tipo_id = tlf.id
        WHERE r.legislacao_pai_id = %s OR r.legislacao_filha_id = %s
    """, (leg_id, leg_id))
    relacoes = {'revogada_por':[],'modificada_por':[],'modifica':[],'revoga':[],'citadas':[],'citada_por':[]}
    for r in (rels or []):
        tipo = (r['tipo_relacao'] or '').lower().strip()
        def mk(prefix):
            return {'id':r[f'legislacao_{prefix}_id'],'numero':r[f'{prefix}_numero'],'ano':r[f'{prefix}_ano'],
                    'tipo_nome':r[f'{prefix}_tipo'],'relacao_id':r['relacao_id'],
                    'ref':f"{r[f'{prefix}_tipo'] or ''} {r[f'{prefix}_numero'] or ''}/{r[f'{prefix}_ano'] or ''}".strip()}
        if tipo in ('revoga','revogacao'):
            if r['legislacao_filha_id']==leg_id: relacoes['revogada_por'].append(mk('pai'))
            else: relacoes['revoga'].append(mk('filha'))
        elif tipo in ('modifica','altera','alteracao','modificacao'):
            if r['legislacao_filha_id']==leg_id: relacoes['modificada_por'].append(mk('pai'))
            else: relacoes['modifica'].append(mk('filha'))
        elif tipo in ('cita','referencia','citacao','menciona'):
            if r['legislacao_pai_id']==leg_id: relacoes['citadas'].append(mk('filha'))
            else: relacoes['citada_por'].append(mk('pai'))
    l['relacoes'] = relacoes
    return jsonify({'success':True,'data':l})

# -- FIX 1: DELETE corrigido -- leg_id + foreign keys + verificacao --
@app.route('/api/legislacoes/<int:leg_id>', methods=['DELETE'])
@editor_required
def api_excluir_legislacao(leg_id):
    leg = qry("SELECT arquivo_url FROM legislacoes WHERE id=%s", (leg_id,), 'one')
    if not leg:
        return jsonify({'success': False, 'error': 'Legislacao nao encontrada'}), 404
    if leg.get('arquivo_url') and r2_disponivel():
        r2_delete(leg['arquivo_url'])
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM legislacao_relacoes WHERE legislacao_pai_id=%s OR legislacao_filha_id=%s", (leg_id, leg_id))
        cur.execute("DELETE FROM alteracoes WHERE legislacao_id=%s", (leg_id,))
        cur.execute("DELETE FROM integracao_atualizacoes WHERE legislacao_id=%s", (leg_id,))
        cur.execute("DELETE FROM legislacoes WHERE id=%s", (leg_id,))
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()
    return jsonify({'success': True})

@app.route('/api/legislacoes/<int:leg_id>/texto')
@login_required
def api_leg_texto(leg_id):
    l = qry("SELECT conteudo_texto, arquivo_nome FROM legislacoes WHERE id=%s", (leg_id,), 'one')
    if not l or not l.get('conteudo_texto'): return jsonify({'success':False,'error':'Texto nao disponivel'}), 404
    return jsonify({'success':True,'data':{'texto':l['conteudo_texto'],'arquivo_nome':l.get('arquivo_nome','')}})

@app.route('/api/legislacoes/<int:leg_id>/documento')
@login_required
def api_leg_documento(leg_id):
    l = qry("SELECT arquivo_url, arquivo_nome FROM legislacoes WHERE id=%s", (leg_id,), 'one')
    if not l: return jsonify({'error':'Nao encontrada'}), 404
    url = l.get('arquivo_url')
    if not url:
        return jsonify({'error':'Documento nao disponivel. Faca upload do arquivo.'}), 404
    if url.startswith('http'):
        return redirect(url)
    if url.startswith('/static/'):
        import os
        file_path = os.path.join(os.path.dirname(__file__), url.lstrip('/'))
        if os.path.isfile(file_path):
            from flask import send_file
            return send_file(file_path, as_attachment=True, download_name=l.get('arquivo_nome') or os.path.basename(file_path))
        return jsonify({'error':'Arquivo local nao encontrado.'}), 404
    if r2_disponivel():
        url_assinada = r2_url_assinada(url, expiracao_seg=3600)
        if url_assinada:
            return redirect(url_assinada)
    return jsonify({'error':'Documento armazenado no R2 mas nao foi possivel gerar URL de acesso.'}), 503

@app.route('/api/legislacoes/<int:leg_id>/monitoramento', methods=['POST'])
@editor_required
def api_toggle_monitoramento(leg_id):
    d      = request.json or {}
    ativar = d.get('ativar', True)
    data_inicio = d.get('data_inicio_monitoramento') or None
    data_fim    = d.get('data_fim_monitoramento') or None
    qry("""UPDATE legislacoes
           SET em_monitoramento=%s,
               data_inicio_monitoramento=%s,
               data_fim_monitoramento=%s,
               ultima_verificacao_monitoramento = NULL
           WHERE id=%s""",
        (ativar, data_inicio, data_fim, leg_id), commit=True, fetch=None)
    return jsonify({'success':True})


@app.route('/api/legislacoes/<int:leg_id>/periodo-monitoramento', methods=['POST'])
@editor_required
def api_atualizar_periodo_monitoramento(leg_id):
    """Atualiza período de monitoramento sem alterar o estado ativo/inativo."""
    d = request.json or {}
    data_inicio = d.get('data_inicio') or None
    data_fim    = d.get('data_fim') or None   # NULL = até hoje
    # Se a data de início mudou para antes da última verificação, resetar verificação
    leg = qry("SELECT data_inicio_monitoramento, ultima_verificacao_monitoramento FROM legislacoes WHERE id=%s",
              (leg_id,), 'one')
    reset_verificacao = False
    if leg and data_inicio:
        from datetime import date as _date_type
        try:
            nova_data = _date_type.fromisoformat(data_inicio) if isinstance(data_inicio, str) else data_inicio
            verif_atual = leg.get('ultima_verificacao_monitoramento')
            inicio_atual = leg.get('data_inicio_monitoramento')
            if inicio_atual and nova_data < inicio_atual:
                reset_verificacao = True
            if verif_atual and nova_data < verif_atual:
                reset_verificacao = True
        except Exception:
            pass
    if reset_verificacao:
        qry("""UPDATE legislacoes
               SET data_inicio_monitoramento=%s, data_fim_monitoramento=%s,
                   ultima_verificacao_monitoramento=NULL
               WHERE id=%s""",
            (data_inicio, data_fim, leg_id), commit=True, fetch=None)
    else:
        qry("""UPDATE legislacoes
               SET data_inicio_monitoramento=%s, data_fim_monitoramento=%s
               WHERE id=%s""",
            (data_inicio, data_fim, leg_id), commit=True, fetch=None)
    return jsonify({'success': True, 'reset_verificacao': reset_verificacao})

# -- Re-upload de arquivo para legislacao existente --
@app.route('/api/legislacoes/<int:leg_id>/upload', methods=['POST'])
@editor_required
def api_reupload_legislacao(leg_id):
    """Faz upload/re-upload de arquivo para uma legislacao existente."""
    leg = qry("SELECT id, arquivo_url FROM legislacoes WHERE id=%s", (leg_id,), 'one')
    if not leg:
        return jsonify({'success': False, 'error': 'Legislacao nao encontrada'}), 404
    arquivo = request.files.get('arquivo')
    if not arquivo:
        return jsonify({'success': False, 'error': 'Nenhum arquivo enviado'}), 400
    nome_arquivo = secure_filename(arquivo.filename)
    arquivo_bytes = arquivo.read()
    texto = extrair_texto_arquivo(arquivo_bytes, nome_arquivo)
    hash_c = hashlib.sha256(arquivo_bytes).hexdigest()
    if leg.get('arquivo_url') and r2_disponivel():
        r2_delete(leg['arquivo_url'])
    url_r2 = None
    if r2_disponivel():
        url_r2 = r2_upload(arquivo_bytes, nome_arquivo, leg_id=leg_id)
    qry("""UPDATE legislacoes
           SET arquivo_nome=%s, arquivo_tipo=%s, arquivo_url=%s,
               conteudo_texto=%s, hash_conteudo=%s, atualizado_em=NOW()
           WHERE id=%s""",
        (nome_arquivo, Path(nome_arquivo).suffix.lower()[1:],
         url_r2, texto, hash_c, leg_id),
        commit=True, fetch=None)
    return jsonify({
        'success': True,
        'arquivo_url': url_r2,
        'texto_extraido': bool(texto),
        'message': 'Arquivo salvo no R2' if url_r2 else 'Arquivo salvo (R2 indisponivel - apenas texto extraido)'
    })

# -- Editar campos da legislacao --
@app.route('/api/legislacoes/<int:leg_id>', methods=['PUT'])
@editor_required
def api_editar_legislacao(leg_id):
    """Atualiza campos editaveis de uma legislacao."""
    leg = qry("SELECT id FROM legislacoes WHERE id=%s", (leg_id,), 'one')
    if not leg:
        return jsonify({'success': False, 'error': 'Legislacao nao encontrada'}), 404
    d = request.json or {}
    if not d:
        return jsonify({'success': False, 'error': 'Nenhum dado enviado'}), 400
    campos_editaveis = {
        'numero', 'ano', 'ementa', 'status', 'estado', 'municipio_nome',
        'tipo_id', 'assunto_id', 'palavras_chave', 'data_publicacao',
        'esfera', 'em_monitoramento', 'data_inicio_monitoramento',
        'data_fim_monitoramento', 'url_original', 'observacoes'
    }
    cols = [k for k in d.keys() if k in campos_editaveis]
    if not cols:
        return jsonify({'success': False, 'error': 'Nenhum campo valido para atualizar'}), 400
    vals = []
    for c in cols:
        v = d[c]
        if c == 'ano' and v is not None:
            v = int(v)
        if c == 'em_monitoramento':
            v = bool(v)
        if c == 'palavras_chave' and isinstance(v, list):
            v = json.dumps(v)
        vals.append(v)
    # Atualizar nomes de tipo e assunto se ids mudaram
    extra_sets = []
    extra_vals = []
    if 'tipo_id' in d:
        tipo_row = qry("SELECT nome FROM tipos_legislacao WHERE id=%s", (d['tipo_id'],), 'one') if d['tipo_id'] else None
        extra_sets.append("tipo_nome=%s")
        extra_vals.append(tipo_row['nome'] if tipo_row else None)
    if 'assunto_id' in d:
        assunto_row = qry("SELECT nome FROM assuntos_legislacao WHERE id=%s", (d['assunto_id'],), 'one') if d['assunto_id'] else None
        extra_sets.append("assunto_nome=%s")
        extra_vals.append(assunto_row['nome'] if assunto_row else None)
    set_sql = ', '.join(f"{c}=%s" for c in cols)
    if extra_sets:
        set_sql += ', ' + ', '.join(extra_sets)
    set_sql += ', atualizado_em=NOW()'
    all_vals = vals + extra_vals + [leg_id]
    qry(f"UPDATE legislacoes SET {set_sql} WHERE id=%s", all_vals, commit=True, fetch=None)
    # Retornar legislacao atualizada
    updated = qry("SELECT l.*, COALESCE(tl.nome, l.tipo_nome) as tipo_nome, COALESCE(al.nome, l.assunto_nome) as assunto_nome FROM legislacoes l LEFT JOIN tipos_legislacao tl ON l.tipo_id=tl.id LEFT JOIN assuntos_legislacao al ON l.assunto_id=al.id WHERE l.id=%s", (leg_id,), 'one')
    return jsonify({'success': True, 'data': updated})

# -------------------------------------------------------------------
# API: GESTAO DE ARQUIVOS POR LEGISLACAO (multi-arquivo)
# -------------------------------------------------------------------

MAX_TOTAL_ARQUIVOS_BYTES = 100 * 1024 * 1024  # 100MB por legislacao

@app.route('/api/legislacoes/<int:leg_id>/arquivos', methods=['GET'])
@login_required
def api_listar_arquivos(leg_id):
    """Lista todos os arquivos associados a uma legislacao."""
    leg = qry("SELECT id FROM legislacoes WHERE id=%s", (leg_id,), 'one')
    if not leg:
        return jsonify({'success': False, 'error': 'Legislacao nao encontrada'}), 404
    arquivos = qry("""SELECT id, legislacao_id, nome_arquivo, arquivo_tipo,
                             tamanho_bytes, criado_em, criado_por
                      FROM legislacao_arquivos
                      WHERE legislacao_id=%s ORDER BY criado_em""", (leg_id,))
    total_bytes = sum(a.get('tamanho_bytes', 0) or 0 for a in arquivos)
    # Incluir tambem o arquivo principal da legislacao (legado)
    leg_principal = qry("SELECT arquivo_nome, arquivo_tipo, arquivo_url FROM legislacoes WHERE id=%s AND arquivo_nome IS NOT NULL", (leg_id,), 'one')
    return jsonify({
        'success': True,
        'data': arquivos,
        'arquivo_principal': leg_principal,
        'total_bytes': total_bytes,
        'limite_bytes': MAX_TOTAL_ARQUIVOS_BYTES,
        'espaco_disponivel': MAX_TOTAL_ARQUIVOS_BYTES - total_bytes
    })

@app.route('/api/legislacoes/<int:leg_id>/arquivos', methods=['POST'])
@editor_required
def api_upload_arquivo(leg_id):
    """Upload de um ou mais arquivos para uma legislacao (limite total: 100MB)."""
    leg = qry("SELECT id FROM legislacoes WHERE id=%s", (leg_id,), 'one')
    if not leg:
        return jsonify({'success': False, 'error': 'Legislacao nao encontrada'}), 404

    arquivos = request.files.getlist('arquivos')
    if not arquivos or all(f.filename == '' for f in arquivos):
        return jsonify({'success': False, 'error': 'Nenhum arquivo enviado'}), 400

    # Verificar espaco total disponivel
    usado = qry("SELECT COALESCE(SUM(tamanho_bytes),0) as total FROM legislacao_arquivos WHERE legislacao_id=%s",
                 (leg_id,), 'one')['total']
    espaco_livre = MAX_TOTAL_ARQUIVOS_BYTES - usado

    resultados = []
    erros = []
    for arquivo in arquivos:
        if not arquivo or arquivo.filename == '':
            continue
        nome_arquivo = secure_filename(arquivo.filename)
        arquivo_bytes = arquivo.read()
        tamanho = len(arquivo_bytes)

        if tamanho > espaco_livre:
            erros.append(f"{nome_arquivo}: excede o limite de 100MB (disponivel: {espaco_livre // (1024*1024)}MB)")
            continue

        hash_c = hashlib.sha256(arquivo_bytes).hexdigest()
        texto = extrair_texto_arquivo(arquivo_bytes, nome_arquivo)

        # Upload para R2
        url_r2 = None
        if r2_disponivel():
            r2_key = f"legislacoes/{leg_id}/arquivos/{hash_c[:8]}_{nome_arquivo}"
            url_r2 = r2_upload(arquivo_bytes, nome_arquivo, leg_id=leg_id)

        # Inserir registro
        row = qry("""INSERT INTO legislacao_arquivos
                      (legislacao_id, nome_arquivo, arquivo_tipo, arquivo_url,
                       tamanho_bytes, hash_conteudo, conteudo_texto, criado_por)
                      VALUES (%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id, nome_arquivo, tamanho_bytes, criado_em""",
                   (leg_id, nome_arquivo, Path(nome_arquivo).suffix.lower().lstrip('.'),
                    url_r2, tamanho, hash_c, texto, session['user_id']),
                   'one', commit=True)

        espaco_livre -= tamanho
        resultados.append(row)

    # Atualizar texto concatenado na legislacao (opcional - para busca)
    todos_textos = qry("SELECT conteudo_texto FROM legislacao_arquivos WHERE legislacao_id=%s AND conteudo_texto IS NOT NULL", (leg_id,))
    if todos_textos:
        texto_completo = '\n\n---\n\n'.join(t['conteudo_texto'] for t in todos_textos if t.get('conteudo_texto'))
        if texto_completo:
            qry("UPDATE legislacoes SET conteudo_texto=%s, atualizado_em=NOW() WHERE id=%s",
                (texto_completo, leg_id), commit=True, fetch=None)

    return jsonify({
        'success': True,
        'arquivos': resultados,
        'erros': erros,
        'espaco_disponivel': espaco_livre
    })

@app.route('/api/legislacoes/<int:leg_id>/arquivos/<int:arq_id>', methods=['DELETE'])
@editor_required
def api_excluir_arquivo(leg_id, arq_id):
    """Exclui um arquivo especifico de uma legislacao."""
    arq = qry("SELECT * FROM legislacao_arquivos WHERE id=%s AND legislacao_id=%s", (arq_id, leg_id), 'one')
    if not arq:
        return jsonify({'success': False, 'error': 'Arquivo nao encontrado'}), 404

    # Deletar do R2
    if arq.get('arquivo_url') and r2_disponivel():
        r2_delete(arq['arquivo_url'])

    qry("DELETE FROM legislacao_arquivos WHERE id=%s", (arq_id,), commit=True, fetch=None)

    return jsonify({'success': True, 'message': f"Arquivo '{arq['nome_arquivo']}' excluido"})

@app.route('/api/legislacoes/<int:leg_id>/arquivos/<int:arq_id>/download')
@login_required
def api_download_arquivo(leg_id, arq_id):
    """Gera URL de download para um arquivo especifico."""
    arq = qry("SELECT * FROM legislacao_arquivos WHERE id=%s AND legislacao_id=%s", (arq_id, leg_id), 'one')
    if not arq:
        return jsonify({'error': 'Arquivo nao encontrado'}), 404
    url = arq.get('arquivo_url')
    if not url:
        return jsonify({'error': 'Arquivo sem URL de acesso'}), 404
    if url.startswith('http'):
        return redirect(url)
    if url.startswith('/static/'):
        import os
        file_path = os.path.join(os.path.dirname(__file__), url.lstrip('/'))
        if os.path.isfile(file_path):
            from flask import send_file
            return send_file(file_path, as_attachment=True, download_name=arq.get('nome_arquivo') or os.path.basename(file_path))
        return jsonify({'error': 'Arquivo local nao encontrado'}), 404
    if r2_disponivel():
        url_assinada = r2_url_assinada(url, expiracao_seg=3600)
        if url_assinada:
            return redirect(url_assinada)
    return jsonify({'error': 'Nao foi possivel gerar URL de download'}), 503

# -- Arvore genealogica --
@app.route('/api/legislacoes/<int:leg_id>/arvore')
@login_required
def api_leg_arvore(leg_id):
    raiz = qry("SELECT l.*, tl.nome as tipo_nome FROM legislacoes l LEFT JOIN tipos_legislacao tl ON l.tipo_id=tl.id WHERE l.id=%s", (leg_id,), 'one')
    if not raiz: return jsonify({'success':False,'error':'Nao encontrada'}), 404
    relacoes = qry("""
        SELECT r.*, lp.id as pai_id, lp.numero as pai_numero, lp.ano as pai_ano, tl1.nome as pai_tipo,
               lf.id as filha_id, lf.numero as filha_numero, lf.ano as filha_ano, tl2.nome as filha_tipo,
               lf.status as filha_status, lf.municipio_nome, lf.data_publicacao, lf.ementa,
               lf.conteudo_texto, lf.arquivo_url
        FROM legislacao_relacoes r
        JOIN legislacoes lp ON r.legislacao_pai_id = lp.id
        JOIN legislacoes lf ON r.legislacao_filha_id = lf.id
        LEFT JOIN tipos_legislacao tl1 ON lp.tipo_id = tl1.id
        LEFT JOIN tipos_legislacao tl2 ON lf.tipo_id = tl2.id
        WHERE lp.id = %s OR lf.id = %s
    """, (leg_id, leg_id))
    ids_vistos = {leg_id}
    nodes = [{'id': leg_id, **raiz}]
    edges = []
    for r in relacoes:
        for node_id, num, ano, tipo, status in [
            (r['pai_id'], r['pai_numero'], r['pai_ano'], r['pai_tipo'], 'vigente'),
            (r['filha_id'], r['filha_numero'], r['filha_ano'], r['filha_tipo'], r['filha_status'])
        ]:
            if node_id not in ids_vistos:
                ids_vistos.add(node_id)
                nodes.append({'id':node_id,'numero':num,'ano':ano,'tipo_nome':tipo,'status':status,
                               'municipio_nome':r.get('municipio_nome'),'ementa':r.get('ementa'),
                               'data_publicacao':str(r.get('data_publicacao','')) if r.get('data_publicacao') else None,
                               'conteudo_texto':bool(r.get('conteudo_texto')),'arquivo_url':r.get('arquivo_url')})
        edges.append({'id':r['id'],'legislacao_pai_id':r['legislacao_pai_id'],
                      'legislacao_filha_id':r['legislacao_filha_id'],'tipo_relacao':r['tipo_relacao']})
    return jsonify({'success':True,'data':{'nodes':nodes,'edges':edges,'raiz':raiz}})

# -- Relacoes categorizadas de uma legislacao --
@app.route('/api/legislacoes/<int:leg_id>/relacoes')
@login_required
def api_leg_relacoes(leg_id):
    """Retorna relacoes categorizadas: revogada_por, modificada_por, modifica, citadas."""
    leg = qry("SELECT id FROM legislacoes WHERE id=%s", (leg_id,), 'one')
    if not leg:
        return jsonify({'success': False, 'error': 'Legislacao nao encontrada'}), 404

    # Buscar todas relacoes envolvendo esta legislacao
    rels = qry("""
        SELECT r.id as relacao_id, r.tipo_relacao,
               r.legislacao_pai_id, r.legislacao_filha_id,
               lp.numero as pai_numero, lp.ano as pai_ano, tlp.nome as pai_tipo,
               lp.municipio_nome as pai_municipio, lp.status as pai_status,
               lf.numero as filha_numero, lf.ano as filha_ano, tlf.nome as filha_tipo,
               lf.municipio_nome as filha_municipio, lf.status as filha_status
        FROM legislacao_relacoes r
        JOIN legislacoes lp ON r.legislacao_pai_id = lp.id
        JOIN legislacoes lf ON r.legislacao_filha_id = lf.id
        LEFT JOIN tipos_legislacao tlp ON lp.tipo_id = tlp.id
        LEFT JOIN tipos_legislacao tlf ON lf.tipo_id = tlf.id
        WHERE r.legislacao_pai_id = %s OR r.legislacao_filha_id = %s
    """, (leg_id, leg_id))

    resultado = {
        'revogada_por': [],
        'modificada_por': [],
        'modifica': [],
        'revoga': [],
        'citadas': [],
        'citada_por': []
    }

    for r in (rels or []):
        tipo = (r['tipo_relacao'] or '').lower().strip()
        pai_id = r['legislacao_pai_id']
        filha_id = r['legislacao_filha_id']

        def info_leg(prefix):
            return {
                'id': r[f'legislacao_{prefix}_id'],
                'numero': r[f'{prefix}_numero'],
                'ano': r[f'{prefix}_ano'],
                'tipo_nome': r[f'{prefix}_tipo'],
                'municipio_nome': r[f'{prefix}_municipio'],
                'status': r[f'{prefix}_status'],
                'relacao_id': r['relacao_id'],
                'ref': f"{r[f'{prefix}_tipo'] or ''} {r[f'{prefix}_numero'] or ''}/{r[f'{prefix}_ano'] or ''}".strip()
            }

        if tipo in ('revoga', 'revogacao'):
            if filha_id == leg_id:
                # Esta legislacao foi revogada pela pai
                resultado['revogada_por'].append(info_leg('pai'))
            else:
                # Esta legislacao revoga a filha
                resultado['revoga'].append(info_leg('filha'))

        elif tipo in ('modifica', 'altera', 'alteracao', 'modificacao'):
            if filha_id == leg_id:
                # Esta legislacao foi modificada pela pai
                resultado['modificada_por'].append(info_leg('pai'))
            else:
                # Esta legislacao modifica a filha
                resultado['modifica'].append(info_leg('filha'))

        elif tipo in ('cita', 'referencia', 'citacao', 'menciona'):
            if pai_id == leg_id:
                # Esta legislacao cita a filha
                resultado['citadas'].append(info_leg('filha'))
            else:
                # Esta legislacao e citada pela pai
                resultado['citada_por'].append(info_leg('pai'))

    return jsonify({'success': True, 'data': resultado})

@app.route('/api/legislacoes/<int:leg_id>/relacoes', methods=['POST'])
@editor_required
def api_add_relacao(leg_id):
    """Adiciona uma relacao entre legislacoes."""
    d = request.json or {}
    tipo_relacao = d.get('tipo_relacao', '').strip()
    outra_leg_id = d.get('outra_legislacao_id')
    direcao = d.get('direcao', 'pai')  # 'pai' = outra e pai, 'filha' = outra e filha

    if not tipo_relacao or not outra_leg_id:
        return jsonify({'success': False, 'error': 'tipo_relacao e outra_legislacao_id obrigatorios'}), 400

    outra = qry("SELECT id FROM legislacoes WHERE id=%s", (outra_leg_id,), 'one')
    if not outra:
        return jsonify({'success': False, 'error': 'Legislacao relacionada nao encontrada'}), 404

    if direcao == 'pai':
        pai_id, filha_id = outra_leg_id, leg_id
    else:
        pai_id, filha_id = leg_id, outra_leg_id

    # Verificar se ja existe
    existente = qry("""SELECT id FROM legislacao_relacoes
                       WHERE legislacao_pai_id=%s AND legislacao_filha_id=%s AND tipo_relacao=%s""",
                    (pai_id, filha_id, tipo_relacao), 'one')
    if existente:
        return jsonify({'success': False, 'error': 'Relacao ja existe'}), 400

    qry("""INSERT INTO legislacao_relacoes (legislacao_pai_id, legislacao_filha_id, tipo_relacao)
           VALUES (%s, %s, %s)""",
        (pai_id, filha_id, tipo_relacao), commit=True, fetch=None)

    return jsonify({'success': True})

@app.route('/api/legislacoes/<int:leg_id>/relacoes/<int:rel_id>', methods=['DELETE'])
@editor_required
def api_del_relacao(leg_id, rel_id):
    """Remove uma relacao."""
    qry("DELETE FROM legislacao_relacoes WHERE id=%s AND (legislacao_pai_id=%s OR legislacao_filha_id=%s)",
        (rel_id, leg_id, leg_id), commit=True, fetch=None)
    return jsonify({'success': True})

@app.route('/api/legislacoes/busca-rapida')
@login_required
def api_busca_rapida_legislacoes():
    """Busca rapida por numero/ano para vincular relacoes."""
    q = request.args.get('q', '').strip()
    if len(q) < 2:
        return jsonify({'success': True, 'data': []})
    data = qry("""SELECT l.id, l.numero, l.ano, tl.nome as tipo_nome, l.municipio_nome, l.status
                  FROM legislacoes l LEFT JOIN tipos_legislacao tl ON l.tipo_id=tl.id
                  WHERE l.pendente_aprovacao=FALSE AND (
                    l.numero ILIKE %s OR CAST(l.ano AS TEXT) ILIKE %s
                    OR tl.nome ILIKE %s OR l.municipio_nome ILIKE %s
                    OR l.ementa ILIKE %s
                  ) ORDER BY l.ano DESC, l.numero LIMIT 20""",
               (f'%{q}%', f'%{q}%', f'%{q}%', f'%{q}%', f'%{q}%'))
    return jsonify({'success': True, 'data': data})

# -- Busca com IA --
@app.route('/api/legislacoes/buscar-ia', methods=['POST'])
@editor_required
def api_buscar_ia():
    consulta = (request.json or {}).get('consulta','').strip()
    if not consulta: return jsonify({'success':False,'error':'Consulta obrigatoria'}), 400
    leg_id = qry("INSERT INTO buscas_ia (consulta,status,solicitado_por,criado_em) VALUES (%s,'pendente',%s,NOW()) RETURNING id",
                 (consulta, session['user_id']), 'one', commit=True)
    def rodar_busca(bid, consul):
        try:
            resultado = _buscar_legislacao_internet(consul)
            if resultado and resultado.get('url'):
                qry("UPDATE buscas_ia SET status='encontrado',resultado_url=%s,resultado_nome=%s,finalizado_em=NOW() WHERE id=%s",
                    (resultado['url'],resultado.get('titulo',''),bid), commit=True, fetch=None)
                qry("""INSERT INTO legislacoes (municipio_nome,estado,ementa,url_original,origem,pendente_aprovacao,criado_em)
                    VALUES (%s,%s,%s,%s,'busca_ia',TRUE,NOW())""",
                    (resultado.get('municipio'),resultado.get('estado'),resultado.get('titulo',''),resultado['url']), commit=True, fetch=None)
            else:
                qry("UPDATE buscas_ia SET status='nao_encontrado',finalizado_em=NOW() WHERE id=%s", (bid,), commit=True, fetch=None)
        except Exception as e:
            qry("UPDATE buscas_ia SET status='erro',erro=%s,finalizado_em=NOW() WHERE id=%s", (str(e),bid), commit=True, fetch=None)
    threading.Thread(target=rodar_busca, args=(leg_id['id'] if leg_id else 0, consulta), daemon=True).start()
    return jsonify({'success':True,'message':'Busca iniciada. A legislacao sera adicionada aos pendentes quando encontrada.'})

# -- Pendentes de aprovacao --
@app.route('/api/legislacoes/pendentes', methods=['GET'])
@login_required
def api_leg_pendentes():
    data = qry("SELECT b.*, u.nome as solicitante FROM buscas_ia b LEFT JOIN users u ON b.solicitado_por=u.id WHERE b.status IN ('encontrado','pendente') ORDER BY b.criado_em DESC")
    return jsonify({'success':True,'data':data})

@app.route('/api/legislacoes/pendentes/<int:bid>/aprovar', methods=['POST'])
@admin_required
def api_aprovar_leg_pendente(bid):
    busca = qry("SELECT * FROM buscas_ia WHERE id=%s", (bid,), 'one')
    if not busca: return jsonify({'success':False,'error':'Nao encontrado'}), 404
    qry("UPDATE legislacoes SET pendente_aprovacao=FALSE, aprovado_em=NOW(), aprovado_por=%s WHERE url_original=%s AND pendente_aprovacao=TRUE",
        (session['user_id'], busca.get('resultado_url','')), commit=True, fetch=None)
    qry("UPDATE buscas_ia SET status='aprovado' WHERE id=%s", (bid,), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/legislacoes/pendentes/<int:bid>/rejeitar', methods=['POST'])
@admin_required
def api_rejeitar_leg_pendente(bid):
    qry("DELETE FROM legislacoes WHERE url_original=(SELECT resultado_url FROM buscas_ia WHERE id=%s) AND pendente_aprovacao=TRUE", (bid,), commit=True, fetch=None)
    qry("UPDATE buscas_ia SET status='rejeitado' WHERE id=%s", (bid,), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/legislacoes/relacoes-todas', methods=['GET'])
@login_required
def api_relacoes_todas():
    """Retorna todas as relações entre legislações (para árvore genealógica)."""
    try:
        rows = qry("""SELECT lr.legislacao_pai_id, lr.legislacao_filha_id, lr.tipo_relacao,
                              lp.numero as pai_numero, lp.ano as pai_ano,
                              lf.numero as filha_numero, lf.ano as filha_ano
                       FROM legislacao_relacoes lr
                       JOIN legislacoes lp ON lr.legislacao_pai_id = lp.id
                       JOIN legislacoes lf ON lr.legislacao_filha_id = lf.id
                       ORDER BY lp.ano DESC, lf.ano DESC""")
        return jsonify({'success': True, 'data': rows or []})
    except Exception as e:
        return jsonify({'success': True, 'data': []})

# -------------------------------------------------------------------
# API: CONFIGURACOES (tipos, assuntos, email)
# -------------------------------------------------------------------

@app.route('/api/config/tipos-legislacao', methods=['GET'])
@login_required
def api_get_tipos(): return jsonify({'success':True,'data':qry("SELECT * FROM tipos_legislacao ORDER BY nome")})

@app.route('/api/config/tipos-legislacao', methods=['POST'])
@editor_required
def api_criar_tipo():
    d = request.json or {}
    nome = d.get('nome','').strip()
    if not nome: return jsonify({'success':False,'error':'Nome obrigatorio'}), 400
    try:
        qry("INSERT INTO tipos_legislacao (nome,descricao,criado_por) VALUES (%s,%s,%s)",
            (nome, d.get('descricao',''), session['user_id']), commit=True, fetch=None)
        return jsonify({'success':True})
    except: return jsonify({'success':False,'error':'Tipo ja existe'}), 400

@app.route('/api/config/tipos-legislacao/<int:tid>', methods=['DELETE'])
@admin_required
def api_del_tipo(tid):
    qry("DELETE FROM tipos_legislacao WHERE id=%s", (tid,), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/config/assuntos', methods=['GET'])
@login_required
def api_get_assuntos(): return jsonify({'success':True,'data':qry("SELECT * FROM assuntos_legislacao ORDER BY nome")})

@app.route('/api/config/assuntos', methods=['POST'])
@editor_required
def api_criar_assunto():
    d = request.json or {}
    nome = d.get('nome','').strip()
    if not nome: return jsonify({'success':False,'error':'Nome obrigatorio'}), 400
    try:
        qry("INSERT INTO assuntos_legislacao (nome,descricao,criado_por) VALUES (%s,%s,%s)",
            (nome, d.get('descricao',''), session['user_id']), commit=True, fetch=None)
        return jsonify({'success':True})
    except: return jsonify({'success':False,'error':'Assunto ja existe'}), 400

@app.route('/api/config/assuntos/<int:aid>', methods=['DELETE'])
@admin_required
def api_del_assunto(aid):
    qry("DELETE FROM assuntos_legislacao WHERE id=%s", (aid,), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/config/email', methods=['GET'])
@admin_required
def api_config_email():
    return jsonify({'success':True,'data':{'host':os.getenv('EMAIL_HOST',''),'port':os.getenv('EMAIL_PORT',''),'user':os.getenv('EMAIL_USER',''),'from':os.getenv('EMAIL_FROM','')}})

@app.route('/api/config/email/testar', methods=['POST'])
@admin_required
def api_testar_email():
    try:
        enviar_email_generico(session['email'], 'Teste UrbanLex', '<p>E-mail de teste do UrbanLex.</p>')
        return jsonify({'success':True})
    except Exception as e:
        return jsonify({'success':False,'error':str(e)}), 500

# -------------------------------------------------------------------
# API: PARAMETROS URBANISTICOS
# -------------------------------------------------------------------

@app.route('/api/municipios')
@login_required
def api_municipios(): return jsonify(qry("SELECT DISTINCT municipio as nome, estado FROM zonas_urbanas ORDER BY estado, municipio"))

@app.route('/api/zonas/<municipio>')
@login_required
def api_zonas(municipio): return jsonify(qry("SELECT zona, subzona FROM zonas_urbanas WHERE municipio=%s ORDER BY zona, subzona", (municipio,)))

@app.route('/api/zona/<municipio>/<zona>')
@login_required
def api_zona(municipio, zona):
    z = qry("SELECT * FROM zonas_urbanas WHERE municipio=%s AND zona=%s LIMIT 1", (municipio,zona), 'one')
    if not z: return jsonify({'error':'Zona nao encontrada'}), 404
    return jsonify(z)

@app.route('/api/zonas/todas')
@login_required
def api_todas_zonas():
    page = int(request.args.get('page',1)); per_page = int(request.args.get('per_page',100))
    offset = (page-1)*per_page
    total = qry("SELECT COUNT(*) as n FROM zonas_urbanas", fetch='one')['n']
    data = qry("SELECT * FROM zonas_urbanas ORDER BY estado, municipio, zona LIMIT %s OFFSET %s", (per_page, offset))
    return jsonify({'success':True,'data':data,'total':total,'page':page})

@app.route('/api/zona', methods=['POST'])
@editor_required
def api_criar_zona():
    d = request.json or {}
    mun = d.get('municipio','').strip()
    zona = d.get('zona','').strip()
    if not mun or not zona: return jsonify({'success':False,'error':'Municipio e zona obrigatorios'}), 400
    cols = [k for k in d.keys() if k not in ('id','criado_em','atualizado_em','atualizado_por')]
    vals = [d[c] for c in cols]
    cols_str = ','.join(cols); ph = ','.join(['%s']*len(cols))
    upd = ','.join(f"{c}=EXCLUDED.{c}" for c in cols if c not in ('municipio','zona','subzona'))
    if upd:
        upd_full = upd + ",atualizado_em=NOW(),atualizado_por=%s"
    else:
        upd_full = "atualizado_em=NOW(),atualizado_por=%s"
    qry(f"INSERT INTO zonas_urbanas ({cols_str}) VALUES ({ph}) ON CONFLICT (municipio,zona,subzona) DO UPDATE SET {upd_full}",
        vals + [session['user_id']], commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/calcular-area', methods=['POST'])
@login_required
def api_calcular_area():
    try:
        from calculador_area_computavel import calcular_areas_computaveis
        resultado = calcular_areas_computaveis(request.json or {})
        return jsonify({'success':True,'resultado':resultado})
    except Exception as e:
        return jsonify({'success':False,'error':str(e)}), 500

# -------------------------------------------------------------------
# API: MONITORAMENTO
# -------------------------------------------------------------------

@app.route('/api/monitor/municipios', methods=['GET'])
@login_required
def api_monitor_municipios(): return jsonify({'success':True,'data':qry("SELECT * FROM municipios WHERE ativo=TRUE ORDER BY nome")})

@app.route('/api/monitor/municipios', methods=['POST'])
@admin_required
def api_monitor_add_municipio():
    d = request.json or {}
    nome = d.get('nome','').strip()
    if not nome: return jsonify({'success':False,'error':'Nome obrigatorio'}), 400
    row = qry("INSERT INTO municipios (nome,estado,url_diario,tipo_site) VALUES (%s,%s,%s,%s) ON CONFLICT DO NOTHING RETURNING id",
        (nome, d.get('estado',''), d.get('url_diario',''), d.get('tipo_site','generico')), 'one', commit=True)
    mun_id = row['id'] if row else None
    if mun_id and d.get('url_diario'):
        def _detectar(mid):
            try:
                from modulos.scraper_inteligente import detectar_e_salvar_perfil
                detectar_e_salvar_perfil(mid)
            except Exception as e:
                import logging; logging.getLogger(__name__).error(f"Deteccao perfil falhou: {e}")
        threading.Thread(target=_detectar, args=(mun_id,), daemon=True).start()
    return jsonify({'success':True,'id':mun_id})

@app.route('/api/monitor/municipios/<int:mid>', methods=['DELETE'])
@admin_required
def api_monitor_del_municipio(mid):
    qry("UPDATE municipios SET ativo=FALSE WHERE id=%s", (mid,), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/monitor/municipios/<int:mid>/perfil')
@login_required
def api_monitor_perfil(mid):
    perfil = qry("""SELECT p.*, m.nome as municipio_nome,
                           COALESCE(m.url_diario, p.url_base) as url_diario
                    FROM perfis_diario p JOIN municipios m ON m.id=p.municipio_id
                    WHERE p.municipio_id=%s""", (mid,), 'one')
    return jsonify({'success':True,'data':perfil})

@app.route('/api/monitor/municipios/<int:mid>/url', methods=['POST'])
@editor_required
def api_monitor_salvar_url(mid):
    d = request.json or {}
    url = d.get('url', '').strip()
    if not url:
        return jsonify({'success': False, 'error': 'URL não informada'}), 400
    # Salvar em municipios.url_diario
    qry("UPDATE municipios SET url_diario=%s, atualizado_em=NOW() WHERE id=%s",
        (url, mid), commit=True, fetch=None)
    # Também atualizar perfis_diario.url_base se existir
    perfil = qry("SELECT id FROM perfis_diario WHERE municipio_id=%s", (mid,), 'one')
    if perfil:
        qry("UPDATE perfis_diario SET url_base=%s WHERE municipio_id=%s",
            (url, mid), commit=True, fetch=None)
    else:
        qry("""INSERT INTO perfis_diario (municipio_id, url_base, status_deteccao)
               VALUES (%s, %s, 'pendente')""",
            (mid, url), commit=True, fetch=None)
    return jsonify({'success': True, 'message': 'URL salva!'})

@app.route('/api/monitor/municipios/<int:mid>/ibge', methods=['POST'])
@editor_required
def api_monitor_salvar_ibge(mid):
    """Salva o código IBGE do município (necessário para API Querido Diário)."""
    d = request.json or {}
    codigo = str(d.get('codigo_ibge', '')).strip()
    if not codigo or not codigo.isdigit():
        return jsonify({'success': False, 'error': 'Código IBGE inválido (deve ser numérico)'}), 400
    qry("UPDATE municipios SET codigo_ibge=%s WHERE id=%s",
        (codigo, mid), commit=True, fetch=None)
    return jsonify({'success': True, 'message': f'Código IBGE {codigo} salvo!'})

@app.route('/api/monitor/municipios/<int:mid>/diagnostico', methods=['POST'])
@editor_required
def api_monitor_diagnostico(mid):
    """
    Executa uma busca diagnóstica SÍNCRONA e retorna cada passo com detalhes.
    Permite ver exatamente o que o scraper faz (ou não faz).
    """
    import traceback, time
    d = request.json or {}
    dias = int(d.get('dias', 7))  # buscar últimos N dias por padrão
    log_steps = []

    def step(emoji, msg, data=None):
        log_steps.append({'emoji': emoji, 'msg': msg, 'data': data})

    # 1. Município
    mun = qry("SELECT * FROM municipios WHERE id=%s", (mid,), 'one')
    if not mun:
        step('❌', 'Município não encontrado no banco')
        return jsonify({'success': False, 'steps': log_steps})
    step('🏙', f'Município: {mun["nome"]} ({mun.get("estado","?")})')
    step('🔗', f'URL em municipios.url_diario: {mun.get("url_diario") or "NULL ⚠️"}')
    codigo_ibge = mun.get('codigo_ibge', '')
    if not codigo_ibge:
        # Auto-preencher código IBGE
        step('🔍', 'Código IBGE não configurado — buscando automaticamente...')
        try:
            from modulos.scraper_inteligente import _obter_codigo_ibge
            codigo_ibge = _obter_codigo_ibge(mid)
            if codigo_ibge:
                step('✅', f'Código IBGE encontrado e salvo: {codigo_ibge}')
                # Recarregar mun do banco
                mun = qry("SELECT * FROM municipios WHERE id=%s", (mid,), 'one')
            else:
                step('⚠️', 'Código IBGE não encontrado na API do IBGE — Querido Diário indisponível')
        except Exception as e:
            step('⚠️', f'Erro ao buscar código IBGE: {e}')

    if codigo_ibge:
        step('🆔', f'Código IBGE: {codigo_ibge} (Querido Diário API habilitada)')

    # 2. Perfil do diário
    perfil = qry("SELECT * FROM perfis_diario WHERE municipio_id=%s", (mid,), 'one')
    if not perfil:
        step('❌', 'Nenhum perfil_diario encontrado! O scraper não sabe como navegar neste site.')
        step('💡', 'Clique em "Re-detectar Perfil" para que a IA aprenda a navegar no diário.')
        return jsonify({'success': False, 'steps': log_steps})

    step('📋', f'Perfil encontrado (id={perfil["id"]})')
    step('🔗', f'URL em perfis_diario.url_base: {perfil.get("url_base") or "NULL ⚠️"}')
    step('📊', f'Status detecção: {perfil.get("status_deteccao","?")}')
    step('🖥', f'Plataforma: {perfil.get("plataforma","desconhecido")}')
    step('📅', f'Detectado em: {perfil.get("detectado_em","nunca")}')
    step('✅', f'Última execução OK: {perfil.get("ultima_execucao_ok","nunca")}')
    step('❌', f'Falhas consecutivas: {perfil.get("falhas_consecutivas",0)}')

    requer_js = perfil.get('requer_playwright', False)
    requer_captcha = perfil.get('requer_captcha', False)
    requer_login = perfil.get('requer_login', False)
    step('🔧', f'Requer JavaScript: {"Sim" if requer_js else "Não"} | CAPTCHA: {"Sim ⚠️" if requer_captcha else "Não"} | Login: {"Sim ⚠️" if requer_login else "Não"}')

    # 3. Perfil JSON (configuração de navegação)
    pjson = perfil.get('perfil_json') or {}
    if pjson:
        step('⚙️', 'Configuração do perfil (perfil_json):', {
            'url_busca': pjson.get('url_busca', '— não definida'),
            'metodo_busca': pjson.get('metodo_busca', 'GET'),
            'formato_data': pjson.get('formato_data', '?'),
            'parametro_data_inicio': pjson.get('parametro_data_inicio', '— não definido'),
            'parametro_data_fim': pjson.get('parametro_data_fim', '— não definido'),
            'seletor_lista_resultados': pjson.get('seletor_lista_resultados', '— não definido'),
            'seletor_titulo_item': pjson.get('seletor_titulo_item', '— não definido'),
            'seletor_data_item': pjson.get('seletor_data_item', '— não definido'),
            'seletor_link_item': pjson.get('seletor_link_item', '— não definido'),
        })
    else:
        step('⚠️', 'perfil_json está VAZIO — o scraper não tem configuração de navegação!')
        step('💡', 'Clique em "Re-detectar Perfil" para que a IA configure os seletores.')

    # 4. Legislações em monitoramento
    legs = qry("""SELECT id, tipo_nome, numero, ano,
                         data_inicio_monitoramento, data_fim_monitoramento,
                         ultima_verificacao_monitoramento, em_monitoramento
                  FROM legislacoes
                  WHERE (municipio_id=%s OR municipio_nome=%s)
                    AND pendente_aprovacao = FALSE
                  ORDER BY em_monitoramento DESC, ano DESC""",
               (mid, mun['nome']))
    if legs:
        for l in (legs or []):
            status = '🟢 Monitorando' if l.get('em_monitoramento') else '⚪ Inativo'
            titulo = f"{l.get('tipo_nome','Lei')} {l.get('numero','?')}/{l.get('ano','?')}"
            step('📜', f'{titulo} — {status} | De: {l.get("data_inicio_monitoramento","não definido")} | Verificado até: {l.get("ultima_verificacao_monitoramento","nunca")}')
    else:
        step('⚠️', 'Nenhuma legislação cadastrada para este município')

    # 5. Teste real de busca (v5: por legislação)
    from datetime import date as _date, timedelta as _td
    dt_fim = _date.today()
    dt_ini = dt_fim - _td(days=dias)

    # Encontrar primeira legislação monitorada para teste
    leg_teste = None
    if legs:
        for l in legs:
            if l.get('em_monitoramento'):
                leg_teste = l
                break
        if not leg_teste:
            leg_teste = legs[0]  # fallback: primeira lei cadastrada

    if leg_teste:
        leg_titulo = f"{leg_teste.get('tipo_nome','Lei')} {leg_teste.get('numero','?')}/{leg_teste.get('ano','?')}"

        # 5a. Teste direto da API Querido Diário
        codigo_ibge = mun.get('codigo_ibge', '')
        if codigo_ibge:
            step('🌐', f'Testando API Querido Diário (IBGE: {codigo_ibge})...')
            try:
                import requests as _req
                termo_qd = f"{leg_teste.get('tipo_nome','Lei')} {leg_teste.get('numero','')}"
                qd_urls = [
                    'https://queridodiario.ok.org.br/api/gazettes',
                    'https://api.queridodiario.ok.org.br/gazettes',
                ]
                qd_ok = False
                for qd_url in qd_urls:
                    try:
                        t0 = time.time()
                        qd_resp = _req.get(qd_url, params={
                            'territory_ids': codigo_ibge,
                            'querystring': f'"{termo_qd}"',
                            'published_since': dt_ini.strftime('%Y-%m-%d'),
                            'published_until': dt_fim.strftime('%Y-%m-%d'),
                            'size': 3,
                        }, timeout=15, headers={'Accept': 'application/json'})
                        elapsed_qd = round(time.time() - t0, 2)

                        if qd_resp.status_code == 200:
                            qd_data = qd_resp.json()
                            qd_total = qd_data.get('total_gazettes', 0)
                            step('✅', f'QD API OK em {elapsed_qd}s: {qd_total} resultado(s) para "{termo_qd}" ({qd_url.split("//")[1][:30]})')
                            qd_ok = True
                            break
                        else:
                            step('⚠️', f'QD API HTTP {qd_resp.status_code} em {elapsed_qd}s ({qd_url.split("//")[1][:30]})')
                    except _req.exceptions.ConnectionError as e:
                        step('❌', f'QD API conexão recusada: {qd_url.split("//")[1][:30]} — {str(e)[:100]}')
                    except _req.exceptions.Timeout:
                        step('❌', f'QD API timeout: {qd_url.split("//")[1][:30]}')
                    except Exception as e:
                        step('❌', f'QD API erro: {str(e)[:150]}')

                if not qd_ok:
                    step('💡', 'API Querido Diário inacessível neste servidor. O scraper usará fallback (requests/Playwright).')
            except Exception as e:
                step('❌', f'Erro teste QD: {e}')

        # 5b. Busca principal
        step('🚀', f'Busca de teste: "{leg_titulo}" nos últimos {dias} dias...')

        try:
            t0 = time.time()
            from modulos.scraper_inteligente import buscar_publicacoes_legislacao
            leg_info = {
                'tipo_nome': leg_teste.get('tipo_nome', 'Lei'),
                'numero': leg_teste.get('numero', ''),
                'ano': leg_teste.get('ano', ''),
                'ementa': '',
            }
            resultado = buscar_publicacoes_legislacao(mid, leg_info, dt_ini, dt_fim)
            elapsed = round(time.time() - t0, 2)

            step('⏱', f'Busca concluída em {elapsed}s (método: {resultado.get("metodo","?")})')
            step('📊', f'Resultado: sucesso={resultado.get("sucesso")} | total={resultado.get("total",0)}')
            step('💬', f'Mensagem: {resultado.get("mensagem","?")}')
            termos = resultado.get('termos_usados', [])
            if termos:
                step('🔍', f'Termos buscados: {", ".join(termos[:4])}')

            pubs = resultado.get('publicacoes', [])
            if pubs:
                step('📰', f'{len(pubs)} publicação(ões) encontrada(s):')
                for i, p in enumerate(pubs[:10]):
                    step('  📄', f'{i+1}. {p.get("titulo","sem título")[:120]} | Data: {p.get("data","?")} | URL: {(p.get("url","") or "sem link")[:80]}')
                if len(pubs) > 10:
                    step('...', f'(+ {len(pubs)-10} publicações omitidas)')
            else:
                step('📭', 'Nenhuma publicação encontrada para este termo')
                step('💡', 'Possíveis causas: site requer JavaScript (precisa Playwright), busca sem resultados para o período, ou site bloqueando acesso automatizado.')

        except ImportError as e:
            step('❌', f'Módulo scraper não encontrado: {e}')
        except Exception as e:
            tb = traceback.format_exc()
            step('❌', f'ERRO na busca: {str(e)[:300]}')
            step('🔍', f'Traceback: {tb[:500]}')
    else:
        step('⚠️', 'Nenhuma legislação cadastrada — não é possível testar a busca (o scraper v5 busca por nome da lei)')

    # 6. Teste de acesso HTTP direto à URL
    url_teste = pjson.get('url_busca') or perfil.get('url_base') or mun.get('url_diario')
    if url_teste:
        step('🌐', f'Teste de acesso HTTP direto a: {url_teste}')
        try:
            import requests as _req
            t0 = time.time()
            resp = _req.get(url_teste, timeout=15, headers={
                'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36'
            })
            elapsed = round(time.time() - t0, 2)
            step('📡', f'HTTP {resp.status_code} em {elapsed}s | Content-Type: {resp.headers.get("Content-Type","?")} | Tamanho: {len(resp.text)} chars')
            if resp.status_code == 200:
                # Mostrar um trecho do HTML
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(resp.text, 'html.parser')
                title = soup.title.string.strip() if soup.title else 'sem <title>'
                step('📄', f'Título da página: {title}')
                # Contar formulários e links
                forms = soup.find_all('form')
                links = soup.find_all('a', href=True)
                step('🔎', f'Encontrados: {len(forms)} formulário(s), {len(links)} link(s)')
            else:
                step('⚠️', f'Resposta não-200 — site pode estar bloqueando ou URL incorreta')
        except Exception as e:
            step('❌', f'Falha ao acessar URL: {str(e)[:200]}')

    return jsonify({'success': True, 'steps': log_steps})

@app.route('/api/monitor/municipios/<int:mid>/detectar', methods=['POST'])
@editor_required
def api_monitor_detectar_perfil(mid):
    forcar = (request.json or {}).get('forcar', False)
    def _detectar():
        try:
            from modulos.scraper_inteligente import detectar_e_salvar_perfil
            detectar_e_salvar_perfil(mid, forcar_redeteccao=forcar)
        except Exception as e:
            import logging; logging.getLogger(__name__).error(f"Deteccao perfil: {e}")
    threading.Thread(target=_detectar, daemon=True).start()
    return jsonify({'success':True,'message':'Deteccao iniciada. Aguarde alguns instantes.'})

# --- Histórico por município ---
@app.route('/api/monitor/municipios/<int:mid>/historico')
@login_required
def api_monitor_municipio_historico(mid):
    # Histórico detalhado por legislação para este município
    try:
        data = qry("""
            SELECT mll.data, mll.status, mll.sucesso, mll.alteracoes_detectadas,
                   mll.publicacoes_encontradas, mll.publicacoes_analisadas,
                   mll.publicacoes_duplicadas,
                   mll.metodo_busca, mll.mensagem, mll.erro,
                   mll.legislacao_id,
                   COALESCE(l.tipo_nome,'Lei') || ' nº ' || COALESCE(l.numero,'?') || '/' || COALESCE(l.ano::text,'?') as legislacao_titulo,
                   EXTRACT(EPOCH FROM (mll.finalizada_em - mll.iniciada_em))::int as duracao_seg
            FROM monitoramento_legislacao_log mll
            LEFT JOIN legislacoes l ON l.id = mll.legislacao_id
            WHERE mll.municipio_id = %s
            ORDER BY mll.data DESC LIMIT 50
        """, (mid,))
        return jsonify({'success': True, 'data': data or []})
    except Exception:
        try:
            data = qry("""
                SELECT iniciada_em as data, status,
                       CASE WHEN status='concluido' THEN true WHEN status='erro' THEN false ELSE null END as sucesso,
                       alteracoes_detectadas,
                       EXTRACT(EPOCH FROM (finalizada_em - iniciada_em))::int as duracao_seg
                FROM scheduler_execucoes
                ORDER BY iniciada_em DESC LIMIT 30
            """)
            return jsonify({'success': True, 'data': data or []})
        except Exception:
            return jsonify({'success': True, 'data': []})

# --- Executar monitoramento de um município ---
@app.route('/api/monitor/municipios/<int:mid>/executar', methods=['POST'])
@editor_required
def api_monitor_executar_municipio(mid):
    mun = qry("SELECT * FROM municipios WHERE id=%s AND ativo=TRUE", (mid,), 'one')
    if not mun:
        return jsonify({'success': False, 'error': 'Município não encontrado'}), 404
    # Verificar se há legislações em monitoramento com data definida
    legs = qry("""SELECT id, tipo_nome, numero, ano, data_inicio_monitoramento
                  FROM legislacoes
                  WHERE (municipio_id=%s OR municipio_nome=%s)
                    AND em_monitoramento = TRUE
                    AND pendente_aprovacao = FALSE""",
               (mid, mun['nome']))
    if not legs:
        return jsonify({'success': False,
            'error': 'Nenhuma legislação em monitoramento neste município. Ative o monitoramento e defina o período primeiro.'}), 400
    sem_data = [l for l in legs if not l.get('data_inicio_monitoramento')]
    if sem_data:
        nomes = ', '.join([f"{l['tipo_nome']} {l['numero']}/{l['ano']}" for l in sem_data[:3]])
        return jsonify({'success': False,
            'error': f'Defina a data de início do monitoramento para: {nomes}'}), 400
    def _executar(municipio_id, user_id):
        try:
            from modulos.scheduler_integrado import executar_ciclo_completo
            executar_ciclo_completo(disparado_por=user_id, municipio_id=municipio_id)
        except Exception as e:
            import logging; logging.getLogger(__name__).error(f"Monitoramento municipio {municipio_id}: {e}")
    uid = session.get('user_id')
    threading.Thread(target=_executar, args=(mid, uid), daemon=True).start()
    return jsonify({'success': True,
        'message': f'Monitoramento iniciado para {len(legs)} legislação(ões) de {mun["nome"]}'})

# --- Legislações de um município (para modal de perfil) ---
@app.route('/api/monitor/municipios/<int:mid>/legislacoes')
@login_required
def api_monitor_municipio_legislacoes(mid):
    data = qry("""SELECT l.id, l.tipo_nome, l.numero, l.ano, l.ementa,
                         l.em_monitoramento,
                         to_char(l.data_inicio_monitoramento, 'YYYY-MM-DD') as data_inicio_monitoramento,
                         to_char(l.data_fim_monitoramento, 'YYYY-MM-DD') as data_fim_monitoramento,
                         to_char(l.ultima_verificacao_monitoramento, 'YYYY-MM-DD') as ultima_verificacao_monitoramento,
                         l.municipio_nome, l.status,
                         COALESCE(tl.nome, l.tipo_nome) as tipo_display
                  FROM legislacoes l
                  LEFT JOIN tipos_legislacao tl ON l.tipo_id=tl.id
                  WHERE (l.municipio_id=%s OR l.municipio_nome=(SELECT nome FROM municipios WHERE id=%s))
                    AND l.pendente_aprovacao = FALSE
                  ORDER BY l.em_monitoramento DESC, l.ano DESC, l.numero""",
               (mid, mid))
    return jsonify({'success': True, 'data': data or []})

@app.route('/api/monitor/historico')
@login_required
def api_monitor_historico():
    data = qry("""SELECT se.*, u.nome as usuario_nome
                  FROM scheduler_execucoes se
                  LEFT JOIN users u ON se.disparado_por=u.id
                  ORDER BY iniciada_em DESC LIMIT 50""")
    for row in (data or []):
        row.pop('log_erros', None)
    return jsonify({'success':True,'data':data})

@app.route('/api/monitor/historico/<int:exec_id>')
@login_required
def api_monitor_execucao_detalhe(exec_id):
    row = qry("""SELECT se.*, u.nome as usuario_nome
                 FROM scheduler_execucoes se
                 LEFT JOIN users u ON se.disparado_por=u.id
                 WHERE se.id=%s""", (exec_id,), 'one')
    if not row:
        return jsonify({'success':False,'error':'Execucao nao encontrada'}), 404
    return jsonify({'success':True,'data':row})

@app.route('/api/monitor/status')
@login_required
def api_monitor_status_resumo():
    ultima = qry("""SELECT id, iniciada_em, finalizada_em, status,
                           municipios_processados, municipios_ok, municipios_erro,
                           alteracoes_detectadas, erros, email_enviado
                    FROM scheduler_execucoes
                    ORDER BY iniciada_em DESC LIMIT 1""", fetch='one')
    if not ultima:
        return jsonify({'success':True,'data':{'status':'nunca_executou','ultima':None}})
    if ultima['status'] == 'rodando':
        saude = 'executando'
    elif ultima['status'] == 'erro' or (ultima.get('municipios_erro',0) == ultima.get('municipios_processados',1) and ultima.get('municipios_processados',0) > 0):
        saude = 'erro'
    elif ultima.get('municipios_erro', 0) > 0 or ultima.get('erros', 0) > 0:
        saude = 'parcial'
    else:
        saude = 'ok'
    ultima['saude'] = saude
    return jsonify({'success':True,'data':ultima})

# -- FIX 7: Query alteracoes usa COALESCE para municipio_nome --
@app.route('/api/monitor/alteracoes')
@login_required
def api_monitor_alteracoes():
    # Buscar de alteracoes_pendentes (gravado pelo scheduler v4.0)
    data = qry("""SELECT ap.id, ap.tipo_alteracao, ap.descricao, ap.status,
                         ap.detectada_em as data_deteccao,
                         ap.validada_em, ap.observacoes,
                         COALESCE(m.nome, l.municipio_nome) as municipio,
                         COALESCE(tl.nome, l.tipo_nome, 'Lei') || ' nº ' ||
                            COALESCE(l.numero, '?') || '/' || COALESCE(l.ano::text, '?') as legislacao_titulo,
                         CASE WHEN ap.status='aprovado' THEN true ELSE false END as aprovado
                  FROM alteracoes_pendentes ap
                  LEFT JOIN legislacoes l ON ap.legislacao_id = l.id
                  LEFT JOIN tipos_legislacao tl ON l.tipo_id = tl.id
                  LEFT JOIN municipios m ON ap.municipio_id = m.id
                  ORDER BY ap.detectada_em DESC LIMIT 100""")
    # Fallback: se não houver dados em alteracoes_pendentes, tentar tabela alteracoes antiga
    if not data:
        data = qry("""SELECT a.*, l.numero, l.ano,
                             COALESCE(tl.nome, l.tipo_nome, 'Lei') || ' nº ' ||
                                COALESCE(l.numero, '?') || '/' || COALESCE(l.ano::text, '?') as legislacao_titulo,
                             COALESCE(m.nome, l.municipio_nome) as municipio
                      FROM alteracoes a
                      JOIN legislacoes l ON a.legislacao_id=l.id
                      LEFT JOIN tipos_legislacao tl ON l.tipo_id=tl.id
                      LEFT JOIN municipios m ON l.municipio_id=m.id
                      ORDER BY a.data_deteccao DESC LIMIT 100""")
    return jsonify({'success':True,'data':data or []})

@app.route('/api/monitor/scheduler', methods=['GET'])
@login_required
def api_scheduler_get():
    cfg = qry("SELECT * FROM scheduler_config ORDER BY id LIMIT 1", fetch='one')
    return jsonify({'success':True,'data':cfg or {}})

@app.route('/api/monitor/scheduler', methods=['POST'])
@admin_required
def api_scheduler_update():
    d = request.json or {}
    upd, vals = [], []
    for campo, col in [('horario','horario_execucao'), ('status','status')]:
        if d.get(campo) is not None:
            upd.append(f"{col}=%s"); vals.append(d[campo])
    for campo, col in [('debug_ativo','debug_ativo'), ('email_relatorio','email_relatorio')]:
        if campo in d:
            valor = d[campo]
            if isinstance(valor, str): valor = valor.lower() == 'true'
            upd.append(f"{col}=%s"); vals.append(bool(valor))
    if upd:
        vals.append(session['user_id'])
        qry(f"UPDATE scheduler_config SET {','.join(upd)}, atualizado_em=NOW() WHERE id=1", vals, commit=True, fetch=None)
    cfg = qry("SELECT * FROM scheduler_config ORDER BY id LIMIT 1", fetch='one')
    return jsonify({'success':True,'data':cfg or {}})

@app.route('/api/monitor/scheduler/toggle', methods=['POST'])
@admin_required
def api_scheduler_toggle():
    d = request.json or {}
    campo = d.get('campo')
    if campo not in ('debug_ativo', 'email_relatorio'):
        return jsonify({'success':False,'error':'Campo invalido'}), 400
    cfg = qry("SELECT * FROM scheduler_config ORDER BY id LIMIT 1", fetch='one')
    if not cfg:
        return jsonify({'success':False,'error':'Configuracao nao encontrada'}), 404
    novo_valor = not bool(cfg[campo])
    qry(f"UPDATE scheduler_config SET {campo}=%s, atualizado_em=NOW() WHERE id=1",
        (novo_valor,), commit=True, fetch=None)
    return jsonify({'success':True, 'campo': campo, 'valor': novo_valor})

@app.route('/api/monitor/executar', methods=['POST'])
@editor_required
def api_monitor_executar():
    try:
        from modulos.scheduler_integrado import executar_ciclo_completo
        from modulos.bridge_integracao import processar_alteracao_detectada
        uid = session['user_id']
        threading.Thread(
            target=executar_ciclo_completo,
            kwargs={'bridge_callback': processar_alteracao_detectada, 'disparado_por': uid},
            daemon=True
        ).start()
        return jsonify({'success':True,'message':'Execucao iniciada. Acompanhe em Historico.'})
    except Exception as e:
        return jsonify({'success':False,'error':str(e)}), 500

# -------------------------------------------------------------------
# API: INTEGRACOES (fila de parametros)
# -------------------------------------------------------------------

@app.route('/api/integracao/pendentes')
@login_required
def api_integ_pendentes():
    data = qry("""SELECT i.*, l.numero as legislacao_num, l.ano as legislacao_ano, tl.nome as tipo_nome,
        CONCAT(l.numero,'/',l.ano) as legislacao_ref
        FROM integracao_atualizacoes i LEFT JOIN legislacoes l ON i.legislacao_id=l.id
        LEFT JOIN tipos_legislacao tl ON l.tipo_id=tl.id WHERE i.status='pendente' ORDER BY i.criado_em DESC""")
    return jsonify({'success':True,'data':data})

@app.route('/api/integracao/<int:iid>')
@login_required
def api_integ_detalhe(iid):
    i = qry("""SELECT i.*, CONCAT(l.numero,'/',l.ano) as legislacao_ref FROM integracao_atualizacoes i
        LEFT JOIN legislacoes l ON i.legislacao_id=l.id WHERE i.id=%s""", (iid,), 'one')
    if not i: return jsonify({'success':False,'error':'Nao encontrado'}), 404
    if i.get('parametros_json') and isinstance(i['parametros_json'], str):
        try: i['parametros_json'] = json.loads(i['parametros_json'])
        except: pass
    return jsonify({'success':True,'data':i})

@app.route('/api/integracao/<int:iid>/aprovar', methods=['POST'])
@editor_required
def api_integ_aprovar(iid):
    i = qry("SELECT * FROM integracao_atualizacoes WHERE id=%s", (iid,), 'one')
    if not i: return jsonify({'success':False,'error':'Nao encontrado'}), 404
    params = i.get('parametros_json') or {}
    if isinstance(params, str):
        try: params = json.loads(params)
        except: params = {}
    if params and i.get('municipio') and i.get('zona'):
        cols = list(params.keys()) + ['municipio','zona','subzona','atualizado_em','atualizado_por']
        vals = list(params.values()) + [i['municipio'],i['zona'],i.get('subzona',''),datetime.now(),session['user_id']]
        ph = ','.join(['%s']*len(cols))
        upd = ','.join(f"{c}=EXCLUDED.{c}" for c in params.keys()) + ',atualizado_em=EXCLUDED.atualizado_em,atualizado_por=EXCLUDED.atualizado_por'
        qry(f"INSERT INTO zonas_urbanas ({','.join(cols)}) VALUES ({ph}) ON CONFLICT (municipio,zona,subzona) DO UPDATE SET {upd}",
            vals, commit=True, fetch=None)
    qry("UPDATE integracao_atualizacoes SET status='aprovado',revisado_em=NOW(),revisado_por=%s WHERE id=%s",
        (session['user_id'], iid), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/integracao/<int:iid>/rejeitar', methods=['POST'])
@editor_required
def api_integ_rejeitar(iid):
    motivo = (request.json or {}).get('motivo','')
    qry("UPDATE integracao_atualizacoes SET status='rejeitado',revisado_em=NOW(),revisado_por=%s,motivo_rejeicao=%s WHERE id=%s",
        (session['user_id'], motivo, iid), commit=True, fetch=None)
    return jsonify({'success':True})

@app.route('/api/integracao/aprovar-todos', methods=['POST'])
@editor_required
def api_integ_aprovar_todos():
    pendentes = qry("SELECT * FROM integracao_atualizacoes WHERE status='pendente'")
    count = 0
    for i in pendentes:
        params = i.get('parametros_json') or {}
        if isinstance(params, str):
            try: params = json.loads(params)
            except: params = {}
        if params and i.get('municipio') and i.get('zona'):
            cols = list(params.keys()) + ['municipio','zona','subzona','atualizado_em','atualizado_por']
            vals = list(params.values()) + [i['municipio'],i['zona'],i.get('subzona',''),datetime.now(),session['user_id']]
            ph = ','.join(['%s']*len(cols))
            upd = ','.join(f"{c}=EXCLUDED.{c}" for c in params.keys()) + ',atualizado_em=EXCLUDED.atualizado_em'
            qry(f"INSERT INTO zonas_urbanas ({','.join(cols)}) VALUES ({ph}) ON CONFLICT (municipio,zona,subzona) DO UPDATE SET {upd}", vals, commit=True, fetch=None)
        qry("UPDATE integracao_atualizacoes SET status='aprovado',revisado_em=NOW(),revisado_por=%s WHERE id=%s",
            (session['user_id'], i['id']), commit=True, fetch=None)
        count += 1
    return jsonify({'success':True,'count':count})

# -------------------------------------------------------------------
# API: DASHBOARD
# -------------------------------------------------------------------

@app.route('/api/dashboard/stats')
@login_required
def api_dashboard_stats():
    try:
        leg_total = qry("SELECT COUNT(*) as n FROM legislacoes WHERE pendente_aprovacao=FALSE", fetch='one')['n']
        leg_vig   = qry("SELECT COUNT(*) as n FROM legislacoes WHERE status='vigente' AND pendente_aprovacao=FALSE", fetch='one')['n']
        leg_rev   = qry("SELECT COUNT(*) as n FROM legislacoes WHERE status='revogada'", fetch='one')['n']
        leg_pend  = qry("SELECT COUNT(*) as n FROM buscas_ia WHERE status IN ('encontrado','pendente')", fetch='one')['n']
        mun       = qry("SELECT COUNT(*) as n FROM municipios WHERE ativo=TRUE", fetch='one')['n']
        zonas     = qry("SELECT COUNT(*) as n FROM zonas_urbanas", fetch='one')['n']
        mun_zon   = qry("SELECT COUNT(DISTINCT municipio) as n FROM zonas_urbanas", fetch='one')['n']
        integ     = qry("SELECT COUNT(*) as n FROM integracao_atualizacoes WHERE status='pendente'", fetch='one')['n']
        sched     = qry("SELECT * FROM scheduler_config ORDER BY id LIMIT 1", fetch='one') or {}
        return jsonify({'success':True,'data':{
            'total_legislacoes':leg_total,'leg_vigentes':leg_vig,'leg_revogadas':leg_rev,
            'leg_pendentes':leg_pend,'total_municipios':mun,'total_zonas':zonas,
            'total_municipios_zonas':mun_zon,'integ_pendentes':integ,
            'scheduler_status':sched.get('status','desconhecido'),
            'horario_scheduler':sched.get('horario_execucao','02:00'),
            'ultima_execucao':str(sched.get('ultima_execucao','')) if sched.get('ultima_execucao') else None,
            'proxima_execucao':str(sched.get('proxima_execucao','')) if sched.get('proxima_execucao') else None,
        }})
    except Exception as e:
        return jsonify({'success':False,'error':str(e)}), 500

@app.route('/api/dashboard/feed')
@login_required
def api_dashboard_feed():
    rows = qry("""
        SELECT tipo, titulo, descricao, criado_em FROM (
            SELECT 'legislacao_adicionada' as tipo,
                   tipo_nome||' '||COALESCE(numero,'')||'/'||COALESCE(CAST(ano AS TEXT),'') as titulo,
                   municipio_nome as descricao, criado_em
            FROM legislacoes WHERE pendente_aprovacao=FALSE
            UNION ALL
            SELECT 'execucao_robo', 'Execucao do robo: '||COALESCE(status,''),
                   COALESCE(municipios_ok::text,'0')||' municipios processados', iniciada_em
            FROM scheduler_execucoes
            UNION ALL
            SELECT tipo, mensagem as titulo, '' as descricao, criado_em
            FROM feed_atividades
        ) sub ORDER BY criado_em DESC LIMIT 30
    """)
    return jsonify({'success':True,'data':rows})

@app.route('/api/badges')
@login_required
def api_badges():
    leg  = qry("SELECT COUNT(*) as n FROM buscas_ia WHERE status IN ('encontrado','pendente')", fetch='one')['n']
    intg = qry("SELECT COUNT(*) as n FROM integracao_atualizacoes WHERE status='pendente'", fetch='one')['n']
    usr  = qry("SELECT COUNT(*) as n FROM users WHERE ativo=TRUE AND aprovado=FALSE", fetch='one')['n']
    return jsonify({'leg_pendentes':leg,'param_pendentes':intg,'integ_pendentes':intg,'users_pendentes':usr})

# -------------------------------------------------------------------
# SISTEMA: DIAGNOSTICO E HEALTH
# -------------------------------------------------------------------

# -------------------------------------------------------------------
# API v6: FEED DE ATIVIDADES DO AGENTE
# -------------------------------------------------------------------

@app.route('/api/feed')
@login_required
def api_feed_atividades():
    """Feed de atividades do agente autônomo."""
    limit = request.args.get('limit', 50, type=int)
    tipo = request.args.get('tipo', '')
    try:
        if tipo:
            rows = qry("""SELECT id, tipo, mensagem, detalhes, lida, criado_em
                         FROM feed_atividades WHERE tipo=%s
                         ORDER BY criado_em DESC LIMIT %s""", (tipo, limit))
        else:
            rows = qry("""SELECT id, tipo, mensagem, detalhes, lida, criado_em
                         FROM feed_atividades
                         ORDER BY criado_em DESC LIMIT %s""", (limit,))
        return jsonify({'success': True, 'data': rows or []})
    except Exception as e:
        return jsonify({'success': True, 'data': []})

@app.route('/api/feed/<int:fid>/lida', methods=['POST'])
@login_required
def api_feed_marcar_lida(fid):
    qry("UPDATE feed_atividades SET lida=TRUE WHERE id=%s", (fid,), commit=True, fetch=None)
    return jsonify({'success': True})

@app.route('/api/feed/resumo')
@login_required
def api_feed_resumo():
    """Resumo do feed: contagens por tipo."""
    try:
        rows = qry("""SELECT tipo, COUNT(*) as total,
                             COUNT(*) FILTER (WHERE NOT lida) as nao_lidas
                      FROM feed_atividades
                      WHERE criado_em > NOW() - INTERVAL '7 days'
                      GROUP BY tipo ORDER BY total DESC""")
        nao_lidas = qry("SELECT COUNT(*) as n FROM feed_atividades WHERE NOT lida", fetch='one')
        return jsonify({
            'success': True,
            'nao_lidas': nao_lidas['n'] if nao_lidas else 0,
            'por_tipo': rows or []
        })
    except Exception:
        return jsonify({'success': True, 'nao_lidas': 0, 'por_tipo': []})

# -------------------------------------------------------------------
# API v6: INTEGRAÇÃO COM PLATAFORMA EXTERNA
# -------------------------------------------------------------------

@app.route('/api/integracao-plataforma/executar', methods=['POST'])
@editor_required
def api_executar_integracao():
    """Executa integração manual com plataforma externa."""
    try:
        from modulos.integrador_plataforma import executar_integracao_plataforma
        resultado = executar_integracao_plataforma()
        # Log da integração
        qry("""INSERT INTO integracao_log (tipo, municipios_consultados,
                novos_detectados, legislacoes_cadastradas, detalhes, criado_em)
               VALUES ('manual', %s, %s, %s, %s, NOW())""",
            (resultado['municipios_consultados'], resultado['novos_detectados'],
             resultado['legislacoes_cadastradas'], json.dumps(resultado)),
            commit=True, fetch=None)
        return jsonify({'success': True, 'data': resultado})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500

@app.route('/api/integracao-plataforma/config', methods=['GET'])
@editor_required
def api_integracao_config():
    """Retorna configuração atual da integração."""
    return jsonify({
        'success': True,
        'configurada': bool(os.getenv('PLATAFORMA_API_URL')),
        'url': os.getenv('PLATAFORMA_API_URL', '(não configurada)'),
    })

@app.route('/api/integracao-plataforma/historico')
@login_required
def api_integracao_historico():
    rows = qry("SELECT * FROM integracao_log ORDER BY criado_em DESC LIMIT 20")
    return jsonify({'success': True, 'data': rows or []})

# -------------------------------------------------------------------
# API v6: DIÁRIOS OFICIAIS DESCOBERTOS
# -------------------------------------------------------------------

@app.route('/api/diarios-oficiais')
@login_required
def api_listar_diarios():
    rows = qry("""SELECT d.*, m.nome as municipio_nome
                  FROM diarios_oficiais d
                  LEFT JOIN municipios m ON d.municipio_id=m.id
                  ORDER BY d.created_at DESC""")
    return jsonify({'success': True, 'data': rows or []})

# -------------------------------------------------------------------
# API v6: REGRAS AUTOMÁTICAS (teste manual)
# -------------------------------------------------------------------

@app.route('/api/agente/descobrir-legislacoes', methods=['POST'])
@editor_required
def api_descobrir_legislacoes():
    """Descobre legislações urbanísticas de um município."""
    d = request.json or {}
    municipio = d.get('municipio', '')
    estado = d.get('estado', '')
    if not municipio or not estado:
        return jsonify({'success': False, 'error': 'municipio e estado obrigatórios'}), 400
    try:
        from modulos.descobridor_legislacoes import (
            descobrir_legislacoes_municipio, cadastrar_legislacoes_descobertas
        )
        from modulos.descobridor_diario import descobrir_diario

        diario = descobrir_diario(municipio, estado, 'municipal')
        legislacoes = descobrir_legislacoes_municipio(municipio, estado)
        ids = []
        if legislacoes and d.get('cadastrar', False):
            ids = cadastrar_legislacoes_descobertas(legislacoes, municipio, estado,
                                                     ativar_monitoramento=d.get('monitorar', True))
        return jsonify({
            'success': True,
            'diario': diario,
            'legislacoes_encontradas': legislacoes,
            'legislacoes_cadastradas': ids,
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500

# -------------------------------------------------------------------

# -------------------------------------------------------------------
# BUSCADOR DE LEGISLAÇÕES — Página + API
# -------------------------------------------------------------------
def _extrair_texto_arquivo(fpath, fname, ext, logs, tmp, chamar_llm):
    """Extrai texto de um arquivo de acordo com sua extensão."""
    import os, subprocess
    txt = ""
    fname_display = os.path.basename(fname)

    if ext == '.pdf':
        r = subprocess.run(['pdftotext', fpath, '-'], capture_output=True, text=True, timeout=60)
        if r.returncode == 0 and len(r.stdout.strip()) > 100:
            txt = r.stdout
            logs.append({'nivel': 'anexo', 'msg': f'  ✅ {fname_display}: {len(txt)} chars via pdftotext'})
        else:
            logs.append({'nivel': 'aviso', 'msg': f'  ⚠️ {fname_display}: PDF rasterizado — usando Gemini Vision...'})
            try:
                import base64
                subprocess.run(['gs','-dNOPAUSE','-dBATCH','-sDEVICE=png16m','-r300',
                    f'-sOutputFile={tmp}/page_%03d.png', fpath],
                    capture_output=True, timeout=120)
                pages = sorted([x for x in os.listdir(tmp) if x.startswith('page_') and x.endswith('.png')])
                if pages:
                    from google import genai as _gv
                    from google.genai import types as _gv_types
                    client_v = _gv.Client(api_key=os.environ.get('GEMINI_API_KEY',''))
                    _LOTE = 10
                    _txt_partes = []
                    for _li in range(0, len(pages), _LOTE):
                        _lote_pgs = pages[_li:_li+_LOTE]
                        parts = [_gv_types.Part.from_text(text='Extraia todo o texto deste documento municipal brasileiro (parte):')]
                        for pg in _lote_pgs:
                            _pg_path = os.path.join(tmp, pg)
                            with open(_pg_path, 'rb') as fp:
                                parts.append(_gv_types.Part.from_bytes(data=fp.read(), mime_type='image/png'))
                            try: os.remove(_pg_path)
                            except: pass
                        try:
                            import concurrent.futures as _cf
                            _ex = _cf.ThreadPoolExecutor(max_workers=1)
                            _fut = _ex.submit(client_v.models.generate_content, model='gemini-2.5-flash', contents=parts)
                            try:
                                _resp_lote = _fut.result(timeout=120)
                                if _resp_lote.text: _txt_partes.append(_resp_lote.text)
                            except _cf.TimeoutError:
                                logs.append({'nivel': 'aviso', 'msg': f'  ⚠️ {fname_display}: lote {_li//_LOTE+1} timeout (120s) — pulando'})
                                _fut.cancel()
                            finally:
                                _ex.shutdown(wait=False)
                        except Exception as _el:
                            logs.append({'nivel': 'aviso', 'msg': f'  ⚠️ {fname_display}: lote {_li//_LOTE+1} falhou: {str(_el)[:60]}'})
                    txt = '\n'.join(_txt_partes)
                    logs.append({'nivel': 'anexo', 'msg': f'  ✅  {fname_display}: {len(txt)} chars via Gemini Vision ({len(pages)} pgs, lotes de {_LOTE})'})
            except Exception as ev:
                logs.append({'nivel': 'aviso', 'msg': f'  ❌ {fname_display}: Gemini Vision falhou: {str(ev)[:80]}'})

    elif ext in ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'):
        try:
            from google import genai as _gv
            from google.genai import types as _gv_types
            import mimetypes
            mime = mimetypes.guess_type(fpath)[0] or 'image/jpeg'
            client_v = _gv.Client(api_key=os.environ.get('GEMINI_API_KEY',''))
            with open(fpath, 'rb') as fp:
                img_bytes = fp.read()
            parts = [
                _gv_types.Part.from_text(text='Extraia todo o texto visível nesta imagem de documento municipal brasileiro:'),
                _gv_types.Part.from_bytes(data=img_bytes, mime_type=mime)
            ]
            import concurrent.futures as _cf2
            with _cf2.ThreadPoolExecutor(max_workers=1) as _ex2:
                _fut2 = _ex2.submit(client_v.models.generate_content, model='gemini-2.5-flash', contents=parts)
                try:
                    resp_v = _fut2.result(timeout=120)
                    txt = resp_v.text or ""
                except _cf2.TimeoutError:
                    logs.append({'nivel': 'aviso', 'msg': f'  ⚠️ {fname_display}: Gemini Vision timeout (120s)'})
                    txt = ""
            logs.append({'nivel': 'anexo', 'msg': f'  ✅ {fname_display}: {len(txt)} chars via Gemini Vision (imagem)'})
        except Exception as ev:
            logs.append({'nivel': 'aviso', 'msg': f'  ❌ {fname_display}: Gemini Vision falhou: {str(ev)[:80]}'})

    elif ext == '.docx':
        try:
            import docx as _docx
            doc = _docx.Document(fpath)
            txt = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            logs.append({'nivel': 'anexo', 'msg': f'  ✅ {fname_display}: {len(txt)} chars via python-docx'})
        except Exception as ed:
            logs.append({'nivel': 'aviso', 'msg': f'  ❌ {fname_display}: python-docx falhou: {str(ed)[:80]}'})

    elif ext == '.doc':
        try:
            r = subprocess.run(['antiword', fpath], capture_output=True, text=True, timeout=60)
            if r.returncode == 0 and r.stdout.strip():
                txt = r.stdout
                logs.append({'nivel': 'anexo', 'msg': f'  ✅ {fname_display}: {len(txt)} chars via antiword'})
            else:
                logs.append({'nivel': 'aviso', 'msg': f'  ❌ {fname_display}: antiword falhou ({r.stderr[:60]})'})
        except Exception as ed:
            logs.append({'nivel': 'aviso', 'msg': f'  ❌ {fname_display}: antiword erro: {str(ed)[:80]}'})

    elif ext == '.xlsx':
        try:
            import openpyxl
            wb = openpyxl.load_workbook(fpath, data_only=True)
            linhas = []
            for ws in wb.worksheets:
                linhas.append(f"[Aba: {ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    linha = "\t".join(str(c) if c is not None else "" for c in row)
                    if linha.strip():
                        linhas.append(linha)
            txt = "\n".join(linhas)
            logs.append({'nivel': 'anexo', 'msg': f'  ✅ {fname_display}: {len(txt)} chars via openpyxl'})
        except Exception as ex:
            logs.append({'nivel': 'aviso', 'msg': f'  ❌ {fname_display}: openpyxl falhou: {str(ex)[:80]}'})

    elif ext == '.xls':
        try:
            import xlrd
            wb = xlrd.open_workbook(fpath)
            linhas = []
            for ws in wb.sheets():
                linhas.append(f"[Aba: {ws.name}]")
                for rx in range(ws.nrows):
                    linha = "\t".join(str(ws.cell_value(rx, cx)) for cx in range(ws.ncols))
                    if linha.strip():
                        linhas.append(linha)
            txt = "\n".join(linhas)
            logs.append({'nivel': 'anexo', 'msg': f'  ✅ {fname_display}: {len(txt)} chars via xlrd'})
        except Exception as ex:
            logs.append({'nivel': 'aviso', 'msg': f'  ❌ {fname_display}: xlrd falhou: {str(ex)[:80]}'})

    elif ext in ('.txt', '.csv', '.md', '.html', '.htm', '.xml'):
        try:
            for enc in ('utf-8', 'cp1252', 'latin-1'):
                try:
                    with open(fpath, 'r', encoding=enc) as fp:
                        txt = fp.read()
                    logs.append({'nivel': 'anexo', 'msg': f'  ✅ {fname_display}: {len(txt)} chars via leitura texto ({enc})'})
                    break
                except UnicodeDecodeError:
                    continue
        except Exception as et:
            logs.append({'nivel': 'aviso', 'msg': f'  ❌ {fname_display}: leitura texto falhou: {str(et)[:80]}'})

    else:
        logs.append({'nivel': 'aviso', 'msg': f'  ⏭️ {fname_display}: tipo não suportado ({ext})'})

    return txt


@app.route('/buscador/anexos')
@login_required
def buscador_anexos():
    return render_template('buscador_anexos.html',
        active_page='buscador-anexos',
        active_group='buscador')

@app.route('/mapeamento/georef')
@login_required
def mapeamento_georef():
    return render_template('mapeamento_georef.html',
        active_page='mapeamento-georef',
        active_group='buscador')


def _baixar_google_maps(bbox, img_w, img_h, api_key, logs):
    """Baixa mapa do Google Maps Static API para o bbox do municipio."""
    import requests, numpy as np, cv2
    from PIL import Image
    import io, math

    south, north, west, east = bbox
    center_lat = (south + north) / 2
    center_lon = (west + east) / 2

    # Calcular zoom ideal
    def lat_rad(lat):
        sin = math.sin(lat * math.pi / 180)
        rad = math.log((1 + sin) / (1 - sin)) / 2
        return max(min(rad, math.pi), -math.pi) / 2

    def zoom_level(lat_diff, lon_diff, w, h):
        for z in range(21, 0, -1):
            lat_px = 256 * (2**z) * lat_diff / 360
            lon_px = 256 * (2**z) * lon_diff / 360
            if lat_px <= h * 0.9 and lon_px <= w * 0.9:
                return z
        return 10

    zoom = zoom_level(north - south, east - west, img_w, img_h)
    zoom = max(10, min(zoom, 16))
    # Reduzir zoom em 1 para garantir que todo o municipio apareca
    zoom = max(10, zoom - 1)

    # Google Maps Static API com scale=2 para maior resolucao (1280x1280 efetivo)
    size = 640
    url = (
        f"https://maps.googleapis.com/maps/api/staticmap"
        f"?center={center_lat},{center_lon}"
        f"&zoom={zoom}"
        f"&size={size}x{size}"
        f"&scale=2"
        f"&maptype=satellite"
        f"&key={api_key}"
    )
    logs.append({'nivel': 'info', 'msg': f'  Google Maps: zoom={zoom} center=({center_lat:.4f},{center_lon:.4f})'})

    try:
        r = requests.get(url, timeout=15)
        if r.status_code != 200:
            logs.append({'nivel': 'aviso', 'msg': f'  Google Maps erro: {r.status_code}'})
            return None
        img = Image.open(io.BytesIO(r.content)).convert('RGB')
        # Redimensionar para dimensoes da planta
        img = img.resize((img_w, img_h), Image.LANCZOS)
        img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
        logs.append({'nivel': 'ok', 'msg': f'  Google Maps carregado: {img_w}x{img_h}px'})
        return img_cv
    except Exception as e:
        logs.append({'nivel': 'aviso', 'msg': f'  Google Maps falhou: {str(e)[:80]}'})
        return None

@app.route('/api/mapeamento/preparar', methods=['POST'])
@login_required
def api_mapeamento_preparar():
    """Prepara os mapas: converte planta para PNG — retorna job_id para polling."""
    import threading, uuid, os, tempfile
    f = request.files.get('arquivo')
    municipio = request.form.get('municipio', 'Municipio')
    estado = request.form.get('estado', 'XX')
    if not f:
        return jsonify({'success': False, 'error': 'Nenhum arquivo enviado'})

    # Salvar arquivo imediatamente
    tmp = tempfile.mkdtemp()
    fpath = os.path.join(tmp, f.filename)
    f.save(fpath)
    fname = f.filename

    # Verificar se ja existe resultado cacheado para este arquivo
    import hashlib, json as _json
    _cache_key = hashlib.md5(f"{municipio}{estado}{fname}".encode()).hexdigest()[:12]
    _meta_path = f"/var/www/urbanlex/static/downloads/georef_meta_{_cache_key}.json"
    _planta_path = f"/var/www/urbanlex/static/downloads/georef_planta_{_cache_key}.png"
    if os.path.exists(_meta_path) and os.path.exists(_planta_path):
        with open(_meta_path) as _mf:
            _meta = _json.load(_mf)
        # Retornar cache direto sem job
        _job_id = str(uuid.uuid4())[:12]
        _job = {'logs': [{'nivel':'ok','msg':'✅ Usando cache anterior — mapas prontos!'}], 'done': True,
                'result': {'success': True, 'planta_url': f'/static/downloads/georef_planta_{_cache_key}.png',
                           'meta_key': _cache_key, 'bbox': _meta.get('bbox', [])}}
        _buscador_jobs[_job_id] = _job
        return jsonify({'success': True, 'job_id': _job_id})

    job_id = str(uuid.uuid4())[:12]
    job = {'logs': _LogList(job_id, get_db), 'done': False, 'result': None}
    _buscador_jobs[job_id] = job

    def _run():
        try:
            import subprocess, hashlib, cv2, numpy as np
            from modulos.mapeador_zonas import _buscar_osm
            logs = job['logs']

            # Converter PDF para PNG (pagina 1, 150 DPI para display)
            ext = os.path.splitext(fname)[1].lower()
            logs.append({'nivel': 'info', 'msg': f'📄 Convertendo {fname}...'})
            if ext == '.pdf':
                subprocess.run(['gs','-dNOPAUSE','-dBATCH','-sDEVICE=png16m','-r150',
                    '-dFirstPage=1','-dLastPage=1',
                    f'-sOutputFile={tmp}/mapa_%03d.png', fpath],
                    capture_output=True, timeout=120)
                pages = sorted([x for x in os.listdir(tmp) if x.startswith('mapa_') and x.endswith('.png')])
                if not pages:
                    job['result'] = {'success': False, 'error': 'Falha ao converter PDF'}
                    return
                planta_png = os.path.join(tmp, pages[0])
            else:
                planta_png = fpath

            # Carregar e redimensionar para canvas (150 DPI, max 2000px)
            img_display = cv2.imread(planta_png)
            if img_display is None:
                job['result'] = {'success': False, 'error': 'Imagem inválida'}
                return
            h_d, w_d = img_display.shape[:2]
            if w_d > 2000:
                scale = 2000 / w_d
                img_display = cv2.resize(img_display, (int(w_d*scale), int(h_d*scale)), interpolation=cv2.INTER_AREA)
            planta_key = hashlib.md5(f"{municipio}{estado}{fname}".encode()).hexdigest()[:12]
            planta_dest = f"/var/www/urbanlex/static/downloads/georef_planta_{planta_key}.png"
            cv2.imwrite(planta_dest, img_display)
            h, w = img_display.shape[:2]
            logs.append({'nivel': 'info', 'msg': f'✅ Planta carregada: {w}x{h}px'})

            # Converter em 300 DPI para segmentacao
            planta_full_dest = f"/var/www/urbanlex/static/downloads/georef_planta_full_{planta_key}.png"
            if ext == '.pdf':
                logs.append({'nivel': 'info', 'msg': '🔍 Convertendo em alta resolução para segmentação...'})
                subprocess.run(['gs','-dNOPAUSE','-dBATCH','-sDEVICE=png16m','-r300',
                    '-dFirstPage=1','-dLastPage=1',
                    f'-sOutputFile={tmp}/full_%03d.png', fpath],
                    capture_output=True, timeout=180)
                full_pages = sorted([x for x in os.listdir(tmp) if x.startswith('full_') and x.endswith('.png')])
                if full_pages:
                    import shutil
                    shutil.copy(os.path.join(tmp, full_pages[0]), planta_full_dest)
                else:
                    planta_full_dest = planta_dest
            else:
                import shutil
                shutil.copy(fpath, planta_full_dest)

            img_full = cv2.imread(planta_full_dest)
            orig_h, orig_w = img_full.shape[:2] if img_full is not None else (h, w)
            logs.append({'nivel': 'info', 'msg': f'✅ Alta resolução: {orig_w}x{orig_h}px'})

            # Buscar bbox do municipio via OSM
            logs.append({'nivel': 'info', 'msg': f'🌐 Buscando coordenadas de {municipio}/{estado}...'})
            osm_data = _buscar_osm(municipio, estado, logs)
            if not osm_data:
                job['result'] = {'success': False, 'error': 'Município não encontrado no OSM'}
                return

            bbox = (float(osm_data['_south']), float(osm_data['_north']),
                    float(osm_data['_west']), float(osm_data['_east']))

            # Salvar meta
            import json
            meta = {'municipio': municipio, 'estado': estado, 'bbox': list(bbox),
                    'w': w, 'h': h, 'planta_key': planta_key,
                    'planta_path': planta_dest,
                    'planta_full_path': planta_full_dest,
                    'orig_w': orig_w, 'orig_h': orig_h}
            meta_path = f"/var/www/urbanlex/static/downloads/georef_meta_{planta_key}.json"
            with open(meta_path, 'w') as mf:
                json.dump(meta, mf)

            logs.append({'nivel': 'ok', 'msg': '✅ Mapas prontos!'})
            job['result'] = {
                'success': True,
                'planta_url': f'/static/downloads/georef_planta_{planta_key}.png',
                'meta_key': planta_key,
                'bbox': list(bbox)
            }

        except Exception as e:
            import traceback
            job['logs'].append({'nivel': 'erro', 'msg': f'Erro: {str(e)[:200]}'})
            job['result'] = {'success': False, 'error': str(e)[:200]}
        finally:
            job['done'] = True

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({'success': True, 'job_id': job_id})


@app.route('/api/mapeamento/upscale', methods=['POST'])
@login_required
def api_mapeamento_upscale():
    """Melhora resolucao da planta via Real-ESRGAN (IA)."""
    import threading, json, os
    data = request.get_json()
    planta_key = data.get('planta_key')
    if not planta_key:
        return jsonify({'success': False, 'error': 'planta_key ausente'})

    meta_path = f"/var/www/urbanlex/static/downloads/georef_meta_{planta_key}.json"
    if not os.path.exists(meta_path):
        return jsonify({'success': False, 'error': 'Metadados não encontrados'})

    with open(meta_path) as f:
        meta = json.load(f)

    job_id = str(__import__('uuid').uuid4())[:12]
    job = {'logs': _LogList(job_id, get_db), 'done': False, 'result': None}
    _buscador_jobs[job_id] = job

    def _run():
        logs = job['logs']
        try:
            import cv2, numpy as np
            full_path = meta.get('planta_full_path', meta['planta_path'])
            logs.append({'nivel': 'info', 'msg': '🤖 Verificando Real-ESRGAN...'})

            # Tentar instalar realesrgan se nao disponivel
            try:
                from basicsr.archs.rrdbnet_arch import RRDBNet
                from realesrgan import RealESRGANer
                esrgan_ok = True
            except ImportError:
                logs.append({'nivel': 'info', 'msg': '  Instalando Real-ESRGAN (primeira vez)...'})
                import subprocess
                subprocess.run(['pip', 'install', 'realesrgan', '--break-system-packages', '-q'],
                              capture_output=True, timeout=120)
                try:
                    from basicsr.archs.rrdbnet_arch import RRDBNet
                    from realesrgan import RealESRGANer
                    esrgan_ok = True
                except ImportError:
                    esrgan_ok = False

            if esrgan_ok:
                import requests as _req
                model_path = '/var/www/urbanlex/modelos/RealESRGAN_x4plus.pth'
                os.makedirs('/var/www/urbanlex/modelos', exist_ok=True)

                if not os.path.exists(model_path):
                    logs.append({'nivel': 'info', 'msg': '  Baixando modelo Real-ESRGAN (~67MB)...'})
                    url = 'https://github.com/xinntao/Real-ESRGAN/releases/download/v0.1.0/RealESRGAN_x4plus.pth'
                    r = _req.get(url, stream=True, timeout=120)
                    with open(model_path, 'wb') as f:
                        for chunk in r.iter_content(8192):
                            f.write(chunk)
                    logs.append({'nivel': 'ok', 'msg': '  Modelo baixado!'})

                logs.append({'nivel': 'info', 'msg': '  Aplicando super resolucao 4x...'})
                model = RRDBNet(num_in_ch=3, num_out_ch=3, num_feat=64,
                               num_block=23, num_grow_ch=32, scale=4)
                upsampler = RealESRGANer(scale=4, model_path=model_path,
                                         model=model, tile=256, tile_pad=10,
                                         pre_pad=0, half=False)
                img = cv2.imread(full_path, cv2.IMREAD_COLOR)
                if img is None:
                    raise Exception('Imagem nao encontrada')
                output, _ = upsampler.enhance(img, outscale=2)
                upscaled_path = full_path.replace('.png', '_upscaled.png')
                cv2.imwrite(upscaled_path, output)
                oh, ow = output.shape[:2]
                logs.append({'nivel': 'ok', 'msg': f'  Super resolucao concluida: {ow}x{oh}px'})

                # Atualizar meta
                import json as _json
                meta['planta_full_path'] = upscaled_path
                meta['orig_w'] = ow
                meta['orig_h'] = oh
                with open(meta_path, 'w') as f:
                    _json.dump(meta, f)

                job['result'] = {'success': True, 'msg': f'Resolucao melhorada: {ow}x{oh}px'}
            else:
                # Fallback: bicubic upscale 2x com OpenCV
                logs.append({'nivel': 'aviso', 'msg': '  Real-ESRGAN indisponivel, usando bicubic 2x...'})
                img = cv2.imread(full_path)
                h, w = img.shape[:2]
                upscaled = cv2.resize(img, (w*2, h*2), interpolation=cv2.INTER_CUBIC)
                upscaled_path = full_path.replace('.png', '_upscaled.png')
                cv2.imwrite(upscaled_path, upscaled)
                import json as _json
                meta['planta_full_path'] = upscaled_path
                meta['orig_w'] = w*2
                meta['orig_h'] = h*2
                with open(meta_path, 'w') as f:
                    _json.dump(meta, f)
                logs.append({'nivel': 'ok', 'msg': f'  Bicubic 2x aplicado: {w*2}x{h*2}px'})
                job['result'] = {'success': True, 'msg': f'Resolucao melhorada (bicubic): {w*2}x{h*2}px'}

        except Exception as e:
            import traceback
            logs.append({'nivel': 'erro', 'msg': f'Erro: {str(e)[:200]}'})
            logs.append({'nivel': 'erro', 'msg': traceback.format_exc()[:300]})
            job['result'] = {'success': False, 'error': str(e)}
        finally:
            job['done'] = True

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({'success': True, 'job_id': job_id})

@app.route('/api/mapeamento/georef-analisar', methods=['POST'])
@login_required
def api_mapeamento_georef_analisar():
    """Inicia análise com pontos de referência manuais."""
    import threading, uuid, json, os, cv2, numpy as np
    data = request.get_json()
    municipio = data.get('municipio', '')
    estado = data.get('estado', '')
    pontos = data.get('pontos', {})

    job_id = str(uuid.uuid4())[:12]
    job = {'logs': _LogList(job_id, get_db), 'done': False, 'result': None}
    _buscador_jobs[job_id] = job

    def _run():
        try:
            from modulos.mapeador_zonas import _buscar_osm, _renderizar_osm, _segmentar_zonas, _gerar_kml
            logs = job['logs']

            # Encontrar meta
            import glob
            metas = glob.glob('/var/www/urbanlex/static/downloads/georef_meta_*.json')
            meta = None
            for m in sorted(metas, key=os.path.getmtime, reverse=True):
                with open(m) as f:
                    d = json.load(f)
                if d['municipio'] == municipio and d['estado'] == estado:
                    meta = d
                    break
            if not meta:
                logs.append({'nivel': 'erro', 'msg': 'Metadados não encontrados — recarregue os mapas'})
                job['done'] = True
                return

            img_w, img_h = meta['w'], meta['h']
            bbox = tuple(meta['bbox'])
            south, north, west, east = bbox

            # Calcular transformacao afim a partir dos pontos manuais
            logs.append({'nivel': 'info', 'msg': '📐 Calculando transformação com pontos de referência...'})
            # Log dos pontos recebidos para debug
            for n in range(1,5):
                p = pontos.get(str(n),{})
                pp = p.get('p',{}) or {}
                po = p.get('o',{}) or {}
                pxp = f'{pp.get("xp",0):.1f}' if pp else '?'
                pyp = f'{pp.get("yp",0):.1f}' if pp else '?'
                if po and po.get('isLatLng'):
                    ostr = f'lat={po.get("lat",0):.5f},lng={po.get("lng",0):.5f}'
                elif po:
                    ostr = f'{po.get("xp",0):.1f}%,{po.get("yp",0):.1f}%'
                else:
                    ostr = '?'
                logs.append({'nivel': 'info', 'msg': f'  Ponto {n}: planta=({pxp}%,{pyp}%) mapa=({ostr})'})
            # Usar pixels absolutos da imagem original diretamente
            src_list, dst_list = [], []
            south, north, west, east = bbox
            for n in range(1, 11):
                p = pontos.get(str(n)) or {}
                pp = p.get('p') or {}
                po = p.get('o') or {}
                # Pular pontos incompletos
                if not pp or not po:
                    continue
                # Planta: coordenadas em % da imagem
                xp = pp.get('xp') or pp.get('xPct') or 0
                yp = pp.get('yp') or pp.get('yPct') or 0
                sx = float(xp) / 100 * img_w
                sy = float(yp) / 100 * img_h
                src_list.append([sx, sy])
                # Mapa: lat/lon ou % da imagem
                if po.get('isLatLng'):
                    lat = po.get('lat', 0)
                    lng = po.get('lng', 0)
                    dx = (lng - west) / (east - west) * img_w
                    dy = (north - lat) / (north - south) * img_h
                else:
                    xp2 = po.get('xp') or po.get('xPct') or 0
                    yp2 = po.get('yp') or po.get('yPct') or 0
                    dx = float(xp2) / 100 * img_w
                    dy = float(yp2) / 100 * img_h
                dst_list.append([dx, dy])
                logs.append({'nivel': 'info', 'msg': f'  Ponto {n}: planta=({sx:.0f},{sy:.0f})px → mapa=({dx:.0f},{dy:.0f})px'})
            src_pts = np.array(src_list, dtype=np.float32)
            dst_pts = np.array(dst_list, dtype=np.float32)

            # Homografia com 4 pontos — força todos a coincidirem exatamente
            H, mask = cv2.findHomography(src_pts, dst_pts, method=0)
            if H is None:
                logs.append({'nivel': 'aviso', 'msg': '⚠️ Falha ao calcular transformação — tente outros pontos'})
                job['done'] = True
                return

            # Verificar erro de reprojecao
            erros = []
            for i in range(len(src_pts)):
                pt = np.array([[[src_pts[i][0], src_pts[i][1]]]], dtype=np.float32)
                pt_t = cv2.perspectiveTransform(pt, H).reshape(2)
                erro = np.sqrt((pt_t[0]-dst_pts[i][0])**2 + (pt_t[1]-dst_pts[i][1])**2)
                erros.append(erro)
            erro_medio = np.mean(erros)
            logs.append({'nivel': 'ok', 'msg': f'✅ Homografia calculada: erro médio={erro_medio:.1f}px | erros={[f"{e:.0f}" for e in erros]}'})

            def px_to_ll(x, y):
                lon = west + (x / img_w) * (east - west)
                lat = north - (y / img_h) * (north - south)
                return lat, lon

            # Escalar homografia para resolucao original (para segmentacao)
            orig_w = meta.get('orig_w', img_w)
            orig_h = meta.get('orig_h', img_h)
            scale_x = orig_w / img_w
            scale_y = orig_h / img_h
            S_src = np.array([[scale_x, 0, 0], [0, scale_y, 0], [0, 0, 1]], dtype=np.float64)
            S_dst = np.array([[scale_x, 0, 0], [0, scale_y, 0], [0, 0, 1]], dtype=np.float64)
            H_full = S_dst @ H.astype(np.float64) @ np.linalg.inv(S_src)

            def px_to_ll_full(x, y):
                lon = west + (x / orig_w) * (east - west)
                lat = north - (y / orig_h) * (north - south)
                return lat, lon

            logs.append({'nivel': 'info', 'msg': f'  Resolução original: {orig_w}x{orig_h}px (segmentação usará essa resolução)'})

            # Gerar validacao — planta colorida semitransparente sobre fundo neutro
            img_planta = cv2.imread(meta['planta_path'])
            planta_warp = cv2.warpPerspective(img_planta, H, (img_w, img_h))
            # Fundo branco com planta semitransparente
            fundo = np.ones((img_h, img_w, 3), dtype=np.uint8) * 240
            mask = cv2.cvtColor(planta_warp, cv2.COLOR_BGR2GRAY)
            _, mask = cv2.threshold(mask, 10, 255, cv2.THRESH_BINARY)
            mask_3ch = cv2.cvtColor(mask, cv2.COLOR_GRAY2BGR) / 255.0
            val = (fundo.astype(float) * (1 - mask_3ch * 0.7) + planta_warp.astype(float) * mask_3ch * 0.7).astype(np.uint8)

            # Desenhar pontos de referencia na imagem de validacao
            cores_val = [(0,0,255),(0,200,0),(255,100,0),(200,0,200)]
            for n in range(1, 5):
                p = pontos.get(str(n), {})
                pp = p.get('p', {})
                po = p.get('o', {})
                cor = cores_val[n-1]
                # Ponto na planta (transformado para OSM)
                if pp:
                    px_p = int(pp.get('xp', 0) / 100 * img_w)
                    py_p = int(pp.get('yp', 0) / 100 * img_h)
                    # Transformar ponto da planta para OSM via homografia
                    pt = np.array([[[float(px_p), float(py_p)]]], dtype=np.float32)
                    pt_t = cv2.perspectiveTransform(pt, H).reshape(2)
                    tx, ty = int(pt_t[0]), int(pt_t[1])
                    cv2.circle(val, (tx, ty), 18, cor, -1)
                    cv2.circle(val, (tx, ty), 22, (255,255,255), 3)
                    cv2.putText(val, str(n), (tx-7, ty+6), cv2.FONT_HERSHEY_SIMPLEX, 0.8, (255,255,255), 2)
                # Ponto no OSM (posicao original)
                if po:
                    if po.get('isLatLng'):
                        _lat = po.get('lat', 0)
                        _lng = po.get('lng', 0)
                        ox = int((_lng - west) / (east - west) * img_w)
                        oy = int((north - _lat) / (north - south) * img_h)
                    else:
                        ox = int((po.get('xp') or 0) / 100 * img_w)
                        oy = int((po.get('yp') or 0) / 100 * img_h)
                    cv2.circle(val, (ox, oy), 18, cor, 2)
                    cv2.circle(val, (ox, oy), 22, (255,255,255), 3)
                    # Linha ligando ponto planta transformado ao ponto OSM
                    if pp:
                        cv2.line(val, (tx, ty), (ox, oy), cor, 2)
                    cv2.putText(val, str(n), (ox-7, oy+6), cv2.FONT_HERSHEY_SIMPLEX, 0.8, cor, 2)

            val_path = f"/var/www/urbanlex/static/downloads/validacao_{municipio.replace(' ','_')}.png"
            cv2.imwrite(val_path, val)
            job['result'] = {
                'geo_ok': True,
                'validacao_pontos_url': f"/static/downloads/validacao_{municipio.replace(' ','_')}.png"
            }

            logs.append({'nivel': 'ok', 'msg': '✅ Transformação aplicada com sucesso'})

            # Estagio 4: Extrair legenda via Gemini
            logs.append({'nivel': 'info', 'msg': '🔍 Extraindo legenda da planta via Gemini...'})
            from modulos.mapeador_zonas import _extrair_legenda, _segmentar_zonas, _gerar_kml
            import tempfile as _tmp2
            _tmp_dir = _tmp2.mkdtemp()
            legenda = _extrair_legenda(
                meta['planta_full_path'], 
                os.path.basename(meta['planta_full_path']),
                municipio, estado, logs, _tmp_dir
            )
            if not legenda:
                logs.append({'nivel': 'aviso', 'msg': '⚠️ Não foi possível extrair a legenda'})
            else:
                logs.append({'nivel': 'ok', 'msg': f'✅ {len(legenda)} zonas na legenda'})

                # Estagio 5: Segmentacao por cor (usa resolucao original)
                logs.append({'nivel': 'info', 'msg': '🎨 Segmentando zonas por cor (resolução original)...'})
                img_planta_full = cv2.imread(meta['planta_full_path'])
                zonas_geo = _segmentar_zonas(img_planta_full, legenda, H_full, px_to_ll_full, logs)

                if zonas_geo:
                    # Estagio 6: Gerar KML
                    logs.append({'nivel': 'info', 'msg': '📦 Gerando KML...'})
                    kml_path = f"/var/www/urbanlex/static/downloads/zoneamento_{municipio.replace(' ','_')}.kml"
                    _gerar_kml(zonas_geo, municipio, estado, kml_path)
                    job['result']['kml_url'] = f"/static/downloads/zoneamento_{municipio.replace(' ','_')}.kml"
                    job['result']['zonas_ok'] = True
                    job['result']['kml_ok'] = True
                    job['result']['zonas'] = [{'nome': z['nome'], 'descricao': z.get('descricao',''), 'cor': z['cor_hex'], 'area_km2': z.get('area_km2','—')} for z in zonas_geo]
                    logs.append({'nivel': 'ok', 'msg': f'✅ KML gerado com {len(zonas_geo)} polígonos!'})
                else:
                    logs.append({'nivel': 'aviso', 'msg': '⚠️ Nenhuma zona segmentada'})

        except Exception as e:
            import traceback
            logs.append({'nivel': 'erro', 'msg': f'Erro: {str(e)[:200]}'})
        finally:
            job['done'] = True

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({'success': True, 'job_id': job_id})

@app.route('/mapeamento/zonas')
@login_required
def mapeamento_zonas():
    return render_template('mapeamento_zonas.html',
        active_page='mapeamento-zonas',
        active_group='buscador')

@app.route('/api/mapeamento/iniciar', methods=['POST'])
@login_required
def api_mapeamento_iniciar():
    import threading, uuid, os, tempfile
    job_id = str(uuid.uuid4())[:12]
    f = request.files.get('arquivo')
    municipio = request.form.get('municipio', 'Municipio')
    estado = request.form.get('estado', 'XX')
    if not f:
        return jsonify({'success': False, 'error': 'Nenhum arquivo enviado'})
    tmp = tempfile.mkdtemp()
    fpath = os.path.join(tmp, f.filename)
    f.save(fpath)
    job = {'logs': _LogList(job_id, get_db), 'done': False, 'result': None}
    _buscador_jobs[job_id] = job
    def _run():
        try:
            from modulos.mapeador_zonas import mapear_zonas
            mapear_zonas(fpath, f.filename, municipio, estado, job['logs'], job, tmp)
        except Exception as e:
            job['logs'].append({'nivel': 'erro', 'msg': f'Erro: {str(e)[:200]}'})
        finally:
            job['done'] = True
    threading.Thread(target=_run, daemon=True).start()
    return jsonify({'success': True, 'job_id': job_id})

@app.route('/api/buscador/analisar-anexo', methods=['POST'])
@login_required
def api_analisar_anexo():
    import threading, uuid, json as _json
    job_id = str(uuid.uuid4())[:12]
    arquivos_upload = request.files.getlist('arquivo')
    if not arquivos_upload:
        arquivos_upload = [request.files.get('arquivo')] if request.files.get('arquivo') else []
    municipio = request.form.get('municipio', 'Municipio')
    estado = request.form.get('estado', 'XX')
    if not arquivos_upload:
        return jsonify({'success': False, 'error': 'Nenhum arquivo enviado'})
    import os, tempfile
    tmp = tempfile.mkdtemp()
    fpaths = []
    for f in arquivos_upload:
        if f and f.filename:
            fpath = os.path.join(tmp, f.filename)
            f.save(fpath)
            fpaths.append((fpath, f.filename))
    if not fpaths:
        return jsonify({'success': False, 'error': 'Nenhum arquivo válido enviado'})
    # Compatibilidade: fpath/f.filename apontam para o primeiro arquivo
    fpath, _ = fpaths[0]
    f = type('F', (), {'filename': fpaths[0][1]})()
    job = {'logs': _LogList(job_id, get_db), 'done': False, 'result': None}
    _buscador_jobs[job_id] = job
    def _run():
        try:
            logs = job['logs']
            _anx_log_path = f"/var/www/urbanlex/static/downloads/job_{job_id}.jsonl"
            # Sobrescrever lista para persistir logs em arquivo
            import json as _json_anx
            class _PersistList(list):
                def append(self, item):
                    super().append(item)
                    try:
                        with open(_anx_log_path, 'a', encoding='utf-8') as _lf:
                            _lf.write(_json_anx.dumps(item, ensure_ascii=False) + '\n')
                    except Exception:
                        pass
            logs = _PersistList()
            job['logs'] = logs
            def chamar_llm(prompt, logs, label=""):
                try:
                    from modulos.buscador_urbanistico import chamar_llm as _cllm
                    return _cllm(prompt, logs, label)
                except Exception:
                    from google import genai as _gai
                    import os as _os2
                    client = _gai.Client(api_key=_os2.environ.get('GEMINI_API_KEY',''))
                    resp = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                    return resp.text
            # Importar funcao de extracao de anexos
            import sys
            sys.path.insert(0, '/var/www/urbanlex')
            from modulos.buscador_urbanistico import _buscar_leismunicipais
            # Processar arquivos
            logs.append({'nivel': 'ok', 'msg': f'\U0001f4ce {len(fpaths)} arquivo(s) recebido(s)'})
            texto_total = ""

            def _processar_arquivo(fpath2, fname2, logs2):
                ext2 = os.path.splitext(fname2)[1].lower()
                logs2.append({'nivel': 'anexo', 'msg': f'\U0001f4c4 Analisando: {os.path.basename(fname2)}...'})
                txt2 = _extrair_texto_arquivo(fpath2, fname2, ext2, logs2, tmp, chamar_llm)
                if txt2:
                    try:
                        desc = chamar_llm(f"Em uma linha, descreva o assunto deste documento municipal:\n\n{txt2[:2000]}", logs2, f"Desc {os.path.basename(fname2)}")
                        if desc:
                            logs2.append({'nivel': 'anexo', 'msg': f'  \U0001f4cb {os.path.basename(fname2)}: {desc.strip()[:150]}'})
                    except Exception:
                        pass
                return txt2

            for _fp, _fn in fpaths:
                _fext = os.path.splitext(_fn)[1].lower()
                if _fext == '.zip':
                    import zipfile, re as _re_sort
                    logs.append({'nivel': 'anexo', 'msg': f'\U0001f4e6 Descompactando ZIP: {_fn}...'})
                    with zipfile.ZipFile(_fp, 'r') as z:
                        _arquivos_map = {}
                        for _info in z.infolist():
                            if _info.flag_bits & 0x800:
                                _nome_correto = _info.filename
                            else:
                                _raw = _info.filename.encode('cp437')
                                try:
                                    _nome_correto = _raw.decode('utf-8', errors='strict')
                                except UnicodeDecodeError:
                                    try:
                                        _nome_correto = _raw.decode('cp1252')
                                    except Exception:
                                        _nome_correto = _info.filename
                            _arquivos_map[_info.filename] = _nome_correto
                        _raw_names = z.namelist()
                        arquivos = sorted(_raw_names, key=lambda x: [int(c) if c.isdigit() else c.lower() for c in _re_sort.split(r'(\d+)', x)])
                        logs.append({'nivel': 'anexo', 'msg': f'\U0001f4c2 {len(arquivos)} arquivo(s) no ZIP'})
                        z.extractall(tmp)
                    for fname in arquivos:
                        fname_display = _arquivos_map.get(fname, fname)
                        fpath2 = os.path.join(tmp, fname)
                        if not os.path.isfile(fpath2):
                            continue
                        txt = _processar_arquivo(fpath2, fname_display, logs)
                        if txt:
                            texto_total += f"\n\n--- {fname_display} ---\n{txt}"
                else:
                    txt = _processar_arquivo(_fp, _fn, logs)
                    if txt:
                        texto_total += f"\n\n--- {_fn} ---\n{txt}"
            if not texto_total.strip():
                logs.append({'nivel': 'aviso', 'msg': '⚠️ Nao foi possivel extrair texto dos anexos'})
                job['done'] = True
                return
            _BLOCO = 30000
            _OVERLAP = 2000
            altera, revoga, revoga_parcialmente, regulamenta, alterado_por, revogado_por, revogado_parcialmente_por, regulamentado_por, cita = [], [], [], [], [], [], [], [], []
            pos = 0
            bi = 0
            total_blocos = max(1, (len(texto_total) + _BLOCO - 1) // _BLOCO)
            while pos < len(texto_total):
                bi += 1
                bloco = texto_total[pos:pos+_BLOCO]
                prompt = (
                    f"Analise este trecho de anexo legislativo de {municipio}/{estado} (bloco {bi}/{total_blocos}).\n"
                    f"Identifique leis mencionadas e suas relacoes. Use as regras abaixo:\n"
                    f'1. altera: leis que ESTE documento modifica\n'
                    f'2. revoga: leis que ESTE documento revoga COMPLETAMENTE\n'
                    f'3. revoga_parcialmente: leis das quais ESTE documento revoga apenas ARTIGOS ESPECIFICOS\n'
                    f'4. regulamenta: leis que ESTE documento regulamenta\n'
                    f'5. alterado_por: leis POSTERIORES que alteram ESTE documento\n'
                    f'6. revogado_por: leis POSTERIORES que revogam COMPLETAMENTE este documento. NAO usar se revogar apenas artigos.\n'
                    f'7. revogado_parcialmente_por: leis POSTERIORES que revogam apenas ARTIGOS ESPECIFICOS deste documento\n'
                    f'8. regulamentado_por: leis que regulamentam ESTE documento\n'
                    f'9. cita: leis citadas em contexto urbanistico\n'
                    f'REGRA: "ficam revogados os Arts. X e Y da Lei Z" → revoga_parcialmente=[Lei Z], NAO revoga=[Lei Z]\n\n'
                    f'Responda APENAS com JSON:\n'
                    f'{"altera":[],"revoga":[],"revoga_parcialmente":[],"regulamenta":[],"alterado_por":[],"revogado_por":[],"revogado_parcialmente_por":[],"regulamentado_por":[],"cita":[]}\n\n'
                    f'TEXTO:\n{bloco}'
                )
                try:
                    resp = chamar_llm(prompt, logs, f"Bloco {bi}/{total_blocos}")
                    if resp:
                        import re, json as _j
                        resp_c = re.sub(r'^```json\s*|\s*```$', '', resp.strip())
                        # Tentar reparar JSON truncado
                        if resp_c and not resp_c.rstrip().endswith(']') and not resp_c.rstrip().endswith('}'):
                            resp_c = resp_c[:resp_c.rfind('}')+1] + ']' if '{' in resp_c else resp_c
                        d = _j.loads(resp_c)
                        altera = list(set(altera + d.get('altera',[])))
                        revoga = list(set(revoga + d.get('revoga',[])))
                        regulamenta = list(set(regulamenta + d.get('regulamenta',[])))
                        alterado_por = list(set(alterado_por + d.get('alterado_por',[])))
                        revogado_por = list(set(revogado_por + d.get('revogado_por',[])))
                        revogado_parcialmente_por = list(set(revogado_parcialmente_por + d.get('revogado_parcialmente_por',[])))
                        revoga_parcialmente = list(set(revoga_parcialmente + d.get('revoga_parcialmente',[])))
                        regulamentado_por = list(set(regulamentado_por + d.get('regulamentado_por',[])))
                        cita = list(set(cita + d.get('cita',[])))
                        logs.append({'nivel': 'relacao', 'msg': f'[Bloco {bi}] altera={len(altera)} revoga={len(revoga)} cita={len(cita)}'})
                except Exception as eb:
                    logs.append({'nivel': 'aviso', 'msg': f'Erro bloco {bi}: {str(eb)[:80]}'})
                pos += _BLOCO - _OVERLAP
            # Deduplicar e filtrar legislacoes
            import re as _rn
            def _eleg(s):
                return bool(_rn.search(r'\b(lei|decreto|resolucao|portaria|medida|emenda|complementar|LC|LO)\b|\d+[./]\d{4}', s, _rn.IGNORECASE))
            def _norm(s):
                s = _rn.sub(r'Lei Complementar n[ºo°.]*\s*', 'Lei Complementar ', s, flags=_rn.IGNORECASE)
                s = _rn.sub(r'\bLC\s+', 'Lei Complementar ', s, flags=_rn.IGNORECASE)
                return _rn.sub(r'\s+', ' ', s).strip().lower()
            def _dedup(lst):
                seen, out = {}, []
                for x in lst:
                    if _eleg(x) and _norm(x) not in seen:
                        seen[_norm(x)] = True; out.append(x)
                return out
            altera=_dedup(altera); revoga=_dedup(revoga); revoga_parcialmente=_dedup(revoga_parcialmente); regulamenta=_dedup(regulamenta)
            alterado_por=_dedup(alterado_por); revogado_por=_dedup(revogado_por); revogado_parcialmente_por=_dedup(revogado_parcialmente_por)
            regulamentado_por=_dedup(regulamentado_por); cita=_dedup(cita)
            logs.append({'nivel':'relacao','msg':f'Apos deduplicacao: altera={len(altera)} cita={len(cita)} alterado_por={len(alterado_por)}'})
            contexto_citas = locals().get('contexto_citas', [])
            resultado = {'altera': altera, 'revoga': revoga, 'revoga_parcialmente': revoga_parcialmente, 'regulamenta': regulamenta,
                        'alterado_por': alterado_por, 'revogado_por': revogado_por, 'revogado_parcialmente_por': revogado_parcialmente_por,
                        'regulamentado_por': regulamentado_por, 'cita': cita, 'contexto_citas': contexto_citas}
            logs.append({'nivel': 'ok', 'msg': f'✅ Analise concluida!'})
            logs.append({'nivel': 'ok', 'msg': f'  Altera: {altera}'})
            logs.append({'nivel': 'ok', 'msg': f'  Revoga: {revoga}'})
            logs.append({'nivel': 'ok', 'msg': f'  Regulamenta: {regulamenta}'})
            logs.append({'nivel': 'ok', 'msg': f'  Alterado por: {alterado_por}'})
            logs.append({'nivel': 'ok', 'msg': f'  Cita: {cita}'})
            # Analise de contexto das citacoes
            contexto_citas = []
            if cita and texto_total:
                logs.append({'nivel': 'relacao', 'msg': f'\U0001f50e Analisando contexto de {len(cita)} legislacoes em {max(1,len(texto_total)//40000)} bloco(s) de texto...'})
                import re as _re_ctx, json as _jctx
                # Dividir texto em blocos de 40k para cobrir todos os anexos
                _CTX_BLOCO = 40000
                _CTX_OVERLAP = 2000
                _ctx_blocos = []
                _p = 0
                while _p < len(texto_total):
                    _ctx_blocos.append(texto_total[_p:_p+_CTX_BLOCO])
                    _p += _CTX_BLOCO - _CTX_OVERLAP
                # Acumular contexto por lei em dict para mesclar resultados de varios blocos
                _ctx_map = {}  # lei_norm -> {lei, zonas, contexto, lei_principal}
                try:
                    for _bi, _bloco_txt in enumerate(_ctx_blocos):
                        logs.append({'nivel': 'relacao', 'msg': f'  Bloco texto {_bi+1}/{len(_ctx_blocos)} ({len(_bloco_txt)} chars)...'})
                        prompt_ctx = (
                            f"Analise o trecho de texto dos anexos de {municipio}/{estado} abaixo e para cada legislacao citada, identifique:\n"
                            f"1. Em que zona(s) ou subzona(s) ela e referenciada (ex: ZRM-2B, ZR-1, ZOP-2)\n"
                            f"2. O contexto da citacao (parametros urbanisticos, indice construtivo, uso permitido)\n"
                            f"3. Se e a lei principal que rege aquela zona ou apenas referencia secundaria\n\n"
                            f"LEGISLACOES A BUSCAR:\n" + "\n".join(f"- {c}" for c in cita) + "\n\n"
                            f'Responda APENAS com JSON para as leis encontradas neste trecho: [{{\'lei\':\'nome\',\'zonas\':[],\'contexto\':\'descricao\',\'lei_principal\':\'\'}}]\n\n'
                            f"TRECHO {_bi+1}/{len(_ctx_blocos)}:\n{_bloco_txt}"
                        )
                        resp_ctx = chamar_llm(prompt_ctx, logs, f'Contexto bloco {_bi+1}/{len(_ctx_blocos)}')
                        if resp_ctx:
                            try:
                                resp_ctx_c = _re_ctx.sub(r'^```json\s*|\s*```$', '', resp_ctx.strip())
                                _bloco_result = _jctx.loads(resp_ctx_c)
                                if isinstance(_bloco_result, list):
                                    for ct in _bloco_result:
                                        _lei = ct.get('lei','')
                                        if not _lei: continue
                                        _key = _lei.lower().strip()
                                        if _key not in _ctx_map:
                                            _ctx_map[_key] = ct
                                        else:
                                            # Mesclar zonas e contexto
                                            _existing = _ctx_map[_key]
                                            _z_new = ct.get('zonas',[])
                                            _z_ex = _existing.get('zonas',[])
                                            _existing['zonas'] = list(set(_z_ex + _z_new))
                                            if ct.get('contexto') and len(ct.get('contexto','')) > len(_existing.get('contexto','')):
                                                _existing['contexto'] = ct['contexto']
                                            if ct.get('lei_principal'):
                                                _existing['lei_principal'] = ct['lei_principal']
                            except Exception as _ep:
                                logs.append({'nivel': 'aviso', 'msg': f'  Erro parse bloco {_bi+1}: {str(_ep)[:60]}'})
                    contexto_citas = list(_ctx_map.values())
                    for ct in contexto_citas:
                        if ct.get('zonas') or ct.get('contexto'):
                            logs.append({'nivel': 'relacao', 'msg': f"  \U0001f4cd {ct.get('lei','')}: zonas={ct.get('zonas',[])} — {ct.get('contexto','')[:100]}"})
                    logs.append({'nivel': 'relacao', 'msg': f'  \u2705 Contexto total: {len(contexto_citas)} legislacoes'})
                except Exception as _ectx:
                    logs.append({'nivel': 'aviso', 'msg': f'Erro contexto: {str(_ectx)[:80]}'})
            # Analise de zonas e subzonas
            zonas_analise = {'tem_zonas': False, 'agrupamento': '', 'subzonas': []}
            if texto_total:
                logs.append({'nivel': 'relacao', 'msg': '\U0001f5fa\ufe0f Analisando zonas e subzonas do municipio...'})
                import json as _jz, re as _rz
                _ZBL = 40000
                _ZOV = 2000
                _z_blocos = []
                _zp = 0
                while _zp < len(texto_total):
                    _z_blocos.append(texto_total[_zp:_zp+_ZBL])
                    _zp += _ZBL - _ZOV
                _z_map = {}  # subzona_norm -> {subzona, divisao, usos}
                _tem_zonas = False
                _agrupamento = ''
                for _zbi, _zbloco in enumerate(_z_blocos):
                    logs.append({'nivel': 'relacao', 'msg': f'  Zonas bloco {_zbi+1}/{len(_z_blocos)}...'})
                    _zprompt = (
                        f"Analise este trecho dos anexos legislativos de {municipio}/{estado} e responda:\n"
                        f"1. Este trecho menciona zonas ou subzonas urbanisticas do municipio? (sim/nao)\n"
                        f"2. Se sim, as zonas sao agrupadas por alguma divisao administrativa ou de planejamento? Qual?\n"
                        f"3. Liste todas as subzonas mencionadas neste trecho com seus usos permitidos.\n\n"
                        f"Responda APENAS com JSON:\n"
                        f"{{\"tem_zonas\": true/false, \"agrupamento\": \"descricao ou vazio\", \"subzonas\": [{{\"subzona\": \"nome\", \"divisao\": \"divisao administrativa ou vazio\", \"usos\": \"usos permitidos\"}}]}}\n\n"
                        f"TRECHO {_zbi+1}/{len(_z_blocos)}:\n{_zbloco}"
                    )
                    try:
                        _zresp = chamar_llm(_zprompt, logs, f'Zonas bloco {_zbi+1}/{len(_z_blocos)}')
                        if _zresp:
                            _zresp_c = _rz.sub(r'^```json\s*|\s*```$', '', _zresp.strip())
                            _zdata = _jz.loads(_zresp_c)
                            if _zdata.get('tem_zonas'):
                                _tem_zonas = True
                            if _zdata.get('agrupamento') and not _agrupamento:
                                _agrupamento = _zdata['agrupamento']
                            for _sz in (_zdata.get('subzonas') or []):
                                _szn = _sz.get('subzona','').strip()
                                if not _szn: continue
                                _szk = _szn.lower()
                                if _szk not in _z_map:
                                    _z_map[_szk] = _sz
                                else:
                                    if _sz.get('usos') and len(_sz.get('usos','')) > len(_z_map[_szk].get('usos','')):
                                        _z_map[_szk]['usos'] = _sz['usos']
                                    if _sz.get('divisao') and not _z_map[_szk].get('divisao'):
                                        _z_map[_szk]['divisao'] = _sz['divisao']
                    except Exception as _ze:
                        logs.append({'nivel': 'aviso', 'msg': f'  Erro zonas bloco {_zbi+1}: {str(_ze)[:60]}'})
                zonas_analise = {'tem_zonas': _tem_zonas, 'agrupamento': _agrupamento, 'subzonas': list(_z_map.values())}
                logs.append({'nivel': 'ok', 'msg': f'\U0001f5fa\ufe0f Zonas: tem_zonas={_tem_zonas} agrupamento={_agrupamento[:60]} subzonas={len(_z_map)}'})
            resultado['zonas_analise'] = zonas_analise
            job['result'] = resultado
            # Persistir resultado no arquivo JSONL
            try:
                import json as _json_res
                _res_path = f"/var/www/urbanlex/static/downloads/job_{job_id}.jsonl"
                with open(_res_path, 'a', encoding='utf-8') as _rf:
                    _rf.write(_json_res.dumps({'_result': resultado}, ensure_ascii=False) + '\n')
            except Exception: pass
        except Exception as e:
            job['logs'].append({'nivel': 'erro', 'msg': f'Erro: {str(e)[:200]}'})
        finally:
            job['done'] = True
    threading.Thread(target=_run, daemon=True).start()
    return jsonify({'success': True, 'job_id': job_id})

@app.route('/buscador')
@app.route('/buscador/manual')
@app.route('/buscador/historico')
@login_required
def pg_buscador():
    page_map = {'/buscador': 'buscador-painel', '/buscador/manual': 'buscador-manual', '/buscador/historico': 'buscador-historico'}
    return render_template('buscador.html', active_group='buscador', active_page=page_map.get(request.path, 'buscador-painel'))

@app.route('/api/buscador/municipio', methods=['POST'])
@editor_required
def api_buscador_municipio():
    """Inicia busca urbanistica de municipio em background."""
    import uuid, threading
    d = request.json or {}
    mun = d.get("municipio", "").strip()
    est = d.get("estado", "").strip()
    _max_legs = d.get("max_legislacoes", None)
    if _max_legs: _max_legs = int(_max_legs)
    if not mun or not est:
        return jsonify({"success": False, "error": "municipio e estado obrigatorios"}), 400
    # Bloquear jobs concorrentes
    _jc_ativos=[j for jid,j in _buscador_jobs.items() if not j.get("done") and not j.get("cancelled") and j.get("tipo")=="auto"]
    if _jc_ativos:
        return jsonify({"success":False,"error":"job_concorrente"}), 409
    job_id = str(uuid.uuid4())[:8]
    job = {"done": False, "cancelled": False, "logs": _LogList(job_id, get_db), "result": None, "tipo": "auto"}
    _buscador_jobs[job_id] = job

    # Registrar inicio no historico
    hist_id = None
    try:
        _hconn = get_db()
        _hcur = _hconn.cursor()
        _hcur.execute("""INSERT INTO buscas_historico (tipo, municipio, estado, iniciado_em, job_id)
                         VALUES (%s, %s, %s, NOW(), %s) RETURNING id""",
                      ('automatica', mun, est, job_id))
        hist_id = _hcur.fetchone()[0]
        _hconn.commit()
        _hcur.close()
        _hconn.close()
    except Exception as _he:
        pass

    def _run():
        try:
            from modulos.buscador_legislacoes import _chamar_llm as _llm
            from modulos.buscador_urbanistico import buscar_legislacoes_urbanisticas
            def chamar_llm(prompt, logs, label="LLM", max_retries=2):
                return _llm(prompt, logs, label, max_retries)
            _fb_url = None
            try:
                _fb_conn = get_db(); _fb_cur = _fb_conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                _fb_cur.execute("SELECT url FROM municipio_fallback WHERE LOWER(municipio)=LOWER(%s) AND LOWER(estado)=LOWER(%s)", (mun, est))
                _fb_row = _fb_cur.fetchone(); _fb_conn.close()
                if _fb_row: _fb_url = _fb_row['url']
            except: pass
            r = buscar_legislacoes_urbanisticas(mun, est, job["logs"], chamar_llm, fallback_url=_fb_url, max_legislacoes=_max_legs)
            # Expor ZIP e JSON no resultado do job
            job["result"] = {
                "encontradas": r.get("encontradas", []),
                "zip_url": r.get("zip_url"),
                "zip_nome": r.get("zip_nome"),
                "relatorio_url": r.get("relatorio_url"),
                "relatorio_nome": r.get("relatorio_nome"),
                "tabela_url": r.get("tabela_url"),
                "tabela_nome": r.get("tabela_nome"),
                "legislacoes_json": r.get("legislacoes_json", []),
                "custo_usd": r.get("custo_usd"),
                "token_stats": r.get("token_stats"),
                "nao_encontrada": r.get("nao_encontrada", False),
            }
            job["hist_id"] = hist_id
            # Atualizar historico ao concluir
            if hist_id:
                try:
                    _enc = (r or {}).get("encontradas", [])
                    _leg = _enc[0] if _enc else {}
                    _log_txt = "\n".join(l.get("msg","") for l in job["logs"])
                    job["logs"].append({"nivel": "info", "msg": f"[DEBUG] _leg keys: {list(_leg.keys())} | pdf_path: {_leg.get('pdf_path','VAZIO')}"})
                    _sucesso = bool(_enc)
                    _hconn2 = get_db()
                    _hcur2 = _hconn2.cursor()
                    import json as _json_h
                    _anexos = _leg.get("anexos_lm", [])
                    _anexos_json = _json_h.dumps(_anexos) if _anexos else "[]"
                    _hcur2.execute("""UPDATE buscas_historico SET
                        concluido_em=NOW(), sucesso=%s,
                        legislacao_tipo=%s, legislacao_numero=%s,
                        legislacao_ano=%s, legislacao_link=%s,
                        log_texto=%s, pdf_path=%s, anexos_paths=%s,
                        relatorio_path=%s, tabela_path=%s, zip_path=%s
                        WHERE id=%s""",
                        (_sucesso,
                         _leg.get("tipo",""), _leg.get("numero",""),
                         _leg.get("ano",""), _leg.get("link",""),
                         _log_txt, _leg.get("pdf_path",""), _anexos_json,
                         r.get("relatorio_url",""), r.get("tabela_url",""), r.get("zip_url",""),
                         hist_id))
                    _hconn2.commit()
                    _hcur2.close()
                    _hconn2.close()
                except Exception as _he2:
                    job["logs"].append({"nivel": "aviso", "msg": f"Erro ao salvar historico: {str(_he2)[:100]}"})
            # Trigger automatico: organizador de dossie em background
            try:
                _zip_url_fim = r.get("zip_url", "")
                if _zip_url_fim:
                    from modulos.dossie_trigger import disparar_organizador_async
                    disparar_organizador_async(mun, est, _zip_url_fim, get_db, origem='manual', busca_id=hist_id)
            except Exception as _e_trig:
                job["logs"].append({"nivel": "aviso", "msg": f"Trigger dossie nao disparado: {str(_e_trig)[:150]}"})
        except Exception as e:
            job["logs"].append({"nivel": "erro", "msg": f"Erro: {str(e)[:200]}"})
        finally:
            job["done"] = True
            # Persistir resultado no JSONL para recuperacao apos restart
            try:
                import json as _json_auto
                _auto_path = f"/var/www/urbanlex/static/downloads/job_{job_id}.jsonl"
                with open(_auto_path, 'a', encoding='utf-8') as _af:
                    _af.write(_json_auto.dumps({'_result': job.get('result')}, ensure_ascii=False) + '\n')
            except Exception:
                pass
    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"success": True, "job_id": job_id})
@app.route('/api/buscador/lei-especifica', methods=['POST'])
@editor_required
def api_buscador_lei_especifica():
    """Inicia analise completa de uma lei especifica em background."""
    import uuid, threading
    d = request.json or {}
    mun = d.get("municipio", "").strip()
    est = d.get("estado", "").strip().upper()
    tipo = d.get("tipo", "").strip()
    numero = d.get("numero", "").strip()
    ano = d.get("ano", "").strip()
    if not mun or not est or not tipo or not numero or not ano:
        return jsonify({"success": False, "error": "municipio, estado, tipo, numero e ano sao obrigatorios"}), 400
    # Bloquear jobs concorrentes
    _jc_ativos = [j for jid, j in _buscador_jobs.items() if not j.get("done") and not j.get("cancelled")]
    if _jc_ativos:
        return jsonify({"success": False, "error": "job_concorrente"}), 409
    job_id = str(uuid.uuid4())[:8]
    job = {"done": False, "cancelled": False, "logs": _LogList(job_id, get_db), "result": None, "tipo": "especifica"}
    _buscador_jobs[job_id] = job
    hist_id = None
    try:
        _hconn = get_db(); _hcur = _hconn.cursor()
        _hcur.execute("""INSERT INTO buscas_historico (tipo, municipio, estado, iniciado_em, job_id)
                         VALUES (%s, %s, %s, NOW(), %s) RETURNING id""",
                      ('especifica', mun, est, job_id))
        hist_id = _hcur.fetchone()[0]
        _hconn.commit(); _hcur.close(); _hconn.close()
    except Exception: pass
    def _run():
        try:
            from modulos.buscador_legislacoes import _chamar_llm as _llm
            from modulos.buscador_urbanistico import buscar_legislacoes_urbanisticas
            def chamar_llm(prompt, logs, label="LLM", max_retries=2):
                return _llm(prompt, logs, label, max_retries)
            _fb_url = None
            try:
                _fb_conn = get_db(); _fb_cur = _fb_conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                _fb_cur.execute("SELECT url FROM municipio_fallback WHERE LOWER(municipio)=LOWER(%s) AND LOWER(estado)=LOWER(%s)", (mun, est))
                _fb_row = _fb_cur.fetchone(); _fb_conn.close()
                if _fb_row: _fb_url = _fb_row['url']
            except: pass
            # Criar lista de leis com apenas a lei especificada
            from modulos.buscador_urbanistico import buscar_legislacoes_urbanisticas as _bla
            # Chamar diretamente com a lei especificada como unica legislacao
            _legs_override = [{"tipo": tipo, "numero": numero, "ano": ano, "descricao": "", "_nivel": 0, "_pergunta_label": ""}]
            from modulos.buscador_urbanistico import _buscar_leismunicipais, _buscar_fallback1, _buscar_fallback2
            # Usar o fluxo completo mas com lista de leis pre-definida
            import json as _json_esp
            r = buscar_legislacoes_urbanisticas(mun, est, job["logs"], chamar_llm,
                                                fallback_url=_fb_url, max_legislacoes=1,
                                                _legs_override=_legs_override)
            job["result"] = {
                "encontradas": r.get("encontradas", []),
                "zip_url": r.get("zip_url"),
                "zip_nome": r.get("zip_nome"),
                "relatorio_url": r.get("relatorio_url"),
                "relatorio_nome": r.get("relatorio_nome"),
                "tabela_url": r.get("tabela_url"),
                "tabela_nome": r.get("tabela_nome"),
                "nao_encontrada": r.get("nao_encontrada", False),
            }
            job["hist_id"] = hist_id
            if hist_id:
                try:
                    _enc = r.get("encontradas", [])
                    _leg = _enc[0] if _enc else {}
                    _hconn2 = get_db(); _hcur2 = _hconn2.cursor()
                    _hcur2.execute("""UPDATE buscas_historico SET concluido_em=NOW(), sucesso=%s,
                        legislacao_tipo=%s, legislacao_numero=%s, legislacao_ano=%s,
                        zip_path=%s, relatorio_path=%s, tabela_path=%s WHERE id=%s""",
                        (bool(_enc), _leg.get("tipo",""), _leg.get("numero",""), _leg.get("ano",""),
                         r.get("zip_url",""), r.get("relatorio_url",""), r.get("tabela_url",""), hist_id))
                    _hconn2.commit(); _hcur2.close(); _hconn2.close()
                except Exception: pass
            # Trigger automatico: organizador de dossie em background
            try:
                _zip_url_fim = r.get("zip_url", "")
                if _zip_url_fim:
                    from modulos.dossie_trigger import disparar_organizador_async
                    disparar_organizador_async(mun, est, _zip_url_fim, get_db, origem='manual', busca_id=hist_id)
            except Exception as _e_trig:
                job["logs"].append({"nivel": "aviso", "msg": f"Trigger dossie nao disparado: {str(_e_trig)[:150]}"})
        except Exception as e:
            job["logs"].append({"nivel": "erro", "msg": f"Erro: {str(e)[:200]}"})
        finally:
            job["done"] = True
    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"success": True, "job_id": job_id})


@app.route('/api/buscador/upload-legislacao', methods=['POST'])
@editor_required
def api_buscador_upload_legislacao():
    """Upload manual de legislacoes para analise completa."""
    import uuid, threading, json as _j_up, os as _os_up, shutil as _sh_up
    from werkzeug.utils import secure_filename
    metadata_str = request.form.get('metadata', '[]')
    try:
        metadata = _j_up.loads(metadata_str)
    except Exception:
        return jsonify({"success": False, "error": "metadata invalido"}), 400
    if not metadata:
        return jsonify({"success": False, "error": "Nenhuma legislacao informada"}), 400
    primeira = metadata[0]
    mun = primeira.get('municipio', '').strip()
    est = primeira.get('estado', '').strip().upper()
    if not mun or not est:
        return jsonify({"success": False, "error": "municipio e estado sao obrigatorios"}), 400
    # Bloquear jobs concorrentes
    _jc = [j for jid, j in _buscador_jobs.items() if not j.get("done") and not j.get("cancelled")]
    if _jc:
        return jsonify({"success": False, "error": "job_concorrente"}), 409
    job_id = str(uuid.uuid4())[:8]
    upload_base = f"/var/www/urbanlex/static/uploads_manuais/{job_id}"
    _os_up.makedirs(upload_base, exist_ok=True)
    legs_override = []
    arquivos_override = []
    for i, leg_meta in enumerate(metadata):
        tipo = leg_meta.get('tipo', 'Lei').strip()
        numero = leg_meta.get('numero', '').strip()
        ano = leg_meta.get('ano', '').strip()
        leg_dir = _os_up.path.join(upload_base, f"leg_{i}")
        _os_up.makedirs(leg_dir, exist_ok=True)
        files = request.files.getlist(f'files_{i}')
        saved = []
        for f in files:
            if f and f.filename:
                fname = secure_filename(f.filename)
                fpath = _os_up.path.join(leg_dir, fname)
                f.save(fpath)
                saved.append(fpath)
        if not saved:
            return jsonify({"success": False, "error": f"Nenhum arquivo para legislacao {i+1}"}), 400
        legs_override.append({"tipo": tipo, "numero": numero, "ano": ano, "_nivel": 0, "_pergunta_label": ""})
        arquivos_override.append({"tipo": tipo, "numero": numero, "ano": ano, "arquivos": saved})
    job = {"done": False, "cancelled": False, "logs": _LogList(job_id, get_db), "result": None, "tipo": "upload"}
    _buscador_jobs[job_id] = job
    hist_id = None
    try:
        _hc = get_db(); _hcur = _hc.cursor()
        _hcur.execute("INSERT INTO buscas_historico (tipo, municipio, estado, iniciado_em, job_id) VALUES (%s,%s,%s,NOW(),%s) RETURNING id",
                      ('upload', mun, est, job_id))
        hist_id = _hcur.fetchone()[0]
        _hc.commit(); _hcur.close(); _hc.close()
    except Exception: pass
    def _run():
        try:
            from modulos.buscador_legislacoes import _chamar_llm as _llm
            from modulos.buscador_urbanistico import buscar_legislacoes_urbanisticas
            def chamar_llm(p, logs, label="LLM", max_retries=2): return _llm(p, logs, label, max_retries)
            r = buscar_legislacoes_urbanisticas(mun, est, job["logs"], chamar_llm,
                max_legislacoes=len(legs_override), _legs_override=legs_override,
                _arquivos_override=arquivos_override)
            job["result"] = {"encontradas": r.get("encontradas", []),
                "zip_url": r.get("zip_url"), "zip_nome": r.get("zip_nome"),
                "relatorio_url": r.get("relatorio_url"), "relatorio_nome": r.get("relatorio_nome"),
                "tabela_url": r.get("tabela_url"), "tabela_nome": r.get("tabela_nome"),
                "nao_encontrada": r.get("nao_encontrada", False)}
            job["hist_id"] = hist_id
            if hist_id:
                try:
                    _enc = r.get("encontradas", []); _leg = _enc[0] if _enc else {}
                    _hc2 = get_db(); _hcur2 = _hc2.cursor()
                    _hcur2.execute("UPDATE buscas_historico SET concluido_em=NOW(), sucesso=%s, legislacao_tipo=%s, legislacao_numero=%s, legislacao_ano=%s, zip_path=%s, relatorio_path=%s, tabela_path=%s WHERE id=%s",
                        (bool(_enc), _leg.get("tipo",""), _leg.get("numero",""), _leg.get("ano",""),
                         r.get("zip_url",""), r.get("relatorio_url",""), r.get("tabela_url",""), hist_id))
                    _hc2.commit(); _hcur2.close(); _hc2.close()
                except Exception: pass
            try:
                _zip_fim = r.get("zip_url","")
                if _zip_fim:
                    from modulos.dossie_trigger import disparar_organizador_async
                    disparar_organizador_async(mun, est, _zip_fim, get_db, origem='upload', busca_id=hist_id)
            except Exception as _et:
                job["logs"].append({"nivel":"aviso","msg":f"Trigger dossie nao disparado: {str(_et)[:150]}"})
        except Exception as e:
            job["logs"].append({"nivel":"erro","msg":f"Erro upload: {str(e)[:200]}"})
        finally:
            job["done"] = True
    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"success": True, "job_id": job_id})

@app.route('/api/buscador/manual', methods=['POST'])
@editor_required
def api_buscador_manual():
    """Busca manual por campos."""
    d = request.json or {}
    try:
        from modulos.buscador_legislacoes import busca_manual
        resultado = busca_manual(d)
        if isinstance(resultado, dict):
            return jsonify({
                'success': True,
                'legislacoes': resultado.get('legislacoes', []),
                'erro': resultado.get('erro'),
                'logs': resultado.get('logs', []),
                'sugestao': resultado.get('sugestao'),
                'fontes': resultado.get('fontes', []),
                'textos_fontes': resultado.get('textos_fontes', []),
            })
        else:
            return jsonify({'success': True, 'legislacoes': resultado, 'logs': []})
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        return jsonify({
            'success': False,
            'error': str(e)[:300],
            'logs': [
                {'nivel': 'erro', 'msg': f'Exceção no servidor: {str(e)[:200]}'},
                {'nivel': 'erro', 'msg': f'Traceback: {tb[-300:]}'},
            ]
        }), 500


# ── Cache IBGE: lista completa de municípios brasileiros ──
_ibge_municipios = []  # [{'nome': 'Natal', 'estado': 'RN'}, ...]
_ibge_carregado = False

def _carregar_ibge():
    """Baixa a lista de todos os municípios do IBGE (5570) e cacheia em memória + arquivo."""
    global _ibge_municipios, _ibge_carregado
    if _ibge_carregado:
        return

    cache_path = os.path.join(os.path.dirname(__file__), 'municipios_ibge.json')

    # Tentar carregar do cache local primeiro
    if os.path.exists(cache_path):
        try:
            with open(cache_path, 'r', encoding='utf-8') as f:
                _ibge_municipios = json.load(f)
            if len(_ibge_municipios) > 5000:
                _ibge_carregado = True
                print(f'[IBGE] {len(_ibge_municipios)} municípios carregados (cache local)')
                return
        except Exception:
            pass

    # Baixar da API do IBGE
    try:
        import requests
        resp = requests.get(
            'https://servicodados.ibge.gov.br/api/v1/localidades/municipios?orderBy=nome',
            timeout=15
        )
        if resp.status_code == 200:
            data = resp.json()
            _ibge_municipios = []
            for m in data:
                try:
                    uf = m['microrregiao']['mesorregiao']['UF']['sigla']
                    _ibge_municipios.append({'nome': m['nome'], 'estado': uf})
                except (KeyError, TypeError):
                    pass
            _ibge_carregado = True
            print(f'[IBGE] {len(_ibge_municipios)} municípios carregados (API)')

            # Salvar cache local
            try:
                with open(cache_path, 'w', encoding='utf-8') as f:
                    json.dump(_ibge_municipios, f, ensure_ascii=False)
                print(f'[IBGE] Cache salvo em {cache_path}')
            except Exception:
                pass
        else:
            print(f'[IBGE] Erro HTTP {resp.status_code}')
    except Exception as e:
        print(f'[IBGE] Falha ao carregar da API: {str(e)[:100]}')

# Carregar IBGE em background no startup
threading.Thread(target=_carregar_ibge, daemon=True).start()

# ── Buscador: jobs em background com polling de logs ──
_buscador_jobs = {}  # job_id -> {'logs': _LogList(job_id, get_db), 'result': None, 'done': False, 'ts': time.time()}
_fila_pausada = False  # pausar worker apos cancelamento
from modulos.fila_worker import iniciar_worker as _iniciar_fila_worker
from modulos.fila_extracao_worker import iniciar_worker_extracao as _iniciar_fila_extracao_worker
from modulos.pipeline_extracao_lei import enfileirar_extracao as _enfileirar_extracao, buscar_consolidado as _buscar_consolidado
from modulos.log_persistente import LogList as _LogList

def _cleanup_old_jobs():
    """Remove jobs com mais de 2 horas."""
    now = time.time()
    expired = [k for k, v in _buscador_jobs.items() if now - v['ts'] > 7200]
    for k in expired:
        _buscador_jobs.pop(k, None)

def _cleanup_chromium_orfaos():
    """Mata processos Chromium orfaos quando nao ha jobs ativos."""
    import subprocess as _sp
    while True:
        try:
            time.sleep(300)
            jobs_ativos = any(not j.get('done') for j in _buscador_jobs.values())
            if not jobs_ativos:
                r = _sp.run(['pgrep', '-f', 'chromium'], capture_output=True, text=True)
                if r.stdout.strip():
                    _sp.run(['pkill', '-9', '-f', 'chromium'], capture_output=True)
        except Exception:
            pass

def _monitor_memoria():
    """Monitora memoria a cada 30s. Se >90% mata Chromium e cancela jobs ativos."""
    import subprocess as _sp
    while True:
        try:
            time.sleep(30)
            with open('/proc/meminfo') as _f:
                lines = {l.split(':')[0]: int(l.split()[1]) for l in _f if ':' in l}
            total = lines.get('MemTotal', 1)
            disponivel = lines.get('MemAvailable', total)
            uso_pct = 100 * (1 - disponivel / total)
            if uso_pct > 90:
                for _j in _buscador_jobs.values():
                    if not _j.get('done'):
                        _j['cancelled'] = True
                        _j['done'] = True
                        _j['logs'].append({'nivel': 'erro', 'msg': f'Job cancelado: memoria critica ({uso_pct:.1f}% usada)'})
                _sp.run(['pkill', '-9', '-f', 'chromium'], capture_output=True)
        except Exception:
            pass

import threading as _th_ck
_th_ck.Thread(target=_cleanup_chromium_orfaos, daemon=True).start()
_th_ck.Thread(target=_monitor_memoria, daemon=True).start()



@app.route('/api/admin/info-sistema')
@login_required
def api_info_sistema():
    import subprocess, requests as _req_gh
    from datetime import datetime
    import pytz
    try:
        commit_local = subprocess.check_output(['git', 'log', '--oneline', '-1'], cwd='/var/www/urbanlex').decode().strip()
        since_raw = subprocess.check_output(['systemctl', 'show', 'urbanlex', '--property=ActiveEnterTimestamp']).decode().strip().replace('ActiveEnterTimestamp=','')
        sp_tz = pytz.timezone('America/Sao_Paulo')
        since_dt = datetime.strptime(since_raw, '%a %Y-%m-%d %H:%M:%S %Z')
        since_sp = since_dt.replace(tzinfo=pytz.utc).astimezone(sp_tz).strftime('%d/%m %H:%M')
    except:
        commit_local = 'desconhecido'
        since_sp = 'desconhecido'
    try:
        gh = _req_gh.get('https://api.github.com/repos/costaerico1978/urbanlex/commits/main', timeout=5).json()
        commit_remote = gh['sha'][:7] + ' ' + gh['commit']['message'].split('\n')[0]
        atualizado = commit_local.split(' ')[0] == gh['sha'][:7]
    except:
        commit_remote = 'desconhecido'
        atualizado = True
    return jsonify({'commit': commit_local, 'since': since_sp, 'commit_remoto': commit_remote, 'atualizado': atualizado})

@app.route('/api/admin/atualizar-worker', methods=['POST'])
@login_required
def api_atualizar_worker():
    import subprocess, threading
    def _update():
        import time
        time.sleep(1)
        subprocess.run(['bash', '-c', 'cd /var/www/urbanlex && git pull && systemctl restart urbanlex'])
    threading.Thread(target=_update, daemon=True).start()
    return jsonify({'success': True, 'msg': 'Atualizando e reiniciando em 1s...'})

@app.route('/api/admin/reiniciar-worker', methods=['POST'])
@login_required
def api_reiniciar_worker():
    import subprocess, threading
    def _restart():
        import time
        time.sleep(1)
        subprocess.run(['pkill', '-9', '-f', 'chromium'], capture_output=True)
        subprocess.run(['pkill', '-9', '-f', 'gunicorn'], capture_output=True)
        time.sleep(2)
        subprocess.run(['bash', '-c', 'cd /var/www/urbanlex && git pull'], capture_output=True)
        subprocess.run(['systemctl', 'start', 'urbanlex'])
    threading.Thread(target=_restart, daemon=True).start()
    return jsonify({'success': True, 'msg': 'Worker reiniciando em 1s...'})

@app.route('/api/buscador/especifica/estado')
@login_required
def api_buscador_especifica_estado():
    """Retorna estado da ultima busca especifica - server-side, sem sessionStorage."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        # 1. Tem busca especifica EM ANDAMENTO?
        cur.execute("""SELECT job_id, municipio, estado, legislacao_tipo,
                              legislacao_numero, legislacao_ano, iniciado_em
                       FROM buscas_historico
                       WHERE tipo='especifica' AND job_id IS NOT NULL
                         AND concluido_em IS NULL
                       ORDER BY iniciado_em DESC LIMIT 1""")
        rodando = cur.fetchone()
        if rodando:
            cur.close(); conn.close()
            return jsonify({
                'success': True, 'modo': 'em_andamento',
                'job_id': rodando['job_id'],
                'municipio': rodando['municipio'], 'estado': rodando['estado'],
                'legislacao_tipo': rodando['legislacao_tipo'],
                'legislacao_numero': rodando['legislacao_numero'],
                'legislacao_ano': rodando['legislacao_ano'],
            })
        # 2. Tem ULTIMA busca CONCLUIDA recente (ultimas 24h)?
        cur.execute("""SELECT id, job_id, municipio, estado, legislacao_tipo,
                              legislacao_numero, legislacao_ano,
                              iniciado_em, concluido_em,
                              zip_path, relatorio_path, tabela_path
                       FROM buscas_historico
                       WHERE tipo='especifica' AND concluido_em IS NOT NULL
                         AND sucesso = true
                       ORDER BY concluido_em DESC LIMIT 1""")
        ultima = cur.fetchone()
        cur.close(); conn.close()
        if ultima:
            def _to_url(p):
                if not p: return None
                # Path no banco pode estar como 'static/...' ou '/static/...'
                p = str(p).lstrip('/')
                return '/' + p if not p.startswith('http') else p
            return jsonify({
                'success': True, 'modo': 'ultima_concluida',
                'id': ultima['id'], 'job_id': ultima['job_id'],
                'municipio': ultima['municipio'], 'estado': ultima['estado'],
                'legislacao_tipo': ultima['legislacao_tipo'],
                'legislacao_numero': ultima['legislacao_numero'],
                'legislacao_ano': ultima['legislacao_ano'],
                'iniciado_em': ultima['iniciado_em'].isoformat() if ultima['iniciado_em'] else None,
                'concluido_em': ultima['concluido_em'].isoformat() if ultima['concluido_em'] else None,
                'zip_url': _to_url(ultima['zip_path']),
                'relatorio_url': _to_url(ultima['relatorio_path']),
                'tabela_url': _to_url(ultima['tabela_path']),
            })
        return jsonify({'success': True, 'modo': None})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:200]})


@app.route('/api/buscador/job-ativo')
@login_required
def api_buscador_job_ativo():
    """Retorna job ativo mais recente — para reconexao apos fechar browser."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""SELECT job_id, municipio, estado, iniciado_em, tipo
                       FROM buscas_historico
                       WHERE job_id IS NOT NULL AND concluido_em IS NULL
                       ORDER BY iniciado_em DESC LIMIT 1""")
        r = cur.fetchone()
        cur.close()
        conn.close()
        # Verificar tambem fila_buscas (worker server-side)
        if not r:
            cur3 = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur3.execute("SELECT job_id, municipio, estado FROM fila_buscas WHERE status='rodando' LIMIT 1")
            rf = cur3.fetchone(); cur3.close()
            if rf and rf['job_id']:
                job_id = rf['job_id']
                job = _buscador_jobs.get(job_id)
                h2 = None
                try:
                    c2b=get_db(); cu2b=c2b.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                    cu2b.execute("SELECT id FROM buscas_historico WHERE job_id=%s LIMIT 1",(job_id,))
                    h2=cu2b.fetchone(); cu2b.close(); c2b.close()
                except: pass
                _tipo_fila = None
                if job and isinstance(job, dict):
                    _tipo_fila = job.get('tipo')
                if not _tipo_fila:
                    _tipo_fila = 'auto'  # jobs da fila do worker sao do tipo 'auto'
                return jsonify({'ativo': True, 'job_id': job_id, 'municipio': rf['municipio'], 'estado': rf['estado'], 'hist_id': h2['id'] if h2 else None, 'tipo': _tipo_fila})
            return jsonify({'ativo': False})
        job_id = r['job_id']
        job = _buscador_jobs.get(job_id)
        if not job or job.get('done'):
            # Verificar se ha arquivo de log persistido (worker reiniciou)
            import json as _json_at
            _log_path_at = f"/var/www/urbanlex/static/downloads/job_{job_id}.jsonl"
            if os.path.exists(_log_path_at):
                try:
                    with open(_log_path_at, 'r', encoding='utf-8') as _lf_at:
                        _logs_at = [_json_at.loads(l) for l in _lf_at if l.strip()]
                    # Verificar se a ultima linha indica conclusao
                    _last = _logs_at[-1] if _logs_at else {}
                    if not _last.get('msg','').startswith('CONCLUIDO'):
                        _tipo_rec = r.get('tipo') or 'manual'
                        return jsonify({'ativo': True, 'job_id': job_id, 'municipio': r['municipio'], 'estado': r['estado'], 'hist_id': None, 'recuperado': True, 'tipo': _tipo_rec})
                except Exception:
                    pass
            return jsonify({'ativo': False})
        # Buscar hist_id para botao de download
        conn2 = get_db()
        cur2 = conn2.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur2.execute("SELECT id FROM buscas_historico WHERE job_id=%s LIMIT 1", (job_id,))
        h = cur2.fetchone()
        cur2.close()
        conn2.close()
        # Pegar tipo: prioridade memoria > banco > fallback
        _tipo_resp = None
        if job and isinstance(job, dict):
            _tipo_resp = job.get('tipo')
        if not _tipo_resp:
            _tipo_resp = r.get('tipo') or 'manual'
        return jsonify({
            'ativo': True,
            'job_id': job_id,
            'municipio': r['municipio'],
            'estado': r['estado'],
            'hist_id': h['id'] if h else None,
            'tipo': _tipo_resp
        })
    except Exception as e:
        return jsonify({'ativo': False})

try:
    import subprocess as _git_sp
    _APP_COMMIT = _git_sp.check_output(['git','rev-parse','--short','HEAD'], cwd='/var/www/urbanlex', stderr=_git_sp.DEVNULL).decode().strip()
except:
    _APP_COMMIT = 'desconhecido'

@app.route('/api/versao')
def api_versao():
    return jsonify({'commit': _APP_COMMIT, 'success': True})

@app.route('/api/buscador/jobs-ativos')
def api_buscador_jobs_ativos():
    """Verifica se ha jobs de busca ativos — usado pelo webhook de deploy."""
    ativos = any(not j.get('done') for j in _buscador_jobs.values())
    return jsonify({'ativos': ativos})


@app.route('/api/buscador/cancelar', methods=['POST'])
@login_required
def api_buscador_cancelar():
    """Cancela um job de busca ativo."""
    job_id = (request.json or {}).get('job_id', '')
    job = _buscador_jobs.get(job_id)
    if job and not job.get('done'):
        job['cancelled'] = True
        job['done'] = True
        job['logs'].append({'nivel': 'aviso', 'msg': '⚠️ Busca cancelada pelo usuário'})
        # Matar processos Chromium para liberar recursos
        import subprocess
        subprocess.run(['pkill', '-9', '-f', 'chromium'], capture_output=True)
    return jsonify({'success': True})

@app.route('/api/buscador/manual-stream', methods=['POST'])
@editor_required
def api_buscador_manual_start():
    """Inicia busca em background e retorna job_id para polling."""
    import uuid
    _cleanup_old_jobs()

    # Cancelar jobs ativos anteriores
    for _jid, _jv in list(_buscador_jobs.items()):
        if not _jv.get('done') and not _jv.get('cancelled'):
            _jv['cancelled'] = True
            _jv['logs'].append({'nivel': 'aviso', 'msg': '⚠️ Busca cancelada — nova busca iniciada'})
            try:
                import os as _os_del
                _old_log = f"/var/www/urbanlex/static/downloads/job_{_jid}.jsonl"
                if _os_del.path.exists(_old_log): _os_del.remove(_old_log)
            except Exception: pass


    d = request.json or {}
    job_id = str(uuid.uuid4())[:12]
    job = {'logs': _LogList(job_id, get_db), 'result': None, 'done': False, 'ts': time.time(), 'log_cursor': 0, 'cancelled': False, 'tipo': 'manual'}
    _buscador_jobs[job_id] = job
    # Registrar inicio no historico
    _hist_id_manual = None
    try:
        _mun_m = d.get('municipio', '').strip()
        _est_m = d.get('estado', '').strip()
        _hconn_m = get_db()
        _hcur_m = _hconn_m.cursor()
        _hcur_m.execute("""INSERT INTO buscas_historico (tipo, municipio, estado, iniciado_em, job_id)
                           VALUES (%s, %s, %s, NOW(), %s) RETURNING id""",
                        ('manual', _mun_m, _est_m, job_id))
        _hist_id_manual = _hcur_m.fetchone()[0]
        _hconn_m.commit()
        _hcur_m.close()
        _hconn_m.close()
        job['hist_id'] = _hist_id_manual
    except Exception:
        pass

    def log_cb(entry):
        job['logs'].append(entry)
        # Persistir log em arquivo para sobreviver restart do worker
        try:
            import json as _json_log
            _log_path = f"/var/www/urbanlex/static/downloads/job_{job_id}.jsonl"
            with open(_log_path, 'a', encoding='utf-8') as _lf:
                _lf.write(_json_log.dumps(entry, ensure_ascii=False) + '\n')
        except Exception:
            pass
        # Interromper thread se job foi cancelado
        if job.get('cancelled'):
            raise InterruptedError('Job cancelado')
        # Se a IA confirmou um PDF, salvar resultado parcial imediatamente
        # Isso garante que mesmo se o worker reiniciar após, o resultado já está no job
        msg = entry.get('msg', '')
        if entry.get('nivel') == 'ok' and 'PDF extraído' in msg and 'pág' in msg:
            try:
                job['_partial_confirmed'] = True
            except Exception:
                pass

    def run():
        try:
            try:
                from modulos.buscador_legislacoes import busca_manual
            except Exception as import_err:
                import traceback
                job['logs'].append({'nivel': 'erro', 'msg': f'❌ Erro ao importar buscador: {str(import_err)[:200]}'})
                job['logs'].append({'nivel': 'erro', 'msg': f'Traceback: {traceback.format_exc()[-400:]}'})
                job['result'] = {'success': False, 'erro': f'Import error: {str(import_err)[:200]}', 'legislacoes': []}
                job['done'] = True
                return

            # Renovar sessão FlareSolverr antes de cada busca
            try:
                import requests as _req_fs
                _old_sid = os.environ.get('FLARESOLVERR_SESSION', '')
                if _old_sid:
                    _req_fs.post('http://localhost:8191/v1', json={'cmd': 'sessions.destroy', 'session': _old_sid}, timeout=10)
                _r_new = _req_fs.post('http://localhost:8191/v1', json={'cmd': 'sessions.create'}, timeout=10)
                _new_sid = _r_new.json().get('session', '')
                if _new_sid:
                    os.environ['FLARESOLVERR_SESSION'] = _new_sid
                    import subprocess as _sp
                    _sp.run(['sed', '-i', f's/FLARESOLVERR_SESSION=.*/FLARESOLVERR_SESSION={_new_sid}/', '/var/www/urbanlex/.env'], capture_output=True)
                    job['logs'].append({'nivel': 'info', 'msg': f'🔄 Sessão FlareSolverr renovada: {_new_sid[:8]}...'})
            except Exception as _e_fs:
                job['logs'].append({'nivel': 'aviso', 'msg': f'⚠️ Erro ao renovar sessão FS: {str(_e_fs)[:60]}'})

            job['logs'].append({'nivel': 'info', 'msg': '🔧 Thread: chamando busca_manual...'})
            result = busca_manual(d, log_callback=log_cb)
            job['logs'].append({'nivel': 'info', 'msg': f'🔧 Thread: busca_manual retornou. Tipo: {type(result).__name__}, keys: {list(result.keys()) if isinstance(result, dict) else "N/A"}'})
            if isinstance(result, dict):
                # Guardar textos COMPLETOS no job (não vão pro frontend)
                textos_completos = result.get('textos_fontes', [])
                job['_textos_completos'] = textos_completos

                # Preservar PDFs baixados — copiar de /tmp/ para localização permanente
                pdf_downloads = []
                for tf in textos_completos:
                    if tf.get('pdf_path') and os.path.isfile(tf['pdf_path']):
                        try:
                            import shutil
                            pdf_dir = os.path.join(os.path.dirname(__file__), 'static', 'downloads')
                            os.makedirs(pdf_dir, exist_ok=True)
                            # Nome baseado no job + fonte
                            nome_fonte = tf.get('nome', 'fonte').replace(' ', '_').replace('/', '_')[:30]
                            nome_pdf = f'{job_id}_{nome_fonte}.pdf'
                            dest = os.path.join(pdf_dir, nome_pdf)
                            shutil.copy2(tf['pdf_path'], dest)
                            tf['pdf_download_url'] = f'/static/downloads/{nome_pdf}'
                            pdf_downloads.append({'nome': tf.get('nome', ''), 'url': tf['pdf_download_url']})
                        except Exception as e_cp:
                            job['logs'].append({'nivel': 'aviso', 'msg': f'⚠️ Erro ao copiar PDF: {str(e_cp)[:40]}'})

                # Gerar PDF a partir do texto para fontes sem pdf_path (ex: LeisMunicipais)
                for tf in textos_completos:
                    # Usar PDF nativo S3 se disponível
                    if tf.get('_fonte') == 'leismunicipais' and tf.get('pdf_nativo_s3_path') and os.path.isfile(tf['pdf_nativo_s3_path']):
                        tf['pdf_path'] = tf['pdf_nativo_s3_path']
                        tf['pdf_download_url'] = tf['pdf_nativo_s3']
                        job['logs'].append({'nivel': 'ok', 'msg': '🎯 LeisMunicipais: usando PDF nativo S3'})
                        continue
                    if tf.get('_fonte') == 'leismunicipais' and tf.get('texto') and not tf.get('pdf_path'):
                        try:
                            from reportlab.lib.pagesizes import A4
                            from reportlab.lib.styles import getSampleStyleSheet
                            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                            from reportlab.lib.units import mm
                            pdf_dir = os.path.join(os.path.dirname(__file__), 'static', 'downloads')
                            os.makedirs(pdf_dir, exist_ok=True)
                            nome_pdf = f'{job_id}_LeisMunicipais.pdf'
                            dest = os.path.join(pdf_dir, nome_pdf)
                            leg_info = result.get('legislacoes', [{}])[0] if result.get('legislacoes') else {}
                            titulo = str(leg_info.get('tipo_legislacao', '')) + ' N. ' + str(leg_info.get('numero', '')) + '/' + str(leg_info.get('ano', ''))
                            linhas = []
                            if tf.get('html_lei'):
                                # WeasyPrint: renderiza HTML com formatação completa
                                _css = 'body{font-family:Arial,sans-serif;font-size:10pt;margin:20mm;} p{margin:4px 0;text-align:justify;}'
                                _body = '<h2>' + titulo.strip() + '</h2><p><i>Fonte: LeisMunicipais.com.br (texto consolidado)</i></p>' + tf['html_lei']
                                _html = '<html><head><meta charset="utf-8"><style>' + _css + '</style></head><body>' + _body + '</body></html>'
                                import tempfile as _tf_wk, subprocess as _sp_wk
                                with _tf_wk.NamedTemporaryFile(suffix='.html', mode='w', encoding='utf-8', delete=False) as _tmp_wk:
                                    _tmp_wk.write(_html); _tmp_wk_path = _tmp_wk.name
                                _sp_wk.run(['wkhtmltopdf','--encoding','utf-8','--quiet', _tmp_wk_path, dest], capture_output=True, timeout=60)
                                import os as _os_wk; _os_wk.unlink(_tmp_wk_path)
                                linhas = ['html_lei']
                            else:
                                # ReportLab: texto puro com quebras por marcadores legais
                                import re as _re
                                doc_rl = SimpleDocTemplate(dest, pagesize=A4,
                                    leftMargin=15*mm, rightMargin=15*mm, topMargin=15*mm, bottomMargin=15*mm)
                                styles = getSampleStyleSheet()
                                story = []
                                if titulo.strip('N. /'):
                                    story.append(Paragraph('<b>' + titulo.strip() + '</b>', styles['Title']))
                                    story.append(Spacer(1, 6*mm))
                                story.append(Paragraph('<i>Fonte: LeisMunicipais.com.br (texto consolidado)</i>', styles['Normal']))
                                story.append(Spacer(1, 4*mm))
                                _txt = tf['texto']
                                _txt = _re.sub(r'(Art[.]|CAPITULO|TITULO|SECAO|ANEXO)', lambda m: '\n' + m.group(0), _txt)
                                linhas = [l.strip() for l in _txt.splitlines() if l.strip()]
                                for linha in linhas:
                                    safe = linha.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                                    story.append(Paragraph(safe, styles['Normal']))
                                doc_rl.build(story)
                            tf['pdf_path'] = dest
                            tf['pdf_download_url'] = f'/static/downloads/{nome_pdf}'
                            pdf_downloads.append({'nome': '\U0001f4d6 LeisMunicipais (texto)', 'url': tf['pdf_download_url']})
                            job['logs'].append({'nivel': 'ok', 'msg': f'\U0001f4c4 LeisMunicipais: PDF gerado ({len(linhas)} linhas)'})
                        except Exception as e_rl:
                            job['logs'].append({'nivel': 'aviso', 'msg': f'\u26a0\ufe0f Erro ao gerar PDF LeisMunicipais: {str(e_rl)[:80]}'})

                if pdf_downloads:
                    job['logs'].append({'nivel': 'ok', 'msg': f'\U0001f4e5 {len(pdf_downloads)} PDF(s) disponivel(is) para download'})

                # Truncar textos para preview no frontend (5000 chars)
                textos_preview = []
                for tf in textos_completos:
                    preview = dict(tf)
                    txt = preview.get('texto', '')
                    preview['texto_preview'] = txt[:5000] + ('...' if len(txt) > 5000 else '')
                    preview['texto_chars'] = len(txt)
                    preview.pop('texto', None)  # Não enviar texto completo pro frontend
                    preview.pop('pdf_path', None)  # Não expor path do servidor
                    # Manter pdf_download_url se existir
                    textos_preview.append(preview)

                job['result'] = {
                    'success': not result.get('error'),
                    'legislacoes': result.get('legislacoes', []),
                    'erro': result.get('erro') or result.get('error'),
                    'fontes': result.get('fontes', []),
                    'textos_fontes': textos_preview,
                    'sugestao': result.get('sugestao'),
                    'pdf_downloads': pdf_downloads,
                }
            else:
                job['result'] = {'success': True, 'legislacoes': result or []}
        except Exception as e:
            import traceback
            tb_str = traceback.format_exc()
            err_msg = f'Exceção: {str(e)[:200]}'
            # Adicionar ao log stream (aparece no frontend em tempo real)
            job['logs'].append({'nivel': 'erro', 'msg': err_msg})
            job['logs'].append({'nivel': 'erro', 'msg': f'Traceback: {tb_str[-300:]}'})
            job['result'] = {
                'success': False,
                'erro': str(e)[:300],
                'error': str(e)[:300],
                'legislacoes': [],
                'logs': [{'nivel': 'erro', 'msg': err_msg}]
            }
        job['done'] = True
        job['hist_id'] = _hist_id_manual
        # Salvar resultado no historico
        if _hist_id_manual:
            try:
                _legs_m = (job.get('result') or {}).get('legislacoes', [])
                _leg_m = _legs_m[0] if _legs_m else {}
                _log_m = "\n".join(l.get("msg","") for l in job["logs"])
                _hconn_m2 = get_db()
                _hcur_m2 = _hconn_m2.cursor()
                import json as _json_m2
                _pdf_m = _leg_m.get("pdf_nativo_s3") or _leg_m.get("pdf_path") or ""
                _anx_m = _leg_m.get("anexos_lm") or []
                _anx_m_json = _json_m2.dumps(_anx_m) if _anx_m else "[]"
                _hcur_m2.execute("""UPDATE buscas_historico SET
                    concluido_em=NOW(), sucesso=%s,
                    legislacao_tipo=%s, legislacao_numero=%s,
                    legislacao_ano=%s, legislacao_link=%s,
                    log_texto=%s, pdf_path=%s, anexos_paths=%s
                    WHERE id=%s""",
                    (bool(_legs_m),
                     _leg_m.get("tipo",""), _leg_m.get("numero",""),
                     _leg_m.get("ano",""), _leg_m.get("url_fonte","") or _leg_m.get("link",""),
                     _log_m, _pdf_m, _anx_m_json, _hist_id_manual))
                _hconn_m2.commit()
                _hcur_m2.close()
                _hconn_m2.close()
            except Exception:
                pass

    t = threading.Thread(target=run, daemon=True)
    t.start()
    return jsonify({'success': True, 'job_id': job_id})


@app.route('/api/buscador/job/<job_id>')
@login_required
def api_buscador_job_poll(job_id):
    """Polling: retorna logs novos + resultado quando pronto."""
    job = _buscador_jobs.get(job_id)
    if not job:
        # Tentar recuperar logs do banco (log_persistente)
        try:
            from modulos.log_persistente import carregar_logs, contar_logs
            _cursor_req = int(request.args.get('cursor', 0))
            _logs_db = carregar_logs(job_id, get_db, _cursor_req)
            if _logs_db:
                _total_db = contar_logs(job_id, get_db)
                _hist_id_db = None
                try:
                    _c2=get_db(); _cu2=_c2.cursor()
                    _cu2.execute("SELECT id FROM buscas_historico WHERE job_id=%s LIMIT 1",(job_id,))
                    _h2=_cu2.fetchone()
                    if _h2: _hist_id_db=_h2[0]
                    _cu2.close(); _c2.close()
                except: pass
                # Verificar se job ainda esta rodando na fila
                _ainda_rodando = False
                try:
                    _c3=get_db(); _cu3=_c3.cursor()
                    _cu3.execute("SELECT status FROM fila_buscas WHERE job_id=%s LIMIT 1",(job_id,))
                    _r3=_cu3.fetchone()
                    if _r3: _ainda_rodando = _r3[0] == 'rodando'
                    _cu3.close(); _c3.close()
                except: pass
                # Buscar tipo do job em buscas_historico
                _tipo_job = None
                try:
                    _c4=get_db(); _cu4=_c4.cursor()
                    _cu4.execute("SELECT tipo FROM buscas_historico WHERE job_id=%s LIMIT 1",(job_id,))
                    _r4=_cu4.fetchone()
                    if _r4: _tipo_job = _r4[0]
                    _cu4.close(); _c4.close()
                except: pass
                return jsonify({'success':True,'logs':[{'nivel':l['nivel'],'msg':l['msg']} for l in _logs_db],'cursor':_cursor_req+len(_logs_db),'done': not _ainda_rodando,'hist_id':_hist_id_db,'result':None,'tipo':_tipo_job})
        except Exception as _edb: pass
        # Tentar recuperar logs do arquivo persistido
        import json as _json_rec
        _log_path = f"/var/www/urbanlex/static/downloads/job_{job_id}.jsonl"
        if os.path.exists(_log_path):
            try:
                with open(_log_path, 'r', encoding='utf-8') as _lf:
                    _logs_rec = [_json_rec.loads(l) for l in _lf if l.strip()]
                cursor = int(request.args.get('cursor', 0))
                # Buscar hist_id no banco
                _hist_id_rec = None
                try:
                    _conn_rec = get_db()
                    _cur_rec = _conn_rec.cursor()
                    _cur_rec.execute("SELECT id FROM buscas_historico WHERE job_id=%s LIMIT 1", (job_id,))
                    _h_rec = _cur_rec.fetchone()
                    if _h_rec: _hist_id_rec = _h_rec[0]
                    _cur_rec.close()
                    _conn_rec.close()
                except Exception:
                    pass
                # Recuperar resultado se existir
                _result_rec = None
                for _lr in _logs_rec:
                    if '_result' in _lr:
                        _result_rec = _lr['_result']
                _logs_rec_clean = [l for l in _logs_rec if '_result' not in l]
                _job_done = _result_rec is not None
                return jsonify({'success': True, 'logs': _logs_rec_clean[cursor:], 'cursor': len(_logs_rec_clean), 'done': _job_done, 'recuperado': True, 'hist_id': _hist_id_rec, 'result': _result_rec})
            except Exception:
                pass
        return jsonify({'success': False, 'error': 'Job não encontrado', 'done': False, 'logs': [], 'cursor': 0})

    cursor = int(request.args.get('cursor', 0))
    new_logs = job['logs'][cursor:]
    next_cursor = len(job['logs'])

    resp = {
        'success': True,
        'logs': new_logs,
        'cursor': next_cursor,
        'done': job['done'],
    }

    resp['hist_id'] = job.get('hist_id')  # Sempre retornar hist_id
    resp['tipo'] = job.get('tipo', 'manual')
    if job['done']:
        resp['result'] = job['result']
        resp['job_id'] = job_id  # Frontend precisa pro cadastrar

    return jsonify(resp)


@app.route('/api/buscador/verificar-duplicatas', methods=['POST'])
@editor_required
def api_buscador_verificar_duplicatas():
    d = request.json or {}
    legislacoes = d.get('legislacoes', [])
    municipio = d.get('municipio', '')
    duplicatas = []
    try:
        conn = get_db()
        cur = conn.cursor()
        for leg in legislacoes:
            tipo = (leg.get('tipo_legislacao') or leg.get('tipo') or '').strip()
            numero = (leg.get('numero') or '').strip()
            ano = str(leg.get('ano') or '').strip()
            sql = 'SELECT id, tipo_nome, numero, ano, municipio_nome, url_original FROM legislacoes WHERE LOWER(tipo_nome)=LOWER(%s) AND numero=%s AND ano=%s AND LOWER(municipio_nome)=LOWER(%s) LIMIT 1'
            cur.execute(sql, (tipo, numero, ano, municipio))
            row = cur.fetchone()
            if row:
                duplicatas.append({'id': row[0], 'tipo': row[1], 'numero': row[2], 'ano': row[3], 'municipio': row[4], 'url_cadastrada': row[5] if len(row) > 5 else ''})
        cur.close()
        conn.close()
    except Exception as e:
        return jsonify({'duplicatas': [], 'error': str(e)[:200]})
    return jsonify({'duplicatas': duplicatas})

@app.route('/api/buscador/cadastrar', methods=['POST'])
@editor_required
def api_buscador_cadastrar():
    """Cadastra legislações selecionadas na biblioteca, com texto da fonte escolhida."""
    d = request.json or {}
    legislacoes = d.get('legislacoes', [])
    municipio = d.get('municipio', '')
    estado = d.get('estado', '')
    monitorar = d.get('monitorar', True)
    substituir = d.get('substituir', False)
    job_id = d.get('job_id', '')

    if not legislacoes:
        return jsonify({'success': False, 'error': 'Nenhuma legislação para cadastrar'}), 400

    # Recuperar textos completos do cache do job (se disponível)
    textos_cache = {}
    pdf_cache = {}
    anexos_cache = {}
    if job_id and job_id in _buscador_jobs:
        job = _buscador_jobs[job_id]
        for tf in job.get('_textos_completos', []):
            url = tf.get('url', '')
            textos_cache[url] = tf.get('texto', '')
            if tf.get('pdf_path'):
                pdf_cache[url] = tf['pdf_path']
            if tf.get('pdf_download_url'):
                pdf_cache[url + '_dl'] = tf['pdf_download_url']
            if tf.get('anexos_lm'):
                anexos_cache[url] = tf['anexos_lm']

    # Enriquecer cada legislação com o texto e PDF da fonte escolhida
    for leg in legislacoes:
        fonte_url = leg.get('fonte_selecionada_url', '') or leg.get('url_fonte', '')
        if fonte_url and fonte_url in textos_cache:
            leg['texto_integral'] = textos_cache[fonte_url]
            leg['url_texto_fonte'] = fonte_url
        if fonte_url and fonte_url in pdf_cache:
            leg['pdf_path'] = pdf_cache[fonte_url]
            leg['pdf_download_url'] = pdf_cache.get(fonte_url + '_dl', '')
        if fonte_url and fonte_url in anexos_cache:
            leg['anexos_lm'] = anexos_cache[fonte_url]
        # Incluir anexos LeisMunicipais independente da fonte selecionada
        if not leg.get('anexos_lm') and anexos_cache:
            for _anx_list in anexos_cache.values():
                if _anx_list:
                    leg['anexos_lm'] = _anx_list
                    break

    try:
        from modulos.buscador_legislacoes import cadastrar_resultados
        # Passar substituir para cada legislação
        if substituir:
            for leg in legislacoes:
                leg['substituir'] = True
        ids = cadastrar_resultados(legislacoes, municipio, estado, monitorar)
        return jsonify({'success': True, 'cadastradas': len(ids), 'ids': ids})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500

@app.route('/api/buscador/varrer-referencias', methods=['POST'])
@editor_required
def api_buscador_varrer_refs():
    """Varre a biblioteca buscando referências a legislações não cadastradas."""
    try:
        from modulos.buscador_legislacoes import varrer_referencias
        refs = varrer_referencias()
        return jsonify({'success': True, 'referencias': refs})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500

@app.route('/api/buscador/fila')
@login_required
def api_buscador_fila():
    """Retorna histórico/fila de buscas."""
    try:
        from modulos.buscador_legislacoes import listar_fila
        return jsonify({'success': True, 'data': listar_fila()})
    except Exception as e:
        return jsonify({'success': True, 'data': []})


@app.route('/api/buscador/historico')
@login_required
def api_buscador_historico():
    """Retorna historico de buscas da tabela buscas_historico."""
    try:
        limit = int(request.args.get('limit', 50))
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""SELECT id, tipo, municipio, estado,
                              iniciado_em, concluido_em, sucesso,
                              job_id,
                              legislacao_tipo, legislacao_numero,
                              legislacao_ano, legislacao_link,
                              pdf_path, anexos_paths,
                              relatorio_path, tabela_path, zip_path
                       FROM buscas_historico
                       ORDER BY iniciado_em DESC LIMIT %s""", (limit,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        data = []
        for r in (rows or []):
            # Calcular status
            _job_id = r.get('job_id') if 'job_id' in r else None
            if r['concluido_em']:
                _status = 'concluida'
            elif _job_id and _buscador_jobs.get(_job_id) and not _buscador_jobs[_job_id].get('done'):
                _status = 'em_andamento'
            else:
                # Verificar arquivo JSONL recente
                import os as _os_h, time as _time_h
                _log_path_h = f"/var/www/urbanlex/static/downloads/job_{_job_id}.jsonl" if _job_id else ""
                if _job_id and _os_h.path.exists(_log_path_h) and (_time_h.time() - _os_h.path.getmtime(_log_path_h)) < 600:
                    _status = 'em_andamento'
                else:
                    _status = 'interrompida'
            data.append({
                'id': r['id'],
                'tipo': r['tipo'],
                'municipio': r['municipio'],
                'estado': r['estado'],
                'iniciado_em': r['iniciado_em'].isoformat() if r['iniciado_em'] else None,
                'concluido_em': r['concluido_em'].isoformat() if r['concluido_em'] else None,
                'sucesso': r['sucesso'],
                'status': _status,
                'legislacao_tipo': r['legislacao_tipo'],
                'legislacao_numero': r['legislacao_numero'],
                'legislacao_ano': r['legislacao_ano'],
                'legislacao_link': r['legislacao_link'],
                'pdf_path': r['pdf_path'],
                'anexos_paths': r['anexos_paths'],
                'relatorio_path': r['relatorio_path'],
                'tabela_path': r['tabela_path'],
                'zip_path': r['zip_path'],
            })
        return jsonify({'success': True, 'data': data})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:200]})

@app.route('/api/buscador/historico/log/<int:hist_id>')
@login_required
def api_buscador_historico_log(hist_id):
    """Retorna log de uma busca especifica."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT log_texto, municipio, estado FROM buscas_historico WHERE id=%s", (hist_id,))
        r = cur.fetchone()
        cur.close()
        conn.close()
        if not r:
            return jsonify({'success': False, 'error': 'Não encontrado'}), 404
        return jsonify({'success': True, 'log': r['log_texto'] or '', 'municipio': r['municipio'], 'estado': r['estado']})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:200]})

@app.route('/api/buscador/historico/log/<int:hist_id>/download')
@login_required
def api_buscador_historico_log_download(hist_id):
    """Download do log de uma busca."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT log_texto, municipio, estado, iniciado_em, tipo, legislacao_tipo, legislacao_numero, legislacao_ano FROM buscas_historico WHERE id=%s", (hist_id,))
        r = cur.fetchone()
        cur.close()
        conn.close()
        if not r:
            return "Não encontrado", 404
        from flask import Response
        import re as _re_log, datetime as _dt_log
        _tipo_busca = (r.get('tipo') or 'automatica').lower()
        _est = (r.get('estado') or 'XX').upper()
        _mun = (r.get('municipio') or 'municipio').replace(' ', '_')
        _dt = r.get('iniciado_em')
        if _dt:
            import datetime as _dtmod
            _dt_brt = _dt + _dtmod.timedelta(hours=-3)
            _dt_str = _dt_brt.strftime('%d_%m_%Y_%Hh%M')
        else:
            _dt_str = 'sem_data'
        if _tipo_busca == 'manual':
            if r.get('legislacao_numero'):
                _leg_tipo = (r.get('legislacao_tipo') or 'Lei').replace(' ', '_')
                _leg_num = r.get('legislacao_numero') or '0'
                _leg_ano = r.get('legislacao_ano') or '0000'
                nome = f"Log_busca_manual_{_est}_{_mun}_{_leg_tipo}_{_leg_num}_{_leg_ano}_data_busca_{_dt_str}.txt"
            else:
                nome = f"Log_busca_manual_{_est}_{_mun}_data_busca_{_dt_str}.txt"
        else:
            nome = f"Log_busca_automatica_{_est}_{_mun}_data_busca_{_dt_str}.txt"
        nome = _re_log.sub(r'[^a-zA-Z0-9_\-.]', '', nome)
        # Se log vazio no banco, ler do job ativo em memoria
        _log_txt = r['log_texto'] or ''
        if not _log_txt:
            for _jv in _buscador_jobs.values():
                if _jv.get('hist_id') == hist_id:
                    _log_txt = "\n".join(l.get("msg","") for l in _jv.get("logs", []))
                    break
        return Response(_log_txt, mimetype='text/plain',
                       headers={'Content-Disposition': f'attachment; filename="{nome}"'})
    except Exception as e:
        return str(e), 500

@app.route('/api/buscador/historico/download/<int:hist_id>')
@login_required
def api_buscador_historico_download(hist_id):
    """Gera ZIP master com relatorio PDF + tabela PDF + zip das legislacoes."""
    import zipfile, io, json as _json_z, os as _os_z
    from flask import send_file
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT municipio, estado, sucesso, pdf_path, anexos_paths, legislacao_numero, legislacao_ano, relatorio_path, tabela_path, zip_path FROM buscas_historico WHERE id=%s", (hist_id,))
        r = cur.fetchone()
        cur.close()
        conn.close()
        if not r:
            return "Não encontrado", 404
        if not r.get('sucesso'):
            return jsonify({'success': False, 'error': 'A busca nao foi concluida com sucesso — pacote indisponivel'}), 400
        base_dir = "/var/www/urbanlex"
        mun = (r.get('municipio') or 'municipio').replace(' ','_')
        est = (r.get('estado') or 'XX').replace(' ','_')
        zip_buffer = io.BytesIO()
        arquivos_incluidos = []
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            def _add(url_or_path, nome_interno):
                if not url_or_path: return
                full = base_dir + url_or_path if url_or_path.startswith('/static') else url_or_path
                if _os_z.path.exists(full):
                    zf.write(full, nome_interno)
                    arquivos_incluidos.append(nome_interno)
            _add(r.get('relatorio_path'), 'relatorio.pdf')
            _add(r.get('tabela_path'), 'tabela_legislacoes.pdf')
            _add(r.get('zip_path'), 'legislacoes.zip')
        if not arquivos_incluidos:
            return jsonify({'success': False, 'error': 'Nenhum arquivo encontrado para esta busca'}), 404
        zip_buffer.seek(0)
        nome_zip = f"busca_{mun}_{est}_id{hist_id}.zip"
        return send_file(zip_buffer, mimetype='application/zip', as_attachment=True, download_name=nome_zip)
    except Exception as e:
        return str(e), 500

@app.route('/api/buscador/historico/apagar', methods=['POST'])
@editor_required
def api_buscador_historico_apagar():
    """Apaga buscas selecionadas do historico."""
    try:
        ids = (request.json or {}).get('ids', [])
        if not ids:
            return jsonify({'success': False, 'error': 'Nenhum ID informado'})
        conn = get_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM buscas_historico WHERE id = ANY(%s)", (ids,))
        deleted = cur.rowcount
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'success': True, 'apagados': deleted})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:200]})

@app.route('/api/buscador/autocomplete-municipios')
@login_required
def api_buscador_autocomplete():
    """Autocomplete de municípios — IBGE (5570) + DB (com sites salvos)."""
    q = request.args.get('q', '').strip()
    estado = request.args.get('estado', '').strip()
    if len(q) < 2:
        return jsonify({'municipios': []})

    # Garantir IBGE carregado
    if not _ibge_carregado:
        _carregar_ibge()

    q_lower = q.lower()

    # 1) Buscar no IBGE (em memória — instantâneo)
    ibge_matches = []
    for m in _ibge_municipios:
        if estado and m['estado'] != estado:
            continue
        if q_lower in m['nome'].lower():
            ibge_matches.append({'nome': m['nome'], 'estado': m['estado'], 'tem_sites': False})
            if len(ibge_matches) >= 50:
                break

    # 2) Buscar no DB (municípios com sites salvos)
    db_com_sites = set()
    try:
        pattern = f'%{q}%'
        params = [pattern]
        estado_filter = ''
        if estado:
            estado_filter = ' AND estado = %s'
            params.append(estado)

        sql = f"""SELECT DISTINCT municipio_nome as nome, estado
                  FROM municipio_sites_referencia
                  WHERE municipio_nome ILIKE %s {estado_filter}
                  LIMIT 50"""
        rows = qry(sql, params)
        for r in (rows or []):
            db_com_sites.add((r['nome'], r['estado']))
    except Exception:
        pass

    # 3) Merge: marcar quais têm sites
    seen = {}
    for m in ibge_matches:
        key = (m['nome'], m['estado'])
        m['tem_sites'] = key in db_com_sites
        seen[key] = m

    # Adicionar do DB que não vieram do IBGE (caso IBGE não carregou)
    for (nome, uf) in db_com_sites:
        key = (nome, uf)
        if key not in seen:
            seen[key] = {'nome': nome, 'estado': uf, 'tem_sites': True}

    # 4) Ordenar: matches que COMEÇAM com a query primeiro, depois os demais
    result = sorted(seen.values(), key=lambda x: (
        0 if x['nome'].lower().startswith(q_lower) else 1,
        x['estado'],
        x['nome']
    ))[:15]

    return jsonify({'municipios': result})


@app.route('/api/buscador/sites-referencia', methods=['GET'])
@login_required
def api_buscador_sites_get():
    """Carrega fontes prioritárias salvas para um município."""
    municipio = request.args.get('municipio', '').strip()
    estado = request.args.get('estado', '').strip()
    if not municipio:
        return jsonify({'success': False, 'error': 'Município obrigatório'})

    try:
        row = qry("""SELECT url_diario, urls_extras FROM municipio_sites_referencia
                     WHERE municipio_nome ILIKE %s AND (estado = %s OR estado = '' OR %s = '')
                     ORDER BY CASE WHEN estado = %s THEN 0 ELSE 1 END
                     LIMIT 1""", (municipio, estado, estado, estado), 'one')
        if row:
            # Compatibilidade: migrar url_diario + urls_extras → fontes_prioritarias
            fontes = list(row.get('urls_extras') or [])
            url_diario_old = (row.get('url_diario') or '').strip()
            if url_diario_old and url_diario_old not in fontes:
                fontes.insert(0, url_diario_old)
            return jsonify({'success': True, 'sites': {
                'fontes_prioritarias': fontes[:3]
            }})
        return jsonify({'success': True, 'sites': None})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:200]})


@app.route('/api/buscador/sites-referencia', methods=['POST'])
@login_required
def api_buscador_sites_post():
    """Salva fontes prioritárias para um município."""
    d = request.get_json(force=True)
    municipio = (d.get('municipio') or '').strip()
    estado_raw = (d.get('estado') or '').strip()
    fontes = d.get('fontes_prioritarias') or []
    if isinstance(fontes, str):
        fontes = [f.strip() for f in fontes.split('\n') if f.strip()]
    fontes = [f for f in fontes[:3] if f]

    if not municipio:
        return jsonify({'success': False, 'error': 'Município obrigatório'})

    # Normalizar estado: apenas UF válida de 2 letras; "Todos"/vazio/qualquer outro → ''
    estado = estado_raw.upper() if (len(estado_raw) == 2 and estado_raw.isalpha()) else ''

    try:
        # Upsert por nome (ILIKE) — evita duplicatas com estado diferente
        updated = qry("""UPDATE municipio_sites_referencia
                         SET urls_extras = %s, estado = %s, url_diario = '', atualizado_em = NOW()
                         WHERE municipio_nome ILIKE %s""",
                      (fontes, estado, municipio), commit=True, fetch=None)
        # Se não existia nenhum registro com esse nome, inserir
        exists = qry("SELECT 1 FROM municipio_sites_referencia WHERE municipio_nome ILIKE %s",
                     (municipio,), fetch='one')
        if not exists:
            qry("""INSERT INTO municipio_sites_referencia (municipio_nome, estado, url_diario, urls_extras, atualizado_em)
                   VALUES (%s, %s, '', %s, NOW())""",
                (municipio, estado, fontes), commit=True, fetch=None)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:200]})

# -------------------------------------------------------------------

# -- Diagnostico R2 (CORRIGIDO: decorator ativo + sem chars especiais) --
@app.route('/api/diagnostico/r2')
@admin_required
def api_diagnostico_r2():
    info = {
        'r2_importado': _R2_IMPORTADO,
        'r2_disponivel': r2_disponivel(),
        'R2_ACCOUNT_ID': bool(os.getenv('R2_ACCOUNT_ID')),
        'R2_ACCESS_KEY': bool(os.getenv('R2_ACCESS_KEY')),
        'R2_SECRET_KEY': bool(os.getenv('R2_SECRET_KEY')),
        'R2_BUCKET_NAME': os.getenv('R2_BUCKET_NAME', ''),
    }
    if info['r2_disponivel']:
        try:
            url_teste = r2_upload(b'teste urbanlex', 'teste_diagnostico.txt', leg_id=0)
            info['upload_teste'] = url_teste
            info['upload_ok'] = bool(url_teste)
            if url_teste:
                r2_delete(url_teste)
                info['delete_ok'] = True
        except Exception as e:
            info['upload_erro'] = str(e)
            info['upload_ok'] = False
    return jsonify(info)

@app.route('/health')
def health():
    try:
        qry("SELECT 1", fetch='one')
        return jsonify({'status':'ok','db':'conectado','version':'3.5'})
    except Exception as e:
        return jsonify({'status':'erro','error':str(e)}), 500

# -- FIX 4: Rotas diagnostico agora exigem token SECRET_KEY --
@app.route('/setup-banco-agora')
def setup_banco():
    token = request.args.get('token', '')
    if token != os.getenv('SECRET_KEY', ''):
        return 'Acesso negado. Use ?token=SUA_SECRET_KEY', 403
    try:
        with open('schema_final.sql', 'r') as f:
            sql = f.read()
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute(sql)
        conn.commit()
        cur.close()
        conn.close()
        return 'Banco inicializado com sucesso!'
    except Exception as e:
        return f'Erro: {str(e)}'

@app.route('/aprovar-admin-agora')
def aprovar_admin():
    token = request.args.get('token', '')
    if token != os.getenv('SECRET_KEY', ''):
        return 'Acesso negado. Use ?token=SUA_SECRET_KEY', 403
    try:
        admin_email = os.environ.get('ADMIN_EMAIL')
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("UPDATE users SET ativo=TRUE, aprovado=TRUE, role='admin' WHERE email=%s", (admin_email,))
        conn.commit()
        cur.close()
        conn.close()
        return f'Admin {admin_email} atualizado! <a href="/login">Fazer login</a>'
    except Exception as e:
        return f'Erro: {str(e)}'

# -------------------------------------------------------------------
# ROTAS ADICIONAIS (FIX 5: movidas para ANTES do if __name__)
# -------------------------------------------------------------------

@app.route('/parametros/nova-zona')
@login_required
def pagina_nova_zona():
    return render_template('parametros.html', active_page='nova-zona', active_group='parametros',
                           abrir_modal='nova_zona', **tmpl_ctx())

@app.route('/api/zona/<int:zona_id>', methods=['PUT'])
@editor_required
def api_atualizar_zona(zona_id):
    d = request.json or {}
    if not d:
        return jsonify({'success': False, 'error': 'Nenhum dado enviado'}), 400
    campos_proibidos = {'id', 'criado_em', 'municipio', 'zona', 'subzona'}
    cols = [k for k in d.keys() if k not in campos_proibidos]
    if not cols:
        return jsonify({'success': False, 'error': 'Nenhum campo valido para atualizar'}), 400
    set_sql = ', '.join(f"{c} = %s" for c in cols)
    vals = [d[c] for c in cols] + [datetime.now(), session['user_id'], zona_id]
    qry(f"UPDATE zonas_urbanas SET {set_sql}, atualizado_em = %s, atualizado_por = %s WHERE id = %s",
        vals, commit=True, fetch=None)
    return jsonify({'success': True})

@app.route('/api/zonas')
@login_required
def api_zonas_query():
    municipio = request.args.get('municipio')
    zona      = request.args.get('zona')
    estado    = request.args.get('estado')
    where = []
    params = []
    if municipio: where.append("municipio ILIKE %s"); params.append(f"%{municipio}%")
    if zona:      where.append("zona ILIKE %s");      params.append(f"%{zona}%")
    if estado:    where.append("estado ILIKE %s");    params.append(f"%{estado}%")
    where_sql = ('WHERE ' + ' AND '.join(where)) if where else ''
    data = qry(f"SELECT * FROM zonas_urbanas {where_sql} ORDER BY estado, municipio, zona LIMIT 200", params)
    return jsonify({'success': True, 'data': data, 'total': len(data)})

@app.route('/api/parametros/importar', methods=['POST'])
@editor_required
def api_parametros_importar():
    arquivo = request.files.get('arquivo')
    if not arquivo:
        return jsonify({'success': False, 'error': 'Nenhum arquivo enviado'}), 400
    if not PANDAS_OK:
        return jsonify({'success': False, 'error': 'Pandas nao instalado'}), 500
    try:
        df = pd.read_excel(io.BytesIO(arquivo.read()))
        df.columns = [c.strip().lower().replace(' ', '_') for c in df.columns]
        ok = erro = 0
        for _, row in df.iterrows():
            try:
                d = {k: (None if pd.isna(v) else v) for k, v in row.items()}
                municipio = str(d.get('municipio', '')).strip()
                zona      = str(d.get('zona', '')).strip()
                if not municipio or not zona:
                    erro += 1; continue
                subzona = str(d.get('subzona', '') or '')
                cols  = [k for k in d if d[k] is not None]
                vals  = [d[c] for c in cols]
                ph    = ', '.join(['%s'] * len(cols))
                upd   = ', '.join(f"{c}=EXCLUDED.{c}" for c in cols if c not in ('municipio','zona','subzona'))
                if upd:
                    qry(f"INSERT INTO zonas_urbanas ({','.join(cols)}) VALUES ({ph}) "
                        f"ON CONFLICT (municipio,zona,subzona) DO UPDATE SET {upd}, atualizado_em=NOW()",
                        vals, commit=True, fetch=None)
                else:
                    qry(f"INSERT INTO zonas_urbanas ({','.join(cols)}) VALUES ({ph}) ON CONFLICT DO NOTHING",
                        vals, commit=True, fetch=None)
                ok += 1
            except Exception as e:
                erro += 1
        return jsonify({'success': True, 'ok': ok, 'erro': erro,
                        'message': f'{ok} zonas importadas, {erro} erros'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/parametros/exportar')
@login_required
def api_parametros_exportar():
    if not PANDAS_OK:
        return jsonify({'success': False, 'error': 'Pandas nao instalado'}), 500
    try:
        data = qry("SELECT * FROM zonas_urbanas ORDER BY estado, municipio, zona, subzona")
        df = pd.DataFrame(data)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Parametros')
        output.seek(0)
        return send_file(output, as_attachment=True,
                         download_name=f'urbanlex_parametros_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx',
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/integracao/aprovadas')
@login_required
def api_integ_aprovadas():
    data = qry("""SELECT i.*, CONCAT(l.numero,'/',l.ano) as legislacao_ref, u.nome as revisor_nome
        FROM integracao_atualizacoes i
        LEFT JOIN legislacoes l ON i.legislacao_id=l.id
        LEFT JOIN users u ON i.revisado_por=u.id
        WHERE i.status='aprovado' ORDER BY i.revisado_em DESC LIMIT 100""")
    return jsonify({'success': True, 'data': data})

@app.route('/api/integracao/rejeitadas')
@login_required
def api_integ_rejeitadas():
    data = qry("""SELECT i.*, CONCAT(l.numero,'/',l.ano) as legislacao_ref, u.nome as revisor_nome
        FROM integracao_atualizacoes i
        LEFT JOIN legislacoes l ON i.legislacao_id=l.id
        LEFT JOIN users u ON i.revisado_por=u.id
        WHERE i.status='rejeitado' ORDER BY i.revisado_em DESC LIMIT 100""")
    return jsonify({'success': True, 'data': data})

# -------------------------------------------------------------------
# INICIALIZACAO (FIX 6: inicializar() no nivel do modulo)
# -------------------------------------------------------------------

def inicializar():
    print("UrbanLex v3.5 iniciando...")
    # Criar tabela de arquivos se nao existe
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS legislacao_arquivos (
                id SERIAL PRIMARY KEY,
                legislacao_id INTEGER NOT NULL REFERENCES legislacoes(id) ON DELETE CASCADE,
                nome_arquivo VARCHAR(500) NOT NULL,
                arquivo_tipo VARCHAR(50),
                arquivo_url TEXT,
                tamanho_bytes BIGINT DEFAULT 0,
                hash_conteudo VARCHAR(64),
                conteudo_texto TEXT,
                criado_em TIMESTAMP DEFAULT NOW(),
                criado_por INTEGER
            )
        """)
        # Indice para busca por legislacao
        cur.execute("""
            CREATE INDEX IF NOT EXISTS idx_leg_arquivos_leg_id
            ON legislacao_arquivos(legislacao_id)
        """)
        conn.commit()
        conn.close()
        print("Tabela legislacao_arquivos verificada/criada")
    except Exception as e:
        print(f"Aviso criacao tabela arquivos: {e}")
    # Criar tabela de log detalhado por legislação
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS monitoramento_legislacao_log (
                id SERIAL PRIMARY KEY,
                execucao_id INTEGER REFERENCES scheduler_execucoes(id) ON DELETE CASCADE,
                legislacao_id INTEGER REFERENCES legislacoes(id) ON DELETE CASCADE,
                municipio_id INTEGER REFERENCES municipios(id) ON DELETE CASCADE,
                data TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                iniciada_em TIMESTAMP,
                finalizada_em TIMESTAMP,
                status VARCHAR(20) DEFAULT 'ok',
                sucesso BOOLEAN DEFAULT TRUE,
                publicacoes_encontradas INTEGER DEFAULT 0,
                publicacoes_analisadas INTEGER DEFAULT 0,
                alteracoes_detectadas INTEGER DEFAULT 0,
                publicacoes_duplicadas INTEGER DEFAULT 0,
                metodo_busca VARCHAR(30),
                url_acessada TEXT,
                mensagem TEXT,
                erro TEXT
            )
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_mll_legislacao ON monitoramento_legislacao_log(legislacao_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_mll_municipio ON monitoramento_legislacao_log(municipio_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_mll_data ON monitoramento_legislacao_log(data DESC)")
        # Adicionar coluna ultima_verificacao_monitoramento se não existir
        try:
            cur.execute("ALTER TABLE legislacoes ADD COLUMN IF NOT EXISTS ultima_verificacao_monitoramento DATE")
        except Exception:
            pass
        try:
            cur.execute("ALTER TABLE legislacoes ADD COLUMN IF NOT EXISTS data_fim_monitoramento DATE")
        except Exception:
            pass
        # v5: coluna para código IBGE (necessária para API Querido Diário)
        try:
            cur.execute("ALTER TABLE municipios ADD COLUMN IF NOT EXISTS codigo_ibge VARCHAR(10)")
        except Exception:
            pass
        # ── v6: Agente Autônomo ──
        try:
            cur.execute("""CREATE TABLE IF NOT EXISTS diarios_oficiais (
                id SERIAL PRIMARY KEY, esfera VARCHAR(20) DEFAULT 'municipal',
                uf VARCHAR(2), municipio_id INTEGER, nome VARCHAR(200),
                url_principal VARCHAR(500), url_busca VARCHAR(500),
                tipo_plataforma VARCHAR(50), metodo_busca JSONB,
                codigo_ibge VARCHAR(10), verificado_em TIMESTAMP,
                funcionando BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP DEFAULT NOW(), updated_at TIMESTAMP DEFAULT NOW())""")
            cur.execute("""CREATE TABLE IF NOT EXISTS feed_atividades (
                id SERIAL PRIMARY KEY, tipo VARCHAR(50) NOT NULL,
                mensagem TEXT NOT NULL, detalhes JSONB,
                lida BOOLEAN DEFAULT FALSE, criado_em TIMESTAMP DEFAULT NOW())""")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_feed_criado ON feed_atividades(criado_em DESC)")
            cur.execute("""CREATE TABLE IF NOT EXISTS alteracoes_parametros (
                id SERIAL PRIMARY KEY, legislacao_id INTEGER, legislacao_alteradora_id INTEGER,
                publicacao_data DATE, parametro VARCHAR(100), zona VARCHAR(50),
                valor_anterior TEXT, valor_novo TEXT, confianca DECIMAL(3,2),
                aprovado_usuario BOOLEAN, aprovado_em TIMESTAMP,
                created_at TIMESTAMP DEFAULT NOW())""")
            cur.execute("""CREATE TABLE IF NOT EXISTS agente_memoria (
                id SERIAL PRIMARY KEY, diario_id INTEGER, url_base VARCHAR(500),
                passos JSONB NOT NULL, sucesso_count INTEGER DEFAULT 1,
                falha_count INTEGER DEFAULT 0, ultimo_uso TIMESTAMP DEFAULT NOW(),
                created_at TIMESTAMP DEFAULT NOW(), updated_at TIMESTAMP DEFAULT NOW())""")
            cur.execute("""CREATE TABLE IF NOT EXISTS integracao_log (
                id SERIAL PRIMARY KEY, tipo VARCHAR(50), municipios_consultados INTEGER DEFAULT 0,
                novos_detectados INTEGER DEFAULT 0, legislacoes_cadastradas INTEGER DEFAULT 0,
                detalhes JSONB, status VARCHAR(20) DEFAULT 'concluido',
                criado_em TIMESTAMP DEFAULT NOW())""")
            cur.execute("ALTER TABLE legislacoes ADD COLUMN IF NOT EXISTS municipio_id INTEGER")
            cur.execute("ALTER TABLE legislacoes ADD COLUMN IF NOT EXISTS origem VARCHAR(50) DEFAULT 'manual'")
            cur.execute("ALTER TABLE legislacoes ADD COLUMN IF NOT EXISTS texto_integral TEXT")
            cur.execute("ALTER TABLE legislacoes ADD COLUMN IF NOT EXISTS url_texto_fonte TEXT")
            # v6.x — Sites de referência por município (buscador)
            cur.execute("""CREATE TABLE IF NOT EXISTS municipio_sites_referencia (
                id SERIAL PRIMARY KEY,
                municipio_nome VARCHAR(200) NOT NULL,
                estado VARCHAR(2) NOT NULL DEFAULT '',
                url_diario TEXT DEFAULT '',
                urls_extras TEXT[] DEFAULT '{}',
                criado_em TIMESTAMP DEFAULT NOW(),
                atualizado_em TIMESTAMP DEFAULT NOW(),
                UNIQUE(municipio_nome, estado))""")
        except Exception as e:
            print(f"Aviso migration v6: {e}")
        conn.commit()
        conn.close()
        print("Tabela monitoramento_legislacao_log verificada/criada")
    except Exception as e:
        print(f"Aviso criacao tabela log legislacao: {e}")
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("""CREATE TABLE IF NOT EXISTS municipio_fallback (
            id SERIAL PRIMARY KEY,
            municipio VARCHAR(200) NOT NULL,
            estado VARCHAR(2) NOT NULL,
            url TEXT NOT NULL,
            criado_em TIMESTAMP DEFAULT NOW(),
            atualizado_em TIMESTAMP DEFAULT NOW(),
            UNIQUE(municipio, estado)
        )""")
        # Adicionar colunas novas se nao existirem
        try:
            cur.execute("ALTER TABLE municipio_fallback ADD COLUMN IF NOT EXISTS lm_nao_catalogado BOOLEAN DEFAULT FALSE")
            cur.execute("ALTER TABLE municipio_fallback ADD COLUMN IF NOT EXISTS fonte_funcionou VARCHAR(20)")
            conn.commit()
        except: pass
        conn.commit(); conn.close()
        print("Tabela municipio_fallback verificada/criada")
    except Exception as e:
        print(f"Aviso municipio_fallback: {e}")
    try:
        conn2 = get_db(); cur2 = conn2.cursor()
        cur2.execute("ALTER TABLE fila_buscas ADD COLUMN IF NOT EXISTS fallback_url_override TEXT")
        conn2.commit(); cur2.close(); conn2.close()
    except: pass
    if SCHEDULER_OK:
        try: iniciar_scheduler(); print("Scheduler iniciado")
        except Exception as e: print(f"Scheduler: {e}")

@app.route('/api/municipio/fallback', methods=['GET'])
@login_required
def api_municipio_fallback_get():
    mun = request.args.get('municipio','').strip()
    est = request.args.get('estado','').strip()
    if not mun or not est: return jsonify({'url': None})
    try:
        conn = get_db(); cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT url FROM municipio_fallback WHERE LOWER(municipio)=LOWER(%s) AND LOWER(estado)=LOWER(%s)", (mun, est))
        row = cur.fetchone(); conn.close()
        return jsonify({'url': row['url'] if row else None})
    except Exception as e:
        return jsonify({'url': None, 'error': str(e)})

@app.route('/api/municipio/fallback', methods=['POST'])
@login_required
def api_municipio_fallback_post():
    d = request.get_json() or {}
    mun = d.get('municipio','').strip(); est = d.get('estado','').strip(); url = d.get('url','').strip()
    if not mun or not est or not url: return jsonify({'success': False, 'error': 'Campos obrigatorios'})
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("INSERT INTO municipio_fallback (municipio, estado, url, atualizado_em) VALUES (%s,%s,%s,NOW()) ON CONFLICT (municipio, estado) DO UPDATE SET url=%s, atualizado_em=NOW()", (mun, est, url, url))
        conn.commit(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/municipio/fallback', methods=['DELETE'])
@login_required
def api_municipio_fallback_delete():
    d = request.get_json() or {}
    mun = d.get('municipio','').strip(); est = d.get('estado','').strip()
    if not mun or not est: return jsonify({'success': False})
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("DELETE FROM municipio_fallback WHERE LOWER(municipio)=LOWER(%s) AND LOWER(estado)=LOWER(%s)", (mun, est))
        conn.commit(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

inicializar()


# ══════════════════════════════════════════════════════════════════
# DIAGNÓSTICO TEMPORÁRIO — remover depois
# ══════════════════════════════════════════════════════════════════
@app.route('/api/diag/do')
@admin_required
def diag_do():
    """Diagnóstico: investiga por que o buscador não acha legislação no DO."""
    import io, contextlib

    output = io.StringIO()
    with contextlib.redirect_stdout(output):
        try:
            _run_diag_do()
        except Exception as e:
            import traceback
            print(f"\n❌ ERRO: {e}")
            print(traceback.format_exc())

    text = output.getvalue()
    return f'<pre style="font-family:monospace;font-size:13px;white-space:pre-wrap;background:#111;color:#0f0;padding:20px">{text}</pre>'


def _run_diag_do():
    import requests, json, re, time as _time

    print("🔍 DIAGNÓSTICO DO DIÁRIO OFICIAL DO RIO DE JANEIRO")
    print(f"   Data: {_time.strftime('%Y-%m-%d %H:%M:%S')}")

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/120.0.0.0',
        'Accept': 'application/json, text/html, */*',
        'Referer': 'https://doweb.rio.rj.gov.br/',
    }
    base = 'https://doweb.rio.rj.gov.br'

    # ── PARTE 1: API REST ──
    print("\n" + "="*60)
    print("PARTE 1: API REST do DO")
    print("="*60)

    endpoints = [
        '/apifront/portal/edicoes',
        '/apifront/portal/edicoes/busca',
        '/apifront/portal/edicoes?dataInicial=16/01/2024&dataFinal=19/01/2024',
        '/apifront/portal/edicoes?q=lei+complementar+270',
        '/api/edicoes',
    ]
    for ep in endpoints:
        try:
            r = requests.get(base + ep, headers=headers, timeout=10)
            ct = r.headers.get('Content-Type', '?')
            print(f"\n{'✅' if r.status_code == 200 else '❌'} {r.status_code} {ep}")
            print(f"   CT: {ct[:40]}")
            body = r.text[:300].replace('\n', ' ')
            print(f"   Body: {body[:250]}")
            if 'json' in ct and r.status_code == 200:
                try:
                    data = r.json()
                    if isinstance(data, list):
                        print(f"   📊 {len(data)} items. Keys: {list(data[0].keys()) if data and isinstance(data[0], dict) else '?'}")
                        for item in data[:3]:
                            print(f"   → {json.dumps(item, ensure_ascii=False)[:250]}")
                    elif isinstance(data, dict):
                        print(f"   📊 keys: {list(data.keys())}")
                except: pass
        except Exception as e:
            print(f"\n❌ {ep}: {str(e)[:100]}")

    # ── PARTE 2: Playwright ──
    print("\n" + "="*60)
    print("PARTE 2: Playwright — Formulário DO")
    print("="*60)

    try:
        from playwright.sync_api import sync_playwright
        import shutil, glob

        exec_path = os.environ.get('PLAYWRIGHT_CHROMIUM_PATH', '')
        if not exec_path:
            for cn in ['chromium', 'chromium-browser']:
                p = shutil.which(cn)
                if p: exec_path = p; break
        if not exec_path:
            nps = glob.glob('/nix/store/*/bin/chromium')
            if nps: exec_path = nps[0]

        print(f"Chromium: {exec_path or 'NÃO ENCONTRADO'}")

        with sync_playwright() as pw:
            browser = pw.chromium.launch(
                headless=True, executable_path=exec_path,
                args=['--no-sandbox', '--disable-dev-shm-usage', '--disable-gpu',
                      '--single-process', '--no-zygote']
            )
            ctx = browser.new_context(
                viewport={'width': 1280, 'height': 900},
                user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/120.0.0.0'
            )
            page = ctx.new_page()

            # Interceptar chamadas de API
            api_calls = []
            def on_req(req):
                u = req.url
                if '/api' in u.lower() or 'busca' in u.lower() or 'edic' in u.lower():
                    api_calls.append({'method': req.method, 'url': u, 'post': req.post_data[:300] if req.post_data else None})
            page.on('request', on_req)

            api_resps = []
            def on_resp(resp):
                u = resp.url
                if '/api' in u.lower() or 'busca' in u.lower() or 'edic' in u.lower():
                    api_resps.append({'status': resp.status, 'url': u, 'ct': resp.headers.get('content-type','?')})
            page.on('response', on_resp)

            print("\n📌 Abrindo DO...")
            page.goto(base + '/', wait_until='networkidle', timeout=30000)
            _time.sleep(2)
            print(f"   URL: {page.url}")
            print(f"   Title: {page.title()}")

            # Campos
            fields = page.evaluate('''() => {
                const info = [];
                document.querySelectorAll('input, select, textarea, button').forEach(el => {
                    info.push({
                        tag: el.tagName, type: el.type||'', id: el.id||'', name: el.name||'',
                        placeholder: el.placeholder||'', visible: el.offsetParent !== null,
                        label: el.labels?.[0]?.textContent?.trim()?.substring(0,50)||'',
                    });
                });
                return info;
            }''')
            print(f"\n📋 {len(fields)} campos:")
            for f in fields:
                v = '👁️' if f['visible'] else '🔒'
                print(f"   {v} <{f['tag'].lower()} type='{f['type']}' id='{f['id']}'> ph='{f['placeholder']}' label='{f['label']}'")

            # Preencher
            print("\n✏️ Preenchendo...")
            for fid, val in [('dataBuscaInicial', '16/01/2024'), ('dataBuscaFinal', '19/01/2024'), ('input2', 'Lei Complementar 270')]:
                try:
                    el = page.query_selector(f'#{fid}')
                    if el:
                        el.click(); _time.sleep(0.2)
                        el.press('Control+a'); el.press('Delete')
                        if '/' in val:
                            el.type(val, delay=50)
                        else:
                            el.fill(val)
                        el.evaluate('el => { el.dispatchEvent(new Event("input",{bubbles:true})); el.dispatchEvent(new Event("change",{bubbles:true})); el.dispatchEvent(new Event("blur",{bubbles:true})); }')
                        _time.sleep(0.3)
                        page.keyboard.press('Escape')
                        print(f"   ✅ {fid} = '{val}'")
                    else:
                        print(f"   ❌ #{fid} não encontrado")
                except Exception as e:
                    print(f"   ❌ {fid}: {str(e)[:60]}")

            _time.sleep(1)
            page.mouse.click(10, 10)

            # Valores após preenchimento
            vals = page.evaluate('''() => {
                const ids = ['dataBuscaInicial', 'dataBuscaFinal', 'input2'];
                const r = {};
                ids.forEach(id => { const el = document.getElementById(id); r[id] = el ? el.value : 'N/A'; });
                return r;
            }''')
            print(f"\n🔍 Valores APÓS preenchimento:")
            for k, v in vals.items():
                status = '✅' if v else '⚠️ VAZIO'
                print(f"   #{k} = '{v}' {status}")

            # Submit
            print("\n🖱️ Submetendo...")
            api_calls.clear(); api_resps.clear()
            submitted = False
            for sel in ['button:has-text("Buscar")', 'button:has-text("Pesquisar")', 'button[type="submit"]', 'button.btn-primary', 'input[type="submit"]']:
                try:
                    btn = page.query_selector(sel)
                    if btn and btn.is_visible():
                        print(f"   Botão: {sel}")
                        btn.click()
                        submitted = True
                        break
                except: continue
            if not submitted:
                page.keyboard.press('Enter')
                print("   Enter pressionado")

            try: page.wait_for_load_state('networkidle', timeout=15000)
            except: pass
            _time.sleep(3)

            # API calls
            print(f"\n📡 {len(api_calls)} API calls:")
            for c in api_calls[:10]:
                print(f"   {c['method']} {c['url'][:120]}")
                if c['post']: print(f"      POST: {c['post'][:200]}")

            print(f"\n📡 {len(api_resps)} API responses:")
            for r in api_resps[:10]:
                print(f"   {r['status']} {r['url'][:120]}")

            # Página resultado
            print(f"\n📄 URL após submit: {page.url}")
            body = (page.inner_text('body') or '')[:2000]
            lines = [l.strip() for l in body.split('\n') if l.strip() and len(l.strip()) > 2]
            print(f"   Body ({len(body)} chars), {len(lines)} linhas significativas:")
            for l in lines[:30]:
                print(f"      {l[:120]}")

            # Todos os links
            links = page.evaluate('''() => {
                const r = [];
                document.querySelectorAll('a[href]').forEach(a => {
                    r.push({t: a.textContent?.trim()?.substring(0,80)||'', h: a.href||'', v: a.offsetParent!==null});
                });
                return r;
            }''')
            print(f"\n🔗 {len(links)} links:")
            for i, l in enumerate(links):
                v = '👁️' if l['v'] else '🔒'
                print(f"   {v} [{i}] '{l['t'][:50]}' → {l['h'][:80]}")

            browser.close()

    except Exception as e:
        import traceback
        print(f"❌ Playwright erro: {e}")
        print(traceback.format_exc()[-500:])

    # ── PARTE 3: BuscaFácil ──
    print("\n" + "="*60)
    print("PARTE 3: BuscaFácil")
    print("="*60)
    try:
        from playwright.sync_api import sync_playwright
        import shutil, glob
        exec_path = os.environ.get('PLAYWRIGHT_CHROMIUM_PATH', '')
        if not exec_path:
            for cn in ['chromium','chromium-browser']:
                p = shutil.which(cn)
                if p: exec_path = p; break
        if not exec_path:
            nps = glob.glob('/nix/store/*/bin/chromium')
            if nps: exec_path = nps[0]

        with sync_playwright() as pw:
            br = pw.chromium.launch(headless=True, executable_path=exec_path,
                args=['--no-sandbox','--disable-dev-shm-usage','--disable-gpu','--single-process','--no-zygote'])
            pg = br.new_context(viewport={'width':1280,'height':900},
                user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/120.0.0.0').new_page()

            print("\n📌 Abrindo BuscaFácil...")
            try: pg.goto('https://www2.rio.rj.gov.br/smu/buscafacil/', wait_until='networkidle', timeout=30000)
            except: pg.goto('https://www2.rio.rj.gov.br/smu/buscafacil/', wait_until='domcontentloaded', timeout=15000)
            _time.sleep(3)

            print(f"   URL: {pg.url}")
            print(f"   Title: {pg.title()}")
            html = pg.content()
            print(f"   HTML: {len(html)} chars")
            print(f"   Preview:\n{html[:800]}")

            frames = pg.frames
            print(f"\n📋 {len(frames)} frames:")
            for i, f in enumerate(frames):
                print(f"   [{i}] name='{f.name}' url='{f.url[:80]}'")
                try:
                    n = f.evaluate('() => document.querySelectorAll("input,select").length')
                    h = f.evaluate('() => document.documentElement.outerHTML.substring(0,300)')
                    print(f"       inputs: {n}")
                    print(f"       HTML: {h[:200]}")
                except Exception as e:
                    print(f"       Erro: {str(e)[:80]}")

            br.close()
    except Exception as e:
        print(f"❌ BuscaFácil erro: {str(e)[:200]}")

    print("\n" + "="*60)
    print("FIM DO DIAGNÓSTICO")
    print("="*60)


# Google Maps key
import os as _os
if not _os.environ.get('GOOGLE_MAPS_KEY'):
    _os.environ['GOOGLE_MAPS_KEY'] = 'AIzaSyCuiZTfrnvUC-1X_suD3w6iGVyT_bhdVpQ'



@app.route('/api/buscador/job-atual')
@login_required
def api_buscador_job_atual():
    """Retorna job rodando agora com logs do banco — para reconexao apos refresh."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT job_id, municipio, estado FROM fila_buscas WHERE status='rodando' ORDER BY iniciado_em DESC LIMIT 1")
        fila = cur.fetchone()
        cur.close(); conn.close()
        if not fila or not fila['job_id']:
            return jsonify({'ativo': False})
        job_id = fila['job_id']
        cursor_req = int(request.args.get('cursor', 0))
        from modulos.log_persistente import carregar_logs, contar_logs
        logs = carregar_logs(job_id, get_db, cursor_req)
        total = contar_logs(job_id, get_db)
        job = _buscador_jobs.get(job_id)
        done = job['done'] if job else False
        hist_id = job.get('hist_id') if job else None
        if not hist_id:
            try:
                c2=get_db(); cu2=c2.cursor()
                cu2.execute("SELECT id FROM buscas_historico WHERE job_id=%s LIMIT 1",(job_id,))
                h=cu2.fetchone(); cu2.close(); c2.close()
                if h: hist_id=h[0]
            except: pass
        return jsonify({'ativo': True, 'job_id': job_id, 'municipio': fila['municipio'], 'estado': fila['estado'], 'logs': [{'nivel':l['nivel'],'msg':l['msg']} for l in logs], 'cursor': cursor_req + len(logs), 'total': total, 'done': done, 'hist_id': hist_id})
    except Exception as e:
        return jsonify({'ativo': False, 'error': str(e)})


@app.route('/api/buscador/ultimo-concluido')
@login_required
def api_buscador_ultimo_concluido():
    """Retorna o ultimo job concluido com URLs dos arquivos."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""SELECT job_id, municipio, estado, zip_path, relatorio_path, tabela_path
                       FROM buscas_historico
                       WHERE concluido_em IS NOT NULL AND sucesso=true
                       ORDER BY concluido_em DESC LIMIT 1""")
        r = cur.fetchone()
        cur.close(); conn.close()
        if not r: return jsonify({'success': False})
        return jsonify({
            'success': True,
            'job_id': r['job_id'],
            'municipio': r['municipio'],
            'estado': r['estado'],
            'zip_url': r['zip_path'],
            'relatorio_url': r['relatorio_path'],
            'tabela_url': r['tabela_path']
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


# ── Dossiê Municipais ─────────────────────────────────────────────────────────

@app.route('/api/dossie/municipios')
@login_required
def api_dossie_municipios():
    """Lista todos os municípios do dossiê."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT dm.id, dm.municipio, dm.estado, dm.origem, dm.criado_em, dm.max_legislacoes,
                   bh.concluido_em as ultima_busca, bh.sucesso,
                   bh.zip_path, bh.relatorio_path, bh.tabela_path, bh.job_id,
                   COALESCE(
                       (SELECT COUNT(*) FROM buscas_logs bl
                        WHERE bl.job_id=bh.job_id AND bl.nivel='tabela'
                        AND bl.msg LIKE '%encontrada%'),
                       (SELECT CAST(NULLIF(regexp_replace(
                           substring(bh2.log_texto FROM '\\((\\d+) legisla'), '[^0-9]', '', 'g'
                       ), '') AS INTEGER)
                        FROM buscas_historico bh2
                        WHERE bh2.municipio=dm.municipio AND bh2.estado=dm.estado
                        AND bh2.sucesso=true ORDER BY bh2.concluido_em DESC LIMIT 1),
                       0
                   ) as total_legislacoes,
                   (SELECT status FROM fila_buscas fb
                    WHERE LOWER(fb.municipio)=LOWER(dm.municipio)
                    AND LOWER(fb.estado)=LOWER(dm.estado)
                    AND fb.status IN ('rodando','aguardando')
                    ORDER BY fb.id DESC LIMIT 1) as fila_status
            FROM dossie_municipios dm
            LEFT JOIN buscas_historico bh ON LOWER(bh.municipio)=LOWER(dm.municipio)
                AND LOWER(bh.estado)=LOWER(dm.estado)
                AND bh.concluido_em IS NOT NULL AND bh.sucesso=true
                AND bh.concluido_em = (
                    SELECT MAX(concluido_em) FROM buscas_historico
                    WHERE municipio=dm.municipio AND estado=dm.estado AND sucesso=true
                )
            ORDER BY dm.municipio ASC
        """)
        rows = cur.fetchall()
        cur.close(); conn.close()
        result = []
        for r in rows:
            leg_status = 'none'
            if r['fila_status'] in ('rodando','aguardando'):
                leg_status = 'pending'
            elif r['sucesso']:
                leg_status = 'ok'
            result.append({
                'id': r['id'],
                'municipio': r['municipio'],
                'estado': r['estado'],
                'origem': r['origem'],
                'ultima_busca': r['ultima_busca'].strftime('%d/%m/%Y') if r['ultima_busca'] else None,
                'total_legislacoes': r['total_legislacoes'] or 0,
                'max_legislacoes': r['max_legislacoes'],  # None = sem limite
                'zip_url': r['zip_path'],
                'relatorio_url': r['relatorio_path'],
                'tabela_url': r['tabela_path'],
                'job_id': r['job_id'],
                'leg_status': leg_status,
            })
        return jsonify({'success': True, 'municipios': result})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/dossie/municipio-dossies/<int:mun_id>')
@login_required
def api_dossie_municipio_dossies(mun_id):
    """Lista todos os dossies (buscas_historico) de um municipio."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT municipio, estado FROM dossie_municipios WHERE id=%s", (mun_id,))
        mun = cur.fetchone()
        if not mun:
            return jsonify({'success': False, 'error': 'municipio nao encontrado'}), 404
        cur.execute("""
            SELECT bh.id, bh.job_id, bh.iniciado_em, bh.concluido_em, bh.sucesso, bh.tipo,
                   bh.zip_path, bh.relatorio_path, bh.tabela_path,
                   COALESCE(
                       (SELECT COUNT(*) FROM buscas_logs bl WHERE bl.job_id=bh.job_id
                        AND bl.nivel='tabela' AND bl.msg LIKE '%%encontrada%%'),
                       0
                   ) as total_legs
            FROM buscas_historico bh
            WHERE LOWER(bh.municipio)=LOWER(%s) AND LOWER(bh.estado)=LOWER(%s)
            ORDER BY bh.iniciado_em DESC
        """, (mun['municipio'], mun['estado']))
        rows = cur.fetchall()
        cur.close(); conn.close()
        dossies = []
        for r in rows:
            dossies.append({
                'id': r['id'],
                'job_id': r['job_id'],
                'iniciado_em': r['iniciado_em'].strftime('%d/%m/%Y %H:%M') if r['iniciado_em'] else None,
                'concluido_em': r['concluido_em'].strftime('%d/%m/%Y %H:%M') if r['concluido_em'] else None,
                'sucesso': r['sucesso'],
                'tipo': r['tipo'] or 'manual',
                'zip_url': r['zip_path'],
                'relatorio_url': r['relatorio_path'],
                'tabela_url': r['tabela_path'],
                'total_legislacoes': r['total_legs'] or 0,
            })
        return jsonify({'success': True, 'municipio': mun['municipio'], 'estado': mun['estado'], 'dossies': dossies})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/dossie/legislacoes-do-municipio/<int:mun_id>')
@login_required
def api_dossie_legislacoes_do_municipio(mun_id):
    """Lista legislacoes organizadas (pastas em /static/dossies/) de um municipio."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT id, dossie_id, legislacao_label, legislacao_meta, categoria,
                   pasta_path, pdf_concatenado_path, n_paginas, total_arquivos,
                   arquivos_originais, arquivos_falhas, duplicados_removidos,
                   anexos_citados, anexos_faltantes,
                   criado_em, atualizado_em
            FROM dossie_legislacoes_pasta
            WHERE dossie_id=%s
            ORDER BY criado_em DESC, legislacao_label
        """, (mun_id,))
        rows = cur.fetchall()
        cur.close(); conn.close()
        
        legislacoes = []
        for r in rows:
            meta = r['legislacao_meta'] or {}
            arquivos = r['arquivos_originais'] or []
            falhas = r['arquivos_falhas'] or []
            
            # Detecta corpo da lei pelos arquivos (primeiro PDF que casa com o label)
            label_lower = (r['legislacao_label'] or '').lower()
            arquivos_marcados = []
            for arq in arquivos:
                nome = arq.get('nome', '')
                # Classifica: 'corpo' se nome bate com o label, 'anexo' caso contrario
                eh_corpo = nome.lower().startswith(label_lower) and arq.get('tipo_detectado') == 'pdf'
                arquivos_marcados.append({
                    'nome': nome,
                    'tipo_detectado': arq.get('tipo_detectado'),
                    'tamanho': arq.get('tamanho', 0),
                    'conversao_ok': arq.get('conversao_ok', False),
                    'foi_convertido': arq.get('foi_convertido', False),
                    'motivo': arq.get('motivo'),
                    'classificacao': 'corpo' if eh_corpo else 'anexo',
                })
            
            # PDF disponivel para download (URL relativa)
            pdf_url = None
            if r['pdf_concatenado_path']:
                pdf_url = r['pdf_concatenado_path'].replace('/var/www/urbanlex', '')
            # ZIP concat_catalogo disponivel para download
            zip_url = None
            _pasta = r.get('pasta_path') or ''
            if _pasta:
                import glob as _glob
                _zips = _glob.glob(_os.path.join(_pasta, '*_concat_catalogo.zip'))
                if _zips:
                    zip_url = _zips[0].replace('/var/www/urbanlex', '')
                elif r['pdf_concatenado_path'] and _os.path.exists(_os.path.join(_pasta, 'etapa4_catalogacao.json')):
                    try:
                        from modulos.preparar_legislacao import gerar_zip as _gz
                        _label = r.get('legislacao_label') or 'lei'
                        _zip_path = _gz(
                            r['pdf_concatenado_path'],
                            _os.path.join(_pasta, 'etapa4_catalogacao.json'),
                            _label, _pasta
                        )
                        zip_url = _zip_path.replace('/var/www/urbanlex', '')
                    except Exception as _ez:
                        logger.warning(f"gerar_zip dossie falhou: {_ez}")
            
            legislacoes.append({
                'id': r['id'],
                'label': r['legislacao_label'],
                'busca_id': busca_id,
                'categoria': r['categoria'] or '',
                'tipo': meta.get('tipo', ''),
                'numero': meta.get('numero', ''),
                'ano': meta.get('ano', ''),
                'descricao': meta.get('descricao', ''),
                'link': meta.get('link', ''),
                'n_paginas': r['n_paginas'] or 0,
                'total_arquivos': r['total_arquivos'] or 0,
                'duplicados_removidos': r['duplicados_removidos'] or 0,
                'pdf_url': pdf_url,
                'zip_url': zip_url,
                'arquivos': arquivos_marcados,
                'falhas': falhas,
                'anexos_citados': r['anexos_citados'] or [],
                'anexos_faltantes': r['anexos_faltantes'] or [],
                'busca_id': r['busca_historico_id'],
                'criado_em': r['criado_em'].strftime('%d/%m/%Y %H:%M') if r['criado_em'] else None,
                'atualizado_em': r['atualizado_em'].strftime('%d/%m/%Y %H:%M') if r['atualizado_em'] else None,
            })
        
        return jsonify({'success': True, 'legislacoes': legislacoes})
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'traceback': traceback.format_exc()[-500:]})


@app.route('/api/dossie/legislacoes-do-dossie/<int:busca_id>')
@login_required
def api_dossie_legislacoes_do_dossie(busca_id):
    """Lista legislacoes organizadas de UMA BUSCA específica (dossie)."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT id, dossie_id, busca_historico_id, legislacao_label, legislacao_meta, categoria,
                   pasta_path, pdf_concatenado_path, n_paginas, total_arquivos,
                   arquivos_originais, arquivos_falhas, duplicados_removidos,
                   anexos_citados, anexos_faltantes,
                   criado_em, atualizado_em
            FROM dossie_legislacoes_pasta
            WHERE busca_historico_id=%s
            ORDER BY criado_em DESC, legislacao_label
        """, (busca_id,))
        rows = cur.fetchall()
        cur.close(); conn.close()
        
        legislacoes = []
        for r in rows:
            meta = r['legislacao_meta'] or {}
            arquivos = r['arquivos_originais'] or []
            falhas = r['arquivos_falhas'] or []
            
            label_lower = (r['legislacao_label'] or '').lower()
            arquivos_marcados = []
            for arq in arquivos:
                nome = arq.get('nome', '')
                eh_corpo = nome.lower().startswith(label_lower) and arq.get('tipo_detectado') == 'pdf'
                arquivos_marcados.append({
                    'nome': nome,
                    'tipo_detectado': arq.get('tipo_detectado'),
                    'tamanho': arq.get('tamanho', 0),
                    'conversao_ok': arq.get('conversao_ok', False),
                    'foi_convertido': arq.get('foi_convertido', False),
                    'motivo': arq.get('motivo'),
                    'classificacao': 'corpo' if eh_corpo else 'anexo',
                    'origem': arq.get('origem', 'busca'),
                })
            
            pdf_url = None
            if r['pdf_concatenado_path']:
                pdf_url = r['pdf_concatenado_path'].replace('/var/www/urbanlex', '')
            
            legislacoes.append({
                'id': r['id'],
                'dossie_id_municipio': r['dossie_id'],
                'busca_historico_id': r['busca_historico_id'],
                'label': r['legislacao_label'],
                'categoria': r['categoria'] or '',
                'tipo': meta.get('tipo', ''),
                'numero': meta.get('numero', ''),
                'ano': meta.get('ano', ''),
                'descricao': meta.get('descricao', ''),
                'link': meta.get('link', ''),
                'n_paginas': r['n_paginas'] or 0,
                'total_arquivos': r['total_arquivos'] or 0,
                'duplicados_removidos': r['duplicados_removidos'] or 0,
                'pdf_url': pdf_url,
                'arquivos': arquivos_marcados,
                'falhas': falhas,
                'anexos_citados': r['anexos_citados'] or [],
                'anexos_faltantes': r['anexos_faltantes'] or [],
                'busca_id': r['busca_historico_id'],
                'criado_em': r['criado_em'].strftime('%d/%m/%Y %H:%M') if r['criado_em'] else None,
                'atualizado_em': r['atualizado_em'].strftime('%d/%m/%Y %H:%M') if r['atualizado_em'] else None,
            })
        
        return jsonify({'success': True, 'busca_id': busca_id, 'legislacoes': legislacoes})
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'traceback': traceback.format_exc()[-500:]})



@app.route('/api/dossie/enviar-gerador', methods=['POST'])
@login_required
def api_dossie_enviar_gerador():
    """
    Envia legislacao(oes) do dossie pro Gerador de Planilha (pipeline 8 etapas).
    
    Reutiliza artifacts do dossie (pdf_concatenado, corpo.pdf, anexos.pdf)
    pulando as Etapas 1-3 do pipeline. Economia: ~30s por legislacao.
    
    Body:
      busca_id (int):              ID da busca historica
      legislacao_label (str, opcional): se fornecido, processa SO essa legislacao
                                          se omitido, processa TODAS as legislacoes
                                          do dossie
    
    Retorna:
      {
        success: bool,
        enfileirados: [{label, fila_id}],
        erros: [{label, erro}],
        total_enfileirados: int
      }
    """
    try:
        data = request.get_json(silent=True) or {}
        busca_id = data.get('busca_id')
        legislacao_label = (data.get('legislacao_label') or '').strip() or None
        
        if not busca_id:
            return jsonify({'success': False, 'error': 'busca_id obrigatorio'}), 400
        
        try:
            busca_id = int(busca_id)
        except:
            return jsonify({'success': False, 'error': 'busca_id invalido'}), 400
        
        # Busca info do dossie + legislacoes
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Pega municipio/estado da busca
        cur.execute("""SELECT id, municipio, estado FROM buscas_historico WHERE id=%s""",
                    (busca_id,))
        bh = cur.fetchone()
        if not bh:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': f'busca_id {busca_id} nao encontrada'}), 404
        
        municipio = bh['municipio']
        estado = bh['estado']
        
        # Pega legislacoes do dossie (filtradas se label especifico)
        if legislacao_label:
            cur.execute("""SELECT id, dossie_id, legislacao_label, pasta_path
                           FROM dossie_legislacoes_pasta
                           WHERE busca_historico_id=%s AND legislacao_label=%s""",
                        (busca_id, legislacao_label))
        else:
            cur.execute("""SELECT id, dossie_id, legislacao_label, pasta_path
                           FROM dossie_legislacoes_pasta
                           WHERE busca_historico_id=%s
                           ORDER BY legislacao_label""",
                        (busca_id,))
        legs = cur.fetchall()
        cur.close(); conn.close()
        
        if not legs:
            return jsonify({'success': False, 
                            'error': f'nenhuma legislacao encontrada para busca {busca_id}' +
                                    (f' com label {legislacao_label}' if legislacao_label else '')}), 404
        
        # Pra cada legislacao: prepara work_dir + enfileira
        from modulos.dossie_para_gerador import preparar_work_dir_pipeline
        from modulos.pipeline_extracao_lei import enfileirar_extracao
        
        enfileirados = []
        erros = []
        
        # Pega nome do usuario pra logs
        user_name = (session.get('nome') or session.get('email') or 'sistema')
        
        for leg in legs:
            label = leg['legislacao_label']
            dossie_id = leg['dossie_id']
            pasta_path = leg['pasta_path']
            
            # 1. Prepara work_dir copiando arquivos do dossie pro pipeline
            try:
                result_prep = preparar_work_dir_pipeline(
                    dossie_id=dossie_id,
                    busca_historico_id=busca_id,
                    legislacao_label=label,
                    get_db=get_db,
                )
                if not result_prep:
                    erros.append({'label': label, 'erro': 'falha preparando work_dir'})
                    continue
                work_dir, zip_path_efetivo = result_prep if isinstance(result_prep, tuple) else (result_prep, os.path.join(result_prep, 'tudo.pdf'))
            except Exception as e:
                erros.append({'label': label, 'erro': f'preparacao: {str(e)[:200]}'})
                continue
            
            # 2. Adiciona como 'compilado' no Gerador de Planilha
            #    (NAO enfileira pipeline ainda - operador escolhe na tela do Gerador)
            try:
                import uuid as _u_c
                novo_job_id = str(_u_c.uuid4())[:8]
                
                conn = get_db(); cur = conn.cursor()
                cur.execute('''
                    INSERT INTO gerador_compilados
                      (job_id, municipio, estado, zip_path, tabela_path, criado_por, criado_por_nome, legislacao_label, busca_historico_id)
                    VALUES (%s, %s, %s, %s, NULL, %s, %s, %s, %s)
                    RETURNING id
                ''', (novo_job_id, municipio, estado, zip_path_efetivo,
                      session.get('user_id'), user_name, label, busca_id))
                comp_row = cur.fetchone()
                conn.commit(); cur.close(); conn.close()
                
                if comp_row:
                    enfileirados.append({
                        'label': label, 
                        'compilado_id': comp_row[0],
                        'job_id': novo_job_id,
                    })
                else:
                    erros.append({'label': label, 'erro': 'falha INSERT em gerador_compilados'})
            except Exception as e:
                erros.append({'label': label, 'erro': f'adicionando compilado: {str(e)[:200]}'})
        
        return jsonify({
            'success': True,
            'enfileirados': enfileirados,
            'erros': erros,
            'total_enfileirados': len(enfileirados),
            'municipio': municipio,
            'estado': estado,
        })
        
    except Exception as e:
        import traceback
        logger.error(f"erro enviar-gerador: {traceback.format_exc()[:1000]}")
        return jsonify({'success': False, 'error': str(e)[:200]}), 500


@app.route('/api/dossie/anexo/upload', methods=['POST'])
@editor_required
def api_dossie_anexo_upload():
    """Recebe upload manual de anexo faltante. Salva em pasta upload_pendente/."""
    try:
        dossie_id = request.form.get('dossie_id', '').strip()
        legislacao_label = request.form.get('legislacao_label', '').strip()
        refere_a = (request.form.get('refere_a') or '').strip()  # opcional: "Anexo 1.4"
        busca_id_form = request.form.get('busca_historico_id', '').strip()
        arquivo = request.files.get('arquivo')
        
        if not dossie_id or not legislacao_label or not arquivo:
            return jsonify({'success': False, 'error': 'dossie_id, legislacao_label e arquivo obrigatorios'}), 400
        
        try:
            dossie_id = int(dossie_id)
        except:
            return jsonify({'success': False, 'error': 'dossie_id invalido'}), 400
        
        # Sanitiza
        from werkzeug.utils import secure_filename
        nome_seguro = secure_filename(arquivo.filename) or 'upload.bin'
        if not nome_seguro:
            return jsonify({'success': False, 'error': 'nome de arquivo invalido'}), 400
        
        # Resolve busca_historico_id: do form ou query banco pela legislacao
        busca_id = None
        try:
            if busca_id_form:
                busca_id = int(busca_id_form)
            else:
                # Procura no banco pela legislacao mais recente desse dossie+label
                _c_u = get_db(); _cu_u = _c_u.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                _cu_u.execute("""SELECT busca_historico_id FROM dossie_legislacoes_pasta 
                                 WHERE dossie_id=%s AND legislacao_label=%s 
                                 ORDER BY criado_em DESC LIMIT 1""",
                              (dossie_id, legislacao_label))
                _row_u = _cu_u.fetchone()
                _cu_u.close(); _c_u.close()
                if _row_u and _row_u['busca_historico_id']:
                    busca_id = _row_u['busca_historico_id']
        except Exception:
            pass
        
        # Cria pasta upload_pendente (com busca_id se disponivel)
        import os as _os_u, hashlib as _h_u
        if busca_id:
            pasta = f'/var/www/urbanlex/static/dossies/{dossie_id}/busca_{busca_id}/{legislacao_label}/upload_pendente'
        else:
            # Fallback legacy
            pasta = f'/var/www/urbanlex/static/dossies/{dossie_id}/{legislacao_label}/upload_pendente'
        _os_u.makedirs(pasta, exist_ok=True)
        
        # Salva arquivo (evita colisao)
        destino = _os_u.path.join(pasta, nome_seguro)
        if _os_u.path.exists(destino):
            base, ext = _os_u.path.splitext(nome_seguro)
            i = 1
            while _os_u.path.exists(_os_u.path.join(pasta, f'{base}_{i}{ext}')):
                i += 1
            nome_seguro = f'{base}_{i}{ext}'
            destino = _os_u.path.join(pasta, nome_seguro)
        
        arquivo.save(destino)
        
        # Calcula MD5
        h = _h_u.md5()
        with open(destino, 'rb') as f:
            for chunk in iter(lambda: f.read(65536), b''):
                h.update(chunk)
        md5_hash = h.hexdigest()
        tamanho = _os_u.path.getsize(destino)
        
        # Insere no banco
        conn = get_db()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO dossie_anexos_uploads 
              (dossie_id, legislacao_label, arquivo_path, nome_original, refere_a, 
               tamanho_bytes, md5_hash, aplicado, criado_por)
            VALUES (%s, %s, %s, %s, %s, %s, %s, FALSE, %s)
            RETURNING id
        """, (dossie_id, legislacao_label, destino, arquivo.filename, refere_a or None,
              tamanho, md5_hash, session.get('user_id')))
        upload_id = cur.fetchone()[0]
        conn.commit(); cur.close(); conn.close()
        
        return jsonify({
            'success': True,
            'upload_id': upload_id,
            'nome_arquivo': nome_seguro,
            'tamanho': tamanho,
            'refere_a': refere_a,
        })
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()[-300:]}), 500


@app.route('/api/dossie/anexo/uploads-pendentes/<int:dossie_id>/<legislacao_label>')
@login_required
def api_dossie_anexo_uploads_pendentes(dossie_id, legislacao_label):
    """Lista uploads pendentes (aplicado=FALSE) de uma legislacao."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT id, nome_original, refere_a, tamanho_bytes, criado_em
            FROM dossie_anexos_uploads
            WHERE dossie_id=%s AND legislacao_label=%s AND aplicado=FALSE
            ORDER BY criado_em DESC
        """, (dossie_id, legislacao_label))
        rows = cur.fetchall()
        cur.close(); conn.close()
        
        pendentes = [{
            'id': r['id'],
            'nome_original': r['nome_original'],
            'refere_a': r['refere_a'] or '',
            'tamanho': r['tamanho_bytes'] or 0,
            'criado_em': r['criado_em'].strftime('%d/%m/%Y %H:%M') if r['criado_em'] else None,
        } for r in rows]
        
        return jsonify({'success': True, 'pendentes': pendentes})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/dossie/anexo/upload-remover/<int:upload_id>', methods=['DELETE'])
@editor_required
def api_dossie_anexo_upload_remover(upload_id):
    """Remove um upload pendente (antes de aplicar)."""
    try:
        import os as _os_d
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT arquivo_path, aplicado FROM dossie_anexos_uploads WHERE id=%s", (upload_id,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': 'nao encontrado'}), 404
        if row['aplicado']:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': 'upload ja aplicado, nao pode ser removido'}), 400
        
        # Remove arquivo
        try:
            if _os_d.path.exists(row['arquivo_path']):
                _os_d.remove(row['arquivo_path'])
        except Exception:
            pass
        
        # Remove do banco
        cur.execute("DELETE FROM dossie_anexos_uploads WHERE id=%s", (upload_id,))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/dossie/anexo/aplicar-uploads', methods=['POST'])
@editor_required
def api_dossie_anexo_aplicar_uploads():
    """
    Aplica os uploads pendentes de uma legislacao:
      1. Pega o ZIP original do dossie (buscas_historico)
      2. Cria ZIP TEMPORARIO com anexos extras (uploads pendentes)
      3. Re-roda organizador + Etapa 4.5 com novo ZIP
      4. Marca uploads como aplicado=TRUE
    """
    try:
        d = request.json or {}
        dossie_id = d.get('dossie_id')
        legislacao_label = d.get('legislacao_label', '').strip()
        
        if not dossie_id or not legislacao_label:
            return jsonify({'success': False, 'error': 'dossie_id e legislacao_label obrigatorios'}), 400
        
        # Pega uploads pendentes
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT id, arquivo_path, nome_original
            FROM dossie_anexos_uploads
            WHERE dossie_id=%s AND legislacao_label=%s AND aplicado=FALSE
        """, (dossie_id, legislacao_label))
        pendentes = cur.fetchall()
        
        if not pendentes:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': 'nenhum upload pendente'}), 400
        
        # Pega o ZIP original do dossie (mais recente)
        cur.execute("""
            SELECT bh.zip_path
            FROM buscas_historico bh
            JOIN dossie_municipios dm ON LOWER(bh.municipio)=LOWER(dm.municipio) AND LOWER(bh.estado)=LOWER(dm.estado)
            WHERE dm.id=%s AND bh.sucesso=TRUE AND bh.zip_path IS NOT NULL AND bh.zip_path != ''
            ORDER BY bh.iniciado_em DESC LIMIT 1
        """, (dossie_id,))
        zip_row = cur.fetchone()
        cur.close()
        
        if not zip_row or not zip_row['zip_path']:
            conn.close()
            return jsonify({'success': False, 'error': 'ZIP original do dossie nao encontrado'}), 404
        
        zip_orig_url = zip_row['zip_path']
        zip_orig_path = '/var/www/urbanlex' + zip_orig_url if zip_orig_url.startswith('/static') else zip_orig_url
        
        if not os.path.exists(zip_orig_path):
            conn.close()
            return jsonify({'success': False, 'error': f'ZIP nao existe no disco: {zip_orig_path}'}), 404
        
        # Dispara em thread separada (nao bloqueia a UI)
        import threading as _th_u
        def _run_aplicar():
            try:
                _aplicar_uploads_worker(dossie_id, legislacao_label, zip_orig_path, pendentes)
            except Exception as e:
                import traceback
                logger.error(f"[aplicar uploads dossie {dossie_id}] EXCECAO: {e}\n{traceback.format_exc()[-500:]}")
        
        _th_u.Thread(target=_run_aplicar, daemon=True, name=f'aplicar-uploads-{dossie_id}-{legislacao_label}').start()
        
        conn.close()
        return jsonify({
            'success': True,
            'message': f'{len(pendentes)} upload(s) sendo aplicado(s) em background. Aguarde alguns minutos e recarregue a pagina.',
            'pendentes': len(pendentes),
        })
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()[-300:]}), 500


def _aplicar_uploads_worker(dossie_id, legislacao_label, zip_orig_path, pendentes):
    """
    Worker que aplica uploads pendentes:
      1. Adiciona arquivos extras ao /static/dossies/<id>/<label>/ (concatena junto)
      2. Re-roda organizador + Etapa 4.5
      3. Marca uploads como aplicado=TRUE
    
    Estratégia simples: copia os arquivos pra pasta da legislação (extras/),
    depois re-roda o organizador que ja lida com ZIPs recursivos. O organizador
    nao le essa pasta extras/ ainda — pra simplificar agora, vou faz inline:
      a) Copia uploads pra uma pasta temporaria
      b) Combina com pdf_concatenado.pdf atual
      c) Atualiza o pdf_concatenado.pdf
      d) Re-roda Etapa 4.5 sobre o PDF novo
    """
    import os as _os_w, shutil as _sh_w, logging as _log_w
    logger_w = _log_w.getLogger(__name__)
    logger_w.info(f"[aplicar dossie {dossie_id} {legislacao_label}] iniciando ({len(pendentes)} pendente(s))")
    
    # Resolve a pasta: primeiro tenta achar via banco (com busca_id), depois fallback legacy
    leg_dir = None
    try:
        _c_w_path = get_db(); _cu_w_path = _c_w_path.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        _cu_w_path.execute("""SELECT pasta_path FROM dossie_legislacoes_pasta 
                              WHERE dossie_id=%s AND legislacao_label=%s 
                              ORDER BY criado_em DESC LIMIT 1""",
                           (dossie_id, legislacao_label))
        _row_path = _cu_w_path.fetchone()
        _cu_w_path.close(); _c_w_path.close()
        if _row_path and _row_path['pasta_path'] and _os_w.path.exists(_row_path['pasta_path']):
            leg_dir = _row_path['pasta_path']
    except Exception:
        pass
    
    # Fallback legacy se nao achou no banco
    if not leg_dir:
        candidatos = [
            f'/var/www/urbanlex/static/dossies/{dossie_id}/{legislacao_label}',  # legacy
        ]
        # Tambem tenta encontrar a busca mais recente em /static/dossies/<id>/busca_*/
        import glob as _glob_w
        for d in _glob_w.glob(f'/var/www/urbanlex/static/dossies/{dossie_id}/busca_*/{legislacao_label}'):
            candidatos.insert(0, d)
        for c in candidatos:
            if _os_w.path.exists(c):
                leg_dir = c
                break
    
    if not leg_dir:
        logger_w.error(f"[aplicar dossie {dossie_id}] pasta nao encontrada para {legislacao_label}")
        return
    
    pdf_atual = _os_w.path.join(leg_dir, 'pdf_concatenado.pdf')
    if not _os_w.path.exists(pdf_atual):
        logger_w.error(f"[aplicar dossie {dossie_id}] PDF concatenado nao existe: {pdf_atual}")
        return
    
    logger_w.info(f"[aplicar dossie {dossie_id}] usando pasta: {leg_dir}")
    
    # Converte cada upload pra PDF
    from modulos.conversor_pdf import converter_para_pdf, concatenar_pdfs, identificar_tipo
    
    tmp_dir = _os_w.path.join(leg_dir, '_tmp_aplicar')
    _os_w.makedirs(tmp_dir, exist_ok=True)
    
    novos_pdfs = []
    arquivos_aplicados = []
    falhas_aplicar = []
    
    for p in pendentes:
        arquivo_path = p['arquivo_path']
        nome_orig = p['nome_original']
        if not _os_w.path.exists(arquivo_path):
            falhas_aplicar.append({'id': p['id'], 'nome': nome_orig, 'motivo': 'arquivo nao existe no disco'})
            continue
        tipo = identificar_tipo(arquivo_path)
        pdf_gerado = converter_para_pdf(arquivo_path, tmp_dir)
        if pdf_gerado:
            novos_pdfs.append(pdf_gerado)
            arquivos_aplicados.append({'id': p['id'], 'nome': nome_orig, 'tipo': tipo})
        else:
            falhas_aplicar.append({'id': p['id'], 'nome': nome_orig, 'motivo': f'falha conversao ({tipo})'})
    
    if not novos_pdfs:
        logger_w.error(f"[aplicar dossie {dossie_id}] nenhum PDF gerado dos uploads")
        # Limpa tmp
        try:
            _sh_w.rmtree(tmp_dir, ignore_errors=True)
        except: pass
        return
    
    # Concatena PDF atual + novos PDFs
    pdf_novo = _os_w.path.join(tmp_dir, 'pdf_concatenado_novo.pdf')
    lista_pra_concat = [pdf_atual] + novos_pdfs
    ok = concatenar_pdfs(lista_pra_concat, pdf_novo)
    
    if not ok or not _os_w.path.exists(pdf_novo):
        logger_w.error(f"[aplicar dossie {dossie_id}] falha concatenando")
        try:
            _sh_w.rmtree(tmp_dir, ignore_errors=True)
        except: pass
        return
    
    # Substitui o pdf_concatenado.pdf
    try:
        _sh_w.copy2(pdf_novo, pdf_atual)
        logger_w.info(f"[aplicar dossie {dossie_id}] PDF atualizado: {pdf_atual}")
    except Exception as e:
        logger_w.error(f"[aplicar dossie {dossie_id}] erro substituindo PDF: {e}")
        try:
            _sh_w.rmtree(tmp_dir, ignore_errors=True)
        except: pass
        return
    
    # Atualiza metadados no banco
    try:
        import pypdf, json as _json_w
        n_paginas = len(pypdf.PdfReader(pdf_atual).pages)
        
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Pega arquivos_originais atual e adiciona os novos
        cur.execute("""
            SELECT arquivos_originais FROM dossie_legislacoes_pasta 
            WHERE dossie_id=%s AND legislacao_label=%s
        """, (dossie_id, legislacao_label))
        row = cur.fetchone()
        arq_atual = (row['arquivos_originais'] or []) if row else []
        
        for a in arquivos_aplicados:
            arq_atual.append({
                'nome': a['nome'],
                'tipo_detectado': a['tipo'],
                'tamanho': 0,
                'conversao_ok': True,
                'foi_convertido': (a['tipo'] != 'pdf'),
                'origem': 'upload_manual',
            })
        
        # Pega o busca_historico_id pra UPDATE
        cur.execute("""SELECT busca_historico_id FROM dossie_legislacoes_pasta 
                       WHERE dossie_id=%s AND legislacao_label=%s 
                       ORDER BY criado_em DESC LIMIT 1""",
                    (dossie_id, legislacao_label))
        _row_bh = cur.fetchone()
        bh_id = _row_bh['busca_historico_id'] if _row_bh else None
        
        if bh_id:
            cur.execute("""
                UPDATE dossie_legislacoes_pasta
                SET n_paginas=%s, 
                    total_arquivos=%s,
                    arquivos_originais=%s::jsonb,
                    atualizado_em=NOW()
                WHERE busca_historico_id=%s AND legislacao_label=%s
            """, (n_paginas, len(arq_atual), _json_w.dumps(arq_atual), bh_id, legislacao_label))
        else:
            cur.execute("""
                UPDATE dossie_legislacoes_pasta
                SET n_paginas=%s, 
                    total_arquivos=%s,
                    arquivos_originais=%s::jsonb,
                    atualizado_em=NOW()
                WHERE dossie_id=%s AND legislacao_label=%s
            """, (n_paginas, len(arq_atual), _json_w.dumps(arq_atual), dossie_id, legislacao_label))
        
        # Marca uploads como aplicado
        for a in arquivos_aplicados:
            cur.execute("UPDATE dossie_anexos_uploads SET aplicado=TRUE, aplicado_em=NOW() WHERE id=%s", (a['id'],))
        
        conn.commit(); cur.close(); conn.close()
        logger_w.info(f"[aplicar dossie {dossie_id}] banco atualizado ({n_paginas} pgs, {len(arq_atual)} arquivos)")
    except Exception as e:
        import traceback
        logger_w.error(f"[aplicar dossie {dossie_id}] erro atualizando banco: {e}\n{traceback.format_exc()[-300:]}")
    
    # Re-roda Etapa 4.5 com o PDF atualizado
    try:
        from modulos.etapa_45 import detectar_anexos_citados
        import json as _json_w2
        
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT arquivos_originais FROM dossie_legislacoes_pasta 
            WHERE dossie_id=%s AND legislacao_label=%s
        """, (dossie_id, legislacao_label))
        row = cur.fetchone()
        cur.close()
        arquivos = (row['arquivos_originais'] or []) if row else []
        label_lower = legislacao_label.lower()
        anexos_baixados = [a for a in arquivos if not (a.get('nome', '').lower().startswith(label_lower) and a.get('tipo_detectado') == 'pdf')]
        
        logger_w.info(f"[aplicar dossie {dossie_id}] re-rodando Etapa 4.5...")
        res = detectar_anexos_citados(legislacao_label, pdf_atual, anexos_baixados)
        
        if res.get('sucesso'):
            cur = conn.cursor()
            # Pega busca_historico_id pra acertar a linha certa
            cur.execute("""SELECT busca_historico_id FROM dossie_legislacoes_pasta 
                           WHERE dossie_id=%s AND legislacao_label=%s 
                           ORDER BY criado_em DESC LIMIT 1""",
                        (dossie_id, legislacao_label))
            _bh_row = cur.fetchone()
            _bh_id = _bh_row[0] if _bh_row else None
            
            if _bh_id:
                cur.execute("""
                    UPDATE dossie_legislacoes_pasta
                    SET anexos_citados=%s::jsonb,
                        anexos_faltantes=%s::jsonb,
                        atualizado_em=NOW()
                    WHERE busca_historico_id=%s AND legislacao_label=%s
                """, (
                    _json_w2.dumps(res.get('anexos_citados', [])),
                    _json_w2.dumps(res.get('anexos_faltantes', [])),
                    _bh_id, legislacao_label
                ))
            else:
                cur.execute("""
                    UPDATE dossie_legislacoes_pasta
                    SET anexos_citados=%s::jsonb,
                        anexos_faltantes=%s::jsonb,
                        atualizado_em=NOW()
                    WHERE dossie_id=%s AND legislacao_label=%s
                """, (
                    _json_w2.dumps(res.get('anexos_citados', [])),
                    _json_w2.dumps(res.get('anexos_faltantes', [])),
                    dossie_id, legislacao_label
                ))
            conn.commit(); cur.close()
            logger_w.info(f"[aplicar dossie {dossie_id}] Etapa 4.5 atualizada: {len(res.get('anexos_faltantes', []))} faltantes")
        conn.close()
    except Exception as e:
        import traceback
        logger_w.error(f"[aplicar dossie {dossie_id}] erro Etapa 4.5: {e}\n{traceback.format_exc()[-300:]}")
    
    # Limpa tmp
    try:
        _sh_w.rmtree(tmp_dir, ignore_errors=True)
    except: pass
    
    logger_w.info(f"[aplicar dossie {dossie_id} {legislacao_label}] CONCLUIDO")


@app.route('/api/prompts-salvos')
@login_required
def api_prompts_salvos_listar():
    try:
        # Filtra apenas o prompt ativo do pipeline end-to-end
        rows = qry("SELECT id, nome, arquivo_origem, tamanho_bytes, criado_em FROM prompts_salvos ORDER BY criado_em DESC")
        for r in rows:
            if r.get('criado_em'):
                r['criado_em'] = r['criado_em'].strftime('%d/%m/%Y %H:%M') if hasattr(r['criado_em'], 'strftime') else str(r['criado_em'])
        return jsonify({'success': True, 'data': rows or []})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/prompts-salvos/<int:pid>')
@login_required
def api_prompts_salvos_get(pid):
    try:
        rows = qry("SELECT id, nome, conteudo, metadata FROM prompts_salvos WHERE id=%s", (pid,))
        if not rows:
            return jsonify({'success': False, 'error': 'nao encontrado'}), 404
        r = rows[0]
        # metadata pode vir como dict (psycopg ja parseou) ou string
        if r.get('metadata') and isinstance(r['metadata'], str):
            try:
                import json as _json_g
                r['metadata'] = _json_g.loads(r['metadata'])
            except Exception:
                r['metadata'] = {}
        return jsonify({'success': True, 'data': r})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/prompts-salvos', methods=['POST'])
@login_required
def api_prompts_salvos_upload():
    try:
        criados = []
        for f in request.files.getlist('arquivos'):
            if not f or not f.filename: continue
            nome = f.filename
            ext = nome.lower().rsplit('.', 1)[-1] if '.' in nome else ''
            try:
                if ext == 'txt':
                    conteudo = f.read().decode('utf-8', errors='ignore')
                elif ext == 'pdf':
                    import tempfile as _t, subprocess as _sp, os as _o
                    with _t.NamedTemporaryFile(suffix='.pdf', delete=False) as _tf:
                        f.save(_tf.name); _path = _tf.name
                    try:
                        _r = _sp.run(['pdftotext', _path, '-'], capture_output=True, text=True, timeout=30)
                        conteudo = _r.stdout if _r.returncode == 0 else ''
                    finally:
                        try: _o.unlink(_path)
                        except: pass
                elif ext in ('doc', 'docx'):
                    import tempfile as _t, os as _o
                    with _t.NamedTemporaryFile(suffix='.'+ext, delete=False) as _tf:
                        f.save(_tf.name); _path = _tf.name
                    try:
                        if ext == 'docx':
                            import docx as _docx
                            _d = _docx.Document(_path)
                            conteudo = '\n'.join(p.text for p in _d.paragraphs)
                        else:
                            import subprocess as _sp
                            _r = _sp.run(['antiword', _path], capture_output=True, text=True, timeout=30)
                            conteudo = _r.stdout if _r.returncode == 0 else ''
                    finally:
                        try: _o.unlink(_path)
                        except: pass
                else:
                    continue
                conn = get_db(); cur = conn.cursor()
                # Extrair metadata YAML do prompt (se existir)
                try:
                    from modulos.gerador_hibrido import extrair_metadata_yaml
                    _meta = extrair_metadata_yaml(conteudo)
                except Exception:
                    _meta = {}
                import json as _json_meta
                cur.execute("INSERT INTO prompts_salvos (nome, conteudo, arquivo_origem, tamanho_bytes, criado_por, metadata) VALUES (%s,%s,%s,%s,%s,%s::jsonb) RETURNING id",
                    (nome.rsplit('.',1)[0][:200], conteudo.strip(), nome, len(conteudo.encode('utf-8')), session.get('user_id'), _json_meta.dumps(_meta)))
                pid = cur.fetchone()[0]
                conn.commit(); cur.close(); conn.close()
                criados.append({'id': pid, 'nome': nome})
            except Exception as _ef:
                continue
        return jsonify({'success': True, 'criados': criados})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/prompts-salvos/excluir', methods=['POST'])
@login_required
def api_prompts_salvos_excluir():
    try:
        ids = (request.json or {}).get('ids', [])
        if not ids:
            return jsonify({'success': False, 'error': 'ids obrigatorio'}), 400
        conn = get_db(); cur = conn.cursor()
        cur.execute("DELETE FROM prompts_salvos WHERE id = ANY(%s)", (ids,))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/planilhas-base', methods=['GET'])
@login_required
def api_planilhas_base_listar():
    try:
        from datetime import timedelta as _td
        # Filtra apenas a planilha base mais recente (v3.1)
        rows = qry("""SELECT id, nome, tamanho_bytes, criado_em, criado_por_nome, arquivo_path
                       FROM planilhas_base ORDER BY criado_em DESC""")
        for r in rows:
            if r.get('criado_em') and hasattr(r['criado_em'], 'strftime'):
                r['criado_em'] = (r['criado_em'] - _td(hours=3)).strftime('%d/%m/%Y %H:%M')
        return jsonify({'success': True, 'data': rows or []})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/planilhas-base', methods=['POST'])
@login_required
def api_planilhas_base_upload():
    try:
        f = request.files.get('arquivo')
        if not f or not f.filename:
            return jsonify({'success': False, 'error': 'arquivo obrigatorio'}), 400
        nome = f.filename
        if not nome.lower().endswith(('.xlsx','.xls','.xlsm')):
            return jsonify({'success': False, 'error': 'arquivo deve ser .xlsx'}), 400
        # Salvar arquivo
        import os as _o, time as _t
        ts = int(_t.time())
        safe_name = ''.join(ch if ch.isalnum() or ch in '._-' else '_' for ch in nome)[:100]
        arq_path = f'/var/www/urbanlex/static/planilhas_base/{ts}_{safe_name}'
        f.save(arq_path)
        size = _o.path.getsize(arq_path)
        # Pegar nome do usuario
        uid = session.get('user_id')
        unome = ''
        try:
            urows = qry("SELECT nome FROM users WHERE id=%s", (uid,))
            if urows: unome = urows[0].get('nome') or ''
        except Exception: pass
        # Inserir
        conn = get_db(); cur = conn.cursor()
        cur.execute("""INSERT INTO planilhas_base (nome, arquivo_path, tamanho_bytes, criado_por, criado_por_nome)
                       VALUES (%s,%s,%s,%s,%s) RETURNING id""",
                    (nome[:200], arq_path, size, uid, unome))
        new_id = cur.fetchone()[0]
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True, 'id': new_id})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/planilhas-base/<int:pid>/download')
@login_required
def api_planilhas_base_download(pid):
    try:
        rows = qry("SELECT nome, arquivo_path FROM planilhas_base WHERE id=%s", (pid,))
        if not rows:
            return jsonify({'success': False, 'error': 'nao encontrado'}), 404
        from flask import send_file as _sf
        return _sf(rows[0]['arquivo_path'], as_attachment=True, download_name=rows[0]['nome'])
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/planilhas-base/<int:pid>', methods=['DELETE'])
@login_required
def api_planilhas_base_excluir(pid):
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'apenas admin pode excluir'}), 403
        rows = qry("SELECT arquivo_path FROM planilhas_base WHERE id=%s", (pid,))
        if rows and rows[0].get('arquivo_path'):
            try:
                import os as _o2
                _o2.remove(rows[0]['arquivo_path'])
            except Exception: pass
        conn = get_db(); cur = conn.cursor()
        cur.execute("DELETE FROM planilhas_base WHERE id=%s", (pid,))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/gerador/compilados', methods=['GET'])
@login_required
def api_gerador_compilados_listar():
    """Lista todos os compilados (visivel para todos os usuarios)."""
    try:
        rows = qry("""
            SELECT gc.id, gc.job_id, gc.municipio, gc.estado, gc.zip_path, gc.tabela_path,
                   gc.criado_em, gc.criado_por, gc.criado_por_nome,
                   gc.legislacao_label, gc.busca_historico_id,
                   dlp.legislacao_meta::text AS legislacao_meta_json,
                   bh.iniciado_em AS dossie_criado_em
            FROM gerador_compilados gc
            LEFT JOIN dossie_legislacoes_pasta dlp
                ON dlp.busca_historico_id = gc.busca_historico_id
                AND dlp.legislacao_label = gc.legislacao_label
            LEFT JOIN buscas_historico bh
                ON bh.id = gc.busca_historico_id
            ORDER BY gc.municipio, gc.estado, gc.busca_historico_id NULLS LAST, gc.criado_em DESC
        """)
        from datetime import timedelta as _td
        import json as _json
        for r in rows:
            if r.get('criado_em') and hasattr(r['criado_em'], 'strftime'):
                r['criado_em'] = (r['criado_em'] - _td(hours=3)).strftime('%d/%m/%Y %H:%M')
            if r.get('dossie_criado_em') and hasattr(r['dossie_criado_em'], 'strftime'):
                r['dossie_criado_em'] = (r['dossie_criado_em'] - _td(hours=3)).strftime('%d/%m/%Y %H:%M')
            # Parseia legislacao_meta (tipo, numero, ano, descricao)
            if r.get('legislacao_meta_json'):
                try: r['legislacao_meta'] = _json.loads(r['legislacao_meta_json'])
                except: r['legislacao_meta'] = {}
                r.pop('legislacao_meta_json', None)
            else:
                r['legislacao_meta'] = {}
        return jsonify({'success': True, 'data': rows or []})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/gerador/compilados', methods=['POST'])
@login_required
def api_gerador_compilados_adicionar():
    """Adiciona um compilado vindo do dossie. Visivel para todos."""
    try:
        data = request.json or {}
        municipio = data.get('municipio'); estado = data.get('estado')
        zip_path = data.get('zip_path') or data.get('zip')
        tabela_path = data.get('tabela_path') or data.get('tabela')
        job_id = data.get('job_id') or ''
        if not municipio or not estado:
            return jsonify({'success': False, 'error': 'municipio e estado obrigatorios'}), 400
        # Quem esta criando
        uid = session.get('user_id')
        unome = ''
        try:
            urows = qry("SELECT nome FROM users WHERE id=%s", (uid,))
            if urows: unome = urows[0].get('nome') or ''
        except Exception: pass
        conn = get_db(); cur = conn.cursor()
        cur.execute("""
            INSERT INTO gerador_compilados (job_id, municipio, estado, zip_path, tabela_path, criado_por, criado_por_nome)
            VALUES (%s,%s,%s,%s,%s,%s,%s) RETURNING id
        """, (job_id, municipio, estado, zip_path, tabela_path, uid, unome))
        new_id = cur.fetchone()[0]
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True, 'id': new_id})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/gerador/compilados/<int:cid>', methods=['DELETE'])
@login_required
def api_gerador_compilados_excluir(cid):
    """Exclui compilado + apaga cache do pipeline se for o ultimo do municipio."""
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'apenas admin pode excluir'}), 403
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT municipio, estado, zip_path FROM gerador_compilados WHERE id=%s", (cid,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': 'nao encontrado'}), 404
        municipio = row['municipio']
        estado = row['estado']
        cur.execute("DELETE FROM gerador_compilados WHERE id=%s", (cid,))
        cur.execute("SELECT COUNT(*) as n FROM gerador_compilados WHERE municipio=%s AND estado=%s", (municipio, estado))
        outros = cur.fetchone()['n']
        # Apaga cache /pipelines/Slug_UF/leg_<md5>/ da legislacao deste compilado
        cache_apagado = False
        try:
            import os as _os_c, shutil as _sh_c
            from modulos.pipeline_extracao_lei import _slug_municipio, PIPELINES_BASE_DIR, calcular_md5_zip
            slug = _slug_municipio(municipio, estado)
            mun_dir = _os_c.path.join(PIPELINES_BASE_DIR, slug)
            zip_path_comp = row.get('zip_path')
            if zip_path_comp:
                md5 = calcular_md5_zip(zip_path_comp)
                if md5:
                    leg_dir = _os_c.path.join(mun_dir, f"leg_{md5[:12]}")
                    if _os_c.path.isdir(leg_dir):
                        try: _sh_c.rmtree(leg_dir); cache_apagado = True
                        except Exception as e_c: logger.warning(f"falha apagando {leg_dir}: {e_c}")
            # Remove pasta-mae do municipio se ficou vazia (sem compilados e sem dossies)
            if outros == 0:
                cur.execute("SELECT COUNT(*) as n FROM buscas_historico WHERE municipio=%s AND estado=%s", (municipio, estado))
                dossies = cur.fetchone()['n']
                if dossies == 0 and _os_c.path.isdir(mun_dir):
                    try:
                        if not _os_c.listdir(mun_dir):
                            _os_c.rmdir(mun_dir)
                    except Exception: pass
        except Exception as e_pipe:
            logger.warning(f"erro limpando pipeline: {e_pipe}")
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True, 'cache_pipeline_apagado': cache_apagado, 'outros_compilados': outros})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]})

@app.route('/api/gerador/test-preencher', methods=['GET'])
@login_required
def api_gerador_test_preencher():
    """Teste sintetico: preenche uma planilha com linhas mock incluindo dict/list/None
    para validar se o fix de conversao de tipos funciona."""
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'apenas admin'}), 403
        # Pegar a planilha base mais recente
        rows = qry("SELECT id, arquivo_path, nome FROM planilhas_base ORDER BY criado_em DESC LIMIT 1")
        if not rows:
            return jsonify({'success': False, 'error': 'nenhuma planilha base cadastrada'}), 400
        template_path_orig = rows[0]['arquivo_path']
        nome_base = rows[0]['nome']
        # Copiar para temp
        import shutil as _sh, os as _o, uuid as _u, openpyxl as _xl
        from datetime import datetime as _dt
        job_id = str(_u.uuid4())[:8]
        out_path = f'/var/www/urbanlex/static/downloads/teste_preencher_{job_id}.xlsx'
        _sh.copy(template_path_orig, out_path)
        wb = _xl.load_workbook(out_path)
        ws = wb.active
        # Achar header
        _header_row_idx = 1
        _col_por_header = {}
        for ri, row in enumerate(ws.iter_rows(values_only=True), start=1):
            _vals = [str(v).strip() if v is not None else '' for v in row]
            if sum(1 for v in _vals if v) >= 10:
                _header_row_idx = ri
                for ci, v in enumerate(_vals, start=1):
                    if v: _col_por_header[v] = ci
                break
        if not _col_por_header:
            return jsonify({'success': False, 'error': 'header nao encontrado'}), 400
        # Pular sub-linha de header (EN)
        _start_row_test = _header_row_idx + 1
        for ri in range(_header_row_idx + 1, _header_row_idx + 5):
            row_vals = [ws.cell(row=ri, column=col).value for col in range(1, min(ws.max_column+1, 20))]
            row_vals = [str(v).strip() if v else '' for v in row_vals]
            if sum(1 for v in row_vals if v) >= 5:
                _start_row_test = ri + 1
            else:
                break
        # Linhas mock — usar os 3 primeiros headers
        headers_lst = list(_col_por_header.keys())
        h1 = headers_lst[0] if len(headers_lst) > 0 else 'col1'
        h2 = headers_lst[1] if len(headers_lst) > 1 else 'col2'
        h3 = headers_lst[2] if len(headers_lst) > 2 else 'col3'
        linhas_mock = [
            {h1: 'string normal', h2: 42, h3: 'NI'},
            {h1: {'valor_aninhado': 'isso é um dict'}, h2: 'string', h3: ''},  # dict
            {h1: ['lista', 'de', 'valores'], h2: None, h3: 'NI'},  # list e None
            {h1: 'teste final', h2: True, h3: 3.14},
        ]
        for i, linha in enumerate(linhas_mock):
            for header, val in linha.items():
                col = _col_por_header.get(header)
                if col:
                    if val is None:
                        _v = ''
                    elif isinstance(val, (dict, list)):
                        import json as _jcell
                        _v = _jcell.dumps(val, ensure_ascii=False)
                    else:
                        _v = val
                    ws.cell(row=_start_row_test+i, column=col, value=_v)
        wb.save(out_path)
        return jsonify({
            'success': True,
            'arquivo_url': f'/static/downloads/teste_preencher_{job_id}.xlsx',
            'arquivo_nome': f'teste_preencher_{job_id}.xlsx',
            'linhas': len(linhas_mock),
            'mensagem': 'Teste OK — incluiu dict, list, None, int, bool, float, string'
        })
    except Exception as e:
        import traceback as _tb
        return jsonify({'success': False, 'error': str(e), 'trace': _tb.format_exc()[-500:]})

@app.route('/api/gerador/conflitos/<job_id>')
@login_required
def api_gerador_conflitos(job_id):
    """Lista conflitos detectados durante uma execucao do gerador."""
    try:
        rows = qry("""SELECT zona, coluna, lei_vencedora, valor_vencedor,
                          lei_perdedora, valor_perdedor, motivo, detectado_em
                   FROM gerador_conflitos_log
                   WHERE job_id=%s
                   ORDER BY detectado_em ASC""", (job_id,))
        conflitos = []
        for r in rows:
            conflitos.append({
                'zona': r.get('zona'),
                'coluna': r.get('coluna'),
                'lei_vencedora': r.get('lei_vencedora'),
                'valor_vencedor': r.get('valor_vencedor'),
                'lei_perdedora': r.get('lei_perdedora'),
                'valor_perdedor': r.get('valor_perdedor'),
                'motivo': r.get('motivo'),
                'detectado_em': r.get('detectado_em').strftime('%d/%m/%Y %H:%M') if r.get('detectado_em') else ''
            })
        return jsonify({'success': True, 'job_id': job_id, 'conflitos': conflitos, 'total': len(conflitos)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/gerador/preview-prompts')
@login_required
def api_gerador_preview_prompts():
    """Retorna os prompts da arquitetura PDF-driven (6 passadas) ja com metadata aplicado."""
    try:
        from modulos.gerador_hibrido import (
            DEFAULT_METADATA,
            prompt_passada_0_catalogacao_avancada,
            prompt_passada_1_inventario,
            prompt_passada_2_pdf_driven_principal,
            prompt_passada_2_pdf_driven_verificacao,
            prompt_passada_3_validacao,
            extrair_metadata_yaml
        )
        prompt_id = request.args.get('prompt_id', '').strip()
        prompt_text = ''
        meta = dict(DEFAULT_METADATA)
        if prompt_id:
            try:
                rows = qry("SELECT conteudo, metadata FROM prompts_salvos WHERE id=%s", (int(prompt_id),))
                if rows:
                    prompt_text = rows[0].get('conteudo','') or ''
                    md = rows[0].get('metadata')
                    if md:
                        if isinstance(md, str):
                            import json as _j
                            md = _j.loads(md)
                        if isinstance(md, dict):
                            meta.update(md)
            except Exception: pass
        if not prompt_text:
            prompt_text = '[Conteudo do prompt selecionado pelo usuario]'
        # Placeholders
        nomes_exemplo = ['LC_270_2024.pdf', 'LC_281_2025.pdf', 'Errata_LC_281.pdf', '...']
        lei_id_ex = 'LC 281/2025'
        zonas_ex = [{'nome_canonico': 'ZRM2', 'unidade_territorial': 'AP-2.1', 'leis_aplicaveis': ['LC 281/2025']}]
        headers_ex = ['Pais', 'Estado', 'Municipio', 'Zona Urbana', 'Subzona Urbana', 'Area_Planejamento', '<...>']
        instrucao_rev_ex = '\n=== INSTRUCAO DE REVOGACOES PARCIAIS ===\nA lei "LC 270/2024" sofreu as seguintes revogacoes parciais:\n  - Art. 47 revogado por LC 281/2025\n  - Tabela XV do Anexo II revogada por LC 281/2025 (apenas para AP-1, AP-2.1)\n'
        estado_ex = {'AP-1||zrm2-a||a||||': ['Pais', 'Estado', 'Zona Urbana', 'Coeficiente de aproveitamento basico']}
        json_consol_ex = '<JSON_CONSOLIDADO>'
        # Gerar
        p0 = prompt_passada_0_catalogacao_avancada(nomes_exemplo, meta)
        p1 = prompt_passada_1_inventario(prompt_text, meta)
        p15_explicacao = (
            'PASSADA 1.5 — MATRIZ DE VIGENCIA (sem chamada IA)\n\n'
            'Backend Python que processa a saida da P0 e calcula:\n'
            '1. Quais leis foram revogadas TOTALMENTE (filtradas)\n'
            '2. Revogacoes PARCIAIS detectadas (escopo: dispositivo + geografia + uso)\n'
            '3. Ordem de processamento dos PDFs (hierarquia: LC > LO > Decreto > Errata)\n'
            '4. Para cada PDF, lista de revogacoes a aplicar antes de ler o conteudo\n\n'
            'Esta passada nao consome tokens. Resultado e usado nas instrucoes da P2.'
        )
        p2a = prompt_passada_2_pdf_driven_principal(prompt_text, lei_id_ex, zonas_ex, headers_ex, instrucao_rev_ex, estado_ex, meta)
        p2b = prompt_passada_2_pdf_driven_verificacao(prompt_text, lei_id_ex, {'linhas': []}, zonas_ex, headers_ex, instrucao_rev_ex, meta)
        p3 = prompt_passada_3_validacao(json_consol_ex, zonas_ex, meta)
        return jsonify({
            'success': True,
            'metadata_versao': meta.get('versao', 0),
            'arquitetura': 'pdf-driven-dupla-passagem',
            'p0': p0,
            'p1': p1,
            'p15': p15_explicacao,
            'p2a': p2a,
            'p2b': p2b,
            'p3': p3
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/gerador/ultimo-job')
@login_required
def api_gerador_ultimo_job():
    """Retorna o ultimo job de gerador (visivel para todos)."""
    try:
        rows = qry("SELECT valor FROM gerador_estado WHERE chave='ultimo_job' LIMIT 1")
        if not rows or not rows[0].get('valor'):
            return jsonify({'success': True, 'job_id': None})
        return jsonify({'success': True, 'job_id': rows[0]['valor']})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/processamento/<int:proc_id>/refs_externas')
@login_required
def api_processamento_refs_externas(proc_id):
    """
    Retorna o relatorio das referencias externas de um processamento (JSON gerado).
    Para cada ref, retorna o status colorido (vermelho/amarelo/verde) + detalhes
    pra alimentar o accordion da aba JSON Gerados.
    """
    try:
        import psycopg2
        from modulos.resolver_referencias_externas import relatorio_refs

        # Le o resultado_json do processamento
        conn = psycopg2.connect(os.environ['DATABASE_URL'])
        cur = conn.cursor()
        cur.execute("""
            SELECT id, municipio, estado, legislacao_label, processado_em,
                   resultado_json, sucesso
            FROM legislacao_processamentos
            WHERE id = %s
        """, (proc_id,))
        row = cur.fetchone()
        if not row:
            cur.close()
            conn.close()
            return jsonify({'success': False, 'error': f'processamento {proc_id} nao encontrado'}), 404

        id_, municipio, estado, label, processado_em, resultado_json, sucesso = row

        if not sucesso:
            cur.close()
            conn.close()
            return jsonify({
                'success': True,
                'processamento_id': id_,
                'legislacao_label': label,
                'municipio': municipio,
                'estado': estado,
                'refs_externas': [],
                'resumo': {'total': 0, 'vermelho': 0, 'amarelo': 0, 'verde': 0},
                'observacao': 'processamento sem sucesso',
            })

        # Gera relatorio (consulta legislacoes + legislacao_processamentos)
        relatorio = relatorio_refs(resultado_json, conn)
        cur.close()
        conn.close()

        # Serializa: converte datas e remove campos pesados
        def _serializar(item):
            r = dict(item)
            # Converte parametros_pendentes (set ja foi convertido pra lista no modulo)
            # detalhes_banco pode ter data_publicacao como objeto date
            for chave in ('detalhes_banco', 'detalhes_processamento'):
                if r.get(chave):
                    d = dict(r[chave])
                    for k, v in list(d.items()):
                        if hasattr(v, 'isoformat'):
                            d[k] = v.isoformat()
                    r[chave] = d
            return r

        refs_serializadas = [_serializar(it) for it in relatorio]

        resumo = {
            'total': len(refs_serializadas),
            'vermelho': sum(1 for r in refs_serializadas if r['status'] == 'vermelho'),
            'amarelo':  sum(1 for r in refs_serializadas if r['status'] == 'amarelo'),
            'verde':    sum(1 for r in refs_serializadas if r['status'] == 'verde'),
        }

        return jsonify({
            'success': True,
            'processamento_id': id_,
            'legislacao_label': label,
            'municipio': municipio,
            'estado': estado,
            'processado_em': processado_em.isoformat() if processado_em else None,
            'refs_externas': refs_serializadas,
            'resumo': resumo,
        })
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'traceback': traceback.format_exc()}), 500


@app.route('/api/legislacao/<int:legislacao_id>/extrair-info', methods=['POST'])
@login_required
def api_legislacao_extrair_info(legislacao_id):
    """
    Fase B.5: dispara extracao do pipeline pra uma lei externa que ja esta
    no dossie (status amarelo). Procura o dossie da lei + enfileira pipeline.

    Returns: {success, compilado_id, fila_id, label, mensagem}
    """
    try:
        from modulos.dossie_para_gerador import preparar_work_dir_pipeline
        from modulos.pipeline_extracao_lei import enfileirar_extracao as _enf

        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

        # 1. Pega a lei - tenta primeiro em legislacoes, depois fallback em dossie_legislacoes_pasta
        dossie_pasta_id = request.args.get('dossie_pasta_id', type=int)
        tipo = numero = ano = municipio = estado = None
        if legislacao_id and legislacao_id > 0:
            cur.execute("""
                SELECT id, tipo_nome, numero, ano, municipio_nome, estado, ementa
                FROM legislacoes
                WHERE id = %s
            """, (legislacao_id,))
            leg = cur.fetchone()
            if leg:
                tipo = leg['tipo_nome']
                numero = str(leg['numero'])
                ano = str(leg['ano'])
                municipio = leg['municipio_nome']
                estado = leg['estado']
        # Fallback: se nao acha em legislacoes, tenta dossie_pasta_id
        if not tipo and dossie_pasta_id:
            cur.execute("""
                SELECT dlp.legislacao_meta, bh.municipio, bh.estado
                FROM dossie_legislacoes_pasta dlp
                JOIN buscas_historico bh ON bh.id = dlp.busca_historico_id
                WHERE dlp.id = %s
            """, (dossie_pasta_id,))
            row_pasta = cur.fetchone()
            if row_pasta:
                meta = row_pasta['legislacao_meta'] or {}
                tipo = meta.get('tipo') or meta.get('tipo_nome')
                numero = str(meta.get('numero') or '')
                ano = str(meta.get('ano') or '')
                municipio = row_pasta['municipio']
                estado = row_pasta['estado']
        if not tipo:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': f'legislacao_id={legislacao_id} nem dossie_pasta_id={dossie_pasta_id} retornaram lei valida'}), 404
        label_user = f"{tipo} {numero}/{ano}"

        # 2. Procura no dossie via match com legislacao_meta
        # Usa comparacao case-insensitive no tipo_nome
        cur.execute("""
            SELECT dlp.id AS dlp_id, dlp.dossie_id, dlp.legislacao_label,
                   dlp.pasta_path, dlp.busca_historico_id,
                   bh.municipio, bh.estado
            FROM dossie_legislacoes_pasta dlp
            JOIN buscas_historico bh ON bh.id = dlp.busca_historico_id
            WHERE LOWER(bh.municipio) = LOWER(%s)
              AND UPPER(bh.estado) = UPPER(%s)
              AND LOWER(dlp.legislacao_meta->>'tipo') = LOWER(%s)
              AND dlp.legislacao_meta->>'numero' = %s
              AND dlp.legislacao_meta->>'ano' = %s
            ORDER BY dlp.id DESC
            LIMIT 1
        """, (municipio, estado, tipo, numero, ano))
        row = cur.fetchone()

        if not row:
            cur.close(); conn.close()
            return jsonify({
                'success': False,
                'error': f'{label_user} nao encontrada no dossie do municipio {municipio}/{estado}. ' +
                         f'Voce precisa primeiro buscar essa lei no Buscador e organizar no Dossie.',
                'codigo': 'NO_DOSSIE'
            }), 404

        cur.close(); conn.close()

        legislacao_label = row['legislacao_label']
        dossie_id = row['dossie_id']
        busca_id = row['busca_historico_id']

        # 3. Prepara work_dir
        work_dir = preparar_work_dir_pipeline(
            dossie_id=dossie_id,
            busca_historico_id=busca_id,
            legislacao_label=legislacao_label,
            get_db=get_db,
        )
        if not work_dir:
            return jsonify({'success': False, 'error': 'falha preparando work_dir do dossie'}), 500

        zip_path_efetivo = os.path.join(work_dir, 'tudo.pdf')
        if not os.path.exists(zip_path_efetivo):
            return jsonify({'success': False, 'error': f'arquivo {zip_path_efetivo} nao existe'}), 500

        # 4. INSERT em gerador_compilados
        import uuid as _u_c
        novo_job_id = str(_u_c.uuid4())[:8]
        user_name = (session.get('nome') or session.get('email') or 'sistema')

        conn = get_db(); cur = conn.cursor()
        cur.execute("""
            INSERT INTO gerador_compilados
              (job_id, municipio, estado, zip_path, tabela_path, criado_por,
               criado_por_nome, legislacao_label, busca_historico_id)
            VALUES (%s, %s, %s, %s, NULL, %s, %s, %s, %s)
            RETURNING id
        """, (novo_job_id, municipio, estado, zip_path_efetivo,
              session.get('user_id'), user_name, legislacao_label, busca_id))
        compilado_id = cur.fetchone()[0]
        conn.commit(); cur.close(); conn.close()

        # 5. enfileirar extracao DIRETAMENTE (B.5 opcao B = nao precisa do operador clicar)
        fila_id = _enf(
            zip_path=zip_path_efetivo,
            municipio=municipio,
            estado_uf=estado,
            legislacao_id=legislacao_id,
            usar_cache=True,
            consolidar_apos=True,
            ordem=0,
            criado_por=session.get('user_id'),
            legislacao_label=legislacao_label,
        )

        if not fila_id:
            return jsonify({
                'success': False,
                'error': 'gerador_compilado criado mas falhou ao enfileirar pipeline',
                'compilado_id': compilado_id
            }), 500

        return jsonify({
            'success': True,
            'compilado_id': compilado_id,
            'fila_id': fila_id,
            'label': label_user,
            'legislacao_label': legislacao_label,
            'mensagem': f'{label_user} enfileirada no pipeline (fila_id={fila_id}). Aguarde concluir e recarregue refs.'
        })

    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'traceback': traceback.format_exc()[:500]}), 500

@app.route('/api/api-keys/<servico>', methods=['POST'])
@login_required
def api_api_keys_salvar(servico):
    """Salva chave de API no .env e atualiza os.environ. Apenas admin."""
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'apenas admin'}), 403
        chave = (request.json or {}).get('chave', '').strip()
        if not chave:
            return jsonify({'success': False, 'error': 'chave vazia'}), 400
        env_var_map = {
            'gemini-flash': 'GEMINI_API_KEY',
            'gemini-pro': 'GEMINI_API_KEY',
            'claude-sonnet': 'ANTHROPIC_API_KEY',
            'claude-opus': 'ANTHROPIC_API_KEY',
            'gemini': 'GEMINI_API_KEY',
            'claude': 'ANTHROPIC_API_KEY',
        }
        env_var = env_var_map.get(servico)
        if not env_var:
            return jsonify({'success': False, 'error': f'servico desconhecido: {servico}'}), 400
        # Atualizar runtime
        os.environ[env_var] = chave
        # Persistir no .env
        env_path = '/var/www/urbanlex/.env'
        try:
            if os.path.exists(env_path):
                with open(env_path, 'r') as f:
                    linhas = f.readlines()
            else:
                linhas = []
            achou = False
            for i, l in enumerate(linhas):
                if l.startswith(env_var + '=') or l.startswith('export ' + env_var + '='):
                    linhas[i] = f'{env_var}={chave}\n'
                    achou = True
                    break
            if not achou:
                linhas.append(f'{env_var}={chave}\n')
            with open(env_path, 'w') as f:
                f.writelines(linhas)
        except Exception as _ee:
            return jsonify({'success': False, 'error': f'erro persistindo: {_ee}'})
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/api-keys/<servico>', methods=['DELETE'])
@login_required
def api_api_keys_remover(servico):
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'apenas admin'}), 403
        env_var_map = {
            'gemini-flash': 'GEMINI_API_KEY',
            'gemini-pro': 'GEMINI_API_KEY',
            'claude-sonnet': 'ANTHROPIC_API_KEY',
            'claude-opus': 'ANTHROPIC_API_KEY',
            'gemini': 'GEMINI_API_KEY',
            'claude': 'ANTHROPIC_API_KEY',
        }
        env_var = env_var_map.get(servico)
        if not env_var:
            return jsonify({'success': False, 'error': 'servico desconhecido'}), 400
        os.environ.pop(env_var, None)
        env_path = '/var/www/urbanlex/.env'
        try:
            if os.path.exists(env_path):
                with open(env_path, 'r') as f:
                    linhas = f.readlines()
                novas = [l for l in linhas if not l.startswith(env_var + '=') and not l.startswith('export ' + env_var + '=')]
                with open(env_path, 'w') as f:
                    f.writelines(novas)
        except Exception: pass
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/api-keys')
@login_required
def api_api_keys():
    """Lista status das API keys configuradas para uso no gerador."""
    try:
        import os as _os_k
        servicos = [
            ('gemini-flash', 'GEMINI_API_KEY', 'Gemini 2.5 Flash'),
            ('gemini-pro', 'GEMINI_API_KEY', 'Gemini 2.5 Pro'),
            ('gemini-hibrido', 'GEMINI_API_KEY', 'Gemini Hibrido (Pro+Flash)'),
            ('gemini-pro-ocr', 'GEMINI_API_KEY', 'OCR + Gemini Pro'),
            ('triagem-ocr-pro-sonnet', 'GEMINI_API_KEY', 'Triagem + OCR + Pro + Sonnet'),
            ('claude-sonnet', 'ANTHROPIC_API_KEY', 'Claude Sonnet 4.5'),
            ('claude-opus', 'ANTHROPIC_API_KEY', 'Claude Opus 4.7'),
        ]
        data = []
        for srv, env_var, label in servicos:
            key = _os_k.environ.get(env_var, '') or ''
            mascarada = ''
            if key and len(key) > 12:
                mascarada = key[:6] + '****' + key[-4:]
            data.append({
                'servico': srv,
                'label': label,
                'configurada': bool(key),
                'mascarada': mascarada,
            })
        return jsonify({'success': True, 'data': data})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/dossie/dossie-detalhe/<int:dossie_id>')
@login_required
def api_dossie_dossie_detalhe(dossie_id):
    """Retorna detalhes de um dossie (busca_historico) individual."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT id, municipio, estado, zip_path, relatorio_path, tabela_path, iniciado_em FROM buscas_historico WHERE id=%s", (dossie_id,))
        r = cur.fetchone()
        cur.close(); conn.close()
        if not r:
            return jsonify({'success': False, 'error': 'nao encontrado'}), 404
        return jsonify({
            'success': True,
            'id': r['id'],
            'municipio': r['municipio'],
            'estado': r['estado'],
            'zip_url': r['zip_path'],
            'relatorio_url': r['relatorio_path'],
            'tabela_url': r['tabela_path'],
            'iniciado_em': r['iniciado_em'].strftime('%d/%m/%Y %H:%M') if r['iniciado_em'] else None,
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/dossie/legislacao/<int:leg_id>', methods=['DELETE'])
@login_required
def api_dossie_apagar_legislacao_dossie(leg_id):
    """Apaga uma legislacao individual do dossie + pasta no disco + cache pipeline."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute('SELECT id, pasta_path, busca_historico_id, legislacao_label FROM dossie_legislacoes_pasta WHERE id=%s', (leg_id,))
        leg = cur.fetchone()
        if not leg:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': 'nao encontrado'}), 404
        import os as _os_l, shutil as _sh_l, hashlib as _hl
        # 1. Calcula MD5 do pdf_concatenado pra achar cache do pipeline
        cache_apagado = False
        if leg['pasta_path'] and _os_l.path.isdir(leg['pasta_path']):
            concat = _os_l.path.join(leg['pasta_path'], 'pdf_concatenado.pdf')
            if _os_l.path.exists(concat):
                try:
                    with open(concat, 'rb') as _f: md5 = _hl.md5(_f.read()).hexdigest()[:12]
                    cur.execute('SELECT municipio, estado FROM buscas_historico WHERE id=%s', (leg['busca_historico_id'],))
                    bh = cur.fetchone()
                    if bh:
                        from modulos.pipeline_extracao_lei import _slug_municipio, PIPELINES_BASE_DIR
                        slug = _slug_municipio(bh['municipio'], bh['estado'])
                        cache_dir = _os_l.path.join(PIPELINES_BASE_DIR, slug, 'leg_' + md5)
                        if _os_l.path.isdir(cache_dir):
                            try: _sh_l.rmtree(cache_dir); cache_apagado = True
                            except Exception as e: logger.warning('falha apagando cache ' + str(cache_dir) + ': ' + str(e))
                        cur.execute('DELETE FROM gerador_compilados WHERE municipio=%s AND estado=%s AND zip_path LIKE %s', (bh['municipio'], bh['estado'], '%' + 'leg_' + md5 + '%'))
                        cur.execute('DELETE FROM legislacao_processamentos WHERE municipio=%s AND estado=%s AND zip_md5=%s', (bh['municipio'], bh['estado'], md5))
                except Exception as e: logger.warning('erro limpando cache: ' + str(e))
            # 2. Apaga pasta do dossie
            try: _sh_l.rmtree(leg['pasta_path'])
            except Exception as e: logger.warning('falha apagando pasta ' + str(leg['pasta_path']) + ': ' + str(e))
        # 3. Apaga linha em dossie_legislacoes_pasta
        cur.execute('DELETE FROM dossie_legislacoes_pasta WHERE id=%s', (leg_id,))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True, 'cache_pipeline_apagado': cache_apagado})
    except Exception as e:
        import traceback
        logger.error('erro apagar legislacao ' + str(leg_id) + ': ' + traceback.format_exc()[:1000])
        return jsonify({'success': False, 'error': str(e)[:300]}), 500


@app.route('/api/dossie/dossie/<int:dossie_id>', methods=['DELETE'])
@login_required
def api_dossie_apagar_dossie(dossie_id):
    """Apaga dossie + arquivos + pastas /static/dossies + cache /static/pipelines."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT job_id, zip_path, relatorio_path, tabela_path, municipio, estado FROM buscas_historico WHERE id=%s", (dossie_id,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': 'nao encontrado'}), 404
        municipio = row['municipio']
        estado = row['estado']
        import os as _os_d, shutil as _sh_d
        # 1. Apaga arquivos individuais
        for _p in (row['zip_path'], row['relatorio_path'], row['tabela_path']):
            if _p:
                _full = _p.lstrip('/')
                if not _full.startswith('var/www'): _full = '/var/www/urbanlex/' + _full
                try: _os_d.unlink(_full)
                except: pass
        # 2. Apaga pastas /static/dossies/<mun>/busca_<id>/<label>/ via dossie_legislacoes_pasta
        try:
            cur.execute("SELECT pasta_path FROM dossie_legislacoes_pasta WHERE busca_historico_id=%s", (dossie_id,))
            for p in cur.fetchall():
                if p['pasta_path'] and _os_d.path.isdir(p['pasta_path']):
                    try: _sh_d.rmtree(p['pasta_path'])
                    except Exception as e_p: logger.warning(f"falha apagando {p['pasta_path']}: {e_p}")
            cur.execute("DELETE FROM dossie_legislacoes_pasta WHERE busca_historico_id=%s", (dossie_id,))
        except Exception as e_dlp:
            logger.warning(f"sem dossie_legislacoes_pasta ou erro: {e_dlp}")
        # 3. Verifica se ainda tem outros dossies pro mesmo municipio
        cur.execute("SELECT COUNT(*) as n FROM buscas_historico WHERE municipio=%s AND estado=%s AND id != %s", (municipio, estado, dossie_id))
        outros = cur.fetchone()['n']
        # 4. Se for o ultimo dossie -> apaga cache do pipeline + gerador_compilados
        cache_apagado = False
        if outros == 0:
            try:
                from modulos.pipeline_extracao_lei import _slug_municipio, PIPELINES_BASE_DIR
                slug = _slug_municipio(municipio, estado)
                cache_dir = _os_d.path.join(PIPELINES_BASE_DIR, slug)
                if _os_d.path.isdir(cache_dir):
                    try: _sh_d.rmtree(cache_dir); cache_apagado = True
                    except Exception as e_c: logger.warning(f"falha apagando cache {cache_dir}: {e_c}")
                cur.execute("DELETE FROM gerador_compilados WHERE municipio=%s AND estado=%s", (municipio, estado))
            except Exception as e_pipe:
                logger.warning(f"erro limpando pipeline: {e_pipe}")
        # 5. Apaga logs e buscas_historico
        if row['job_id']:
            cur.execute("DELETE FROM buscas_logs WHERE job_id=%s", (row['job_id'],))
        cur.execute("DELETE FROM buscas_historico WHERE id=%s", (dossie_id,))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True, 'cache_pipeline_apagado': cache_apagado, 'outros_dossies_municipio': outros})
    except Exception as e:
        import traceback
        logger.error(f"erro apagar dossie {dossie_id}: {traceback.format_exc()[:1000]}")
        return jsonify({'success': False, 'error': str(e)[:300]})

@app.route('/api/dossie/municipio/<municipio>/<estado>')
@login_required
def api_dossie_municipio_detalhe(municipio, estado):
    """Retorna detalhe completo de um município incluindo logs."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT bh.job_id, bh.concluido_em, bh.sucesso,
                   bh.zip_path, bh.relatorio_path, bh.tabela_path, bh.log_texto
            FROM buscas_historico bh
            WHERE LOWER(bh.municipio)=LOWER(%s) AND LOWER(bh.estado)=LOWER(%s)
            AND bh.sucesso=true
            ORDER BY bh.concluido_em DESC LIMIT 1
        """, (municipio, estado))
        bh = cur.fetchone()
        logs = []
        if bh and bh['job_id']:
            cur.execute("""
                SELECT nivel, msg FROM buscas_logs
                WHERE job_id=%s ORDER BY cursor ASC
            """, (bh['job_id'],))
            logs = [dict(r) for r in cur.fetchall()]
        cur.close(); conn.close()
        return jsonify({
            'success': True,
            'job_id': bh['job_id'] if bh else None,
            'ultima_busca': bh['concluido_em'].strftime('%d/%m/%Y %H:%M') if bh and bh['concluido_em'] else None,
            'zip_url': bh['zip_path'] if bh else None,
            'relatorio_url': bh['relatorio_path'] if bh else None,
            'tabela_url': bh['tabela_path'] if bh else None,
            'logs': logs,
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


# ── Integração Landly ──────────────────────────────────────────────────────────

@app.route('/api/integracao/landly/config', methods=['GET','POST'])
@login_required
def api_integracao_landly_config():
    """Salvar ou buscar configuração da integração Landly."""
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    if request.method == 'POST':
        d = request.json or {}
        cur.execute("SELECT id FROM integracao_landly LIMIT 1")
        exists = cur.fetchone()
        api_key = d.get('api_key','')
        if api_key and not api_key.startswith('••'):
            from cryptography.fernet import Fernet
            _fk = os.environ.get('SECRET_KEY','urbanlex_secret_key_32chars_here').encode()[:32].ljust(32,b'0')
            import base64
            _fkey = base64.urlsafe_b64encode(_fk)
            api_key_enc = Fernet(_fkey).encrypt(api_key.encode()).decode()
        else:
            api_key_enc = None
        if exists:
            if api_key_enc:
                cur.execute("""UPDATE integracao_landly SET api_url=%s, api_key_enc=%s,
                    agendamento_ativo=%s, horario_1=%s, horario_2=%s,
                    max_legislacoes=%s, atualizado_em=NOW() WHERE id=%s""",
                    (d.get('api_url'), api_key_enc, d.get('agendamento_ativo',False),
                     d.get('horario_1'), d.get('horario_2') or None,
                     d.get('max_legislacoes') or None, exists['id']))
            else:
                cur.execute("""UPDATE integracao_landly SET api_url=%s,
                    agendamento_ativo=%s, horario_1=%s, horario_2=%s,
                    max_legislacoes=%s, atualizado_em=NOW() WHERE id=%s""",
                    (d.get('api_url'), d.get('agendamento_ativo',False),
                     d.get('horario_1'), d.get('horario_2') or None,
                     d.get('max_legislacoes') or None, exists['id']))
        else:
            cur.execute("""INSERT INTO integracao_landly
                (api_url, api_key_enc, agendamento_ativo, horario_1, horario_2, max_legislacoes)
                VALUES (%s,%s,%s,%s,%s,%s)""",
                (d.get('api_url'), api_key_enc, d.get('agendamento_ativo',False),
                 d.get('horario_1'), d.get('horario_2') or None,
                 d.get('max_legislacoes') or None))
        conn.commit()
        cur.close(); conn.close()
        # Atualizar cron job
        try:
            _atualizar_cron_landly(d.get('horario_1'), d.get('horario_2'), d.get('agendamento_ativo', False))
        except Exception as _ce: print(f'Cron update error: {_ce}')
        return jsonify({'success': True})
    else:
        cur.execute("""SELECT api_url, agendamento_ativo, horario_1, horario_2,
            max_legislacoes,
            CASE WHEN api_key_enc IS NOT NULL THEN '••••••••••••' ELSE NULL END as api_key_display
            FROM integracao_landly LIMIT 1""")
        r = cur.fetchone()
        cur2 = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur2.execute("""SELECT id, executado_em, total_municipios, novos_municipios, status, erro, municipios_snapshot, novos_snapshot, log_linhas
            FROM integracao_landly_sync ORDER BY executado_em DESC LIMIT 10""")
        syncs = []
        for s in cur2.fetchall():
            syncs.append({
                'executado_em': (s['executado_em'] - __import__('datetime').timedelta(hours=3)).strftime('%d/%m %H:%M'),
                'total_municipios': s['total_municipios'],
                'novos_municipios': s['novos_municipios'],
                'status': s['status'],
                'erro': s['erro'],
                'municipios_snapshot': s['municipios_snapshot'] or [],
                'novos_snapshot': s['novos_snapshot'] or [],
                'log_linhas': s['log_linhas'] or [],
                'id': s['id'],
            })
        cur.close(); cur2.close(); conn.close()
        return jsonify({'success': True, 'config': dict(r) if r else {}, 'syncs': syncs})


@app.route('/api/integracao/landly/testar', methods=['POST'])
@login_required
def api_integracao_landly_testar():
    """Testa conexão com a API Landly."""
    d = request.json or {}
    url = d.get('api_url','').rstrip('/')
    key = d.get('api_key','')
    try:
        import requests as _req
        _headers = {'X-API-Key': key} if key and not key.startswith('••') else {}
        r = _req.get(url, headers=_headers, timeout=10, stream=True)
        r.close()
        if r.status_code == 200:
            return jsonify({'success': True, 'msg': 'Conexão OK'})
        return jsonify({'success': False, 'msg': f'HTTP {r.status_code}'})
    except Exception as e:
        return jsonify({'success': False, 'msg': str(e)})


@app.route('/api/integracao/landly/sincronizar', methods=['POST'])
@login_required
def api_integracao_landly_sincronizar():
    """Executa sincronização manual com Landly."""
    import threading
    def _run():
        _executar_sync_landly()
    threading.Thread(target=_run, daemon=True).start()
    return jsonify({'success': True, 'msg': 'Sincronização iniciada'})


def _atualizar_cron_landly(horario_1, horario_2, ativo):
    """Atualiza cron jobs para sync Landly."""
    import subprocess, re
    script = '/var/www/urbanlex/sync_landly.py'
    log = '/var/log/urbanlex-landly-sync.log'
    # Ler crontab atual
    result = subprocess.run(['crontab', '-l'], capture_output=True, text=True)
    cron = result.stdout if result.returncode == 0 else ''
    # Remover linhas anteriores do landly
    cron = '\n'.join(l for l in cron.splitlines() if 'sync_landly.py' not in l)
    if ativo and horario_1:
        h, m = horario_1.split(':')
        cron += f'\n{m} {h} * * * /usr/bin/python3 {script} >> {log} 2>&1'
        if horario_2:
            h2, m2 = horario_2.split(':')
            cron += f'\n{m2} {h2} * * * /usr/bin/python3 {script} >> {log} 2>&1'
    cron = cron.strip() + '\n'
    proc = subprocess.run(['crontab', '-'], input=cron, capture_output=True, text=True)
    if proc.returncode == 0:
        print(f'Cron Landly atualizado: ativo={ativo} h1={horario_1} h2={horario_2}')
    else:
        print(f'Cron erro: {proc.stderr}')

def _executar_sync_landly():
    """Executa a sincronização com a API Landly."""
    import requests as _req
    from cryptography.fernet import Fernet
    import base64
    from datetime import datetime, timezone, timedelta
    _logs = []
    def _log(nivel, msg):
        br = datetime.now(timezone.utc) - timedelta(hours=3)
        _logs.append({'ts': br.strftime('%H:%M:%S'), 'nivel': nivel, 'msg': msg})
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT api_url, api_key_enc, max_legislacoes FROM integracao_landly LIMIT 1")
        cfg = cur.fetchone()
        if not cfg or not cfg['api_url']:
            cur.close(); conn.close()
            return
        api_key = None
        if cfg['api_key_enc']:
            _fk = os.environ.get('SECRET_KEY','urbanlex_secret_key_32chars_here').encode()[:32].ljust(32,b'0')
            _fkey = base64.urlsafe_b64encode(_fk)
            try: api_key = Fernet(_fkey).decrypt(cfg['api_key_enc'].encode()).decode()
            except: pass
        url = cfg['api_url'].rstrip('/')
        _log('info', f'Iniciando sincronização com Landly...')
        _log('info', f'Conectando à API: {url}')
        headers = {'X-API-Key': api_key} if api_key else {}
        r = _req.get(url, headers=headers, timeout=30)
        data = r.json()
        # Suporta lista direta ou objeto com campo 'properties'
        if isinstance(data, list):
            props = data
        else:
            props = data.get('properties', data.get('municipios', []))
        municipios_landly = list({(p['municipio'], p['estado']) for p in props if p.get('municipio') and p.get('estado')})
        _log('ok', f'✓ {len(props)} propriedades recebidas da API')
        _log('info', f'Verificando municípios existentes no dossiê...')
        cur.execute("SELECT municipio, estado FROM dossie_municipios")
        existentes = {(r['municipio'], r['estado']) for r in cur.fetchall()}
        novos = [(m, e) for m, e in municipios_landly if (m, e) not in existentes]
        _log('ok', f'{len(novos)} municípios novos detectados')
        for mun, est in novos:
            _log('info', f'Adicionando ao dossiê: {mun}/{est}')
            cur.execute("INSERT INTO dossie_municipios (municipio, estado, origem) VALUES (%s,%s,'integracao') ON CONFLICT DO NOTHING", (mun, est))
        import json
        _log('ok', f'✓ Sincronização concluída — {len(municipios_landly)} totais, {len(novos)} novos')
        try:
            with open('/tmp/landly_sync_done', 'w') as _f: _f.write(str(__import__('time').time()))
        except: pass
        cur.execute("""INSERT INTO integracao_landly_sync
            (total_municipios, novos_municipios, status, municipios_snapshot, novos_snapshot, log_linhas)
            VALUES (%s,%s,'sucesso',%s,%s,%s)""",
            (len(municipios_landly), len(novos),
             json.dumps([{'municipio':m,'estado':e} for m,e in municipios_landly]),
             json.dumps([{'municipio':m,'estado':e} for m,e in novos]),
             json.dumps(_logs)))
        conn.commit()
        cur.close(); conn.close()
    except Exception as ex:
        _log('erro', f'Erro: {str(ex)[:200]}')
        try:
            conn2 = get_db(); cur2 = conn2.cursor()
            cur2.execute("INSERT INTO integracao_landly_sync (status, erro, log_linhas) VALUES ('erro',%s,%s)", (str(ex), json.dumps(_logs)))
            conn2.commit(); cur2.close(); conn2.close()
        except: pass

@app.route('/gerador-planilha')
@login_required
def gerador_planilha():
    return render_template('gerador_planilha.html', active_page='gerador-planilha', active_group='gerador', **tmpl_ctx())

@app.route('/api/gerador/iniciar', methods=['POST'])
@editor_required
def api_gerador_iniciar():
    import uuid, threading
    job_id = str(uuid.uuid4())[:8]
    template_file = request.files.get('template')
    planilha_base_id = request.form.get('planilha_base_id', '').strip()
    compilados = json.loads(request.form.get('compilados', '[]'))
    prompt = request.form.get('prompt', '')
    ia_provider = request.form.get('ia_provider', 'gemini-pro')
    prompt_id_form = (request.form.get('prompt_id') or request.form.get('prompt_salvo_id') or '').strip()
    if not compilados:
        return jsonify({'success': False, 'error': 'compilados obrigatorios'}), 400
    if not template_file and not planilha_base_id:
        return jsonify({'success': False, 'error': 'planilha base obrigatoria (anexar ou escolher da lista)'}), 400
    import os as _os_gp
    template_path = f'/var/www/urbanlex/static/downloads/template_{job_id}.xlsx'
    if template_file:
        template_file.save(template_path)
    else:
        # Copiar do banco
        try:
            _rows = qry("SELECT arquivo_path FROM planilhas_base WHERE id=%s", (int(planilha_base_id),))
            if not _rows:
                return jsonify({'success': False, 'error': 'planilha base nao encontrada'}), 404
            import shutil as _sh_gp
            _sh_gp.copy(_rows[0]['arquivo_path'], template_path)
        except Exception as _ep:
            return jsonify({'success': False, 'error': f'erro carregando planilha base: {_ep}'}), 500
    job = {'done': False, 'cancelled': False, 'logs': _LogList(job_id, get_db), 'result': None}
    _buscador_jobs[job_id] = job
    # Registrar como ultimo job (substitui antigo)
    try:
        _conn = get_db(); _cur = _conn.cursor()
        _cur.execute("INSERT INTO gerador_estado (chave, valor, atualizado_em) VALUES ('ultimo_job',%s,now()) ON CONFLICT (chave) DO UPDATE SET valor=EXCLUDED.valor, atualizado_em=now()", (job_id,))
        # Limpar logs antigos de outros jobs de gerador (manter so esse)
        _cur.execute("DELETE FROM buscas_logs WHERE job_id != %s AND job_id IN (SELECT job_id FROM buscas_historico WHERE tipo='gerador')", (job_id,))
        _conn.commit(); _cur.close(); _conn.close()
    except Exception as _eu: pass

    def _run():
        try:
            import zipfile as _zf, tempfile as _tmp, base64 as _b64, openpyxl as _xl, re as _re_gp, shutil as _sh
            from pathlib import Path
            from modulos.gerador_hibrido import (
                extrair_json, DEFAULT_METADATA,
                prompt_passada_0_catalogacao_avancada,
                prompt_passada_1_inventario,
                prompt_passada_2_pdf_driven_principal,
                prompt_passada_2_pdf_driven_verificacao,
                prompt_passada_3_validacao,
                merge_resultado_no_estado,
                estado_para_resumo_para_prompt,
                estado_para_planilha_final,
            )
            from modulos.multi_ia import chamar_ia, chamar_ia_com_blocos, montar_client, info_modelo, adequar_pdfs_para_janela, resolver_ia_para_fase
            from modulos.vigencia import (
                ordenar_pdfs_por_prioridade,
                calcular_matriz_vigencia,
                gerar_instrucao_revogacao_para_pdf,
                filtrar_pdfs_revogados_totalmente,
            )
            _ia = ia_provider or 'gemini-pro'
            info_ia = info_modelo(_ia)
            job['logs'].append({'nivel':'info','msg':f'IA: {_ia} ({info_ia["modelo"]})'})
            client = montar_client(_ia)
            # Cache de clients para modo hibrido (cada fase pode usar modelo diferente)
            _clients_cache = {_ia: client}
            def _resolver_ia_e_client(_label_fase):
                _ia_resolvido = resolver_ia_para_fase(_ia, _label_fase)
                if _ia_resolvido not in _clients_cache:
                    _clients_cache[_ia_resolvido] = montar_client(_ia_resolvido)
                    job['logs'].append({'nivel':'info','msg':f'  [hibrido] criando client {_ia_resolvido} para fase {_label_fase}'})
                return _ia_resolvido, _clients_cache[_ia_resolvido]
            if client is None:
                raise Exception(f'IA {_ia} sem chave configurada')
            _metadata = dict(DEFAULT_METADATA)
            try:
                if prompt_id_form:
                    _mrows = qry("SELECT metadata FROM prompts_salvos WHERE id=%s", (int(prompt_id_form),))
                    if _mrows and _mrows[0].get('metadata'):
                        _md = _mrows[0]['metadata']
                        if isinstance(_md, str):
                            import json as _jm
                            _md = _jm.loads(_md)
                        if isinstance(_md, dict):
                            _metadata.update(_md)
                            job['logs'].append({'nivel':'info','msg':f'Metadata v{_metadata.get("versao",0)}'})
                else:
                    from modulos.gerador_hibrido import extrair_metadata_yaml
                    _ext = extrair_metadata_yaml(prompt)
                    if _ext: _metadata.update(_ext)
            except Exception as _emm:
                job['logs'].append({'nivel':'aviso','msg':f'Erro metadata: {_emm}'})
            todos_anexos = []
            # Log resumo da configuracao do job (essencial pra diagnostico)
            try:
                _info_prompt = '?'
                if prompt_id_form:
                    _pn = qry("SELECT nome, length(conteudo) as bytes FROM prompts_salvos WHERE id=%s", (int(prompt_id_form),))
                    if _pn:
                        _info_prompt = f'{_pn[0]["nome"]} (id={prompt_id_form}, {_pn[0]["bytes"]/1024:.0f}KB, versao={_metadata.get("versao","?")})'
                else:
                    _info_prompt = f'inline ({len(prompt or "")} chars, versao={_metadata.get("versao","?")})'
                _info_planilha = _os_gp.path.basename(template_path)
                if planilha_base_id:
                    _bn = qry("SELECT nome FROM planilhas_base WHERE id=%s", (int(planilha_base_id),))
                    if _bn:
                        _info_planilha = f'{_bn[0]["nome"]} (id={planilha_base_id})'
                _info_comp = ', '.join([f'{c.get("municipio","?")}/{c.get("estado","??")}' for c in compilados])
                job['logs'].append({'nivel':'info','msg':'=== CONFIGURACAO DO JOB ==='})
                job['logs'].append({'nivel':'info','msg':f'  Prompt:    {_info_prompt}'})
                job['logs'].append({'nivel':'info','msg':f'  Planilha:  {_info_planilha}'})
                job['logs'].append({'nivel':'info','msg':f'  IA:        {_ia} ({info_ia["modelo"]})'})
                job['logs'].append({'nivel':'info','msg':f'  Compilados: {len(compilados)} ({_info_comp})'})
                job['logs'].append({'nivel':'info','msg':'==========================='})
            except Exception as _einfo:
                job['logs'].append({'nivel':'aviso','msg':f'Erro logando config: {_einfo}'})
            job['logs'].append({'nivel':'info','msg':'Iniciando arquitetura PDF-driven'})
            for comp in compilados:
                mun = comp.get('municipio','?'); est = comp.get('estado','??')
                if comp.get('zip_on') and comp.get('zip'):
                    zpath = comp['zip']
                    if zpath.startswith('/static/'): zpath = '/var/www/urbanlex' + zpath
                    if _os_gp.path.exists(zpath):
                        # Extracao recursiva: detecta tipo por magic bytes (cobre arquivos
                        # sem extensao e ZIPs aninhados, ex.: anexo do LeisMunicipais e' ZIP)
                        import io as _io_gp_x, re as _re_gp_x
                        def _saneia_nome(s):
                            s = _re_gp_x.sub(r'[\\/:*?"<>|]+', '_', s.strip())
                            return s.strip('_') or 'arquivo'
                        def _extrair_pdfs_rec(blob, origem='arquivo', prof=0):
                            achados = []
                            if prof > 3 or len(blob) < 8: return achados
                            if blob.startswith(b'%PDF'):
                                nome = origem if origem.lower().endswith('.pdf') else origem + '.pdf'
                                achados.append((nome, blob))
                            elif blob.startswith(b'PK\x03\x04'):
                                try:
                                    with _zf.ZipFile(_io_gp_x.BytesIO(blob)) as inner:
                                        for inner_n in inner.namelist():
                                            if inner_n.endswith('/') or inner_n.startswith('__MACOSX'): continue
                                            try: inner_d = inner.read(inner_n)
                                            except Exception: continue
                                            inner_base = _saneia_nome(_os_gp.path.basename(inner_n))
                                            achados.extend(_extrair_pdfs_rec(inner_d, inner_base, prof+1))
                                except Exception: pass
                            return achados
                        try:
                            qtd_zips_aninhados = 0
                            qtd_pdf_extraidos = 0
                            def _pasta_lei_de_path(p):
                                # Se o path passa por '/anexos/', pasta_lei e tudo ANTES.
                                # Isso resiste a nomes de arquivo que contem '/' (ex.:
                                # 'Lei Complementar No 148/2023 - Xangri-la-RS').
                                if not p: return ''
                                idx = p.lower().find('/anexos/')
                                if idx >= 0:
                                    return p[:idx]
                                parts = p.split('/')
                                return '/'.join(parts[:-1])
                            with _zf.ZipFile(zpath) as zf:
                                for fname in zf.namelist():
                                    if fname.endswith('/') or fname.startswith('__MACOSX'): continue
                                    base = _saneia_nome(_os_gp.path.basename(fname))
                                    try: data = zf.read(fname)
                                    except Exception: continue
                                    if len(data) < 100: continue
                                    if data.startswith(b'PK\x03\x04'): qtd_zips_aninhados += 1
                                    pasta_lei_calc = _pasta_lei_de_path(fname)
                                    pdfs = _extrair_pdfs_rec(data, base)
                                    for nome_pdf, data_pdf in pdfs:
                                        todos_anexos.append({'title': f'{mun}/{est}: {nome_pdf}', 'data_b64': _b64.standard_b64encode(data_pdf).decode(), 'nome_arquivo': nome_pdf, 'pasta_lei': pasta_lei_calc, 'caminho_zip': fname})
                                        qtd_pdf_extraidos += 1
                            if qtd_zips_aninhados:
                                job['logs'].append({'nivel':'info','msg':f'  Compilado {mun}/{est}: {qtd_pdf_extraidos} PDFs ({qtd_zips_aninhados} de ZIPs aninhados)'})
                        except Exception as _ezi:
                            job['logs'].append({'nivel':'erro','msg':f'Erro ZIP: {_ezi}'})
                if comp.get('tabela_on') and comp.get('tabela'):
                    tpath = comp['tabela']
                    if tpath.startswith('/static/'): tpath = '/var/www/urbanlex' + tpath
                    if _os_gp.path.exists(tpath):
                        try:
                            with open(tpath, 'rb') as f2:
                                data = f2.read()
                                todos_anexos.append({'title': f'Tabela {mun}/{est}', 'data_b64': _b64.standard_b64encode(data).decode(), 'nome_arquivo': _os_gp.path.basename(tpath)})
                        except Exception: pass
            if not todos_anexos:
                raise Exception('Nenhum PDF valido')
            _kb_total = sum(len(_b64.b64decode(a['data_b64'])) for a in todos_anexos) / 1024
            job['logs'].append({'nivel':'info','msg':f'Total: {len(todos_anexos)} arquivo(s), {_kb_total/1024:.1f}MB'})

            # ====== PRE-PASSAGEM: NORMALIZACAO PDF (Plano Y v2 - Etapa 1) ======
            # Se IA selecionada usa pipeline 'triagem_ocr_pro_sonnet',
            # normaliza PDFs com OCRmyPDF para garantir texto extraivel
            # mantendo layout visual original (PDF searchable).
            _info_ia_pre = info_modelo(_ia)
            if _info_ia_pre.get('pipeline') == 'triagem_ocr_pro_sonnet':
                job['logs'].append({'nivel':'info','msg':'======= PRE-PASSAGEM: NORMALIZACAO PDF ======='})
                try:
                    from modulos.pdf_normalizador import normalizar_pdf as _norm_pdf
                    import tempfile as _tmpnorm, base64 as _b64norm
                    _pre_aplicado = 0
                    _pre_skip = 0
                    _pre_falhou = 0
                    _pre_t0 = time.time()
                    for _pi, _pdf_item in enumerate(todos_anexos, 1):
                        try:
                            _nome_pdf = _pdf_item.get('nome_arquivo') or f'pdf_{_pi}'
                            with _tmpnorm.NamedTemporaryFile(suffix='.pdf', delete=False, prefix='urb_norm_') as _tf_norm:
                                _tf_norm.write(_b64norm.b64decode(_pdf_item['data_b64']))
                                _tf_norm_path = _tf_norm.name
                            _norm_res = _norm_pdf(_tf_norm_path)
                            if _norm_res.get('ocr_aplicado'):
                                _new_path = _norm_res['path_final']
                                with open(_new_path, 'rb') as _f_new:
                                    _new_data = _f_new.read()
                                _pdf_item['data_b64'] = _b64norm.standard_b64encode(_new_data).decode()
                                _ocr_res = _norm_res.get('ocr_resultado') or {}
                                _diag = _norm_res.get('diagnostico') or {}
                                _pre_aplicado += 1
                                job['logs'].append({
                                    'nivel': 'info',
                                    'msg': f'  {_nome_pdf[:50]}: OCR aplicado ({_diag.get("paginas_imagem", "?")} pag-img de {_diag.get("total_paginas", "?")} | {_ocr_res.get("tempo_ms", 0)}ms)'
                                })
                            elif _norm_res.get('erro'):
                                _pre_falhou += 1
                                job['logs'].append({
                                    'nivel': 'aviso',
                                    'msg': f'  {_nome_pdf[:50]}: falha normalizacao - {_norm_res["erro"][:100]}'
                                })
                            else:
                                _pre_skip += 1
                        except Exception as _e_pre:
                            _pre_falhou += 1
                            job['logs'].append({
                                'nivel': 'aviso',
                                'msg': f'  Erro normalizando PDF {_pi}: {type(_e_pre).__name__}: {str(_e_pre)[:100]}'
                            })
                    _pre_dt = int(time.time() - _pre_t0)
                    job['logs'].append({
                        'nivel': 'info',
                        'msg': f'Pre-passagem concluida em {_pre_dt}s: {_pre_aplicado} OCR aplicado, {_pre_skip} ja searchable, {_pre_falhou} falhas'
                    })
                except Exception as _e_pre_setup:
                    job['logs'].append({
                        'nivel': 'aviso',
                        'msg': f'Pre-passagem falhou no setup: {type(_e_pre_setup).__name__}: {str(_e_pre_setup)[:200]} - continuando sem normalizacao'
                    })

                # ====== TRIAGEM FLASH (Plano Y v2 - Etapa 3) ======
                # 1a passada: classifica paginas por tipo e mapeia zonas->paginas.
                # Resultado fica em job['triagem'] para uso pela 2a passada (Pro filtrada).
                try:
                    from modulos.triagem_flash import triagem_anexos as _triagem_anexos
                    from modulos.triagem_flash import filtrar_anexos_por_zonas as _filtrar_por_zonas
                    _ia_flash = 'gemini-flash'
                    _cli_flash = montar_client(_ia_flash)
                    _job_id_safe = str(job.get('id') or job.get('job_id') or 'unknown').replace('/', '_')
                    _triagem_persist = f'/var/www/urbanlex/static/debug/triagem/{_job_id_safe}.json'
                    _triagem_res = _triagem_anexos(
                        pdfs_list=todos_anexos,
                        client=_cli_flash,
                        ia_id=_ia_flash,
                        logs=job['logs'],
                        persistir_em=_triagem_persist,
                    )
                    job['triagem'] = _triagem_res
                    
                    # ====== MODO DEV: filtrar zonas via URBANLEX_DEV_ZONAS ======
                    import os as _os_dev
                    _dev_zonas_str = (_os_dev.environ.get('URBANLEX_DEV_ZONAS') or '').strip()
                    if _dev_zonas_str:
                        _dev_zonas = [z.strip().upper() for z in _dev_zonas_str.split(',') if z.strip()]
                        if _dev_zonas:
                            _antes = len(todos_anexos)
                            _filtrados = _filtrar_por_zonas(todos_anexos, _dev_zonas, _triagem_res.get('indice_zona_paginas', {}))
                            if _filtrados and len(_filtrados) < _antes:
                                todos_anexos = _filtrados
                                # Reconstroi anexos_principais ainda nao foi feito, vai ser recalculado abaixo
                                job['logs'].append({
                                    'nivel': 'info',
                                    'msg': f'  [MODO DEV] URBANLEX_DEV_ZONAS={_dev_zonas}: filtrado de {_antes} para {len(todos_anexos)} PDFs'
                                })
                            else:
                                job['logs'].append({
                                    'nivel': 'aviso',
                                    'msg': f'  [MODO DEV] URBANLEX_DEV_ZONAS={_dev_zonas} nao encontrou PDFs relevantes; usando todos'
                                })
                except Exception as _e_triagem:
                    job['logs'].append({
                        'nivel': 'aviso',
                        'msg': f'Triagem falhou: {type(_e_triagem).__name__}: {str(_e_triagem)[:200]} - continuando sem triagem'
                    })
                    job['triagem'] = None
            # Para P0/P1/P3 (catalogacao, inventario, validacao) usar SO os PDFs
            # principais (corpo das leis), NAO os anexos volumosos. Os anexos
            # ainda vao pra P2 vinculados a sua lei via mapa_anexos_extras.
            anexos_principais = [a for a in todos_anexos if '/anexos/' not in (a.get('caminho_zip') or '').lower()]
            if len(anexos_principais) < len(todos_anexos):
                _kb_p = sum(len(_b64.b64decode(a['data_b64'])) for a in anexos_principais) / 1024
                job['logs'].append({'nivel':'info','msg':f'  P0/P1/P3 usarao {len(anexos_principais)} PDF(s) principais ({_kb_p/1024:.1f}MB) — anexos so na P2'})
            else:
                anexos_principais = todos_anexos
            # adequar_pdfs_para_janela agora e feito automaticamente dentro de chamar_ia_com_blocos
            wb = _xl.load_workbook(template_path)
            ws = wb.active
            _headers = []
            _header_row_idx = 1
            for ri, row in enumerate(ws.iter_rows(values_only=True), start=1):
                _vals = [str(v).strip() if v is not None else '' for v in row]
                if sum(1 for v in _vals if v) >= 10:
                    _header_row_idx = ri
                    _headers = [v for v in _vals if v]
                    break
            job['logs'].append({'nivel':'info','msg':f'Planilha: {len(_headers)} colunas (header linha {_header_row_idx})'})
            job['logs'].append({'nivel':'info','msg':'======= P0 CATALOGACAO ======='})
            catalogacao = []
            try:
                _p0 = prompt_passada_0_catalogacao_avancada([a['nome_arquivo'] for a in anexos_principais], _metadata)
                _ia_f, _cli_f = _resolver_ia_e_client('P0')
                _r0 = chamar_ia_com_blocos(_cli_f, _ia_f, _p0, anexos_principais, job['logs'], 'P0', chave_agregar='arquivos')
                _j0 = extrair_json(_r0)
                if _j0 and 'arquivos' in _j0:
                    catalogacao = _j0['arquivos']
                    job['logs'].append({'nivel':'ok','msg':f'P0: {len(catalogacao)} catalogado(s)'})
                else:
                    raise Exception('P0: JSON invalido')
            except Exception as _e0:
                job['logs'].append({'nivel':'erro','msg':f'P0: {_e0}'})
                raise
            job['logs'].append({'nivel':'info','msg':'======= P1.5 MATRIZ VIGENCIA ======='})
            matriz = calcular_matriz_vigencia(catalogacao)
            cat_filtrada = filtrar_pdfs_revogados_totalmente(catalogacao, matriz)
            cat_ordenada = ordenar_pdfs_por_prioridade(cat_filtrada)
            job['logs'].append({'nivel':'info','msg':f'Revogadas: {len(matriz["leis_revogadas_totalmente"])}, Parciais: {len(matriz["revogacoes_parciais"])}'})
            job['logs'].append({'nivel':'info','msg':'======= P1 INVENTARIO ======='})
            zonas_canonicas = []
            try:
                _p1 = prompt_passada_1_inventario(prompt, _metadata)
                _ia_f, _cli_f = _resolver_ia_e_client('P1')
                _r1 = chamar_ia_com_blocos(_cli_f, _ia_f, _p1, anexos_principais, job['logs'], 'P1', chave_agregar=_metadata['chave_inventario'])
                _j1 = extrair_json(_r1)
                if _j1 and _metadata['chave_inventario'] in _j1:
                    zonas_canonicas = _j1[_metadata['chave_inventario']] or []
                    job['logs'].append({'nivel':'ok','msg':f'P1: {len(zonas_canonicas)} zonas'})
                else:
                    raise Exception('P1: JSON invalido')
            except Exception as _e1:
                job['logs'].append({'nivel':'erro','msg':f'P1: {_e1}'})
                raise
            if not zonas_canonicas:
                raise Exception('P1 nao identificou zonas')
            job['logs'].append({'nivel':'info','msg':f'======= P2 PDF-DRIVEN ({len(cat_ordenada)} PDFs) ======='})
            estado_planilha = {}
            conflitos_log = []
            mapa_nome_anexo = {a['nome_arquivo']: a for a in todos_anexos}
            mapa_anexos_extras = {}
            for _a in todos_anexos:
                _cz = (_a.get('caminho_zip') or '').lower()
                if '/anexos/' in _cz:
                    _pl = _a.get('pasta_lei', '')
                    if _pl:
                        mapa_anexos_extras.setdefault(_pl, []).append(_a)
            for idx_pdf, item_cat in enumerate(cat_ordenada, 1):
                if job.get('cancelled'):
                    job['logs'].append({'nivel':'aviso','msg':'Cancelado'})
                    break
                lei_id = item_cat.get('identificacao', '?')
                nome_arq = item_cat.get('nome_arquivo', '?')
                anexo_pdf = mapa_nome_anexo.get(nome_arq)
                if not anexo_pdf:
                    job['logs'].append({'nivel':'aviso','msg':f'PDF {nome_arq} nao encontrado'})
                    continue
                _pasta_da_lei = anexo_pdf.get('pasta_lei', '')
                _anexos_extras = mapa_anexos_extras.get(_pasta_da_lei, [])
                _chave_agreg = _metadata.get('chave_zona_individual', 'linhas')
                # Batching D (refinado): dividir anexos por TAMANHO (MB) e quantidade.
                # Cap de seguranca: ~8MB e <=40 PDFs por batch (folga vs Gemini 5min).
                # Se a lei principal sozinha for grande, chamar_ia_com_blocos comprime/divide.
                _BATCH_BYTES = 6 * 1024 * 1024
                _BATCH_QTD   = 24
                def _tam_pdf_b64(_p):
                    try: return len(_b64.b64decode(_p['data_b64']))
                    except Exception: return 1024 * 100
                _tam_lei = _tam_pdf_b64(anexo_pdf)
                _batches = []
                if not _anexos_extras:
                    _batches = [[]]
                else:
                    _atual = []
                    _atual_bytes = 0
                    for _ax in _anexos_extras:
                        _t = _tam_pdf_b64(_ax)
                        # se nao cabe e ja tem algo no batch, fecha o batch atual
                        if _atual and (_atual_bytes + _t > _BATCH_BYTES or len(_atual) >= _BATCH_QTD):
                            _batches.append(_atual)
                            _atual = []
                            _atual_bytes = 0
                        _atual.append(_ax)
                        _atual_bytes += _t
                    if _atual: _batches.append(_atual)
                _qt_anexos = len(_anexos_extras)
                _mb_lei = _tam_lei / (1024*1024)
                _mb_anexos = sum(_tam_pdf_b64(a) for ax in _batches for a in ax) / (1024*1024)
                if len(_batches) > 1:
                    job['logs'].append({'nivel':'info','msg':f'--- {idx_pdf}/{len(cat_ordenada)}: {lei_id} (lei {_mb_lei:.1f}MB + {_qt_anexos} anexos {_mb_anexos:.1f}MB divididos em {len(_batches)} batches) ---'})
                elif _qt_anexos:
                    job['logs'].append({'nivel':'info','msg':f'--- {idx_pdf}/{len(cat_ordenada)}: {lei_id} (lei {_mb_lei:.1f}MB + {_qt_anexos} anexo(s) {_mb_anexos:.1f}MB) ---'})
                else:
                    job['logs'].append({'nivel':'info','msg':f'--- {idx_pdf}/{len(cat_ordenada)}: {lei_id} (lei {_mb_lei:.1f}MB, sem anexos) ---'})
                instrucao_rev = gerar_instrucao_revogacao_para_pdf(lei_id, matriz)
                resumo_estado = estado_para_resumo_para_prompt(estado_planilha)
                # Helper: tenta enviar [lei] + anexos. Se 504, subdivide anexos pela metade.
                # Recursivo ate caber. Retorna lista de JSONs (um por sub-batch que conseguiu).
                # Estrategia condicional por modelo (gemini=pdf_nativo, claude=texto_lei_principal)
                from modulos.multi_ia import info_modelo as _info_modelo_strat
                _estrat = _info_modelo_strat(_ia).get('estrategia_pdf', 'pdf_nativo')
                job['logs'].append({'nivel':'info','msg':f'  [DEBUG] estrategia ativa: {_estrat} (ia={_ia})'})
                _texto_lei_cache = {'val': None}  # cache do texto extraido (1 por loop)

                def _extrair_texto_lei(_pdf_b64):
                    """Extrai texto da lei via pdftotext. Retorna '' se PDF escaneado/falhar."""
                    if _texto_lei_cache['val'] is not None:
                        return _texto_lei_cache['val']
                    try:
                        import base64, tempfile, subprocess
                        # debug: tipo e estrutura
                        if isinstance(_pdf_b64, dict):
                            _ks = list(_pdf_b64.keys())
                            job['logs'].append({'nivel':'info','msg':f'  [DEBUG_TXT] dict com keys={_ks}'})
                            _b64_str = _pdf_b64.get('data_b64') or _pdf_b64.get('data') or _pdf_b64.get('content') or ''
                        elif isinstance(_pdf_b64, bytes):
                            job['logs'].append({'nivel':'info','msg':f'  [DEBUG_TXT] bytes ({len(_pdf_b64)} bytes), assumindo PDF binario'})
                            _bin = _pdf_b64
                            _b64_str = None  # sinaliza que ja esta binario
                        else:
                            job['logs'].append({'nivel':'info','msg':f'  [DEBUG_TXT] tipo={type(_pdf_b64).__name__}, len={len(str(_pdf_b64))}'})
                            _b64_str = _pdf_b64
                        if _b64_str is not None:
                            _bin = base64.b64decode(_b64_str) if _b64_str else b''
                        job['logs'].append({'nivel':'info','msg':f'  [DEBUG_TXT] _bin len={len(_bin)} bytes'})
                        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as _tf:
                            _tf.write(_bin)
                            _tmp_path = _tf.name
                        _r = subprocess.run(['pdftotext', '-layout', _tmp_path, '-'],
                                            capture_output=True, timeout=60)
                        os.unlink(_tmp_path)
                        _txt = _r.stdout.decode('utf-8', errors='replace')
                        _texto_lei_cache['val'] = _txt
                        return _txt
                    except Exception as _exe:
                        job['logs'].append({'nivel':'aviso','msg':f'  pdftotext falhou: {str(_exe)[:120]}'})
                        _texto_lei_cache['val'] = ''
                        return ''

                def _chamar_com_subdivisao(_lei_pdf, _anexos_lista, _label_base, _prompt_para_ia, _profundidade=0):
                    _resultados = []
                    # Modo OCR: extrai texto/tabelas estruturadas dos anexos antes de mandar para IA
                    _info_modelo_atual = info_modelo(_ia)
                    _usa_ocr = _info_modelo_atual.get('pre_processamento') == 'ocr_tabelas'
                    if _usa_ocr:
                        try:
                            from modulos.ocr_tabelas import processar_pdf as _ocr_pdf
                            import base64 as _b64ocr, tempfile as _tmpocr, os as _osocr, json as _json_ocr
                            _markdown_anexos = []
                            _falhas_ocr = 0
                            _diagnostico_batch = []  # diagnostico detalhado por anexo
                            for _ai, _ax in enumerate(_anexos_lista or [], 1):
                                _nome_anexo = _ax.get('nome_arquivo') or _ax.get('title') or f'anexo_{_ai}'
                                _diag_anexo = {'idx': _ai, 'nome': _nome_anexo}
                                try:
                                    # Salva PDF temporariamente para processar
                                    with _tmpocr.NamedTemporaryFile(suffix='.pdf', delete=False) as _tf:
                                        _tf.write(_b64ocr.b64decode(_ax['data_b64']))
                                        _tf_path = _tf.name
                                    _r_ocr = _ocr_pdf(_tf_path)
                                    _osocr.unlink(_tf_path)
                                    # Captura diagnostico rico
                                    _diag_anexo.update({
                                        'metodo': _r_ocr.get('metodo'),
                                        'overall': _r_ocr.get('overall'),
                                        'chars': _r_ocr.get('chars', 0),
                                        'paginas_total': _r_ocr.get('paginas_total', 0),
                                        'paginas_ok': _r_ocr.get('paginas_ok', 0),
                                        'paginas_falhou': _r_ocr.get('paginas_falhou', 0),
                                        'tempo_ms': _r_ocr.get('tempo_ms', 0),
                                        'erro_geral': _r_ocr.get('erro_geral'),
                                        'tamanho_bytes': _r_ocr.get('tamanho_bytes', 0),
                                        'paginas_detalhe': _r_ocr.get('paginas_detalhe', []),
                                    })
                                    _conteudo = _r_ocr.get('conteudo_md', '') or ''
                                    if _conteudo and len(_conteudo) >= 100:
                                        _markdown_anexos.append(f'\n=== ANEXO: {_nome_anexo} (metodo: {_r_ocr.get("metodo")}) ===\n{_conteudo}\n=== FIM ANEXO ===\n')
                                        _diag_anexo['usado_no_prompt'] = True
                                    else:
                                        _falhas_ocr += 1
                                        _diag_anexo['usado_no_prompt'] = False
                                except Exception as _eocr:
                                    _falhas_ocr += 1
                                    _diag_anexo.update({
                                        'metodo': 'exception',
                                        'overall': 'falhou',
                                        'erro_geral': f'{type(_eocr).__name__}: {str(_eocr)[:200]}',
                                        'usado_no_prompt': False,
                                    })
                                _diagnostico_batch.append(_diag_anexo)
                            # Persiste diagnostico do batch em JSON
                            try:
                                _dbg_dir_ocr = '/var/www/urbanlex/static/debug/ocr'
                                _osocr.makedirs(_dbg_dir_ocr, exist_ok=True)
                                _dbg_label_safe = _label_base.replace('/', '_').replace('.', '_')
                                _dbg_file = f'{_dbg_dir_ocr}/{job_id}_{_dbg_label_safe}.json'
                                with open(_dbg_file, 'w') as _df_ocr:
                                    _json_ocr.dump({
                                        'job_id': job_id,
                                        'label': _label_base,
                                        'total_anexos': len(_anexos_lista or []),
                                        'sucesso': len(_markdown_anexos),
                                        'falhou': _falhas_ocr,
                                        'anexos': _diagnostico_batch,
                                    }, _df_ocr, indent=2, ensure_ascii=False)
                            except Exception:
                                pass
                            # Log detalhado: lista PDFs problematicos com motivo
                            _falhas_resumo = []
                            for _d in _diagnostico_batch:
                                if not _d.get('usado_no_prompt'):
                                    _motivo = _d.get('erro_geral') or _d.get('overall', '?')
                                    _falhas_resumo.append(f"{_d.get('nome', '?')[:40]}: {_motivo[:80]}")
                            if _markdown_anexos:
                                # Lei via texto + anexos via Markdown estruturado, sem PDF binario
                                _texto_lei = _extrair_texto_lei(_lei_pdf) if _estrat == 'texto_lei_principal' else ''
                                _prefixo_lei = (f'=== TEXTO DA LEI PRINCIPAL ===\n{_texto_lei}\n=== FIM DO TEXTO DA LEI ===\n\n' if _texto_lei and len(_texto_lei) >= 1000 else '')
                                _prompt_efetivo = (_prefixo_lei + ''.join(_markdown_anexos) + '\n\n' + _prompt_para_ia)
                                _pdfs_call = [] if _texto_lei and len(_texto_lei) >= 1000 else [_lei_pdf]
                                job['logs'].append({'nivel':'info','msg':f'  {_label_base}: OCR processou {len(_markdown_anexos)}/{len(_anexos_lista)} anexos como Markdown ({_falhas_ocr} falharam)'})
                                # Detalhe das falhas (ate 5 primeiras para nao poluir log)
                                if _falhas_resumo:
                                    for _fr in _falhas_resumo[:5]:
                                        job['logs'].append({'nivel':'aviso','msg':f'    OCR FALHOU: {_fr}'})
                                    if len(_falhas_resumo) > 5:
                                        job['logs'].append({'nivel':'aviso','msg':f'    ... +{len(_falhas_resumo)-5} outras falhas (ver JSON: {job_id}_{_dbg_label_safe}.json)'})
                            else:
                                # Todos falharam -> fallback para PDF nativo
                                job['logs'].append({'nivel':'aviso','msg':f'  {_label_base}: OCR falhou em todos os {len(_anexos_lista)} anexos, fallback para PDF nativo'})
                                for _fr in _falhas_resumo[:3]:
                                    job['logs'].append({'nivel':'aviso','msg':f'    OCR FALHOU: {_fr}'})
                                _usa_ocr = False
                        except Exception as _e_setup_ocr:
                            job['logs'].append({'nivel':'aviso','msg':f'  {_label_base}: erro setup OCR ({_e_setup_ocr}), usando PDF nativo'})
                            _usa_ocr = False
                    if not _usa_ocr:
                        # Fluxo original (sem OCR)
                        if _estrat == 'texto_lei_principal':
                            _texto = _extrair_texto_lei(_lei_pdf)
                            if _texto and len(_texto) >= 1000:
                                _pdfs_call = list(_anexos_lista or [])
                                _prompt_efetivo = f'=== TEXTO DA LEI PRINCIPAL ===\n{_texto}\n=== FIM DO TEXTO DA LEI ===\n\n' + _prompt_para_ia
                            else:
                                job['logs'].append({'nivel':'info','msg':f'  {_label_base}: lei sem texto extraivel ({len(_texto)} chars), usando PDF nativo'})
                                _pdfs_call = [_lei_pdf] + (_anexos_lista or [])
                                _prompt_efetivo = _prompt_para_ia
                        else:
                            _pdfs_call = [_lei_pdf] + (_anexos_lista or [])
                            _prompt_efetivo = _prompt_para_ia
                    try:
                        _ia_f_p2, _cli_f_p2 = _resolver_ia_e_client(_label_base)
                        _r = chamar_ia_com_blocos(_cli_f_p2, _ia_f_p2, _prompt_efetivo, _pdfs_call,
                                                   job['logs'], _label_base, chave_agregar=_chave_agreg,
                                                   max_tentativas=2)
                        return [_r]
                    except Exception as _e:
                        _msg = str(_e).lower()
                        _e_timeout = ('deadline' in _msg or '504' in _msg or 'timeout' in _msg)
                        if _e_timeout and len(_anexos_lista) > 1 and _profundidade < 4:
                            _meio = len(_anexos_lista) // 2
                            _h1 = _anexos_lista[:_meio]
                            _h2 = _anexos_lista[_meio:]
                            job['logs'].append({'nivel':'aviso','msg':f'  {_label_base} estourou. Subdividindo {len(_anexos_lista)} anexos em {len(_h1)}+{len(_h2)} (prof. {_profundidade+1})'})
                            _resultados.extend(_chamar_com_subdivisao(_lei_pdf, _h1, f'{_label_base}.sub1', _prompt_para_ia, _profundidade+1))
                            _resultados.extend(_chamar_com_subdivisao(_lei_pdf, _h2, f'{_label_base}.sub2', _prompt_para_ia, _profundidade+1))
                            return _resultados
                        else:
                            job['logs'].append({'nivel':'erro','msg':f'  {_label_base} falhou definitivamente: {str(_e)[:160]}'})
                            return []
                _j2a = None  # ultimo resultado, usado pela verificacao P2b
                _ultimo_r2a = ''  # ultima resposta bruta, usada pelo prompt P2b
                _total_principal = 0
                try:
                    _p2a = prompt_passada_2_pdf_driven_principal(prompt, lei_id, zonas_canonicas, _headers, instrucao_rev, resumo_estado, _metadata)
                    for _bi, _batch_anexos in enumerate(_batches, 1):
                        if job.get('cancelled'): break
                        _label_batch = f'P2.{idx_pdf}a' if len(_batches) == 1 else f'P2.{idx_pdf}a.{_bi}/{len(_batches)}'
                        # Chama com auto-subdivisao em caso de 504
                        _resultados_lista = _chamar_com_subdivisao(anexo_pdf, _batch_anexos, _label_batch, _p2a)
                        for _ri, _r2a in enumerate(_resultados_lista):
                            _ultimo_r2a = _r2a
                            _label_real = _label_batch if len(_resultados_lista) == 1 else f'{_label_batch}.r{_ri+1}'
                            try:
                                _dbg_dir = '/var/www/urbanlex/static/debug'
                                _os_gp.makedirs(_dbg_dir, exist_ok=True)
                                _dbg_path = f'{_dbg_dir}/{job_id}_{_label_real.replace("/","_")}_{nome_arq[:50]}.json'
                                with open(_dbg_path, 'w') as _df:
                                    _df.write(_r2a)
                            except Exception:
                                pass
                            _jb = extrair_json(_r2a)
                            if _jb:
                                _key_count_b = 0
                                try:
                                    _linhas_b = _jb.get(_metadata['chave_zona_individual'], _jb.get('linhas', []))
                                    for _ln in _linhas_b:
                                        if isinstance(_ln, dict):
                                            _key_count_b += sum(1 for v in _ln.values() if v not in (None, '', 'NI'))
                                except: pass
                                _adic_b, _conf_b = merge_resultado_no_estado(_jb, lei_id, nome_arq, estado_planilha, conflitos_log, _metadata)
                                _total_principal += _adic_b
                                job['logs'].append({'nivel':'ok','msg':f'  {_label_real}: IA retornou {_key_count_b} valores, +{_adic_b} celulas, {_conf_b} conflitos'})
                                _j2a = _jb
                    _r2a = _ultimo_r2a
                    job['logs'].append({'nivel':'ok','msg':f'  Principal (total {len(_batches)} batch(es)): +{_total_principal} celulas'})
                    if _j2a:
                        pass  # ja mesclado/logado dentro do loop de batches
                    else:
                        # nenhum batch retornou JSON valido — pula verificacao
                        continue
                except Exception as _e2a:
                    job['logs'].append({'nivel':'erro','msg':f'  Principal: {_e2a}'})
                    continue
                try:
                    _p2b = prompt_passada_2_pdf_driven_verificacao(prompt, lei_id, _j2a, zonas_canonicas, _headers, instrucao_rev, _metadata)
                    _total_verif = 0
                    for _bi, _batch_anexos in enumerate(_batches, 1):
                        if job.get('cancelled'): break
                        _label_b = f'P2.{idx_pdf}b' if len(_batches) == 1 else f'P2.{idx_pdf}b.{_bi}/{len(_batches)}'
                        _resultados_b = _chamar_com_subdivisao(anexo_pdf, _batch_anexos, _label_b, _p2b)
                        for _ri, _r2b in enumerate(_resultados_b):
                            _label_real_b = _label_b if len(_resultados_b) == 1 else f'{_label_b}.r{_ri+1}'
                            try:
                                _dbg_dir = '/var/www/urbanlex/static/debug'
                                _os_gp.makedirs(_dbg_dir, exist_ok=True)
                                _dbg_path = f'{_dbg_dir}/{job_id}_{_label_real_b.replace("/","_")}_{nome_arq[:50]}.json'
                                with open(_dbg_path, 'w') as _df:
                                    _df.write(_r2b)
                            except Exception:
                                pass
                            _j2b = extrair_json(_r2b)
                            if _j2b:
                                _adic_v, _ = merge_resultado_no_estado(_j2b, lei_id, nome_arq, estado_planilha, conflitos_log, _metadata)
                                _total_verif += _adic_v
                                if _adic_v > 0:
                                    job['logs'].append({'nivel':'ok','msg':f'  {_label_real_b}: +{_adic_v}'})
                    if len(_batches) > 1:
                        job['logs'].append({'nivel':'ok','msg':f'  Verificacao (total {len(_batches)} batch(es)): +{_total_verif} celulas'})
                except Exception as _e2b:
                    job['logs'].append({'nivel':'aviso','msg':f'  Verificacao falhou: {_e2b}'})
            job['logs'].append({'nivel':'info','msg':'======= P3 VALIDACAO ======='})
            try:
                import json as _json3
                linhas_finais = estado_para_planilha_final(estado_planilha)
                _consolidado = _json3.dumps({'linhas': linhas_finais}, ensure_ascii=False, indent=2)[:50000]
                _p3 = prompt_passada_3_validacao(_consolidado, zonas_canonicas, _metadata)
                _ia_f, _cli_f = _resolver_ia_e_client('P3')
                _r3 = chamar_ia_com_blocos(_cli_f, _ia_f, _p3, anexos_principais, job['logs'], 'P3', chave_agregar=_metadata['chave_validacao'])
                _j3 = extrair_json(_r3)
                if _j3:
                    faltantes = _j3.get(_metadata['chave_validacao']) or []
                    if faltantes:
                        adic_p3, _ = merge_resultado_no_estado({'linhas': faltantes}, 'P3', 'P3', estado_planilha, conflitos_log, _metadata)
                        job['logs'].append({'nivel':'ok','msg':f'P3: +{adic_p3}'})
            except Exception as _e3:
                job['logs'].append({'nivel':'aviso','msg':f'P3 falhou: {_e3}'})
            if conflitos_log:
                try:
                    conn_c = get_db(); cur_c = conn_c.cursor()
                    for cf in conflitos_log:
                        cur_c.execute("INSERT INTO gerador_conflitos_log (job_id, zona, coluna, lei_vencedora, valor_vencedor, lei_perdedora, valor_perdedor, motivo) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
                            (job_id, cf.get('chave_zona'), cf.get('coluna'), cf.get('lei_vencedora'), str(cf.get('valor_vencedor'))[:500], cf.get('lei_perdedora'), str(cf.get('valor_perdedor'))[:500], cf.get('motivo')))
                    conn_c.commit(); cur_c.close(); conn_c.close()
                    job['logs'].append({'nivel':'info','msg':f'{len(conflitos_log)} conflitos auditados'})
                except Exception as _ec:
                    job['logs'].append({'nivel':'aviso','msg':f'Conflitos: {_ec}'})
            todas_linhas = estado_para_planilha_final(estado_planilha)
            job['logs'].append({'nivel':'info','msg':f'Preenchendo planilha com {len(todas_linhas)} linhas'})
            _col_por_header = {}
            for ri, row in enumerate(ws.iter_rows(values_only=True), start=1):
                _vals = [str(v).strip() if v is not None else '' for v in row]
                if sum(1 for v in _vals if v) >= 10:
                    _header_row_idx = ri
                    for ci, v in enumerate(_vals, start=1):
                        if v: _col_por_header[v] = ci
                    break
            _start_row = _header_row_idx + 1
            for ri in range(_header_row_idx + 1, _header_row_idx + 5):
                row_vals = [ws.cell(row=ri, column=col).value for col in range(1, min(ws.max_column+1, 20))]
                row_vals = [str(v).strip() if v else '' for v in row_vals]
                if sum(1 for v in row_vals if v) >= 5:
                    _start_row = ri + 1
                else:
                    break
            for i, linha in enumerate(todas_linhas):
                if not isinstance(linha, dict): continue
                for header, val in linha.items():
                    col = _col_por_header.get(header)
                    if col:
                        if val is None: _v = ''
                        elif isinstance(val, (dict, list)):
                            import json as _jcell
                            _v = _jcell.dumps(val, ensure_ascii=False)
                        else: _v = val
                        ws.cell(row=_start_row+i, column=col, value=_v)
            from datetime import datetime as _dt, timezone as _tz, timedelta as _td
            _now = _dt.now(_tz(_td(hours=-3)))  # America/Sao_Paulo (UTC-3)
            _ts = _now.strftime('%d%m%Y_%H%M')
            muns = list(set(c['municipio'] for c in compilados))
            _est = compilados[0]['estado'] if compilados else 'XX'
            _mun = muns[0].replace(' ','_') if muns else 'municipio'
            nome_arquivo = f'parametros_urbanos_{_est}_{_mun}_{_ts}.xlsx'
            out_path = f'/var/www/urbanlex/static/downloads/{nome_arquivo}'
            wb.save(out_path)
            job['logs'].append({'nivel':'ok','msg':f'Salvo: {nome_arquivo}'})
            try:
                conn2 = get_db(); cur2 = conn2.cursor()
                for comp in compilados:
                    cur2.execute("INSERT INTO planilhas_geradas (municipio, estado, arquivo_path, arquivo_nome, zonas_count, prompt) VALUES (%s,%s,%s,%s,%s,%s)",
                        (comp['municipio'], comp['estado'], out_path, nome_arquivo, len(todas_linhas), prompt))
                conn2.commit(); cur2.close(); conn2.close()
            except Exception: pass
            job['result'] = {'arquivo_url': f'/static/downloads/{nome_arquivo}', 'arquivo_nome': nome_arquivo, 'zonas': len(todas_linhas)}
        except Exception as e:
            job['logs'].append({'nivel':'erro','msg':f'Erro: {str(e)[:300]}'})
        finally:
            job['done'] = True
            try: _os_gp.unlink(template_path)
            except: pass
    threading.Thread(target=_run, daemon=True).start()
    return jsonify({'success': True, 'job_id': job_id})

@app.route('/api/gerador/job/<job_id>')
@login_required
def api_gerador_job(job_id):
    cursor = int(request.args.get('cursor', 0))
    job = _buscador_jobs.get(job_id)
    if not job:
        # Buscar do banco (job de processo anterior)
        try:
            rows = qry("SELECT nivel, msg, ts FROM buscas_logs WHERE job_id=%s ORDER BY id ASC OFFSET %s", (job_id, cursor))
        except Exception:
            rows = []
        return jsonify({'done': True, 'logs': rows or [], 'cursor': cursor + len(rows or []), 'result': None})
    logs = job['logs'][cursor:]
    return jsonify({'done': job['done'], 'logs': logs, 'cursor': cursor + len(logs), 'result': job.get('result')})

@app.route('/api/gerador/cancelar/<job_id>', methods=['POST'])
@login_required
def api_gerador_cancelar(job_id):
    job = _buscador_jobs.get(job_id)
    if job: job['done'] = True; job['cancelled'] = True
    return jsonify({'success': True})

@app.route('/api/gerador/jsons-gerados/excluir', methods=['POST'])
@login_required
def api_gerador_jsons_gerados_excluir():
    """Apaga registros + pastas leg do disco."""
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'apenas admin'}), 403
        ids = (request.json or {}).get('ids', [])
        if not ids:
            return jsonify({'success': False, 'error': 'nenhum id'}), 400
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute('SELECT id, municipio, estado, output_dir, zip_md5 FROM legislacao_processamentos WHERE id = ANY(%s)', (ids,))
        regs = cur.fetchall()
        import os as _os, shutil as _sh
        from modulos.pipeline_extracao_lei import PIPELINES_BASE_DIR, _slug_municipio
        pastas = 0
        for r in regs:
            target = None
            if r.get('output_dir') and _os.path.isdir(r['output_dir']):
                target = r['output_dir']
            elif r.get('zip_md5') and r.get('municipio') and r.get('estado'):
                slug = _slug_municipio(r['municipio'], r['estado'])
                target = _os.path.join(PIPELINES_BASE_DIR, slug, 'leg_' + r['zip_md5'][:12])
            if target and _os.path.isdir(target):
                try:
                    _sh.rmtree(target); pastas += 1
                except Exception as e:
                    logger.warning('falha apagar ' + str(target) + ': ' + str(e))
        cur.execute('DELETE FROM legislacao_processamentos WHERE id = ANY(%s)', (ids,))
        n = cur.rowcount
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True, 'deletados': n, 'pastas': pastas})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500


@app.route('/api/gerador/json/download/<int:proc_id>')
@login_required
def api_gerador_json_download(proc_id):
    """Baixa resultado_final.json com filename customizado (tipo_numero_ano)."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT lp.id, lp.municipio, lp.estado, lp.legislacao_label, lp.output_dir, lp.processado_em,
                   l.tipo_nome, l.numero AS lei_numero, l.ano AS lei_ano,
                   dlp.legislacao_meta::text AS dlp_meta_json
            FROM legislacao_processamentos lp
            LEFT JOIN legislacoes l ON l.id = lp.legislacao_id
            LEFT JOIN buscas_historico bh
                ON LOWER(bh.municipio) = LOWER(lp.municipio)
                AND bh.estado = lp.estado
            LEFT JOIN dossie_legislacoes_pasta dlp
                ON dlp.busca_historico_id = bh.id
                AND dlp.legislacao_label = lp.legislacao_label
            WHERE lp.id = %s
        """, (proc_id,))
        r = cur.fetchone()
        cur.close(); conn.close()
        if not r:
            return jsonify({'success': False, 'error': 'nao encontrado'}), 404
        # Monta filename (mesma logica de jsons-gerados)
        tipo = (r.get('tipo_nome') or '').replace(' ', '')
        numero = (r.get('lei_numero') or '').replace('/', '-')
        ano = str(r.get('lei_ano') or '')
        if (not tipo or not numero or not ano) and r.get('dlp_meta_json'):
            try:
                import json as _json
                meta = _json.loads(r['dlp_meta_json'])
                if not tipo:
                    t = (meta.get('tipo') or '').strip()
                    if 'Complementar' in t: tipo = 'LC'
                    elif 'Decreto' in t: tipo = 'Dec'
                    elif t: tipo = t.replace(' ', '')
                if not numero: numero = str(meta.get('numero') or '').replace('/', '-').replace('.', '')
                if not ano: ano = str(meta.get('ano') or '')
            except Exception:
                pass
        if (not tipo or not numero or not ano) and r.get('legislacao_label'):
            parts = r['legislacao_label'].split('_')
            if len(parts) >= 3:
                if not tipo: tipo = parts[0]
                if not numero: numero = parts[1]
                if not ano: ano = parts[2]
        tipo = tipo or 'Lei'; numero = numero or '?'; ano = ano or '?'
        mun_safe = (r['municipio'] or '').replace(' ', '-').replace('/', '-')
        estado = r['estado'] or '?'
        dt = r['processado_em']
        data_str = dt.strftime('%d%m%Y_%H%M') if dt else '?'
        filename = f"Extracao_{estado}_{mun_safe}_{tipo}_{numero}_{ano}_{data_str}.json"
        # Path do arquivo
        output_dir = r.get('output_dir') or ''
        json_path = os.path.join(output_dir, 'resultado_final.json')
        if not os.path.exists(json_path):
            return jsonify({'success': False, 'error': 'arquivo nao existe no disco'}), 404
        from flask import send_file
        return send_file(json_path, as_attachment=True, download_name=filename, mimetype='application/json')
    except Exception as e:
        import traceback; logger.error('json_download: ' + traceback.format_exc()[:500])
        return jsonify({'success': False, 'error': str(e)[:200]}), 500


# ════════════════════════════════════════════════════════════
# PREENCHEDOR DE PLANILHA - gera xlsx consolidado por municipio
# ════════════════════════════════════════════════════════════

@app.route('/api/preenchedor/gerar', methods=['POST'])
@login_required
def api_preenchedor_gerar():
    """Recebe lista de jsons_ids + template_id + municipio + estado,
    gera planilha consolidada e salva em planilhas_geradas."""
    try:
        data = request.get_json() or {}
        jsons_ids = data.get('jsons_ids') or []
        template_id = data.get('template_id')   # id em planilhas_base
        municipio = (data.get('municipio') or '').strip()
        estado = (data.get('estado') or '').strip()

        if not jsons_ids:
            return jsonify({'success': False, 'error': 'jsons_ids obrigatorio'}), 400
        if not template_id:
            return jsonify({'success': False, 'error': 'template_id obrigatorio'}), 400
        if not municipio or not estado:
            return jsonify({'success': False, 'error': 'municipio/estado obrigatorios'}), 400

        # Busca planilha base
        rows = qry("SELECT id, nome, arquivo_path FROM planilhas_base WHERE id=%s", (int(template_id),))
        if not rows:
            return jsonify({'success': False, 'error': 'planilha base nao encontrada'}), 404
        template_path = rows[0]['arquivo_path']
        template_nome = rows[0]['nome']
        import os as _os
        if not _os.path.exists(template_path):
            return jsonify({'success': False, 'error': f'arquivo template nao existe: {template_path}'}), 404

        # Detecta versao do template (v3.1, v3.2 etc) pra metadados
        import re as _re
        m = _re.search(r'_v(\d+_\d+)_', template_nome)
        versao_base = m.group(1).replace('_', '.') if m else None

        # Importa o modulo (lazy)
        from modulos import preenchedor_planilha as pp

        # Buffer de log
        logs = []
        def log(m):
            logs.append(m)
            print(f"[preenchedor] {m}")

        # Roda geracao
        resultado = pp.gera_planilha_municipio(
            jsons_ids=[int(j) for j in jsons_ids],
            template_path=template_path,
            municipio=municipio,
            estado=estado,
            get_db_func=get_db,
            log_callback=log,
        )
        if not resultado:
            return jsonify({'success': False, 'error': 'falha ao gerar planilha', 'logs': logs}), 500

        # INSERT em planilhas_geradas
        uid = session.get('user_id')
        unome = ''
        try:
            urows = qry("SELECT nome FROM users WHERE id=%s", (uid,))
            if urows: unome = urows[0].get('nome') or ''
        except Exception: pass

        conn = get_db(); cur = conn.cursor()
        cur.execute("""INSERT INTO planilhas_geradas
            (municipio, estado, arquivo_nome, arquivo_path,
             zonas_count, template_path, criado_por, legislacoes_usadas, tamanho_bytes)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s::jsonb,%s) RETURNING id""",
            (municipio, estado, resultado['filename'], resultado['filepath'],
             resultado['n_zonas'], template_path, uid,
             json.dumps(resultado['jsons_ids']), resultado['tamanho_bytes']))
        new_id = cur.fetchone()[0]
        conn.commit(); cur.close(); conn.close()

        return jsonify({
            'success': True,
            'id': new_id,
            'filename': resultado['filename'],
            'n_zonas': resultado['n_zonas'],
            'n_jsons': resultado['n_jsons_usados'],
            'tamanho_bytes': resultado['tamanho_bytes'],
            'versao_base': versao_base,
            'logs': logs,
        })
    except Exception as e:
        import traceback as _tb
        return jsonify({'success': False, 'error': str(e)[:500], 'tb': _tb.format_exc()[-500:]}), 500


@app.route('/api/planilhas/geradas', methods=['GET'])
@login_required
def api_planilhas_geradas_listar():
    """Lista xlsx gerados, agrupados por municipio."""
    try:
        rows = qry("""
            SELECT pg.id, pg.municipio, pg.estado, pg.arquivo_nome, pg.arquivo_path,
                   pg.zonas_count, pg.tamanho_bytes, pg.criado_em,
                   pg.template_path, pg.legislacoes_usadas,
                   u.nome AS criado_por_nome
              FROM planilhas_geradas pg
              LEFT JOIN users u ON u.id = pg.criado_por
             ORDER BY pg.municipio, pg.estado, pg.criado_em DESC
        """) or []
        from datetime import timedelta as _td
        for r in rows:
            if r.get('criado_em') and hasattr(r['criado_em'], 'strftime'):
                r['criado_em'] = (r['criado_em'] - _td(hours=3)).strftime('%d/%m/%Y %H:%M')
            # versao do template
            tn = (r.get('template_path') or '').split('/')[-1]
            import re as _re
            mm = _re.search(r'_v(\d+_\d+)_', tn)
            r['versao_base'] = mm.group(1).replace('_', '.') if mm else None
        return jsonify({'success': True, 'data': rows})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500


@app.route('/api/planilhas/geradas/<int:pid>', methods=['DELETE'])
@login_required
def api_planilhas_geradas_apagar(pid):
    """Apaga 1 planilha gerada (arquivo + registro)."""
    try:
        rows = qry("SELECT arquivo_path FROM planilhas_geradas WHERE id=%s", (pid,))
        if rows and rows[0].get('arquivo_path'):
            import os as _os
            try: _os.remove(rows[0]['arquivo_path'])
            except Exception: pass
        conn = get_db(); cur = conn.cursor()
        cur.execute("DELETE FROM planilhas_geradas WHERE id=%s", (pid,))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500


@app.route('/api/planilhas/geradas/apagar-lote', methods=['POST'])
@login_required
def api_planilhas_geradas_apagar_lote():
    """Apaga varias planilhas geradas de uma vez."""
    try:
        data = request.get_json() or {}
        ids = data.get('ids') or []
        ids = [int(i) for i in ids if str(i).isdigit()]
        if not ids:
            return jsonify({'success': False, 'error': 'ids vazio'}), 400
        rows = qry("SELECT id, arquivo_path FROM planilhas_geradas WHERE id = ANY(%s)", (ids,)) or []
        import os as _os
        for r in rows:
            try: _os.remove(r['arquivo_path'])
            except Exception: pass
        conn = get_db(); cur = conn.cursor()
        cur.execute("DELETE FROM planilhas_geradas WHERE id = ANY(%s)", (ids,))
        n = cur.rowcount
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True, 'apagados': n})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500


@app.route('/api/planilhas/geradas/<int:pid>/download', methods=['GET'])
@login_required
def api_planilhas_geradas_download(pid):
    """Baixa 1 xlsx."""
    try:
        rows = qry("SELECT arquivo_path, arquivo_nome FROM planilhas_geradas WHERE id=%s", (pid,))
        if not rows:
            return jsonify({'success': False, 'error': 'nao encontrada'}), 404
        from flask import send_file as _sf
        return _sf(rows[0]['arquivo_path'], as_attachment=True, download_name=rows[0]['arquivo_nome'])
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500


@app.route('/api/planilhas/geradas/download-zip', methods=['POST'])
@login_required
def api_planilhas_geradas_download_zip():
    """Recebe lista de ids, junta num zip e baixa."""
    try:
        data = request.get_json() or {}
        ids = data.get('ids') or []
        ids = [int(i) for i in ids if str(i).isdigit()]
        if not ids:
            return jsonify({'success': False, 'error': 'ids vazio'}), 400
        rows = qry("SELECT id, arquivo_path, arquivo_nome FROM planilhas_geradas WHERE id = ANY(%s)", (ids,)) or []
        if not rows:
            return jsonify({'success': False, 'error': 'nenhuma encontrada'}), 404
        import io as _io, zipfile as _zip, os as _os
        buf = _io.BytesIO()
        with _zip.ZipFile(buf, 'w', _zip.ZIP_DEFLATED) as zf:
            for r in rows:
                if r.get('arquivo_path') and _os.path.exists(r['arquivo_path']):
                    zf.write(r['arquivo_path'], arcname=r['arquivo_nome'])
        buf.seek(0)
        from flask import send_file as _sf
        import time as _t
        zipname = f"planilhas_geradas_{_t.strftime('%d%m%Y_%H%M')}.zip"
        return _sf(buf, mimetype='application/zip', as_attachment=True, download_name=zipname)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500


@app.route('/api/gerador/jsons-gerados')
@login_required
def api_gerador_jsons_gerados():
    """Lista JSONs gerados pelo pipeline, agrupados por municipio."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT lp.id, lp.municipio, lp.estado, lp.processado_em,
                   lp.legislacao_id, lp.legislacao_label, lp.output_dir,
                   lp.sucesso, lp.metricas, lp.zip_md5,
                   l.tipo_nome, l.numero as lei_numero, l.ano as lei_ano
            FROM legislacao_processamentos lp
            LEFT JOIN legislacoes l ON l.id = lp.legislacao_id
            WHERE lp.sucesso = TRUE
            ORDER BY lp.municipio ASC, lp.estado ASC, lp.processado_em DESC
        """)
        rows = cur.fetchall()
        cur.close(); conn.close()
        
        # Agrupa por municipio
        grupos = {}
        for r in rows:
            key = f"{r['municipio']}|{r['estado']}"
            if key not in grupos:
                grupos[key] = {'municipio': r['municipio'], 'estado': r['estado'], 'jsons': []}
            
            # Monta nome: tabela legislacoes -> meta do dossie -> label parsing
            tipo = (r.get('tipo_nome') or '').replace(' ', '')
            numero = (r.get('lei_numero') or '').replace('/', '-')
            ano = str(r.get('lei_ano') or '')
            # Fallback 1: meta do dossie
            if (not tipo or not numero or not ano) and r.get('dlp_meta_json'):
                try:
                    import json as _json
                    meta = _json.loads(r['dlp_meta_json'])
                    if not tipo:
                        t = (meta.get('tipo') or '').strip()
                        if 'Complementar' in t: tipo = 'LC'
                        elif 'Decreto' in t: tipo = 'Dec'
                        elif t: tipo = t.replace(' ', '')
                    if not numero:
                        numero = str(meta.get('numero') or '').replace('/', '-').replace('.', '')
                    if not ano:
                        ano = str(meta.get('ano') or '')
                except Exception:
                    pass
            # Fallback 2: parsing do legislacao_label
            if (not tipo or not numero or not ano) and r.get('legislacao_label'):
                parts = r['legislacao_label'].split('_')
                if len(parts) >= 3:
                    if not tipo: tipo = parts[0]
                    if not numero: numero = parts[1]
                    if not ano: ano = parts[2]
            # Defaults
            tipo = tipo or 'Lei'
            numero = numero or '?'
            ano = ano or '?'
            
            mun_safe = (r['municipio'] or '').replace(' ', '-').replace('/', '-')
            estado = r['estado'] or '?'
            
            dt = r['processado_em']
            data_str = dt.strftime('%d%m%Y_%H%M') if dt else '?'
            
            filename = f"Extracao_{estado}_{mun_safe}_{tipo}_{numero}_{ano}_{data_str}.json"
            
            # URL do resultado_final.json
            output_dir = r.get('output_dir') or ''
            web_url = None
            if output_dir and '/static/' in output_dir:
                web_url = output_dir[output_dir.index('/static/'):] + '/resultado_final.json'
            
            metricas = r.get('metricas') or {}
            grupos[key]['jsons'].append({
                'id': r['id'],
                'filename': filename,
                'tipo_nome': r.get('tipo_nome'),
                'numero': r.get('lei_numero'),
                'ano': r.get('lei_ano'),
                'legislacao_label': r.get('legislacao_label'),
                'processado_em': dt.strftime('%d/%m/%Y %H:%M') if dt else None,
                'web_url': web_url,
                'custo_total': metricas.get('custo_total', 0),
                'tempo_total': metricas.get('tempo_total', 0),
                'tokens_in': metricas.get('tokens_in', 0),
                'tokens_out': metricas.get('tokens_out', 0),
                'zip_md5_short': (r.get('zip_md5') or '')[:12],
            })
        
        # Ordena por municipio
        lista = sorted(grupos.values(), key=lambda g: (g['municipio'] or '').lower())
        return jsonify({'success': True, 'grupos': lista, 'total': sum(len(g['jsons']) for g in lista)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:300]}), 500


@app.route('/api/gerador/historico')
@login_required
def api_gerador_historico():
    try:
        conn = get_db(); cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM planilhas_geradas ORDER BY criado_em DESC LIMIT 100")
        rows = cur.fetchall(); cur.close(); conn.close()
        result = []
        for r in rows:
            result.append({'id': r['id'], 'municipio': r['municipio'], 'estado': r['estado'],
                'arquivo_nome': r['arquivo_nome'], 'zonas_count': r['zonas_count'] or 0,
                'criado_em': r['criado_em'].strftime('%d/%m/%Y às %H:%M') if r['criado_em'] else '',
                'arquivo_url': '/static/downloads/' + r['arquivo_nome'] if r['arquivo_nome'] else None,
                'tokens_total': 0,
                'tokens_in': 0,
                'tokens_out': 0,
                'custo_usd': 0.0,
                'tamanho_bytes': 0,
                'ia_modelo': None,
                'ia_provider': None,
                'landly_status': 'nao_enviada',
                'landly_enviado_em': None,
                'landly_erro': None,
                'legislacoes_usadas': []})
        stats = {'total': len(result), 'tokens_total': 0, 'custo_usd_total': 0.0, 'landly_enviadas': 0}
        return jsonify({'success': True, 'planilhas': result, 'stats': stats})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/gerador/excluir', methods=['DELETE'])
@editor_required
def api_gerador_excluir():
    ids = (request.json or {}).get('ids', [])
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("DELETE FROM planilhas_geradas WHERE id = ANY(%s)", (ids,))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/dossie-municipais')
@login_required
def dossie_municipais():
    return render_template('dossie_municipais.html', active_page='dossie-municipais')


@app.route('/api/dossie/municipio/<municipio>/<estado>', methods=['DELETE'])
@login_required
def api_dossie_municipio_deletar(municipio, estado):
    """Apaga todos os dados de um município do dossiê."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        # Buscar arquivos para deletar fisicamente
        cur.execute("""SELECT zip_path, relatorio_path, tabela_path, pdf_path, job_id
            FROM buscas_historico
            WHERE LOWER(municipio)=LOWER(%s) AND LOWER(estado)=LOWER(%s)""",
            (municipio, estado))
        rows = cur.fetchall()
        # Deletar arquivos físicos
        import glob
        for r in rows:
            for fpath in [r['zip_path'], r['relatorio_path'], r['tabela_path'], r['pdf_path']]:
                if fpath:
                    full = os.path.join('/var/www/urbanlex', fpath.lstrip('/'))
                    try:
                        if os.path.exists(full): os.remove(full)
                    except: pass
            # Deletar logs do banco
            if r['job_id']:
                cur.execute("DELETE FROM buscas_logs WHERE job_id=%s", (r['job_id'],))
        # Deletar histórico
        cur.execute("DELETE FROM buscas_historico WHERE LOWER(municipio)=LOWER(%s) AND LOWER(estado)=LOWER(%s)",
            (municipio, estado))
        # Deletar da fila
        cur.execute("DELETE FROM fila_buscas WHERE LOWER(municipio)=LOWER(%s) AND LOWER(estado)=LOWER(%s)",
            (municipio, estado))
        # Deletar do dossiê
        cur.execute("DELETE FROM dossie_municipios WHERE LOWER(municipio)=LOWER(%s) AND LOWER(estado)=LOWER(%s)",
            (municipio, estado))
        conn.commit()
        cur.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/integracao/landly/sync/<int:sync_id>', methods=['DELETE'])
@login_required
def api_landly_sync_deletar(sync_id):
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("DELETE FROM integracao_landly_sync WHERE id=%s", (sync_id,))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/integracao/landly/municipios')
@login_required
def api_landly_municipios():
    """Retorna municipios do dossie com origem integracao para o accordion."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT dm.municipio, dm.estado, dm.max_legislacoes,
                   bh.concluido_em as ultima_busca,
                   (SELECT status FROM fila_buscas fb
                    WHERE LOWER(fb.municipio)=LOWER(dm.municipio)
                    AND LOWER(fb.estado)=LOWER(dm.estado)
                    AND fb.status IN ('rodando','aguardando')
                    ORDER BY fb.id DESC LIMIT 1) as fila_status
            FROM dossie_municipios dm
            LEFT JOIN buscas_historico bh ON LOWER(bh.municipio)=LOWER(dm.municipio)
                AND LOWER(bh.estado)=LOWER(dm.estado) AND bh.sucesso=true
                AND bh.concluido_em=(SELECT MAX(concluido_em) FROM buscas_historico
                    WHERE municipio=dm.municipio AND estado=dm.estado AND sucesso=true)
            WHERE dm.origem='integracao'
            ORDER BY dm.municipio ASC
        """)
        rows = cur.fetchall()
        cur.close(); conn.close()
        result = []
        for r in rows:
            status = 'none'
            if r['fila_status'] in ('rodando','aguardando'): status = 'pending'
            elif r['ultima_busca']: status = 'ok'
            result.append({
                'municipio': r['municipio'],
                'estado': r['estado'],
                'status': status,
                'ultima_busca': r['ultima_busca'].strftime('%d/%m/%Y') if r['ultima_busca'] else None,
                'max_legislacoes': r['max_legislacoes'],
            })
        return jsonify({'success': True, 'municipios': result})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/integracao/landly/ultimo-log')
@login_required
def api_landly_ultimo_log():
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT status, erro, log_linhas, executado_em FROM integracao_landly_sync ORDER BY executado_em DESC LIMIT 1")
        r = cur.fetchone()
        cur.close(); conn.close()
        if not r: return jsonify({'success': True, 'logs': [], 'status': 'idle'})
        return jsonify({
            'success': True,
            'status': r['status'],
            'logs': r['log_linhas'] or [],
            'executado_em': (r['executado_em'] - __import__('datetime').timedelta(hours=3)).strftime('%d/%m %H:%M') if r['executado_em'] else None,
            '_ts': r['executado_em'].isoformat() if r['executado_em'] else None,
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/integracao/landly/eventos')
@login_required
def api_landly_eventos():
    """SSE endpoint para notificar browser quando sync Landly terminar."""
    import time
    def _stream():
        last = None
        yield 'data: connected\n\n'
        for _ in range(300):  # max 10min
            try:
                flag = '/tmp/landly_sync_done'
                if os.path.exists(flag):
                    ts = os.path.getmtime(flag)
                    if last is None:
                        last = ts
                    elif ts != last:
                        last = ts
                        yield f'data: sync_done\n\n'
                else:
                    last = None
            except: pass
            time.sleep(2)
        yield 'data: timeout\n\n'
    return Response(stream_with_context(_stream()),
                   mimetype='text/event-stream',
                   headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

@app.route('/api/hora-servidor')
def api_hora_servidor():
    from datetime import datetime, timezone, timedelta
    utc_now = datetime.now(timezone.utc)
    brasilia = utc_now - timedelta(hours=3)
    return jsonify({'hora': brasilia.strftime('%H:%M'), 'data': brasilia.strftime('%d/%m/%Y')})

@app.route('/api/fila/reordenar', methods=['POST'])
@login_required
def api_fila_reordenar():
    try:
        ids = request.json.get('ids', [])
        conn = get_db(); cur = conn.cursor()
        for ordem, item_id in enumerate(ids):
            cur.execute("UPDATE fila_buscas SET ordem=%s WHERE id=%s AND status='aguardando'", (ordem, item_id))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/fila/item/<int:item_id>', methods=['DELETE'])
@login_required
def api_fila_item_deletar(item_id):
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("DELETE FROM fila_buscas WHERE id=%s AND status='aguardando'", (item_id,))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/fila/pausar', methods=['POST'])
@login_required
def api_fila_pausar():
    import modulos.fila_worker as _fw
    _fw._fila_pausada = True
    return jsonify({'success': True})

@app.route('/api/fila/despausar', methods=['POST'])
@login_required
def api_fila_despausar():
    import modulos.fila_worker as _fw
    _fw._fila_pausada = False
    return jsonify({'success': True})

@app.route('/api/fila/adicionar', methods=['POST'])
@login_required
def api_fila_adicionar():
    data = request.get_json()
    municipio = (data.get('municipio') or '').strip()
    estado = (data.get('estado') or '').strip()
    max_legs = data.get('max_legislacoes')
    if not municipio or not estado:
        return jsonify({'success':False,'error':'municipio e estado obrigatorios'}),400
    try:
        conn=get_db(); cur=conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT COALESCE(MAX(ordem),0) as max FROM fila_buscas WHERE status IN ('aguardando','rodando')")
        max_ordem=cur.fetchone()['max']
        cur.execute("INSERT INTO fila_buscas (municipio,estado,max_legislacoes,ordem) VALUES (%s,%s,%s,%s) RETURNING id",(municipio,estado,max_legs,max_ordem+1))
        new_id=cur.fetchone()['id']; conn.commit(); cur.close(); conn.close()
        return jsonify({'success':True,'id':new_id})
    except Exception as e:
        return jsonify({'success':False,'error':str(e)}),500

@app.route('/api/fila/listar')
@login_required
def api_fila_listar():
    try:
        conn=get_db(); cur=conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT id,municipio,estado,max_legislacoes,fallback_url_override,status,job_id,criado_em,iniciado_em,concluido_em,erro,origem FROM fila_buscas ORDER BY ordem ASC,criado_em ASC")
        rows=cur.fetchall(); cur.close(); conn.close()
        result=[]
        for r in rows:
            item=dict(r)
            for k in ['criado_em','iniciado_em','concluido_em']:
                if item.get(k): item[k]=item[k].isoformat()
            result.append(item)
        return jsonify({'success':True,'fila':result})
    except Exception as e:
        return jsonify({'success':False,'error':str(e)}),500

@app.route('/api/fila/remover/<int:fila_id>', methods=['DELETE'])
@login_required
def api_fila_remover(fila_id):
    try:
        conn=get_db(); cur=conn.cursor()
        cur.execute("DELETE FROM fila_buscas WHERE id=%s AND status='aguardando'",(fila_id,))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success':True})
    except Exception as e:
        return jsonify({'success':False,'error':str(e)}),500

@app.route('/api/fila/editar/<int:fila_id>', methods=['PATCH'])
@login_required
def api_fila_editar(fila_id):
    data = request.json or {}
    fallback_url = data.get('fallback_url_override') or None
    max_legs = data.get('max_legislacoes')
    if max_legs is not None:
        try: max_legs = int(max_legs)
        except: max_legs = None
    try:
        conn=get_db(); cur=conn.cursor()
        cur.execute("UPDATE fila_buscas SET fallback_url_override=%s, max_legislacoes=%s WHERE id=%s AND status='aguardando'",(fallback_url, max_legs, fila_id))
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success':True})
    except Exception as e:
        return jsonify({'success':False,'error':str(e)}),500

@app.route('/api/fila/limpar', methods=['POST'])
@login_required
def api_fila_limpar():
    try:
        conn=get_db(); cur=conn.cursor()
        cur.execute("DELETE FROM fila_buscas WHERE status='aguardando'")
        conn.commit(); cur.close(); conn.close()
        return jsonify({'success':True})
    except Exception as e:
        return jsonify({'success':False,'error':str(e)}),500

@app.route('/api/fila/cancelar-rodando', methods=['POST'])
@login_required
def api_fila_cancelar_rodando():
    try:
        conn=get_db(); cur=conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT job_id FROM fila_buscas WHERE status='rodando' LIMIT 1")
        row=cur.fetchone()
        if row and row['job_id']:
            job_id = row['job_id']
            if job_id in _buscador_jobs:
                _buscador_jobs[job_id]['done'] = True
        cur.execute("UPDATE fila_buscas SET status='cancelado',concluido_em=NOW() WHERE status='rodando'")
        conn.commit(); cur.close(); conn.close()
        import subprocess; subprocess.Popen(['pkill', '-9', '-f', 'chromium'])
        return jsonify({'success':True})
    except Exception as e:
        return jsonify({'success':False,'error':str(e)}),500

# ═══════════════════════════════════════════════════════════════════════════════
# ROTAS REST — FILA DE EXTRAÇÃO (pipeline_extracao_lei)
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/api/fila_extracao/adicionar', methods=['POST'])
@login_required
def api_fila_extracao_adicionar():
    """Adiciona item na fila_extracao. Body: {zip_path, municipio, estado, ...}"""
    data = request.get_json() or {}
    zip_path = (data.get('zip_path') or '').strip()
    municipio = (data.get('municipio') or '').strip()
    estado = (data.get('estado') or '').strip()
    
    if not zip_path or not municipio or not estado:
        return jsonify({'success': False, 'error': 'zip_path, municipio e estado obrigatorios'}), 400
    
    if not os.path.exists(zip_path):
        return jsonify({'success': False, 'error': f'ZIP nao encontrado: {zip_path}'}), 400
    
    try:
        item_id = _enfileirar_extracao(
            zip_path=zip_path,
            municipio=municipio,
            estado_uf=estado,
            legislacao_id=data.get('legislacao_id'),
            usar_cache=data.get('usar_cache', True),
            consolidar_apos=data.get('consolidar_apos', True),
            ordem=data.get('ordem', 0),
            criado_por=session.get('user_id'),
        )
        if item_id:
            return jsonify({'success': True, 'id': item_id})
        return jsonify({'success': False, 'error': 'falhou ao enfileirar'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/fila_extracao/ultima')
@login_required
def api_fila_extracao_ultima():
    """Retorna a fila mais recente (qualquer status)."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""SELECT id, municipio, estado, status, job_id, progresso_atual,
                              erro_etapa, erro_msg, iniciado_em, concluido_em, criado_em,
                              legislacao_label
                       FROM fila_extracao ORDER BY id DESC LIMIT 1""")
        item = cur.fetchone()
        cur.close(); conn.close()
        if not item:
            return jsonify({'success': True, 'fila': None})
        return jsonify({
            'success': True,
            'fila': {
                'fila_id': item['id'],
                'municipio': item['municipio'],
                'estado': item['estado'],
                'status': item['status'],
                'job_id': item['job_id'],
                'progresso_atual': item['progresso_atual'],
                'erro_etapa': item['erro_etapa'],
                'erro_msg': item['erro_msg'],
                'iniciado_em': item['iniciado_em'].isoformat() if item['iniciado_em'] else None,
                'concluido_em': item['concluido_em'].isoformat() if item['concluido_em'] else None,
                'criado_em': item['criado_em'].isoformat() if item['criado_em'] else None,
                'legislacao_label': item['legislacao_label'],
                'ativa': item['status'] in ('aguardando', 'rodando'),
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:200]}), 500


@app.route('/api/fila_extracao/listar')
@login_required
def api_fila_extracao_listar():
    """Lista itens da fila_extracao com status atual."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT id, municipio, estado, zip_path, legislacao_id,
                   usar_cache, consolidar_apos, status, job_id, ordem,
                   processamento_id, progresso_atual, erro_etapa, erro_msg,
                   criado_em, iniciado_em, concluido_em
            FROM fila_extracao
            ORDER BY ordem ASC, criado_em DESC
            LIMIT 100
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
        result = []
        for r in rows:
            item = dict(r)
            for k in ['criado_em', 'iniciado_em', 'concluido_em']:
                if item.get(k):
                    item[k] = item[k].isoformat()
            result.append(item)
        return jsonify({'success': True, 'fila': result})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500



@app.route('/api/fila_extracao/log/<int:fila_id>')
@login_required
def api_fila_extracao_log(fila_id):
    """
    Retorna logs ao vivo de um item da fila_extracao.
    UI usa pra exibir log de execucao do pipeline em tempo real.
    
    Query: cursor=N (retorna logs com cursor > N)
    Returns: {status, job_id, progresso_atual, done, logs:[{ts, nivel, msg, cursor}]}
    """
    try:
        cursor_after = int(request.args.get('cursor', 0))
    except:
        cursor_after = 0
    
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""SELECT id, municipio, estado, status, job_id, progresso_atual,
                              erro_etapa, erro_msg, iniciado_em, concluido_em
                       FROM fila_extracao WHERE id=%s""", (fila_id,))
        item = cur.fetchone()
        if not item:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': 'item nao encontrado'}), 404
        
        logs = []
        if item['job_id']:
            cur.execute("""SELECT cursor, ts, nivel, msg 
                           FROM buscas_logs 
                           WHERE job_id=%s AND cursor > %s 
                           ORDER BY cursor ASC LIMIT 500""",
                        (item['job_id'], cursor_after))
            for r in cur.fetchall():
                logs.append({
                    'cursor': r['cursor'],
                    'ts': r['ts'],
                    'nivel': r['nivel'],
                    'msg': r['msg'],
                })
        
        cur.close(); conn.close()
        
        done = item['status'] in ('concluido', 'erro', 'cancelado')
        
        return jsonify({
            'success': True,
            'fila_id': fila_id,
            'job_id': item['job_id'],
            'municipio': item['municipio'],
            'estado': item['estado'],
            'status': item['status'],
            'progresso_atual': item['progresso_atual'],
            'erro_etapa': item['erro_etapa'],
            'erro_msg': item['erro_msg'],
            'done': done,
            'logs': logs,
            'next_cursor': logs[-1]['cursor'] if logs else cursor_after,
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:200]}), 500



@app.route('/api/fila_extracao/resultado/<int:fila_id>')
@login_required
def api_fila_extracao_resultado(fila_id):
    """Retorna o resultado_final.json do job + caminho dos arquivos."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT id, municipio, estado, status, job_id FROM fila_extracao WHERE id=%s", (fila_id,))
        item = cur.fetchone()
        cur.close(); conn.close()
        if not item:
            return jsonify({'success': False, 'error': 'item nao encontrado'}), 404
        if item['status'] != 'concluido':
            return jsonify({'success': False, 'error': f'job nao concluido (status={item["status"]})'}), 400
        
        # Determina work_dir do pipeline
        from modulos.pipeline_extracao_lei import _slug_municipio, PIPELINES_BASE_DIR
        slug = _slug_municipio(item['municipio'], item['estado'])
        work_dir = os.path.join(PIPELINES_BASE_DIR, slug)
        
        if not os.path.isdir(work_dir):
            return jsonify({'success': False, 'error': f'work_dir nao existe: {work_dir}'}), 404
        
        # Le resultado_final.json
        resultado_path = os.path.join(work_dir, 'resultado_final.json')
        resultado = None
        if os.path.exists(resultado_path):
            try:
                with open(resultado_path) as f:
                    resultado = json.load(f)
            except Exception as e:
                resultado = {'erro': f'falha lendo: {e}'}
        
        # Lista outros arquivos relevantes
        arquivos = []
        if os.path.isdir(work_dir):
            for f in sorted(os.listdir(work_dir)):
                if f.endswith('.json') or f.endswith('.txt') or f.endswith('.pdf'):
                    full = os.path.join(work_dir, f)
                    arquivos.append({
                        'nome': f,
                        'tamanho': os.path.getsize(full),
                        'url': f.replace('/var/www/urbanlex', '') if '/var/www/urbanlex' in full else '/static/pipelines/' + slug + '/' + f,
                    })
        
        return jsonify({
            'success': True,
            'fila_id': fila_id,
            'municipio': item['municipio'],
            'estado': item['estado'],
            'work_dir_url': '/static/pipelines/' + slug,
            'resultado': resultado,
            'resultado_url': '/static/pipelines/' + slug + '/resultado_final.json' if os.path.exists(resultado_path) else None,
            'arquivos': arquivos,
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:200]}), 500


@app.route('/api/fila_extracao/cancelar/<int:fila_id>', methods=['POST'])
@login_required
def api_fila_extracao_cancelar(fila_id):
    """
    Cancela um item da fila_extracao.
    - Se aguardando: marca como cancelado direto
    - Se rodando: marca como cancelado (worker vai detectar e parar)
    - Se ja terminou: retorna sem fazer nada
    """
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT id, status FROM fila_extracao WHERE id=%s", (fila_id,))
        item = cur.fetchone()
        if not item:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': 'item nao encontrado'}), 404
        
        if item['status'] in ('concluido', 'erro', 'cancelado'):
            cur.close(); conn.close()
            return jsonify({'success': True, 'ja_terminou': True, 'status': item['status']})
        
        cur.execute("""UPDATE fila_extracao 
                       SET status='cancelado', concluido_em=NOW(), 
                           erro_msg='Cancelado pelo operador via UI' 
                       WHERE id=%s""", (fila_id,))
        conn.commit(); cur.close(); conn.close()
        
        return jsonify({'success': True, 'fila_id': fila_id, 'status': 'cancelado'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)[:200]}), 500


@app.route('/api/fila_extracao/<int:item_id>')
@login_required
def api_fila_extracao_detalhe(item_id):
    """Retorna detalhe de 1 item da fila."""
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM fila_extracao WHERE id=%s", (item_id,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        if not row:
            return jsonify({'success': False, 'error': 'nao encontrado'}), 404
        item = dict(row)
        for k in ['criado_em', 'iniciado_em', 'concluido_em']:
            if item.get(k):
                item[k] = item[k].isoformat()
        return jsonify({'success': True, 'item': item})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/fila_extracao/<int:item_id>', methods=['DELETE'])
@login_required
def api_fila_extracao_remover(item_id):
    """Remove item da fila (apenas se status != 'rodando')."""
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT status FROM fila_extracao WHERE id=%s", (item_id,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': 'nao encontrado'}), 404
        if row[0] == 'rodando':
            cur.close(); conn.close()
            return jsonify({'success': False, 'error': 'item esta rodando'}), 400
        cur.execute("DELETE FROM fila_extracao WHERE id=%s", (item_id,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/consolidado/<municipio>/<estado>')
@login_required
def api_consolidado_municipio(municipio, estado):
    """Busca estado consolidado de um municipio."""
    try:
        cons = _buscar_consolidado(municipio, estado)
        if not cons:
            return jsonify({'success': False, 'error': 'nao consolidado'}), 404
        if cons.get('consolidado_em'):
            cons['consolidado_em'] = cons['consolidado_em'].isoformat()
        return jsonify({'success': True, 'consolidado': cons})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/gerador/extrair-dados', methods=['POST'])
@login_required
def api_gerador_extrair_dados():
    """
    Nova rota: substitui /api/gerador/iniciar para usar o pipeline end-to-end.
    Body: compilados=[{municipio,estado,zip_path,...}, ...]
    Enfileira 1 item de extracao por compilado.
    """
    try:
        compilados_raw = request.form.get('compilados') or '[]'
        try:
            compilados = json.loads(compilados_raw)
        except Exception:
            return jsonify({'success': False, 'error': 'compilados JSON invalido'}), 400
        
        if not isinstance(compilados, list) or not compilados:
            return jsonify({'success': False, 'error': 'nenhum compilado selecionado'}), 400
        
        # Toggle: se True, ignora cache MD5 e força reprocessamento
        forcar_repr = (request.form.get('forcar_reprocessamento') or '').strip() in ('1', 'true', 'on', 'yes')
        
        ids_enfileirados = []
        erros = []
        for c in compilados:
            municipio = (c.get('municipio') or '').strip()
            estado = (c.get('estado') or '').strip()
            zip_path = (c.get('zip_path') or '').strip()
            
            if not municipio or not estado or not zip_path:
                erros.append({'municipio': municipio, 'erro': 'campos obrigatorios faltando'})
                continue
            
            if not os.path.exists(zip_path):
                erros.append({'municipio': municipio, 'erro': f'ZIP nao encontrado: {zip_path}'})
                continue
            
            # Resolve legislacao_label a partir do zip_path em gerador_compilados
            leg_label_lookup = None
            try:
                _c2 = get_db().cursor()
                _c2.execute('SELECT legislacao_label FROM gerador_compilados WHERE zip_path=%s LIMIT 1', (zip_path,))
                _r2 = _c2.fetchone()
                if _r2 and _r2[0]: leg_label_lookup = _r2[0]
                _c2.close()
            except Exception as _e_lk: logger.warning('lookup legislacao_label falhou: ' + str(_e_lk))
            item_id = _enfileirar_extracao(
                zip_path=zip_path,
                municipio=municipio,
                estado_uf=estado,
                usar_cache=(not forcar_repr),
                consolidar_apos=True,
                criado_por=session.get('user_id'),
                legislacao_label=leg_label_lookup,
            )
            if item_id:
                ids_enfileirados.append({'municipio': municipio, 'estado': estado, 'fila_id': item_id})
            else:
                erros.append({'municipio': municipio, 'erro': 'falhou ao enfileirar'})
        
        return jsonify({
            'success': True,
            'enfileirados': ids_enfileirados,
            'erros': erros,
            'total_enfileirados': len(ids_enfileirados),
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


_iniciar_fila_worker(app, get_db, _buscador_jobs)
_iniciar_fila_extracao_worker(app, get_db)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.getenv('PORT',5000)), debug=os.getenv('FLASK_ENV')!='production')
